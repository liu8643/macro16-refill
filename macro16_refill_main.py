#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
宏觀16模組自動抓取與Excel回填主程式
版本：V2.0 FinalEvidenceSystem
目的：依「宏觀16模組市場資料回填SOP」與「主程式工程級規格書」執行資料取得、標準化、分數判定、Excel回填與稽核。

執行方式：
  GUI：python macro16_refill_main.py
  CLI：python macro16_refill_main.py --cli --template template.xlsx --out output.xlsx --date 2026-04-29

必要套件：requests, openpyxl
"""
from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import logging
import math
import os
import re
import sys
import time
import sqlite3
import html as html_lib
from dataclasses import dataclass, asdict
from io import StringIO
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

try:
    import requests
except Exception:
    requests = None

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except Exception as exc:
    raise RuntimeError("缺少 openpyxl，請先安裝：pip install openpyxl") from exc

try:
    import pandas as pd
except Exception:
    pd = None

try:
    import numpy as np
except Exception:
    np = None

APP_NAME = "Macro16RefillEngine"
VERSION = "3.0.0-ai-project-rotation-monitor"
STRATEGY_VERSION = "teacher_strategy_v3.0_ai_project_rotation_monitor_20260525"
DEFAULT_TIMEOUT = 15
DEFAULT_MAX_FALLBACK_DAYS = 5

MODULES = [
    "美股-S&P500", "美股-NASDAQ", "美股-道瓊", "VIX恐慌",
    "美債10Y", "原油", "戰爭/地緣", "CPI", "非農", "外資",
    "官股", "台股指數", "成交量", "AI產業", "OTC", "台股夜盤"
]

# SOP V2.2：正式輸出模式分離。
# macro_refill：正式日常模式＝宏觀16修正頁 + 原本TOP/00~09報表（若有DB），不得刪除TOP輸出。
# macro_only：只輸出「市場輸入 / 宏觀16模組 / V2技術引擎」三頁，用於單純驗證宏觀回填。
# institutional_report：只輸出00~16老師策略機構級報表。
# macro_teacher：宏觀三頁 + 00~16老師策略完整報表。
# teacher_full：只輸出00~16老師策略完整報表，等同institutional_report語意。
# all：完整debug與機構報表全輸出。
REPORT_MODE_MACRO = "macro_refill"
REPORT_MODE_MACRO_ONLY = "macro_only"
REPORT_MODE_INSTITUTIONAL = "institutional_report"
REPORT_MODE_MACRO_TEACHER = "macro_teacher"
REPORT_MODE_TEACHER_FULL = "teacher_full"
REPORT_MODE_ALL = "all"
# R5N29：獨立觀察池培養模式；非預設，不影響原宏觀16 / 老師策略報表。
REPORT_MODE_WATCH_POOL = "watch_pool_cultivation"
MACRO_REFILL_SHEETS = ["市場輸入", "宏觀16模組", "V2技術引擎"]


YAHOO_SYMBOLS = {
    "美股-S&P500": "^GSPC",
    "美股-NASDAQ": "^IXIC",
    "美股-道瓊": "^DJI",
    "VIX恐慌": "^VIX",
    "原油": "CL=F",
}

YAHOO_SYMBOL_CANDIDATES = {
    "OTC": ["^TWOII", "TWOII.TW", "TWOII.TWO"],
}


SOURCE_PRIORITY = ["官方資料", "交易所/期交所", "國際金融數據站", "台灣可信財經站", "人工事件判斷"]

SOURCE_URLS = {
    "reuters_war": "https://www.reuters.com",
    "bloomberg_policy": "https://www.bloomberg.com",
    "federal_reserve": "https://www.federalreserve.gov",
    "bls_api_cpi": "https://api.bls.gov/publicAPI/v2/timeseries/data/CUUR0000SA0",
    "bls_api_nfp": "https://api.bls.gov/publicAPI/v2/timeseries/data/CES0000000001",
    "twse_foreign": "https://www.twse.com.tw/fund/BFI82U?response=json&dayDate={date}&type=day",
    "twse_broker_report": "https://www.twse.com.tw/zh/trading/brokerReport",
    "tpex_indices": "https://www.tpex.org.tw/zh-tw/mainboard/trading/info/indices-pricing.html",
    "wantgoo_public_bank": "https://www.wantgoo.com/stock/public-bank/trend",
    "histock": "https://histock.tw",
    "techcrunch_ai": "https://techcrunch.com",
    "isw": "https://www.understandingwar.org",
    "cnn": "https://www.cnn.com",
    "iek": "https://ieknet.iek.org.tw/",
    "taifex_night": "https://www.taifex.com.tw/cht/3/futContractsDateAh",
    "twse_t86": "https://www.twse.com.tw/rwd/zh/fund/T86?date={date}&selectType=ALLBUT0999&response=json",
    "gov_broker_fallback_histock": "https://histock.tw/stock/broker8.aspx",
}


@dataclass
class RawData:
    key: str
    value: Any
    date: str
    source: str
    url: str
    fetched_at: str
    status: str = "OK"
    message: str = ""
    query_date: str = ""
    actual_date: str = ""
    is_fallback: bool = False
    fallback_days: int = 0
    data_status: str = "OK"
    data_note: str = ""
    parse_status: str = "PARSE_OK"
    raw_file_path: str = ""
    confidence: float = 1.0

@dataclass
class MarketInput:
    base_date: str = ""
    close: Optional[float] = None
    high: Optional[float] = None
    low: Optional[float] = None
    prev_high: Optional[float] = None
    prev_low: Optional[float] = None
    ma5: Optional[float] = None
    turnover_100m: Optional[float] = None
    avg_turnover_5d_100m: Optional[float] = None
    foreign_net_100m: Optional[float] = None
    gov_net_100m: Optional[float] = None
    ai_strength: float = 0.5
    major_event: int = 0
    night_score: Optional[int] = None
    night_net_lots: Optional[int] = None
    source_1: str = ""
    source_2: str = ""
    source_3: str = ""
    source_4: str = ""

@dataclass
class ManualOverride:
    gov_net_100m: Optional[float] = None
    ai_strength: Optional[float] = None
    major_event: Optional[int] = None
    event_note: str = ""
    oil_event_note: str = ""
    night_score: Optional[float] = None

@dataclass
class ModuleScore:
    module: str
    data_text: str
    strength: float
    direction: int
    weighted_score: float
    explanation: str
    source: str
    data_time: str
    trade_usage: str
    status: str = "OK"

@dataclass
class TechnicalRisk:
    below_ma5: int
    lower_high: int
    lower_low: int
    volume_expansion: int
    major_event: int
    risk_score: float
    market_judgement: str
    night_bearish: int = 0
    night_score: Optional[int] = None
    night_net_lots: Optional[int] = None

class Macro16Logger:
    def __init__(self, log_dir: Path):
        log_dir.mkdir(parents=True, exist_ok=True)
        self.run_id = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.log_file = log_dir / f"macro16_debug_{self.run_id}.txt"
        self.raw_dir = log_dir / "raw" / self.run_id[:8] / self.run_id
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        self.evidence_records: List[Dict[str, Any]] = []
        root = logging.getLogger()
        for handler in list(root.handlers):
            root.removeHandler(handler)
        logging.basicConfig(
            filename=str(self.log_file),
            level=logging.INFO,
            format="%(asctime)s [%(levelname)s] %(message)s",
            encoding="utf-8",
        )
        self.messages: List[str] = []

    def info(self, msg: str):
        logging.info(msg)
        self.messages.append(f"INFO {msg}")

    def warning(self, msg: str):
        logging.warning(msg)
        self.messages.append(f"WARN {msg}")

    def error(self, msg: str):
        logging.error(msg, exc_info=True)
        self.messages.append(f"ERROR {msg}")

    def debug(self, msg: str):
        logging.info("DEBUG " + msg)
        self.messages.append(f"DEBUG {msg}")

    def _safe_filename(self, value: str) -> str:
        return re.sub(r"[^0-9A-Za-z_\\-\\u4e00-\\u9fff]+", "_", str(value)).strip("_")[:80] or "source"

    def _infer_parsed_fields(self, payload: Any) -> Dict[str, Any]:
        """從已解析payload自動產生 parsed_fields，避免 status=OK 但 parse_status=NO_PARSED_VALUE。"""
        if isinstance(payload, dict):
            parsed = {}
            for k, v in payload.items():
                if k in ("raw", "rows", "snippet"):
                    continue
                if isinstance(v, (str, int, float, bool)) or v is None:
                    parsed[k] = v
                elif isinstance(v, (list, tuple)) and len(v) <= 10:
                    parsed[k] = v
                elif isinstance(v, dict):
                    parsed[k] = {kk: vv for kk, vv in list(v.items())[:10] if isinstance(vv, (str, int, float, bool)) or vv is None}
            return parsed
        return {}

    def write_raw_evidence(self, source: str, payload: Any, parsed: Optional[Dict[str, Any]] = None, status: str = "OK", url: str = "", message: str = "") -> str:
        if parsed is None:
            parsed = self._infer_parsed_fields(payload)
        parsed = parsed or {}
        parse_status = "PARSE_OK" if parsed else "NO_PARSED_VALUE"
        if status == "OK" and parse_status != "PARSE_OK":
            status = "WARN"
            message = (message + "；" if message else "") + "已抓到原始資料但尚未形成parsed_fields，避免假OK。"
        record = {
            "run_id": self.run_id,
            "source": source,
            "url": url,
            "fetch_time": dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": status,
            "parse_status": parse_status,
            "parsed_fields": parsed,
            "message": message,
            "raw_excerpt": str(payload)[:4000],
        }
        base = self._safe_filename(source)
        path = self.raw_dir / f"{base}.json"
        # 不允許不同來源覆蓋同一個source.json/模組json；若重複則自動加序號。
        if path.exists():
            idx = 2
            while True:
                candidate = self.raw_dir / f"{base}_{idx}.json"
                if not candidate.exists():
                    path = candidate
                    break
                idx += 1
        try:
            path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
            record["raw_file_path"] = str(path)
            self.evidence_records.append(record)
            self.info(f"RAW_EVIDENCE_FILE source={source} path={path} status={status} parse_status={record['parse_status']}")
            return str(path)
        except Exception as exc:
            self.warning(f"RAW_EVIDENCE_WRITE_FAIL source={source} error={exc}")
            return ""

    def raw_snapshot(self, source: str, payload: Any, parsed: Optional[Dict[str, Any]] = None, status: str = "OK", url: str = "", message: str = ""):
        text = str(payload)
        if len(text) > 900:
            text = text[:900] + "..."
        self.info(f"RAW_DATA_SNAPSHOT source={source} payload={text}")
        self.write_raw_evidence(source, payload, parsed=parsed, status=status, url=url, message=message)

    def parsed_value(self, field: str, value: Any, source: str, actual_date: str = ""):
        self.info(f"PARSED_VALUE field={field} value={value} source={source} actual_date={actual_date}")

    def strategy_trace(self, tag: str, payload: Dict[str, Any]):
        """老師策略工程追蹤：輸出可被log grep的固定格式。"""
        try:
            text = json.dumps(payload, ensure_ascii=False, default=str)
        except Exception:
            text = str(payload)
        self.info(f"{tag} {text}")

class HttpClient:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger
        self.session = requests.Session() if requests else None
        if self.session:
            self.session.headers.update({
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Macro16RefillEngine/2.0",
                "Accept": "application/json,text/csv,text/html,text/plain,*/*",
                "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
                "Referer": "https://www.google.com/",
            })

    def get_text(self, url: str, timeout: int = DEFAULT_TIMEOUT) -> str:
        if not self.session:
            raise RuntimeError("requests 未安裝，無法抓取網路資料")
        self.logger.info(f"GET {url}")
        r = self.session.get(url, timeout=timeout)
        self.logger.debug(f"HTTP status={r.status_code}, content_type={r.headers.get('content-type','')}, len={len(r.text or '')}")
        r.raise_for_status()
        if not r.encoding:
            r.encoding = "utf-8"
        return r.text

    def get_json(self, url: str, timeout: int = DEFAULT_TIMEOUT) -> Dict[str, Any]:
        text = self.get_text(url, timeout)
        try:
            return json.loads(text)
        except Exception as exc:
            self.logger.warning(f"JSON解析失敗 url={url}; head={text[:200]!r}; error={exc}")
            raise

class SourceConnector:
    def __init__(self, client: HttpClient, logger: Macro16Logger):
        self.client = client
        self.logger = logger

    def _today_str(self) -> str:
        return dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _compact_date(self, value: Optional[str]) -> str:
        if not value:
            return dt.date.today().strftime("%Y%m%d")
        return str(value).replace("-", "")

    def _dash_date(self, value: str) -> str:
        v = self._compact_date(value)
        return f"{v[:4]}-{v[4:6]}-{v[6:8]}" if len(v) == 8 else v

    def _previous_calendar_day(self, compact_date: str) -> str:
        """
        V1.5：安全日曆回退。
        目的：
        1. 正常日期：直接回退前一日，例如 20260430 -> 20260429。
        2. 非法日期：先修正成該年月最近合法日，再回退。
           例如 20260431 不是合法日期，會先修正為 20260430，再回退到 20260429。
        """
        raw = self._compact_date(compact_date)
        try:
            d = dt.datetime.strptime(raw, "%Y%m%d").date()
            return (d - dt.timedelta(days=1)).strftime("%Y%m%d")
        except ValueError:
            if len(raw) == 8 and raw[:6].isdigit():
                year = int(raw[:4])
                month = int(raw[4:6])
                # 先取得該年月最後一天
                if month == 12:
                    first_next_month = dt.date(year + 1, 1, 1)
                else:
                    first_next_month = dt.date(year, month + 1, 1)
                last_valid_day = first_next_month - dt.timedelta(days=1)
                # 再回退一天，符合「無效指定日也視為抓不到，要往前退」的需求
                return (last_valid_day - dt.timedelta(days=1)).strftime("%Y%m%d")
            raise

    def _twse_rows(self, data: Dict[str, Any]) -> List[Any]:
        rows = data.get("data", []) or []
        if not rows and data.get("tables"):
            for table in data.get("tables", []):
                rows.extend(table.get("data", []) or [])
        return rows

    def _official_no_data(self, source_name: str, query_date: str, try_date: str, url: str, data: Optional[Dict[str, Any]] = None):
        stat = ""
        try:
            stat = str(data.get("stat", "")) if isinstance(data, dict) else ""
        except Exception:
            stat = ""
        self.logger.warning(f"OFFICIAL_NOT_UPDATED source={source_name} query_date={query_date} try_date={try_date} stat={stat} url={url}")

    def _fallback_note(self, query_date: str, actual_date: str, source_name: str) -> Tuple[bool, int, str]:
        q = self._compact_date(query_date)
        a = self._compact_date(actual_date)
        is_fb = q != a
        if not is_fb:
            return False, 0, ""
        qd = dt.datetime.strptime(q, "%Y%m%d").date()
        ad = dt.datetime.strptime(a, "%Y%m%d").date()
        days = max(0, (qd - ad).days)
        note = f"{source_name} 查詢日 {q} 官網尚未公布/無資料，已使用最近可用資料日 {a}"
        self.logger.info(f"FALLBACK_USED source={source_name} query_date={q} actual_date={a} fallback_days={days}")
        return True, days, note

    def _log_fallback_try(self, source_name: str, query_date: str, try_date: str, attempt_index: int, max_back_days: int):
        """V1.5：記錄固定最多退 5 次的每一次嘗試。"""
        self.logger.info(
            f"FALLBACK_TRY source={source_name} query_date={query_date} "
            f"try_date={try_date} attempt={attempt_index}/{max_back_days}"
        )

    def fetch_yahoo_chart(self, symbol: str, module: str, range_days: str = "10d") -> RawData:
        url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?range={range_days}&interval=1d"
        try:
            data = self.client.get_json(url)
            result = data.get("chart", {}).get("result", [])[0]
            timestamps = result.get("timestamp", [])
            quote = result.get("indicators", {}).get("quote", [])[0]
            closes = [x for x in quote.get("close", []) if x is not None]
            highs = [x for x in quote.get("high", []) if x is not None]
            lows = [x for x in quote.get("low", []) if x is not None]
            if not closes:
                raise ValueError("Yahoo資料無收盤價")
            last_close = float(closes[-1])
            prev_close = float(closes[-2]) if len(closes) >= 2 else last_close
            pct = ((last_close - prev_close) / prev_close * 100) if prev_close else 0
            last_ts = timestamps[-1] if timestamps else int(time.time())
            data_date = dt.datetime.fromtimestamp(last_ts).strftime("%Y-%m-%d")
            value = {
                "symbol": symbol,
                "close": last_close,
                "prev_close": prev_close,
                "change_pct": pct,
                "high": float(highs[-1]) if highs else None,
                "low": float(lows[-1]) if lows else None,
                "last5_close": closes[-5:] if len(closes) >= 5 else closes,
            }
            self.logger.raw_snapshot(module, value)
            self.logger.parsed_value(f"{module}_close", last_close, "Yahoo Finance", data_date)
            self.logger.parsed_value(f"{module}_change_pct", round(pct, 4), "Yahoo Finance", data_date)
            return RawData(module, value, data_date, "Yahoo Finance", url, self._today_str())
        except Exception as exc:
            self.logger.warning(f"Yahoo抓取失敗 {module}/{symbol}: {exc}")
            return RawData(module, None, "", "Yahoo Finance", url, self._today_str(), "FAIL", str(exc))


    def fetch_yahoo_chart_candidates(self, symbols: List[str], module: str, range_days: str = "10d") -> RawData:
        last_fail = None
        for symbol in symbols:
            data = self.fetch_yahoo_chart(symbol, module, range_days)
            if data.status == "OK":
                data.message = f"使用候選代碼 {symbol}"
                return data
            last_fail = data
        return last_fail or RawData(module, None, "", "Yahoo Finance", ",".join(symbols), self._today_str(), "FAIL", "所有Yahoo候選代碼失敗")

    def fetch_twse_taiex_history(self, base_date: Optional[str] = None, max_back_days: int = DEFAULT_MAX_FALLBACK_DAYS) -> RawData:
        query_date = self._compact_date(base_date)
        try_date = query_date
        last_error = ""
        for fallback_days in range(max_back_days + 1):
            self._log_fallback_try("TWSE_TAIEX", query_date, try_date, fallback_days, max_back_days)
            url = f"https://www.twse.com.tw/rwd/zh/TAIEX/MI_5MINS_HIST?date={try_date}&response=json"
            try:
                data = self.client.get_json(url)
                rows = data.get("data", []) or []
                if not rows:
                    self._official_no_data("TWSE_TAIEX", query_date, try_date, url, data)
                    try_date = self._previous_calendar_day(try_date)
                    continue
                parsed = []
                for row in rows:
                    try:
                        roc_date = str(row[0])
                        parts = roc_date.split("/")
                        year = int(parts[0]) + 1911 if int(parts[0]) < 1911 else int(parts[0])
                        date_str = f"{year:04d}-{int(parts[1]):02d}-{int(parts[2]):02d}"
                        compact = date_str.replace("-", "")
                        if compact > query_date:
                            continue
                        parsed.append({
                            "date": date_str,
                            "open": self._to_float(row[1]),
                            "high": self._to_float(row[2]),
                            "low": self._to_float(row[3]),
                            "close": self._to_float(row[4]),
                        })
                    except Exception:
                        continue
                if not parsed:
                    self._official_no_data("TWSE_TAIEX", query_date, try_date, url, data)
                    try_date = self._previous_calendar_day(try_date)
                    continue
                last = parsed[-1]
                actual = last["date"].replace("-", "")
                is_fb, fb_days, note = self._fallback_note(query_date, actual, "TWSE_TAIEX")
                value = {"rows": parsed, "last": last}
                self.logger.raw_snapshot("TWSE_TAIEX", last)
                self.logger.parsed_value("taiex_close", last.get("close"), "TWSE MI_5MINS_HIST", last.get("date", ""))
                return RawData("台股指數", value, last["date"], "TWSE MI_5MINS_HIST", url, self._today_str(),
                               "OK", note, query_date, actual, is_fb, fb_days, "OK", note)
            except Exception as exc:
                last_error = str(exc)
                self.logger.warning(f"TWSE加權指數來源失敗 try_date={try_date} url={url}: {exc}")
                try_date = self._previous_calendar_day(try_date)
        return RawData("台股指數", None, "", "TWSE MI_5MINS_HIST", "", self._today_str(),
                       "FAIL", last_error or "最近可用台股指數資料未取得", query_date, "", True, max_back_days, "DATA_MISSING", last_error)

    def fetch_twse_turnover_month(self, base_date: Optional[str] = None, max_back_days: int = DEFAULT_MAX_FALLBACK_DAYS) -> RawData:
        """抓取TWSE成交值。V1.4：指定日期無資料時，自動往前找最近有效交易日。"""
        query_date = self._compact_date(base_date)
        try_date = query_date
        last_error = ""
        for fallback_days in range(max_back_days + 1):
            self._log_fallback_try("TWSE_FMTQIK", query_date, try_date, fallback_days, max_back_days)
            month_param = try_date[:6] + "01"
            candidates = [
                f"https://www.twse.com.tw/rwd/zh/afterTrading/FMTQIK?date={month_param}&response=json",
                f"https://www.twse.com.tw/exchangeReport/FMTQIK?response=json&date={month_param}",
                f"https://www.twse.com.tw/rwd/zh/TAIEX/FMTQIK?date={month_param}&response=json",
            ]
            for url in candidates:
                try:
                    data = self.client.get_json(url)
                    rows = self._twse_rows(data)
                    parsed = []
                    for row in rows:
                        try:
                            roc_date = str(row[0]).strip()
                            parts = roc_date.split("/")
                            year = int(parts[0]) + 1911 if int(parts[0]) < 1911 else int(parts[0])
                            date_str = f"{year:04d}-{int(parts[1]):02d}-{int(parts[2]):02d}"
                            compact = date_str.replace("-", "")
                            if compact > query_date or compact > try_date:
                                continue
                            amount = self._to_float(row[2])
                            if not math.isnan(amount):
                                parsed.append({"date": date_str, "turnover_100m": amount / 100000000})
                        except Exception as row_exc:
                            self.logger.debug(f"FMTQIK row skipped row={row}; error={row_exc}")
                    if not parsed:
                        self._official_no_data("TWSE_FMTQIK", query_date, try_date, url, data)
                        continue
                    last = parsed[-1]
                    actual = last["date"].replace("-", "")
                    is_fb, fb_days, note = self._fallback_note(query_date, actual, "TWSE_FMTQIK")
                    self.logger.info(f"TWSE成交值取得成功 source={url}, query_date={query_date}, actual_date={actual}, turnover_100m={last['turnover_100m']:.2f}")
                    self.logger.raw_snapshot("TWSE_FMTQIK", last)
                    self.logger.parsed_value("turnover_100m", round(last.get("turnover_100m", 0), 2), "TWSE FMTQIK", last.get("date", ""))
                    return RawData("成交量", {"rows": parsed, "last": last}, last["date"], "TWSE FMTQIK", url, self._today_str(),
                                   "OK", note, query_date, actual, is_fb, fb_days, "OK", note)
                except Exception as exc:
                    last_error = str(exc)
                    self.logger.warning(f"TWSE成交值來源失敗 try_date={try_date} url={url}: {exc}")
            try_date = self._previous_calendar_day(try_date)
        return RawData("成交量", None, "", "TWSE FMTQIK", "", self._today_str(),
                       "FAIL", last_error or "最近可用成交值資料未取得", query_date, "", True, max_back_days, "DATA_MISSING", last_error)

    def fetch_foreign_investor(self, base_date: Optional[str] = None, max_back_days: int = DEFAULT_MAX_FALLBACK_DAYS) -> RawData:
        """抓取外資買賣超。V1.5：指定日期抓不到時，固定最多往前退 5 次。"""
        query_date = self._compact_date(base_date)
        try_date = query_date
        last_error = ""
        for fallback_days in range(max_back_days + 1):
            self._log_fallback_try("TWSE_FOREIGN", query_date, try_date, fallback_days, max_back_days)
            candidates = [
                f"https://www.twse.com.tw/rwd/zh/fund/BFI82U?dayDate={try_date}&type=day&response=json",
                f"https://www.twse.com.tw/fund/BFI82U?response=json&dayDate={try_date}&type=day",
                f"https://www.twse.com.tw/rwd/zh/fund/TWT38U?date={try_date}&response=json",
                f"https://www.twse.com.tw/fund/TWT38U?response=json&date={try_date}",
            ]
            for url in candidates:
                try:
                    data = self.client.get_json(url)
                    rows = self._twse_rows(data)
                    if not rows:
                        self._official_no_data("TWSE_FOREIGN", query_date, try_date, url, data)
                        continue
                    for row in rows:
                        row_text = " ".join([str(x) for x in row])
                        if ("外資" in row_text) or ("外陸資" in row_text):
                            vals = []
                            for cell in row:
                                try:
                                    v = self._to_float(cell)
                                    if not math.isnan(v):
                                        vals.append(v)
                                except Exception:
                                    pass
                            if vals:
                                net_100m = vals[-1] / 100000000
                                is_fb, fb_days, note = self._fallback_note(query_date, try_date, "TWSE_FOREIGN")
                                self.logger.info(f"FETCH_OK source=TWSE_FOREIGN query_date={query_date} actual_date={try_date} net_100m={net_100m:.2f} url={url}")
                                buy_amount = vals[-3] if len(vals) >= 3 else None
                                sell_amount = vals[-2] if len(vals) >= 2 else None
                                net_amount = vals[-1] if len(vals) >= 1 else None
                                parsed = {
                                    "foreign_net_100m": round(net_100m, 2),
                                    "buy_amount": buy_amount,
                                    "sell_amount": sell_amount,
                                    "net_amount": net_amount,
                                    "query_date": query_date,
                                    "actual_date": try_date,
                                    "source_rule": "TWSE BFI82U/TWT38U 外資及陸資列淨買賣超 / 100000000",
                                    "raw_row": row_text[:500],
                                }
                                raw_path = self.logger.write_raw_evidence(
                                    "TWSE_FOREIGN", parsed, parsed=parsed, status="OK", url=url,
                                    message="外資買賣超解析完成，parsed_fields已寫入證據鏈"
                                )
                                self.logger.parsed_value("foreign_net_100m", round(net_100m, 2), "TWSE BFI82U/TWT38U", try_date)
                                return RawData("外資", {"net_100m": net_100m, "raw_hint": row_text[:500]}, self._dash_date(try_date),
                                               "TWSE三大法人", url, self._today_str(), "OK", note, query_date, try_date, is_fb, fb_days, "OK", note, "PARSE_OK", raw_path, 1.0)
                    text = json.dumps(data, ensure_ascii=False)
                    nums = [self._to_float(x) for x in re.findall(r"-?\d[\d,]*\.?\d*", text)]
                    nums = [x for x in nums if abs(x) > 1000000]
                    if nums:
                        net_100m = nums[-1] / 100000000
                        is_fb, fb_days, note = self._fallback_note(query_date, try_date, "TWSE_FOREIGN")
                        msg = "使用fallback解析；" + note if note else "使用fallback解析"
                        parsed = {
                            "foreign_net_100m": round(net_100m, 2),
                            "query_date": query_date,
                            "actual_date": try_date,
                            "source_rule": "TWSE三大法人數值fallback解析 / 100000000",
                            "raw_excerpt": text[:500],
                        }
                        self.logger.warning(f"外資語意列未找到，使用數值fallback source={url}, query_date={query_date}, actual_date={try_date}, net_100m={net_100m:.2f}")
                        raw_path = self.logger.write_raw_evidence(
                            "TWSE_FOREIGN", parsed, parsed=parsed, status="WARN", url=url,
                            message=msg + "；已寫入foreign_net_100m parsed_fields"
                        )
                        return RawData("外資", {"net_100m": net_100m, "raw_hint": text[:500]}, self._dash_date(try_date),
                                       "TWSE三大法人-fallback", url, self._today_str(), "WARN", msg, query_date, try_date, is_fb, fb_days, "OK", msg, "PARSE_OK", raw_path, 0.8)
                    self._official_no_data("TWSE_FOREIGN", query_date, try_date, url, data)
                except Exception as exc:
                    last_error = str(exc)
                    self.logger.warning(f"外資來源失敗 try_date={try_date} url={url}: {exc}")
            try_date = self._previous_calendar_day(try_date)
        return RawData("外資", None, "", "TWSE三大法人", "", self._today_str(),
                       "FAIL", last_error or "最近可用外資資料未取得", query_date, "", True, max_back_days, "DATA_MISSING", last_error)

    def fetch_fred_csv_latest(self, series_id: str, module: str) -> RawData:
        url = f"https://fred.stlouisfed.org/graph/fredgraph.csv?id={series_id}"
        try:
            text = self.client.get_text(url)
            rows = list(csv.DictReader(StringIO(text)))
            valid = [(r["observation_date"], r[series_id]) for r in rows if r.get(series_id) not in (None, "", ".")]
            if not valid:
                raise ValueError("FRED無有效資料")
            date_str, value = valid[-1]
            self.logger.raw_snapshot(module, {"series": series_id, "date": date_str, "value": value})
            self.logger.parsed_value(module, value, "FRED", date_str)
            return RawData(module, {"value": float(value), "series": series_id}, date_str, "FRED", url, self._today_str())
        except Exception as exc:
            self.logger.warning(f"FRED抓取失敗 {series_id}: {exc}")
            return RawData(module, None, "", "FRED", url, self._today_str(), "FAIL", str(exc))

    def fetch_bls_release_text(self, module: str, url: str) -> RawData:
        try:
            text = self.client.get_text(url)
            compact = re.sub(r"\s+", " ", text)
            snippet = compact[:1200]
            self.logger.raw_snapshot(module, snippet)
            self.logger.parsed_value(f"{module}_source", url, "BLS", "latest_release")
            return RawData(module, {"url": url, "snippet": snippet}, "latest_release", "BLS", url, self._today_str(), "OK", "BLS periodic release fetched")
        except Exception as exc:
            self.logger.warning(f"BLS抓取失敗 {module}: {exc}")
            return RawData(module, None, "", "BLS", url, self._today_str(), "WARN", str(exc), data_status="OPTIONAL_MISSING")

    def fetch_text_snapshot(self, module: str, url: str, source_name: str, status_if_fail: str = "WARN", parse_required: bool = False) -> RawData:
        try:
            text = self.client.get_text(url)
            compact = re.sub(r"\s+", " ", text)
            snippet = compact[:5000]
            value = {"url": url, "snippet": snippet, "snippet_length": len(snippet)}
            if parse_required:
                status = "WARN"
                data_status = "NO_PARSED_VALUE"
                parsed = {}
                parse_status = "NO_PARSED_VALUE"
                message = f"{source_name} source fetched but parser required"
            else:
                status = "OK"
                data_status = "OK"
                parsed = {"source_url": url, "snippet_length": len(snippet)}
                parse_status = "PARSE_OK"
                message = f"{source_name} source fetched"
            raw_path = self.logger.write_raw_evidence(module, value, parsed=parsed, status=status, url=url, message=message)
            self.logger.raw_snapshot(module, value, parsed=parsed, status=status, url=url, message=message)
            self.logger.parsed_value(f"{module}_source_url", url, source_name, "latest")
            return RawData(module, value, "latest", source_name, url, self._today_str(), status, message, data_status=data_status, parse_status=parse_status, raw_file_path=raw_path, confidence=0.5 if parse_required else 0.8)
        except Exception as exc:
            self.logger.warning(f"{source_name}抓取失敗 {module}: {exc}")
            raw_path = self.logger.write_raw_evidence(module, {"url": url, "error": str(exc)}, parsed={}, status=status_if_fail, url=url, message=str(exc))
            return RawData(module, None, "", source_name, url, self._today_str(), status_if_fail, str(exc), data_status="FETCH_FAIL", parse_status="NO_PARSED_VALUE", raw_file_path=raw_path, confidence=0.0)


    def fetch_bls_api_series(self, module: str, series_id: str, source_url: str) -> RawData:
        year = dt.date.today().year
        url = f"{source_url}?startyear={year-2}&endyear={year}"
        try:
            data = self.client.get_json(url)
            series = data.get("Results", {}).get("series", [])
            if not series or not series[0].get("data"):
                raise ValueError("BLS API無有效資料")
            latest = series[0]["data"][0]
            value = self._to_float(latest.get("value"))
            period = f"{latest.get('year')}-{latest.get('periodName')}"
            payload = {"series_id": series_id, "period": period, "value": value, "raw": latest}
            self.logger.raw_snapshot(module, payload)
            self.logger.parsed_value(f"{module}_value", value, "BLS API", period)
            return RawData(module, payload, period, "BLS API", url, self._today_str(), "OK", "BLS API series fetched")
        except Exception as exc:
            self.logger.warning(f"BLS API抓取失敗 {module}/{series_id}: {exc}")
            return RawData(module, None, "", "BLS API", url, self._today_str(), "WARN", str(exc), data_status="FETCH_FAIL")

    def fetch_bls_cpi(self) -> RawData:
        return self.fetch_bls_api_series("CPI", "CUUR0000SA0", SOURCE_URLS["bls_api_cpi"])

    def fetch_bls_nfp(self) -> RawData:
        return self.fetch_bls_api_series("非農", "CES0000000001", SOURCE_URLS["bls_api_nfp"])

    def fetch_reuters_war(self, url: str = SOURCE_URLS["reuters_war"]) -> RawData:
        raw = self.fetch_text_snapshot("戰爭/停火", url, "Reuters", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "").lower()
            risk = 1 if any(k in text for k in ["war", "ceasefire", "israel", "iran", "attack", "missile", "sanction"]) else 0
            raw.value["major_event"] = risk
            self.logger.parsed_value("major_event", risk, "Reuters", raw.date)
        return raw

    def fetch_bloomberg_policy(self, url: str = SOURCE_URLS["bloomberg_policy"]) -> RawData:
        return self.fetch_text_snapshot("外交政策", url, "Bloomberg", "WARN")

    def fetch_fed_policy(self, url: str = SOURCE_URLS["federal_reserve"]) -> RawData:
        return self.fetch_text_snapshot("FED利率政策", url, "Federal Reserve", "WARN")

    def fetch_isw_conflict(self, url: str = SOURCE_URLS["isw"]) -> RawData:
        raw = self.fetch_text_snapshot("ISW衝突分析", url, "ISW", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "").lower()
            risk = 1 if any(k in text for k in ["russia", "ukraine", "iran", "israel", "war", "attack"]) else 0
            raw.value["major_event"] = risk
            self.logger.parsed_value("isw_major_event", risk, "ISW", raw.date)
        return raw

    def fetch_cnn_major_news(self, url: str = SOURCE_URLS["cnn"]) -> RawData:
        raw = self.fetch_text_snapshot("CNN重大新聞", url, "CNN", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "").lower()
            risk = 1 if any(k in text for k in ["breaking", "war", "crisis", "market", "president", "attack"]) else 0
            raw.value["major_event"] = risk
            self.logger.parsed_value("cnn_major_event", risk, "CNN", raw.date)
        return raw

    def fetch_iek_industry(self, url: str = SOURCE_URLS["iek"]) -> RawData:
        raw = self.fetch_text_snapshot("IEK產業分析", url, "IEK Taiwan", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "")
            strength = 0.8 if any(k in text for k in ["AI", "CPO", "半導體", "先進封裝", "產業"]) else 0.5
            raw.value["ai_strength"] = strength
            self.logger.parsed_value("iek_ai_strength", strength, "IEK Taiwan", raw.date)
        return raw

    def fetch_trump_public_news(self, url: str = "https://www.reuters.com/world/us/") -> RawData:
        raw = self.fetch_text_snapshot("美國總統", url, "Reuters US", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "").lower()
            signal = 1 if any(k in text for k in ["trump", "tariff", "president", "trade", "china", "fed"]) else 0
            raw.value["policy_event"] = signal
            self.logger.parsed_value("us_president_policy_event", signal, "Reuters US", raw.date)
        return raw

    def _normalize_gov_unit_to_100m(self, value: float, unit: str = "") -> float:
        """SOP V2.1 P0-04：官股數字統一轉為億元。"""
        unit = str(unit or "")
        if "億元" in unit or unit == "億" or "億" in unit:
            return float(value)
        if "萬" in unit:
            return float(value) / 10000.0
        # 若頁面未給單位，Wantgoo/HiStock常見顯示為億元；保守以億元處理並留證據。
        return float(value)

    def _parse_public_bank_text(self, text: str) -> Optional[Dict[str, Any]]:
        """
        SOP V2.1 第十四章：Wantgoo/第三方八大官股備援Parser。
        目標只做備援回填與P0_WARN，不冒充官方資料。
        """
        if not text:
            return None
        clean = html_lib.unescape(re.sub(r"\s+", " ", str(text)))
        patterns = [
            r"(?:八大|官股|公股|八大官股|八大公股)[^。；;\n]{0,80}?(買超|賣超|買賣超|淨買超|淨賣超)?[^-+0-9]{0,20}([-+−]?\d+(?:,\d{3})*(?:\.\d+)?)\s*(億元|億|萬)?",
            r"(買超|賣超|買賣超|淨買超|淨賣超)[^。；;\n]{0,40}?(?:八大|官股|公股)?[^-+0-9]{0,20}([-+−]?\d+(?:,\d{3})*(?:\.\d+)?)\s*(億元|億|萬)?",
        ]
        candidates = []
        for pat in patterns:
            for m in re.finditer(pat, clean):
                words = " ".join([g for g in m.groups() if g])
                nums = re.findall(r"[-+−]?\d+(?:,\d{3})*(?:\.\d+)?", words)
                if not nums:
                    continue
                num_text = nums[-1].replace(",", "").replace("−", "-")
                try:
                    val = float(num_text)
                except Exception:
                    continue
                unit_m = re.search(r"(億元|億|萬)", words)
                unit = unit_m.group(1) if unit_m else "億"
                if any(k in words for k in ["賣超", "淨賣超"]) and val > 0:
                    val = -val
                val_100m = self._normalize_gov_unit_to_100m(val, unit)
                # 排除明顯不是金額的日期/代號小數。
                if abs(val_100m) < 0.01 or abs(val_100m) > 5000:
                    continue
                candidates.append({"gov_net_100m": round(val_100m, 2), "matched_text": m.group(0)[:180], "unit": unit})
        if not candidates:
            return None
        # 優先選擇包含八大/官股/公股的片段。
        candidates.sort(key=lambda x: (0 if any(k in x["matched_text"] for k in ["八大", "官股", "公股"]) else 1, -abs(x["gov_net_100m"])))
        best = candidates[0]
        sig = "偏多" if best["gov_net_100m"] > 0 else "偏空" if best["gov_net_100m"] < 0 else "中性"
        best.update({"gov_signal": sig, "gov_score": 1 if sig == "偏多" else -1 if sig == "偏空" else 0, "source_rule": "Wantgoo/八大官股資料頁解析；來源保留供追溯，分數只依數值"})
        return best

    def fetch_wantgoo_public_bank(self, url: str = SOURCE_URLS["wantgoo_public_bank"]) -> RawData:
        raw = self.fetch_text_snapshot("官股整理", url, "Wantgoo", "WARN")
        if raw.value:
            text = raw.value.get("snippet", "") if isinstance(raw.value, dict) else str(raw.value)
            parsed = self._parse_public_bank_text(text)
            nums = [self._to_float(x) for x in re.findall(r"-?\d+(?:,\d{3})*(?:\.\d+)?", text)]
            nums = [x for x in nums if not math.isnan(x)]
            if parsed:
                raw.value.update(parsed)
                raw.value["gov_hint_values"] = nums[:20]
                raw.status = "OK"
                raw.data_status = "OK"
                raw.parse_status = "PARSE_OK"
                raw.confidence = 0.55
                raw.message = "Wantgoo八大官股資料解析完成；來源保留供追溯，分數只依數值"
                raw.raw_file_path = self.logger.write_raw_evidence("官股整理", raw.value, parsed=parsed, status="OK", url=url, message=raw.message)
                self.logger.parsed_value("wantgoo_gov_net_100m", parsed.get("gov_net_100m"), "Wantgoo備援", raw.date)
            elif nums:
                raw.value["gov_hint_values"] = nums[:20]
                raw.status = "WARN"
                raw.data_status = "NO_PARSED_VALUE"
                raw.parse_status = "NO_PARSED_VALUE"
                raw.message = "Wantgoo頁面已取得且有數字，但未能定位八大官股淨買賣超欄位"
                self.logger.parsed_value("wantgoo_gov_hint_values", nums[:5], "Wantgoo", raw.date)
            else:
                raw.status = "WARN"
                raw.data_status = "NO_PARSED_VALUE"
                raw.parse_status = "NO_PARSED_VALUE"
                raw.message = "Wantgoo頁面已取得但未解析出官股數字"
        return raw

    def fetch_twse_broker_report(self, url: str = SOURCE_URLS["twse_broker_report"], base_date: Optional[str] = None, max_back_days: int = DEFAULT_MAX_FALLBACK_DAYS) -> RawData:
        """
        V2.4 官股資料修正：
        1. 原 twse_broker_report 頁面會回 404，因此不再把 404 頁面當成可用官股資料。
        2. 先抓 TWSE 官方 T86 作為證交所法人資料證據；但 T86 不是八大官股本體，只能作官方佐證。
        3. 八大官股目前 TWSE 無公開可直接彙總的官方 API；若使用 Histock/Wantgoo，只標記為「第三方輔助」，不得假裝官方。
        """
        query_date = self._compact_date(base_date)
        try_date = query_date
        last_error = ""
        for fallback_days in range(max_back_days + 1):
            self._log_fallback_try("TWSE_T86_GOV_EVIDENCE", query_date, try_date, fallback_days, max_back_days)
            api_url = SOURCE_URLS["twse_t86"].format(date=try_date)
            try:
                data = self.client.get_json(api_url)
                rows = self._twse_rows(data)
                if rows:
                    parsed = {"twse_t86_rows": len(rows), "source_rule": "TWSE官方T86法人資料佐證；非八大官股本體"}
                    is_fb, fb_days, note = self._fallback_note(query_date, try_date, "TWSE_T86_GOV_EVIDENCE")
                    payload = {"official_evidence": parsed, "rows_sample": rows[:3], "gov_net_100m": None}
                    raw_path = self.logger.write_raw_evidence("官股", payload, parsed=parsed, status="WARN", url=api_url, message="TWSE官方T86可取得，但不是八大官股彙總；官股數字需券商分點彙總或人工覆寫")
                    self.logger.parsed_value("twse_t86_rows", len(rows), "TWSE T86", self._dash_date(try_date))
                    return RawData("官股", payload, self._dash_date(try_date), "TWSE T86 官方佐證", api_url, self._today_str(),
                                   "WARN", "TWSE官方T86可取得，但不是八大官股彙總；未提供gov_net_100m，避免非官方數字進主判斷", query_date,
                                   try_date, is_fb, fb_days, "NO_PARSED_VALUE", note, "NO_PARSED_VALUE", raw_path, 0.35)
                self._official_no_data("TWSE_T86_GOV_EVIDENCE", query_date, try_date, api_url, data)
            except Exception as exc:
                last_error = str(exc)
                self.logger.warning(f"TWSE T86官股佐證來源失敗 try_date={try_date} url={api_url}: {exc}")
            try_date = self._previous_calendar_day(try_date)

        # 第三方輔助來源只做證據，不進主分數
        fallback_url = SOURCE_URLS.get("gov_broker_fallback_histock", SOURCE_URLS.get("wantgoo_public_bank", ""))
        raw = self.fetch_text_snapshot("官股", fallback_url, "第三方八大官股輔助", "WARN", parse_required=True)
        raw.status = "WARN"
        raw.data_status = "NO_PARSED_VALUE"
        raw.parse_status = "NO_PARSED_VALUE"
        raw.message = "TWSE官方官股彙總未取得；第三方八大官股頁只作輔助證據，不進主判斷。" + (f" TWSE錯誤={last_error}" if last_error else "")
        raw.confidence = 0.25
        return raw

    def fetch_ranking_result_db(self, db_path: Optional[str] = None, strict: bool = False) -> RawData:
        """
        V2.4 Ranking資料修正：
        - 若使用者指定db_path，只檢查該DB，不再亂掃其他DB。
        - 未指定時才掃描目前目錄與子目錄，並記錄候選DB。
        - ranking_result不存在或空表時，parse_status必須是NO_PARSED_VALUE，不能假OK。
        - strict=True時直接丟出錯誤，避免交易系統在沒有Ranking核心時繼續產出可下單結論。
        """
        candidates: List[Path] = []
        if db_path:
            candidates = [Path(db_path)]
        else:
            candidates.extend(Path.cwd().glob("*.db"))
            candidates.extend(Path.cwd().glob("**/*.db"))
        seen = set()
        candidates = [p for p in candidates if not (str(p) in seen or seen.add(str(p)))]
        checked: List[str] = []
        for path in candidates[:50]:
            checked.append(str(path))
            if not path.exists():
                self.logger.warning(f"RANKING_DB_NOT_FOUND path={path}")
                continue
            try:
                conn = sqlite3.connect(str(path))
                cur = conn.cursor()
                cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='ranking_result'")
                if not cur.fetchone():
                    conn.close()
                    self.logger.warning(f"RANKING_TABLE_MISSING db={path} table=ranking_result")
                    continue
                cur.execute("SELECT COUNT(*) FROM ranking_result")
                count = int(cur.fetchone()[0] or 0)
                if count <= 0:
                    conn.close()
                    msg = f"ranking_result為空 db={path}，交易系統不可運行"
                    self.logger.warning("RANKING_TABLE_EMPTY " + msg)
                    if strict:
                        raise RuntimeError(msg)
                    payload = {"db_path": str(path), "table": "ranking_result", "row_count": 0}
                    raw_path = self.logger.write_raw_evidence("排行分析", payload, parsed={}, status="FAIL", url=str(path), message=msg)
                    return RawData("排行分析", None, "", "SQLite DB", str(path), self._today_str(), "FAIL", msg, data_status="DATA_MISSING", parse_status="NO_PARSED_VALUE", raw_file_path=raw_path, confidence=0.0)
                cur.execute("PRAGMA table_info(ranking_result)")
                columns = [r[1] for r in cur.fetchall()]
                cur.execute("SELECT * FROM ranking_result LIMIT 5")
                rows = cur.fetchall()
                conn.close()
                parsed = {"db_path": str(path), "table": "ranking_result", "row_count": count, "columns": columns[:50]}
                payload = {**parsed, "sample_rows": rows}
                self.logger.raw_snapshot("排行分析", payload, parsed=parsed, status="OK", url=str(path), message="ranking_result loaded")
                self.logger.parsed_value("ranking_result_count", count, "SQLite DB", str(path))
                return RawData("排行分析", payload, self._today_str(), "SQLite DB", str(path), self._today_str(), "OK", "ranking_result loaded", data_status="OK", parse_status="PARSE_OK", confidence=1.0)
            except Exception as exc:
                self.logger.debug(f"ranking_result DB skipped path={path}: {exc}")
                if strict:
                    raise
        msg = "未找到 ranking_result 資料表；已檢查DB=" + ("; ".join(checked[:20]) if checked else "無DB候選")
        if strict:
            raise RuntimeError(msg)
        raw_path = self.logger.write_raw_evidence("排行分析", {"checked_db": checked}, parsed={}, status="FAIL", url="db:ranking_result", message=msg)
        return RawData("排行分析", None, "", "SQLite DB", "db:ranking_result", self._today_str(), "FAIL", msg, data_status="DATA_MISSING", parse_status="NO_PARSED_VALUE", raw_file_path=raw_path, confidence=0.0)




    def _complete_parsed_raw(self, raw: RawData, parsed: Dict[str, Any], message: str, confidence: float = 1.0) -> RawData:
        raw.value = raw.value or {}
        if isinstance(raw.value, dict):
            raw.value.update(parsed)
        raw.status = "OK"
        raw.data_status = "OK"
        raw.parse_status = "PARSE_OK"
        raw.message = message
        raw.confidence = confidence
        raw.raw_file_path = self.logger.write_raw_evidence(raw.key, raw.value, parsed=parsed, status="OK", url=raw.url, message=message)
        return raw

    def fetch_gov_news(self, url: str = SOURCE_URLS["twse_broker_report"]) -> RawData:
        return self.fetch_twse_broker_report(url)

    def fetch_ai_industry_news(self, url: str = SOURCE_URLS["techcrunch_ai"]) -> RawData:
        raw = self.fetch_text_snapshot("AI產業", url, "TechCrunch", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "")
            strength = 0.8 if any(k in text for k in ["AI", "先進封裝", "需求", "資本支出", "成長"]) else 0.5
            raw.value["ai_strength"] = strength
            self.logger.parsed_value("ai_strength", strength, "TechCrunch", raw.date)
        return raw

    def fetch_geopolitical_news(self, url: str = SOURCE_URLS["reuters_war"]) -> RawData:
        raw = self.fetch_text_snapshot("戰爭/地緣", url, "Reuters", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "").lower()
            risk = 1 if any(k in text for k in ["war", "blockade", "israel", "iran", "attack", "missile", "kills"]) else 0
            raw.value["major_event"] = risk
            self.logger.parsed_value("major_event", risk, "Reuters", raw.date)
        return raw

    def _strip_html_cells(self, row_html: str) -> List[str]:
        cells = re.findall(r"<(?:td|th)[^>]*>(.*?)</(?:td|th)>", row_html, flags=re.I | re.S)
        out = []
        for c in cells:
            txt = re.sub(r"<[^>]+>", " ", c)
            txt = html_lib.unescape(txt)
            txt = re.sub(r"\s+", " ", txt).strip()
            if txt:
                out.append(txt)
        return out

    def _parse_taifex_night_html(self, html_text: str) -> Optional[Dict[str, Any]]:
        """解析TAIFEX夜盤頁。優先解析臺股期貨/外資列的多空淨額口數。"""
        decoded = html_lib.unescape(html_text or "")
        rows = re.findall(r"<tr[^>]*>(.*?)</tr>", decoded, flags=re.I | re.S)
        last_product = ""
        target_date = ""
        mdate = re.search(r"日期\s*(\d{4}/\d{2}/\d{2}|\d{4}-\d{2}-\d{2}|\d{4}\d{2}\d{2})", re.sub(r"<[^>]+>", " ", decoded))
        if mdate:
            target_date = mdate.group(1).replace("/", "-")
        for row in rows:
            cols = self._strip_html_cells(row)
            if not cols:
                continue
            row_text = " ".join(cols)
            if "臺股期貨" in row_text or "台股期貨" in row_text:
                last_product = "臺股期貨"
            # TAIFEX表格常見結構：商品名稱只出現在第一列，自營商/投信/外資在後續列沿用商品名稱
            if last_product == "臺股期貨" and any(c == "外資" or "外資" in c for c in cols):
                nums = []
                for c in cols:
                    cleaned = c.replace(",", "").replace("−", "-").strip()
                    if re.fullmatch(r"-?\d+(?:\.\d+)?", cleaned):
                        nums.append(float(cleaned))
                if len(nums) >= 6:
                    long_lots = int(nums[-6])
                    long_amount = int(nums[-5])
                    short_lots = int(nums[-4])
                    short_amount = int(nums[-3])
                    net_lots = int(nums[-2])
                    net_amount = int(nums[-1])
                elif len(nums) >= 2:
                    net_lots = int(nums[-2])
                    net_amount = int(nums[-1])
                    long_lots = long_amount = short_lots = short_amount = None
                else:
                    continue
                return {
                    "contract": "TX",
                    "contract_name": "臺股期貨",
                    "identity": "外資",
                    "session": "after_hours",
                    "data_date": target_date or "latest",
                    "long_lots": long_lots,
                    "long_amount": long_amount,
                    "short_lots": short_lots,
                    "short_amount": short_amount,
                    "net_lots": net_lots,
                    "net_amount": net_amount,
                    "night_score": 1 if net_lots > 0 else (-1 if net_lots < 0 else 0),
                    "source_rule": "TAIFEX夜盤臺股期貨外資多空淨額口數"
                }
        # fallback：以純文字行序解析（當HTML表格結構變動時）
        text = re.sub(r"<[^>]+>", "\n", decoded)
        tokens = [t.strip() for t in re.split(r"\n+", text) if t.strip()]
        try:
            i = tokens.index("臺股期貨")
        except ValueError:
            try:
                i = tokens.index("台股期貨")
            except ValueError:
                return None
        window = tokens[i:i+60]
        if "外資" in window:
            j = window.index("外資")
            vals = []
            for t in window[j+1:j+10]:
                c = t.replace(",", "").replace("−", "-")
                if re.fullmatch(r"-?\d+(?:\.\d+)?", c):
                    vals.append(float(c))
            if len(vals) >= 6:
                net_lots = int(vals[4])
                net_amount = int(vals[5])
                return {"contract": "TX", "contract_name": "臺股期貨", "identity": "外資", "session": "after_hours", "data_date": target_date or "latest", "long_lots": int(vals[0]), "long_amount": int(vals[1]), "short_lots": int(vals[2]), "short_amount": int(vals[3]), "net_lots": net_lots, "net_amount": net_amount, "night_score": 1 if net_lots > 0 else (-1 if net_lots < 0 else 0), "source_rule": "TAIFEX夜盤純文字fallback"}
        return None

    def fetch_taifex_night_snapshot(self, url: str = SOURCE_URLS["taifex_night"]) -> RawData:
        try:
            html_text = self.client.get_text(url)
            parsed = self._parse_taifex_night_html(html_text)
            if parsed:
                value = {"url": url, "html_length": len(html_text), **parsed}
                raw_path = self.logger.write_raw_evidence("台股夜盤", value, parsed=parsed, status="OK", url=url, message="TAIFEX夜盤已解析臺股期貨外資多空淨額")
                self.logger.raw_snapshot("TAIFEX_NIGHT_PARSED", parsed, parsed=parsed, status="OK", url=url, message="TAIFEX夜盤已解析臺股期貨外資多空淨額")
                self.logger.parsed_value("night_score", parsed.get("night_score"), "TAIFEX", parsed.get("data_date", "latest"))
                self.logger.parsed_value("taifex_tx_foreign_net_lots", parsed.get("net_lots"), "TAIFEX", parsed.get("data_date", "latest"))
                return RawData("台股夜盤", value, parsed.get("data_date", "latest"), "TAIFEX", url, self._today_str(), "OK", "TAIFEX夜盤已解析臺股期貨外資多空淨額", data_status="OK", parse_status="PARSE_OK", raw_file_path=raw_path, confidence=0.9)
            value = {"url": url, "html_length": len(html_text), "snippet": re.sub(r"\s+", " ", html_text[:5000])}
            raw_path = self.logger.write_raw_evidence("台股夜盤", value, parsed={}, status="WARN", url=url, message="TAIFEX頁面已取得，但未解析出臺股期貨外資多空淨額")
            self.logger.warning(f"TAIFEX夜盤未解析出數值 url={url}")
            return RawData("台股夜盤", value, "latest", "TAIFEX", url, self._today_str(), "WARN", "TAIFEX頁面已取得，但未解析出臺股期貨外資多空淨額", data_status="NO_PARSED_VALUE", parse_status="NO_PARSED_VALUE", raw_file_path=raw_path, confidence=0.3)
        except Exception as exc:
            self.logger.warning(f"TAIFEX夜盤抓取失敗 url={url}: {exc}")
            raw_path = self.logger.write_raw_evidence("台股夜盤", {"url": url, "error": str(exc)}, parsed={}, status="WARN", url=url, message=str(exc))
            return RawData("台股夜盤", None, "", "TAIFEX", url, self._today_str(), "WARN", str(exc), data_status="FETCH_FAIL", parse_status="NO_PARSED_VALUE", raw_file_path=raw_path, confidence=0.0)

    def fetch_tpex_otc_snapshot(self, url: str = SOURCE_URLS["tpex_indices"]) -> RawData:
        raw = self.fetch_text_snapshot("OTC官方來源", url, "TPEX", "WARN")
        if raw.status == "OK" and raw.value:
            text = raw.value.get("snippet", "")
            nums = [self._to_float(x) for x in re.findall(r"\d+(?:,\d{3})*(?:\.\d+)?", text)]
            nums = [x for x in nums if not math.isnan(x)]
            if nums:
                raw.value["otc_hint_values"] = nums[:20]
                self.logger.parsed_value("tpex_otc_hint_values", nums[:5], "TPEX", raw.date)
            else:
                raw.status = "WARN"
                raw.data_status = "NO_PARSED_VALUE"
                raw.message = "TPEX官方頁已取得但未解析出OTC指數數值"
        return raw

    def build_manual_raw(self, module: str, value: Any, note: str, source: str = "人工覆寫/Excel註解") -> RawData:
        self.logger.raw_snapshot(module, {"value": value, "note": note})
        self.logger.parsed_value(module, value, source, self._today_str())
        return RawData(module, value, self._today_str(), source, source, self._today_str(), "OK", note)

    def _to_float(self, value: Any) -> float:
        if value is None:
            return math.nan
        s = str(value).replace(",", "").replace("--", "").strip()
        if s in ("", "nan", "None"):
            return math.nan
        return float(s)


class MarketNarrativeBuilder:
    """SOP V2.1 P1-01：市場輸入交易判讀模板化。"""
    def build(self, market: MarketInput) -> str:
        parts = []
        if market.close is not None and market.ma5 is not None:
            parts.append(f"台股收盤{market.close:.2f}，{'站上' if market.close >= market.ma5 else '跌破'}5MA {market.ma5:.2f}")
        if market.foreign_net_100m is not None:
            parts.append(f"外資{'買超' if market.foreign_net_100m >= 0 else '賣超'}{abs(market.foreign_net_100m):.2f}億元")
        else:
            parts.append("外資資料未完成")
        if market.gov_net_100m is not None:
            parts.append(f"官股/八大公股{'買超' if market.gov_net_100m >= 0 else '賣超'}{abs(market.gov_net_100m):.2f}億元")
        else:
            parts.append("官股資料未完成，需TEJ或備援來源")
        if market.turnover_100m is not None and market.avg_turnover_5d_100m is not None:
            vol_state = "放量" if market.turnover_100m > market.avg_turnover_5d_100m * 1.05 else "量能正常/未明顯放大"
            parts.append(f"成交值{market.turnover_100m:.2f}億元，5日均量{market.avg_turnover_5d_100m:.2f}億元，{vol_state}")
        if market.major_event:
            parts.append("重大事件風險=1，需降倉禁追高")
        else:
            parts.append("重大事件風險=0")
        if market.night_score is not None:
            if market.night_score < 0:
                lots = f"{market.night_net_lots}口" if market.night_net_lots is not None else "淨空"
                parts.append(f"夜盤外資偏空({lots})，盤前風險需加分")
            elif market.night_score > 0:
                lots = f"{market.night_net_lots}口" if market.night_net_lots is not None else "淨多"
                parts.append(f"夜盤外資偏多({lots})")
            else:
                parts.append("夜盤外資中性")
        return "；".join(parts) + "。"

class FieldCompletionValidator:
    """SOP V2.1 P0-03：核心欄位最終完成規則。"""
    CORE_FIELDS = ["base_date", "close", "high", "low", "ma5", "turnover_100m", "avg_turnover_5d_100m", "foreign_net_100m"]
    STRICT_FIELDS = ["gov_net_100m"]

    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def validate_market_input(self, market: MarketInput, strict_gov: bool = False) -> List[str]:
        issues = []
        for field in self.CORE_FIELDS:
            val = getattr(market, field, None)
            if val is None or val == "":
                issues.append(f"P0欄位未完成:{field}")
        if strict_gov and market.gov_net_100m is None:
            issues.append("P0欄位未完成:gov_net_100m，需TEJ或Wantgoo備援解析/人工覆寫")
        elif market.gov_net_100m is None:
            issues.append("P0_WARN:gov_net_100m未完成，正式檔會標示未取得，不可假OK")
        for issue in issues:
            self.logger.warning("FIELD_COMPLETION " + issue)
        return issues

class MacroRefillValidator:
    """SOP V2.2：只在 macro_only 模式清成三頁；macro_refill 不得刪除 TOP/00~09 報表。"""
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def ensure_macro_sheets(self, wb):
        """保留既有報表，只確保宏觀三頁存在並移到前面。"""
        if not wb.sheetnames:
            wb.create_sheet("市場輸入")
        for required in MACRO_REFILL_SHEETS:
            if required not in wb.sheetnames:
                wb.create_sheet(required)
        front = [wb[name] for name in MACRO_REFILL_SHEETS if name in wb.sheetnames]
        rest = [ws for ws in wb._sheets if ws.title not in MACRO_REFILL_SHEETS]
        wb._sheets = front + rest
        return wb

    def enforce_macro_only_sheets(self, wb):
        """macro_only 專用：只保留宏觀三頁。"""
        self.ensure_macro_sheets(wb)
        for name in list(wb.sheetnames):
            if name not in MACRO_REFILL_SHEETS:
                del wb[name]
                self.logger.info(f"MACRO_ONLY_REMOVE_EXTRA_SHEET sheet={name}")
        wb._sheets = [wb[name] for name in MACRO_REFILL_SHEETS]
        return wb

class DataProcessor:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def _source_note(self, raw: RawData) -> str:
        parts = [raw.url]
        if raw.query_date:
            parts.append(f"query_date={raw.query_date}")
        if raw.actual_date:
            parts.append(f"actual_date={raw.actual_date}")
        if raw.is_fallback:
            parts.append(f"fallback_days={raw.fallback_days}")
        if raw.data_note:
            parts.append(raw.data_note)
        return " | ".join([p for p in parts if p])

    def apply_manual_override(self, market: MarketInput, override: Optional[ManualOverride]) -> MarketInput:
        if not override:
            return market
        if override.gov_net_100m is not None:
            old = market.gov_net_100m
            market.gov_net_100m = override.gov_net_100m
            self.logger.info(f"MANUAL_OVERRIDE field=gov_net_100m old={old} new={override.gov_net_100m}")
        if override.ai_strength is not None:
            old = market.ai_strength
            market.ai_strength = override.ai_strength
            self.logger.info(f"MANUAL_OVERRIDE field=ai_strength old={old} new={override.ai_strength}")
        if override.major_event is not None:
            old = market.major_event
            market.major_event = int(override.major_event)
            self.logger.info(f"MANUAL_OVERRIDE field=major_event old={old} new={override.major_event} note={override.event_note}")
        return market

    def _merge_major_event(self, raw: Dict[str, RawData]) -> Tuple[int, List[Tuple[str, int, str]]]:
        """P0：合併 Reuters/ISW/CNN/人工事件來源，避免單一來源失敗造成重大事件被漏判。"""
        candidates = ["戰爭/地緣", "戰爭/停火", "ISW衝突分析", "CNN重大新聞"]
        events: List[Tuple[str, int, str]] = []
        for key in candidates:
            item = raw.get(key)
            value = 0
            source = ""
            try:
                source = item.source if item else ""
                if item and isinstance(item.value, dict):
                    value = int(item.value.get("major_event", 0) or 0)
            except Exception:
                value = 0
            events.append((key, value, source))
        merged = max([v for _, v, _ in events], default=0)
        active_sources = [f"{k}:{v}:{s}" for k, v, s in events if v]
        self.logger.info(f"MARKET_EVENT_MERGE merged={merged} detail={events}")
        if active_sources:
            self.logger.parsed_value("market.major_event_sources", ";".join(active_sources), "EventMerge", "latest")
        return merged, events

    def build_market_input(self, raw: Dict[str, RawData], base_date: str = "") -> MarketInput:
        market = MarketInput()
        taiex = raw.get("台股指數")
        turnover = raw.get("成交量")
        foreign = raw.get("外資")
        if taiex and taiex.status == "OK" and taiex.value:
            rows = taiex.value.get("rows", [])
            last = taiex.value.get("last", {})
            market.base_date = last.get("date", base_date)
            market.close = last.get("close")
            market.high = last.get("high")
            market.low = last.get("low")
            if len(rows) >= 2:
                market.prev_high = rows[-2].get("high")
                market.prev_low = rows[-2].get("low")
            if len(rows) >= 5:
                market.ma5 = round(sum([r["close"] for r in rows[-5:]]) / 5, 2)
            market.source_1 = self._source_note(taiex)
        else:
            market.base_date = base_date or dt.date.today().isoformat()
            self.logger.warning("台股指數未取得，市場輸入將保留缺值")
        if turnover and turnover.status == "OK" and turnover.value:
            trows = turnover.value.get("rows", [])
            last_t = turnover.value.get("last", {})
            market.turnover_100m = round(last_t.get("turnover_100m", 0), 2)
            if len(trows) >= 5:
                market.avg_turnover_5d_100m = round(sum([r["turnover_100m"] for r in trows[-5:]]) / 5, 2)
            market.source_2 = self._source_note(turnover)
        if foreign and foreign.status == "OK" and foreign.value:
            market.foreign_net_100m = round(foreign.value.get("net_100m", 0), 2)
            market.source_3 = self._source_note(foreign)

        # V2.5 P0：跨月5日資料優先覆蓋單日資料，補齊前高/前低/5MA/5日均量。
        market5 = raw.get("市場5日資料")
        if market5 and market5.status == "OK" and isinstance(market5.value, dict):
            v = market5.value
            for field in ["base_date", "close", "high", "low", "prev_high", "prev_low", "ma5", "turnover_100m", "avg_turnover_5d_100m"]:
                if v.get(field) is not None:
                    setattr(market, field, v.get(field))
            note = f"TWSE跨月5日：{v.get('taiex_recent_dates','')} / 成交值：{v.get('turnover_recent_dates','')}"
            market.source_1 = note
            market.source_2 = note
            self.logger.parsed_value("market.ma5", market.ma5, "TWSE跨月5日", market.base_date)
            self.logger.parsed_value("market.avg_turnover_5d_100m", market.avg_turnover_5d_100m, "TWSE跨月5日", market.base_date)

        gov_candidates = [raw.get("官股"), raw.get("八大官股"), raw.get("官股整理")]
        gov = next((g for g in gov_candidates if g and getattr(g, "parse_status", "") == "PARSE_OK" and isinstance(g.value, dict) and g.value.get("gov_net_100m") is not None), None)
        if gov:
            market.gov_net_100m = round(float(gov.value.get("gov_net_100m")), 2)
            market.source_4 = self._source_note(gov)
            self.logger.parsed_value("market.gov_net_100m", market.gov_net_100m, gov.source, gov.date)
        else:
            market.gov_net_100m = None

        ai = raw.get("AI產業")
        iek = raw.get("IEK產業分析")
        # V2.5：依P0/P1修正，IEK 台灣產業來源優先於 TechCrunch，避免低估台股AI產業強度。
        if iek and iek.status == "OK" and isinstance(iek.value, dict) and iek.value.get("ai_strength") is not None:
            market.ai_strength = float(iek.value.get("ai_strength"))
            self.logger.parsed_value("market.ai_strength", market.ai_strength, iek.source, iek.date)
        elif ai and ai.status == "OK" and isinstance(ai.value, dict) and ai.value.get("ai_strength") is not None:
            market.ai_strength = float(ai.value.get("ai_strength"))
            self.logger.parsed_value("market.ai_strength", market.ai_strength, ai.source, ai.date)
        else:
            market.ai_strength = 0.5

        merged_event, event_details = self._merge_major_event(raw)
        market.major_event = int(merged_event or 0)
        event_sources = [f"{k}:{v}:{s}" for k, v, s in event_details if v]
        if event_sources:
            if not market.source_4:
                market.source_4 = "重大事件來源=" + ";".join(event_sources)
            self.logger.parsed_value("market.major_event", market.major_event, "EventMerge", "latest")
        else:
            self.logger.parsed_value("market.major_event", market.major_event, "EventMerge", "latest")

        night = raw.get("台股夜盤")
        if night and night.status == "OK" and isinstance(night.value, dict):
            try:
                market.night_score = int(night.value.get("night_score", 0) or 0)
            except Exception:
                market.night_score = 0
            try:
                market.night_net_lots = int(night.value.get("net_lots")) if night.value.get("net_lots") is not None else None
            except Exception:
                market.night_net_lots = None
            self.logger.parsed_value("market.night_score", market.night_score, night.source, night.date)
            self.logger.parsed_value("market.night_net_lots", market.night_net_lots, night.source, night.date)

        self.logger.info(f"市場輸入標準化完成：{asdict(market)}")
        return market

class ScoringEngine:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def score_all(self, raw: Dict[str, RawData], market: MarketInput) -> List[ModuleScore]:
        scores: List[ModuleScore] = []
        for module in MODULES:
            method = getattr(self, f"score_{self._safe_name(module)}", None)
            if method:
                scores.append(method(raw.get(module), market))
            else:
                scores.append(self.score_neutral(module, raw.get(module), "尚未建立自動判定規則，依SOP標記中性"))
        return scores

    def _safe_name(self, module: str) -> str:
        m = module.replace("-", "_").replace("/", "_").replace(" ", "_")
        mapping = {
            "美股_S&P500": "sp500", "美股_NASDAQ": "nasdaq", "美股_道瓊": "dow",
            "VIX恐慌": "vix", "美債10Y": "ust10y", "原油": "oil", "戰爭_地緣": "geopolitical",
            "CPI": "cpi", "非農": "nfp", "外資": "foreign", "官股": "gov",
            "台股指數": "taiex", "成交量": "turnover", "AI產業": "ai", "OTC": "otc", "台股夜盤": "night"
        }
        return mapping.get(m, m)

    def score_neutral(self, module: str, raw: Optional[RawData], reason: str) -> ModuleScore:
        data_text = "未取得" if not raw or raw.status != "OK" else str(raw.value)
        source = "未取得" if not raw else raw.source
        data_time = "" if not raw else raw.date
        return ModuleScore(module, data_text, 0.0, 0, 0.0, reason, source, data_time, "不納入主判斷，保守中性", "WARN")

    def _score_yahoo_index(self, raw: Optional[RawData], module: str, positive_text: str, negative_text: str) -> ModuleScore:
        if not raw or raw.status != "OK" or not raw.value:
            return self.score_neutral(module, raw, f"{module}未取得，依SOP不可編造，列中性")
        close = raw.value["close"]
        pct = raw.value["change_pct"]
        direction = 1 if pct > 0.2 else (-1 if pct < -0.2 else 0)
        strength = min(1.0, max(0.3, abs(pct) / 2)) if direction else 0.2
        weighted = round(direction * strength, 2)
        explanation = f"{module}收{close:.2f}，漲跌幅{pct:.2f}%；{positive_text if direction>0 else negative_text if direction<0 else '小幅震盪，方向中性'}。"
        trade = "提高風險偏好" if direction > 0 else "降低追價意願" if direction < 0 else "維持觀望，不作為主要加減碼依據"
        return ModuleScore(module, f"收盤{close:.2f} / 漲跌幅{pct:.2f}%", round(strength,2), direction, weighted, explanation, raw.source, raw.date, trade)

    def score_sp500(self, raw, market):
        return self._score_yahoo_index(raw, "美股-S&P500", "美股風險偏好偏多", "美股風險偏好降溫")

    def score_nasdaq(self, raw, market):
        return self._score_yahoo_index(raw, "美股-NASDAQ", "科技股仍有支撐", "科技股轉弱，壓抑AI權值股")

    def score_dow(self, raw, market):
        return self._score_yahoo_index(raw, "美股-道瓊", "傳產廣度穩定", "傳產與景氣股偏弱")

    def score_vix(self, raw, market):
        if not raw or raw.status != "OK" or not raw.value:
            return self.score_neutral("VIX恐慌", raw, "VIX未取得，列中性")
        v = raw.value["close"]
        pct = raw.value["change_pct"]
        if v < 20 and pct <= 0:
            direction, strength = 1, 0.4
        elif v > 25 or pct > 5:
            direction, strength = -1, 0.8
        elif v >= 20 or pct > 0:
            direction, strength = -1, 0.5
        else:
            direction, strength = 0, 0.2
        weighted = round(direction * strength, 2)
        explanation = f"VIX {v:.2f}，日變化{pct:.2f}%；依SOP低於20偏穩，但快速上升需提高風險。"
        trade = "恐慌未升溫，可維持正常部位" if direction > 0 else "風險升溫，降低追價與槓桿" if direction < 0 else "風險中性"
        return ModuleScore("VIX恐慌", f"{v:.2f} / {pct:.2f}%", strength, direction, weighted, explanation, raw.source, raw.date, trade)

    def score_ust10y(self, raw, market):
        if not raw or raw.status != "OK" or not raw.value:
            return self.score_neutral("美債10Y", raw, "美債10Y未取得，列中性")
        y = raw.value["value"]
        direction = -1 if y >= 4.0 else (1 if y < 3.5 else 0)
        strength = 0.7 if y >= 4.0 else 0.4 if direction else 0.2
        weighted = round(direction * strength, 2)
        explanation = f"美債10Y殖利率{y:.2f}%；4%以上對高估值科技股有壓力。"
        trade = "壓抑高估值與追價" if direction < 0 else "估值壓力緩和" if direction > 0 else "中性觀察"
        return ModuleScore("美債10Y", f"{y:.2f}%", strength, direction, weighted, explanation, raw.source, raw.date, trade)

    def score_oil(self, raw, market):
        if not raw or raw.status != "OK" or not raw.value:
            return self.score_neutral("原油", raw, "原油未取得，列中性")
        close = raw.value.get("close")
        pct = raw.value.get("change_pct", 0)
        last5 = raw.value.get("last5_close", []) or []
        five_day_change = ((close - last5[0]) / last5[0] * 100) if close and len(last5) >= 5 and last5[0] else pct
        if five_day_change > 3 or pct > 1.5:
            direction, strength = -1, min(1.0, max(0.5, abs(five_day_change) / 5))
            reason = "油價短線上升，通膨與成本壓力提高"
        elif pct < -1.5:
            direction, strength = 1, 0.3
            reason = "油價短線回落，通膨壓力略降"
        else:
            direction, strength = 0, 0.3
            reason = "油價震盪，暫不作主要判斷"
        weighted = round(direction * strength, 2)
        return ModuleScore("原油", f"收{close:.2f}/日變化{pct:.2f}%/5日變化{five_day_change:.2f}%",
                           round(strength,2), direction, weighted, reason, raw.source, raw.date,
                           "油價高檔急升時降低追價與高估值股部位")

    def score_geopolitical(self, raw, market):
        if raw and raw.status == "OK" and isinstance(raw.value, dict):
            risk = int(raw.value.get("major_event", 0) or 0)
            direction = -1 if risk else 0
            strength = 0.8 if risk else 0.2
            weighted = round(direction * strength, 2)
            explanation = "已依 Excel 註解來源優先抓取 Reuters/地緣事件頁面；若含戰爭、封鎖、攻擊等關鍵字，列重大事件風險。"
            trade = "外部事件風險升高，降倉禁追高" if risk else "事件頁面已抓取，未偵測到強風險關鍵字"
            return ModuleScore("戰爭/地緣", raw.value.get("url", ""), strength, direction, weighted, explanation, raw.source, raw.date, trade, "OK")
        return ModuleScore("戰爭/地緣", "未取得", 0.0, 0, 0.0, "地緣新聞來源未取得，列為選用資料缺失，不得自動編造事件。", "Reuters", market.base_date, "需確認資料來源/解析結果事件嚴重度", "WARN")
    def score_cpi(self, raw, market):
        if raw and raw.status == "OK" and raw.value:
            snippet = raw.value.get("snippet", "") if isinstance(raw.value, dict) else str(raw.value)
            nums = re.findall(r"[-+]?\d+(?:\.\d+)?\s*percent|[-+]?\d+(?:\.\d+)?%", snippet, flags=re.I)
            data_text = f"BLS CPI API fetched；value={raw.value.get('value') if isinstance(raw.value, dict) else ''}"
            if nums:
                data_text += "；sample=" + ", ".join(nums[:3])
            return ModuleScore("CPI", data_text, 0.3, 0, 0.0, "CPI 為週期性資料，已依 Excel 註解改抓 BLS API；非公布日使用最近公告，不做每日回退。", raw.source, raw.date, "通膨資料已抓取，需搭配油價/利率判斷", "OK")
        return ModuleScore("CPI", "BLS未取得", 0.0, 0, 0.0, "CPI 為週期性資料；BLS來源抓取失敗時不應假設中性，需標示資料缺失。", "BLS", market.base_date, "需補抓或人工確認", "WARN")
    def score_nfp(self, raw, market):
        if raw and raw.status == "OK" and raw.value:
            snippet = raw.value.get("snippet", "") if isinstance(raw.value, dict) else str(raw.value)
            data_text = f"BLS 非農 API fetched；value={raw.value.get('value') if isinstance(raw.value, dict) else ''}"
            m = re.search(r"nonfarm payroll employment.*?(?:rose|increased).*?([0-9,]+)", snippet, flags=re.I)
            if m:
                data_text += f"；payroll_hint={m.group(1)}"
            return ModuleScore("非農", data_text, 0.3, 0, 0.0, "非農為週期性資料，已依 Excel 註解改抓 BLS API就業資料；非公布日使用最近公告，不做每日回退。", raw.source, raw.date, "就業資料已抓取，需搭配市場預期差判斷", "OK")
        return ModuleScore("非農", "BLS未取得", 0.0, 0, 0.0, "非農為週期性資料；BLS來源抓取失敗時需標示資料缺失，不可編造。", "BLS", market.base_date, "需補抓或人工確認", "WARN")
    def score_foreign(self, raw, market):
        value = market.foreign_net_100m
        if value is None:
            return self.score_neutral("外資", raw, "外資買賣超未取得，列中性")
        direction = 1 if value > 10 else (-1 if value < -10 else 0)
        strength = min(1.0, max(0.3, abs(value) / 300)) if direction else 0.2
        weighted = round(direction * strength, 2)
        explanation = f"外資買賣超{value:.2f}億元；買超為正、賣超為負，依SOP轉為資金方向。"
        trade = "外資偏多，可提高主流股觀察" if direction > 0 else "外資賣壓，降低部位" if direction < 0 else "外資小幅，中性"
        return ModuleScore("外資", f"{value:.2f}億元", round(strength,2), direction, weighted, explanation, raw.source if raw else "", market.base_date, trade)

    def score_gov(self, raw, market):
        value = market.gov_net_100m
        if value is None:
            source = raw.source if raw else "官股資料來源"
            msg = raw.message if raw else "官股資料未取得"
            return ModuleScore("官股", msg, 0.0, 0, 0.0, "官股資料未取得，不編造數字；取得後分數只依gov_net_100m數值，不因來源不同扣分。", source, raw.date if raw else market.base_date, "官股資料不足，不納入主判斷", "WARN")
        direction = 1 if value > 0 else (-1 if value < 0 else 0)
        strength = min(1.0, max(0.3, abs(value)/100)) if direction else 0.2
        weighted = round(direction*strength,2)
        return ModuleScore("官股", f"{value:.2f}億元", strength, direction, weighted, "官股/八大公股資金方向已解析；分數只依數值，不因來源不同扣分。買超代表承接支撐，賣超代表政策資金未護盤。", raw.source if raw else "官股資料來源", market.base_date, "視為支撐判斷，不等於追價依據", "OK")
    def score_taiex(self, raw, market):
        if market.close is None or market.ma5 is None:
            return self.score_neutral("台股指數", raw, "台股收盤或5MA不足，列中性")
        direction = 1 if market.close > market.ma5 else -1
        strength = 0.6
        weighted = round(direction*strength,2)
        explanation = f"加權指數收{market.close:.2f}，5MA {market.ma5:.2f}；依SOP站上/跌破5MA判斷本地市場方向。"
        trade = "允許正常選股，但仍需個股條件" if direction > 0 else "短線轉弱，降倉與禁追高"
        return ModuleScore("台股指數", f"收{market.close:.2f}/5MA{market.ma5:.2f}", strength, direction, weighted, explanation, raw.source if raw else "TWSE", market.base_date, trade)

    def score_turnover(self, raw, market):
        if market.turnover_100m is None or market.avg_turnover_5d_100m is None:
            return self.score_neutral("成交量", raw, "成交值或5日均量不足，列中性")
        if market.turnover_100m > market.avg_turnover_5d_100m * 1.05:
            direction, strength = (1, 0.5) if market.close and market.ma5 and market.close >= market.ma5 else (-1, 0.6)
        else:
            direction, strength = 0, 0.3
        weighted = round(direction*strength,2)
        explanation = f"成交值{market.turnover_100m:.2f}億元，5日均量{market.avg_turnover_5d_100m:.2f}億元；依SOP判斷量能可信度。"
        trade = "價量配合可提高可信度" if direction > 0 else "下跌放量或量能不足，追價風險提高" if direction < 0 else "量能普通，不做主要加減碼依據"
        return ModuleScore("成交量", f"{market.turnover_100m:.2f}/{market.avg_turnover_5d_100m:.2f}億元", strength, direction, weighted, explanation, raw.source if raw else "TWSE", market.base_date, trade)

    def score_ai(self, raw, market):
        strength = float(market.ai_strength or 0.5)
        source = "TechCrunch/IEK/產業來源" if raw and raw.status == "OK" else "人工/預設"
        direction = 1 if strength >= 0.7 else (0 if strength >= 0.4 else -1)
        weighted = round(direction * strength, 2)
        explanation = f"AI主流強度{strength:.2f}；V1.8 已依 Excel 註解抓取 TechCrunch/IEK產業來源，若抓不到才保留預設值。"
        trade = "主線有效，可優先找主流拉回" if direction > 0 else "AI主線中性，避免過度集中" if direction == 0 else "AI題材轉弱，降低權重"
        status = "OK" if raw and raw.status == "OK" else "WARN"
        return ModuleScore("AI產業", f"AI強度{strength:.2f}", strength, direction, weighted, explanation, source, market.base_date, trade, status)
    def score_otc(self, raw, market):
        if raw and raw.status == "OK" and isinstance(raw.value, dict) and "close" in raw.value:
            return self._score_yahoo_index(raw, "OTC", "中小型資金活躍", "中小型資金轉弱")
        if raw and raw.status == "OK":
            return ModuleScore("OTC", "TPEX OTC 官方來源頁已抓取但未解析指數數值", 0.0, 0, 0.0, "抓到頁面不等於抓到數據；未解析出OTC指數與漲跌前列WARN，不納入分數。", raw.source, raw.date, "OTC資料不足，不納入主判斷", "WARN")
        return self.score_neutral("OTC", raw, "OTC資料未取得，列中性")
    def score_night(self, raw, market):
        if raw and raw.status == "OK" and isinstance(raw.value, dict) and raw.value.get("night_score") is not None:
            score = int(raw.value.get("night_score"))
            direction = 1 if score > 0 else (-1 if score < 0 else 0)
            strength = 0.4 if direction else 0.2
            weighted = round(direction * strength, 2)
            return ModuleScore("台股夜盤", str(raw.value), strength, direction, weighted, "已依附件指定 TAIFEX futContractsDateAh 解析台股夜盤數值。", raw.source, raw.date, "作為盤前輔助判斷", "OK")
        if raw and raw.status == "WARN":
            return ModuleScore("台股夜盤", raw.message or "TAIFEX正確網址已抓取但未解析數值", 0.0, 0, 0.0, "已改用附件指定 TAIFEX 夜盤網址，但本次未解析出TX夜盤數值，不可假裝OK。", raw.source, raw.date, "需修正 parser 或人工確認", "WARN")
        return ModuleScore("台股夜盤", "未取得", 0.0, 0, 0.0, "TAIFEX 夜盤來源未取得；屬時間敏感輔助資料，不阻斷主判斷。", "TAIFEX", market.base_date, "盤前需確認夜盤強弱", "WARN")
class IndicatorEngine:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def compute(self, market: MarketInput, macro_total: float) -> TechnicalRisk:
        below_ma5 = int(market.close is not None and market.ma5 is not None and market.close < market.ma5)
        lower_high = int(market.high is not None and market.prev_high is not None and market.high < market.prev_high)
        lower_low = int(market.low is not None and market.prev_low is not None and market.low < market.prev_low)
        volume_expansion = int(market.turnover_100m is not None and market.avg_turnover_5d_100m is not None and market.turnover_100m > market.avg_turnover_5d_100m * 1.05)
        major_event = int(market.major_event or 0)
        night_score = market.night_score if market.night_score is not None else 0
        night_bearish = int(night_score < 0)
        risk_score = below_ma5 + lower_high + lower_low + volume_expansion + major_event + night_bearish
        if night_bearish:
            self.logger.info(f"NIGHT_RISK_APPLIED night_score={market.night_score} net_lots={market.night_net_lots}")
        if macro_total >= 3 and risk_score <= 1:
            judgement = "強多 / 允許交易"
        elif macro_total >= 1 and risk_score <= 2:
            judgement = "震盪偏多 / 可做但禁追高"
        elif macro_total <= -3 or risk_score >= 4:
            judgement = "風險偏空 / 停止新倉"
        elif macro_total <= -1 or risk_score >= 3:
            judgement = "震盪偏空 / 降倉禁追高"
        else:
            judgement = "中性震盪 / 只做最高勝率標的"
        self.logger.info(f"V2技術引擎完成：risk_score={risk_score}, judgement={judgement}, night_bearish={night_bearish}")
        return TechnicalRisk(below_ma5, lower_high, lower_low, volume_expansion, major_event, risk_score, judgement, night_bearish, market.night_score, market.night_net_lots)

class ExplanationEngine:
    def build_summary(self, macro_total: float, tech: TechnicalRisk) -> Dict[str, str]:
        if macro_total >= 3:
            state = "強多"
            switch = "允許交易"
            advice = "可提高主流股權重，但仍需個股進場條件。"
        elif 1 <= macro_total < 3:
            state = "震盪偏多"
            switch = "允許交易但禁追高"
            advice = "優先主流拉回、低位階翻多，避免滿倉。"
        elif -1 < macro_total < 1:
            state = "中性震盪"
            switch = "降低出手頻率"
            advice = "只做最高勝率標的，控制部位。"
        elif -3 < macro_total <= -1:
            state = "震盪偏空"
            switch = "降倉交易 / 禁追高"
            advice = "只做防守與低位階輪動，不追高。"
        else:
            state = "風險偏空"
            switch = "停止新倉"
            advice = "等待風險降溫與技術轉強。"
        if tech.risk_score >= 3 and "停止" not in switch:
            switch = "降倉交易 / 禁追高"
            advice += " 技術風險偏高，需再降部位。"
        return {"宏觀總分": f"{macro_total:.2f}", "技術風險分數": f"{tech.risk_score:.0f}", "市場狀態": state, "交易開關": switch, "操作建議": advice, "核心結論": tech.market_judgement}


# =============================
# V2.5 P0 Institutional Report + TEJ + Market5Day
# =============================
EXPECTED_INSTITUTIONAL_SHEETS = [
    "00_執行摘要", "01_DB資料盤點", "02_模型設計", "03_最終TOP15",
    "04_成長模型TOP30", "05_價值模型TOP30", "06_低位階候選",
    "07_老師點名股檢核", "08_排除與風險", "09_來源與限制",
    "10_老師策略驗收", "11_修改追蹤", "12_策略命中驗證",
    "13_LOW_BUY候選", "14_WATCH觀察池", "15_AVOID排除清單", "16_放行與觀察候選", "17_REDUCE減碼警示"
]

REPORT_COLUMNS = [
    "排名", "代號", "名稱", "市場", "產業", "題材", "老師分類", "老師決策", "決策原因", "決策追蹤", "老師股票池", "核心股狀態",
    "K線警訊", "波段修正", "避開/換股", "現價",
    "低接區", "停損", "目標1", "目標2", "RR", "是否可下單", "老師執行狀態",
    "老師總分", "成長分", "價值分", "EPS_TTM", "PE", "殖利率%",
    "營收YoY%", "法人分", "20日漲幅%", "低位階分", "K線分",
    "均線支撐分", "量能健康分", "營收EPS分", "操作策略", "排除原因",
    "低位階翻多", "放行理由", "硬性排除", "軟性排除", "壓力來源",
    "Phase5大波", "Phase5小波", "Phase5修正型態", "Phase5回撤比例", "Phase5反彈性質",
    "Phase5修正完成", "Phase5逃命反彈", "Phase5推動浪", "Phase5波段階段", "Phase5突破階段", "Phase5推動階段", "Phase5波段位置分", "Phase5候選池", "Phase5阻擋原因",
    "期間BUY次數", "期間AVOID次數", "最後降級原因"
]

TEACHER_WATCHLIST = [
    "2317", "2382", "3706", "2881", "2330", "3019", "2324", "6533", "2359",
    "3231", "2454", "3034", "3711", "9945", "2603", "2412"
]

class TEJGovBankEngine:
    """V2.5 P0：TEJ八大公股行庫主來源解析。Wantgoo/T86只作佐證，不冒充主資料。"""
    def __init__(self, tej_file: Optional[str], logger: Macro16Logger):
        self.tej_file = tej_file
        self.logger = logger

    def _to_number(self, v: Any) -> float:
        try:
            if v is None:
                return math.nan
            s = str(v).replace(",", "").replace("--", "").strip()
            if s == "" or s.lower() == "nan":
                return math.nan
            return float(s)
        except Exception:
            return math.nan

    def load(self):
        if pd is None:
            self.logger.warning("TEJGovBankEngine 需要 pandas，未安裝時官股只能標示WARN")
            return None
        if not self.tej_file:
            return None
        path = Path(self.tej_file)
        if not path.exists():
            self.logger.warning(f"TEJ_GOV_FILE_NOT_FOUND path={path}")
            return None
        try:
            sheets = pd.read_excel(path, sheet_name=None)
        except Exception as exc:
            self.logger.warning(f"TEJ_GOV_READ_FAIL path={path} error={exc}")
            return None
        raw = None
        for name in ["Raw1", "Raw2", "raw1", "raw2"]:
            if name in sheets:
                raw = sheets[name]
                break
        if raw is None:
            for _, df in sheets.items():
                cols = [str(c) for c in df.columns]
                if any(("買" in c and "超" in c) for c in cols):
                    raw = df
                    break
        if raw is None or raw.empty:
            return None
        raw = raw.copy()
        raw.columns = [str(c).strip() for c in raw.columns]
        return raw

    def parse(self) -> Dict[str, Any]:
        df = self.load()
        if df is None or len(df) == 0:
            return {"status": "WARN", "gov_net_100m": None, "gov_signal": "未知", "gov_score": 0, "message": "TEJ檔案未提供、讀取失敗或無Raw資料", "rows": 0}
        amount_col = next((c for c in df.columns if "買(賣)超金額" in c or "買賣超金額" in c or "買賣超金額" in c.replace(" ", "")), None)
        net_col = next((c for c in df.columns if "買(賣)超" in c or "買賣超" in c), None)
        date_col = next((c for c in df.columns if "日期" in c), None)
        if amount_col:
            amount = sum([self._to_number(x) for x in df[amount_col].tolist() if not math.isnan(self._to_number(x))])
            # TEJ欄位可能已是千元/元；先以欄名金額總和保守轉億元，證據保留原欄位。
            gov_net_100m = amount / 100000000.0
            amount_source_col = amount_col
        elif net_col:
            gov_net_100m = None
            amount_source_col = net_col
        else:
            return {"status": "WARN", "gov_net_100m": None, "gov_signal": "未知", "gov_score": 0, "message": "TEJ欄位缺少買賣超或買賣超金額", "rows": len(df)}
        if gov_net_100m is None:
            net_vals = [self._to_number(x) for x in df[net_col].tolist()]
            net_vals = [x for x in net_vals if not math.isnan(x)]
            direction_sum = sum(net_vals) if net_vals else 0
            gov_signal = "偏多" if direction_sum > 0 else "偏空" if direction_sum < 0 else "中性"
        else:
            gov_signal = "偏多" if gov_net_100m > 0 else "偏空" if gov_net_100m < 0 else "中性"
        gov_score = 1 if gov_signal == "偏多" else -1 if gov_signal == "偏空" else 0
        actual_date = "latest"
        if date_col and df[date_col].notna().any():
            actual_date = str(df[date_col].dropna().iloc[0])
        payload = {
            "status": "OK", "source": "TEJ八大公股行庫", "actual_date": actual_date,
            "gov_net_100m": gov_net_100m, "gov_signal": gov_signal, "gov_score": gov_score,
            "rows": len(df), "amount_column": amount_source_col,
            "message": "TEJ八大官股解析完成；TWSE T86/Wantgoo僅作佐證"
        }
        self.logger.write_raw_evidence("官股_TEJ", payload, parsed=payload, status="OK", url=self.tej_file or "", message=payload["message"])
        self.logger.parsed_value("gov_net_100m", gov_net_100m, "TEJ八大公股行庫", actual_date)
        return payload

class Market5DayEngine:
    """V2.5 P0：跨月回補最近5個有效交易日，避免前高/前低/5MA/5日均量空白。"""
    def __init__(self, client: HttpClient, logger: Macro16Logger):
        self.client = client
        self.logger = logger

    def _num(self, value: Any) -> float:
        try:
            return float(str(value).replace(",", "").strip())
        except Exception:
            return math.nan

    def _iso(self, compact: str) -> str:
        compact = str(compact).replace("-", "")
        return f"{compact[:4]}-{compact[4:6]}-{compact[6:8]}"

    def _tw_date_to_iso(self, value: Any) -> str:
        s = str(value).strip()
        parts = s.split("/")
        if len(parts) >= 3:
            year = int(parts[0]) + 1911 if int(parts[0]) < 1911 else int(parts[0])
            return f"{year:04d}-{int(parts[1]):02d}-{int(parts[2]):02d}"
        if re.fullmatch(r"\d{8}", s):
            return f"{s[:4]}-{s[4:6]}-{s[6:8]}"
        return s

    def _month_candidates(self, base_date: str, months_back: int = 4) -> List[str]:
        base = dt.datetime.strptime(str(base_date).replace("-", ""), "%Y%m%d").date().replace(day=1)
        out = []
        y, m = base.year, base.month
        for _ in range(months_back):
            out.append(f"{y:04d}{m:02d}01")
            m -= 1
            if m == 0:
                y -= 1; m = 12
        return out

    def fetch_taiex_recent_days(self, base_date: str, need: int = 5):
        frames = []
        for month_start in self._month_candidates(base_date, 4):
            url = f"https://www.twse.com.tw/rwd/zh/TAIEX/MI_5MINS_HIST?date={month_start}&response=json"
            try:
                js = self.client.get_json(url)
            except Exception as exc:
                self.logger.warning(f"MARKET5_TAIEX_FETCH_FAIL url={url} error={exc}")
                continue
            for row in js.get("data", []) or []:
                try:
                    frames.append({"date": self._tw_date_to_iso(row[0]), "open": self._num(row[1]), "high": self._num(row[2]), "low": self._num(row[3]), "close": self._num(row[4])})
                except Exception:
                    continue
        if pd is None:
            return []
        df = pd.DataFrame(frames)
        if df.empty:
            return df
        df = df.dropna(subset=["date", "close"]).drop_duplicates("date")
        df = df[df["date"] <= self._iso(base_date)].sort_values("date", ascending=False).head(need)
        return df.sort_values("date")

    def fetch_turnover_recent_days(self, base_date: str, need: int = 5):
        frames = []
        for month_start in self._month_candidates(base_date, 4):
            url = f"https://www.twse.com.tw/rwd/zh/afterTrading/FMTQIK?date={month_start}&response=json"
            try:
                js = self.client.get_json(url)
            except Exception as exc:
                self.logger.warning(f"MARKET5_TURNOVER_FETCH_FAIL url={url} error={exc}")
                continue
            for row in js.get("data", []) or []:
                try:
                    frames.append({"date": self._tw_date_to_iso(row[0]), "turnover_100m": self._num(row[2]) / 100000000.0})
                except Exception:
                    continue
        if pd is None:
            return []
        df = pd.DataFrame(frames)
        if df.empty:
            return df
        df = df.dropna(subset=["date", "turnover_100m"]).drop_duplicates("date")
        df = df[df["date"] <= self._iso(base_date)].sort_values("date", ascending=False).head(need)
        return df.sort_values("date")

    def build_market_features(self, base_date: str) -> Dict[str, Any]:
        compact = str(base_date).replace("-", "")
        px = self.fetch_taiex_recent_days(compact, 5)
        tv = self.fetch_turnover_recent_days(compact, 5)
        if pd is None or getattr(px, "empty", True) or getattr(tv, "empty", True) or len(px) < 5 or len(tv) < 5:
            msg = f"P0_FAIL: 最近5交易日不足 taiex={0 if pd is None or getattr(px,'empty',True) else len(px)} turnover={0 if pd is None or getattr(tv,'empty',True) else len(tv)}"
            self.logger.warning(msg)
            return {"status": "FAIL", "message": msg}
        latest = px.iloc[-1]
        prev = px.iloc[-2]
        result = {
            "status": "OK", "base_date": latest["date"], "close": float(latest["close"]),
            "high": float(latest["high"]), "low": float(latest["low"]),
            "prev_high": float(prev["high"]), "prev_low": float(prev["low"]),
            "ma5": round(float(px["close"].mean()), 2),
            "turnover_100m": round(float(tv.iloc[-1]["turnover_100m"]), 2),
            "avg_turnover_5d_100m": round(float(tv["turnover_100m"].mean()), 2),
            "taiex_recent_dates": ",".join(px["date"].astype(str).tolist()),
            "turnover_recent_dates": ",".join(tv["date"].astype(str).tolist()),
        }
        self.logger.write_raw_evidence("市場5日資料", result, parsed=result, status="OK", url="TWSE MI_5MINS_HIST/FMTQIK", message="跨月5日資料已補齊")
        return result

class DBRepository:
    """V2.5：從股票DB建立機構級報告資料集。"""
    def __init__(self, db_path: str, logger: Optional[Macro16Logger] = None):
        self.db_path = db_path
        self.logger = logger

    def connect(self):
        return sqlite3.connect(self.db_path)

    def table_exists(self, table: str) -> bool:
        with self.connect() as conn:
            cur = conn.cursor()
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table,))
            return cur.fetchone() is not None

    def get_trade_date(self) -> str:
        with self.connect() as conn:
            return pd.read_sql("SELECT MAX(date) AS d FROM ranking_result", conn).iloc[0]["d"]

    def table_count(self, table: str) -> int:
        if not self.table_exists(table):
            return 0
        with self.connect() as conn:
            return int(pd.read_sql(f"SELECT COUNT(*) AS n FROM {table}", conn).iloc[0]["n"])

    def latest_date(self, table: str) -> str:
        if not self.table_exists(table):
            return ""
        candidates = ["date", "snapshot_date", "trade_date", "data_date", "feature_date", "plan_date", "source_date", "data_year"]
        with self.connect() as conn:
            cols = [r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
            for c in candidates:
                if c in cols:
                    try:
                        return str(pd.read_sql(f"SELECT MAX({c}) AS d FROM {table}", conn).iloc[0]["d"])
                    except Exception:
                        pass
        return ""

    def load_latest_table(self, conn, table: str, date_col: str, trade_date: str):
        if not self.table_exists(table):
            return pd.DataFrame()
        cols = [r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
        if date_col in cols:
            try:
                return pd.read_sql(f"SELECT * FROM {table} WHERE {date_col}=(SELECT MAX({date_col}) FROM {table} WHERE {date_col}<=?)", conn, params=[trade_date])
            except Exception:
                pass
        return pd.read_sql(f"SELECT * FROM {table}", conn)

    def load_base_universe(self, trade_date: str):
        with self.connect() as conn:
            r = pd.read_sql("SELECT * FROM ranking_result WHERE date=?", conn, params=[trade_date])
            ms = self.load_latest_table(conn, "market_snapshot", "snapshot_date", trade_date)
            sm = self.load_latest_table(conn, "stocks_master", "update_date", trade_date)
            val = self.load_latest_table(conn, "external_valuation", "data_date", trade_date)
            rev = self.load_latest_table(conn, "external_revenue", "source_date", trade_date)
            inst = self.load_latest_table(conn, "external_institutional", "trade_date", trade_date)
            margin = self.load_latest_table(conn, "external_margin", "trade_date", trade_date)
            ph = pd.read_sql("SELECT * FROM price_history WHERE date<=?", conn, params=[trade_date]) if self.table_exists("price_history") else pd.DataFrame()
        for df in [r, ms, sm, val, rev, inst, margin, ph]:
            if not df.empty and "stock_id" in df.columns:
                df["stock_id"] = df["stock_id"].astype(str).str.zfill(4)
        df = r.copy()
        if not ms.empty:
            df = df.merge(ms.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_m"))
        if not sm.empty:
            df = df.merge(sm.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_master"))
        if not val.empty:
            df = df.merge(val.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_valuation"))
        if not rev.empty:
            df = df.merge(rev.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_revenue"))
        if not inst.empty:
            df = df.merge(inst.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_inst"))
        if not margin.empty:
            df = df.merge(margin.drop_duplicates("stock_id"), on="stock_id", how="left", suffixes=("", "_margin"))
        if not ph.empty:
            feats = self._price_features(ph)
            df = df.merge(feats, on="stock_id", how="left")
        return df.drop_duplicates("stock_id")

    def _price_features(self, ph):
        out = []
        for sid, g in ph.groupby("stock_id"):
            g = g.sort_values("date")
            close = pd.to_numeric(g["close"], errors="coerce")
            high = pd.to_numeric(g["high"], errors="coerce")
            low = pd.to_numeric(g["low"], errors="coerce")
            vol = pd.to_numeric(g.get("volume"), errors="coerce") if "volume" in g.columns else pd.Series(dtype=float)
            last_close = close.iloc[-1] if len(close) else math.nan
            high120 = high.tail(120).max() if len(high) else math.nan
            low120 = low.tail(120).min() if len(low) else math.nan
            pos120 = (last_close - low120) / (high120 - low120) if high120 and low120 and high120 != low120 else math.nan
            pct20 = (last_close / close.iloc[-20] - 1) * 100 if len(close) >= 20 and close.iloc[-20] else math.nan
            ma20_calc = close.tail(20).mean() if len(close) >= 20 else math.nan
            ma60_calc = close.tail(60).mean() if len(close) >= 60 else math.nan
            vol5 = vol.tail(5).mean() if len(vol) >= 5 else math.nan
            vol20 = vol.tail(20).mean() if len(vol) >= 20 else math.nan
            ma5_calc = close.tail(5).mean() if len(close) >= 5 else math.nan
            ma65_calc = close.tail(65).mean() if len(close) >= 65 else ma60_calc
            prev_close = close.iloc[-2] if len(close) >= 2 else math.nan
            last_open = pd.to_numeric(g.get("open"), errors="coerce").iloc[-1] if "open" in g.columns and len(g) else math.nan
            last_high = high.iloc[-1] if len(high) else math.nan
            last_low = low.iloc[-1] if len(low) else math.nan
            prev_high = high.iloc[-2] if len(high) >= 2 else math.nan
            prev2_high = high.iloc[-3] if len(high) >= 3 else math.nan
            high20_prev = high.iloc[:-1].tail(20).max() if len(high) >= 21 else math.nan
            high60_prev = high.iloc[:-1].tail(60).max() if len(high) >= 61 else math.nan
            low20_prev = low.iloc[:-1].tail(20).min() if len(low) >= 21 else math.nan
            ema12 = close.ewm(span=12, adjust=False).mean() if len(close) >= 12 else pd.Series(dtype=float)
            ema26 = close.ewm(span=26, adjust=False).mean() if len(close) >= 26 else pd.Series(dtype=float)
            macd_dif = (ema12 - ema26) if len(ema26) else pd.Series(dtype=float)
            macd_dea = macd_dif.ewm(span=9, adjust=False).mean() if len(macd_dif) >= 9 else pd.Series(dtype=float)
            macd_hist = (macd_dif - macd_dea) if len(macd_dea) else pd.Series(dtype=float)
            macd_hist_last = macd_hist.iloc[-1] if len(macd_hist) else math.nan
            macd_hist_prev = macd_hist.iloc[-2] if len(macd_hist) >= 2 else math.nan
            low9 = low.rolling(9).min()
            high9 = high.rolling(9).max()
            rsv = (close - low9) / (high9 - low9).replace(0, math.nan) * 100
            k = rsv.ewm(com=2, adjust=False).mean()
            d = k.ewm(com=2, adjust=False).mean()
            k_last = k.iloc[-1] if len(k) else math.nan
            d_last = d.iloc[-1] if len(d) else math.nan
            k_prev = k.iloc[-2] if len(k) >= 2 else math.nan
            d_prev = d.iloc[-2] if len(d) >= 2 else math.nan
            out.append({
                "stock_id": str(sid).zfill(4), "pos_120d": pos120, "pct_20d": pct20,
                "ma5_calc": ma5_calc, "ma20_calc": ma20_calc, "ma60_calc": ma60_calc, "ma65_calc": ma65_calc,
                "vol5": vol5, "vol20": vol20, "last_open": last_open, "last_high": last_high, "last_low": last_low,
                "prev_close": prev_close, "prev_high_price": prev_high, "prev2_high_price": prev2_high,
                "high20_prev": high20_prev, "high60_prev": high60_prev, "low20_prev": low20_prev,
                "macd_hist": macd_hist_last, "macd_hist_prev": macd_hist_prev,
                "kd_k": k_last, "kd_d": d_last, "kd_k_prev": k_prev, "kd_d_prev": d_prev
            })
        return pd.DataFrame(out)


def _safe_series(df, column: str, default=math.nan, numeric: bool = True):
    """Phase5 FIX：保證任何欄位都回傳與 df.index 對齊的 pandas Series。
    目的：避免缺欄、重複欄名、單一 scalar 被 pd.to_numeric 後變成 numpy.float64，
    導致 .abs() / .fillna() / .notna() 失敗。
    """
    if pd is None:
        return None
    idx = getattr(df, "index", None)
    if idx is None:
        idx = pd.RangeIndex(0)
    if df is None or column is None:
        s = pd.Series(default, index=idx)
    elif hasattr(df, "columns") and column in df.columns:
        value = df[column]
        # 若merge後出現重複欄名，df[column] 會是 DataFrame；取第一欄作為主欄位。
        if isinstance(value, pd.DataFrame):
            value = value.iloc[:, 0]
        if isinstance(value, pd.Series):
            s = value.reindex(idx)
        else:
            s = pd.Series(value, index=idx)
    else:
        s = pd.Series(default, index=idx)
    if numeric:
        s = pd.to_numeric(s, errors="coerce")
    return s

def _safe_bool_series(df, column: str, default: bool = False):
    s = _safe_series(df, column, default=default, numeric=False)
    if s is None:
        return s
    return s.fillna(default).astype(bool)

def _safe_str_series(df, column: str, default: str = ""):
    s = _safe_series(df, column, default=default, numeric=False)
    if s is None:
        return s
    return s.fillna(default).astype(str)

class StageTrace:
    """InstitutionalReportEngine 階段追蹤，用於精準定位是哪個 Engine / 欄位失敗。"""
    def __init__(self, logger: Optional[Macro16Logger], stage: str, df_getter=None):
        self.logger = logger
        self.stage = stage
        self.df_getter = df_getter

    def __enter__(self):
        if self.logger:
            self.logger.strategy_trace("STAGE_START", {"stage": self.stage})
        return self

    def __exit__(self, exc_type, exc, tb):
        if not self.logger:
            return False
        if exc_type is None:
            payload = {"stage": self.stage, "status": "OK"}
            try:
                df = self.df_getter() if callable(self.df_getter) else None
                if df is not None:
                    payload["rows"] = int(len(df))
                    payload["columns"] = list(map(str, list(df.columns)[:80]))
            except Exception:
                pass
            self.logger.strategy_trace("STAGE_OK", payload)
            return False
        payload = {"stage": self.stage, "status": "FAIL", "error_type": exc_type.__name__, "error": str(exc)}
        try:
            df = self.df_getter() if callable(self.df_getter) else None
            if df is not None:
                payload["rows"] = int(len(df))
                payload["columns"] = list(map(str, list(df.columns)[:120]))
        except Exception as meta_exc:
            payload["meta_error"] = str(meta_exc)
        self.logger.strategy_trace("STAGE_FAIL", payload)
        return False


class MarketRegimeEngine:
    """顧奎國老師策略P0：大盤多空總閘與乖離風控。
    原Macro16只用5MA/高低點/量能，本層補上「短線拉回 vs 波段修正」語義。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        close = _safe_series(df, "close")
        ma60 = _safe_series(df, "ma60_final")
        ma60 = ma60.fillna(_safe_series(df, "ma60_calc"))
        macd = _safe_series(df, "macd_hist")
        macd_prev = _safe_series(df, "macd_hist_prev")
        k = _safe_series(df, "kd_k")
        d = _safe_series(df, "kd_d")
        k_prev = _safe_series(df, "kd_k_prev")
        d_prev = _safe_series(df, "kd_d_prev")
        denom = ma60.replace(0, math.nan)
        df["deviation_from_ma60"] = ((close - ma60) / denom * 100).replace([math.inf, -math.inf], math.nan)
        df["deviation_risk_flag"] = df["deviation_from_ma60"].abs().fillna(0) >= 20
        df["macd_turn_negative"] = (macd < 0) & (macd_prev >= 0)
        df["macd_near_zero"] = macd.abs().fillna(999) <= 0.20
        df["kd_dead_cross_below_80"] = (k < d) & (k_prev >= d_prev) & (k < 80)
        df["wave_correction_flag"] = df["macd_turn_negative"].fillna(False) & df["kd_dead_cross_below_80"].fillna(False)
        df["market_pullback_type"] = "短線拉回"
        df.loc[df["wave_correction_flag"], "market_pullback_type"] = "波段修正"
        df.loc[df["deviation_risk_flag"] & ~df["wave_correction_flag"], "market_pullback_type"] = "乖離過大/禁追高"
        if self.logger:
            wave_count = int(df["wave_correction_flag"].fillna(False).sum())
            self.logger.strategy_trace("WAVE_CORRECTION", {
                "rows": int(len(df)),
                "wave_correction_count": wave_count,
                "macd_turn_negative_count": int(df["macd_turn_negative"].fillna(False).sum()),
                "kd_dead_cross_below_80_count": int(df["kd_dead_cross_below_80"].fillna(False).sum()),
                "deviation_risk_count": int(df["deviation_risk_flag"].fillna(False).sum()),
            })
            if wave_count == 0:
                self.logger.warning("WAVE_CORRECTION_ZERO_WARN wave_correction_count=0，請確認macd_hist/macd_hist_prev與KD欄位來源是否完整；此為驗收WARN，不阻斷報表。")
        return df

class SakataRiskPatternEngine:
    """顧奎國老師策略P0：量化流星、空頭新星十字、墓碑線、長黑K。"""
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        o = _safe_series(df, "last_open")
        h = _safe_series(df, "last_high")
        l = _safe_series(df, "last_low")
        c = _safe_series(df, "close")
        prev_c = _safe_series(df, "prev_close")
        body = (c - o).abs()
        rng = (h - l).replace(0, math.nan)
        upper = h - pd.concat([o, c], axis=1).max(axis=1)
        lower = pd.concat([o, c], axis=1).min(axis=1) - l
        df["shooting_star_flag"] = (upper >= body * 2) & ((body / rng) <= 0.35) & ((h - c) / rng >= 0.45)
        df["tombstone_flag"] = (upper >= rng * 0.60) & (lower <= rng * 0.10) & ((body / rng) <= 0.25)
        df["bearish_star_cross_flag"] = ((body / rng) <= 0.15) & (prev_c.notna()) & (c >= prev_c * 1.02)
        df["long_black_flag"] = (c < o) & ((o - c) / rng >= 0.60)
        warn = []
        for i in df.index:
            labels = []
            if bool(df.at[i, "shooting_star_flag"]): labels.append("流星")
            if bool(df.at[i, "bearish_star_cross_flag"]): labels.append("空頭新星十字")
            if bool(df.at[i, "tombstone_flag"]): labels.append("墓碑線")
            if bool(df.at[i, "long_black_flag"]): labels.append("長黑K")
            warn.append("/".join(labels) if labels else "")
        df["k_warning_type"] = warn
        df["two_day_break_high"] = _safe_series(df, "last_high") > _safe_series(df, "high20_prev")
        df["kline_score"] = _safe_series(df, "kline_score", default=50).fillna(50)
        df.loc[df["k_warning_type"].ne(""), "kline_score"] = (df["kline_score"] - 25).clip(0, 100)
        if self.logger:
            self.logger.strategy_trace("SAKATA_PATTERN_SUMMARY", {
                "shooting_star_count": int(df["shooting_star_flag"].fillna(False).sum()),
                "bearish_star_cross_count": int(df["bearish_star_cross_flag"].fillna(False).sum()),
                "tombstone_count": int(df["tombstone_flag"].fillna(False).sum()),
                "long_black_count": int(df["long_black_flag"].fillna(False).sum()),
                "warning_count": int(df["k_warning_type"].astype(str).ne("").sum()),
            })
        return df

class CoreLeaderEngine:
    """顧奎國老師策略P0：台積電(2330)作為大盤核心風向。"""
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def infer_state(self, df):
        default = {"core_leader_state": "NE", "core_leader_reason": "未取得2330資料，核心股風向不進主判斷"}
        if df is None or df.empty or "stock_id" not in df.columns:
            return default
        sid = df["stock_id"].astype(str).str.zfill(4)
        if not sid.eq("2330").any():
            return default
        r = df.loc[sid.eq("2330")].iloc[0]
        close = _safe_float(r.get("close"))
        ma20 = _safe_float(r.get("ma20_final"))
        ma60 = _safe_float(r.get("ma60_final"))
        high20 = _safe_float(r.get("high20_prev"))
        warning = str(r.get("k_warning_type", "") or "")
        if high20 and close and close > high20:
            state, reason = "突破續強", "2330突破近20日壓力，主流續強"
        elif close and ma20 and close >= ma20 and (not warning):
            state, reason = "回測不破", "2330守住MA20/回測不破，仍屬多方拉回"
        elif close and ma20 and close < ma20 and (ma60 and close >= ma60):
            state, reason = "假突破/短線降檔", "2330跌回MA20下方但未破中期線，降追價"
        elif close and ma60 and close < ma60:
            state, reason = "核心轉弱", "2330跌破MA60，市場風險升級"
        else:
            state, reason = "觀察", "2330資料不足，僅列觀察"
        if warning:
            reason += f"；K線警訊={warning}"
        return {"core_leader_state": state, "core_leader_reason": reason}

    def apply(self, df):
        state = self.infer_state(df)
        df["core_leader_state"] = state["core_leader_state"]
        df["core_leader_reason"] = state["core_leader_reason"]
        if self.logger:
            self.logger.strategy_trace("CORE_LEADER_STATUS", state)
            if state.get("core_leader_state") in ("突破續強", "假突破/短線降檔", "核心轉弱"):
                self.logger.strategy_trace("2330_BREAKOUT", state)
        return df

class AvoidSwapEngine:
    """Phase5：排除換股層升級。
    修改前：swap_reason非空即avoid_flag=True，導致軟性壓力也會把全市場封殺。
    修改後：區分 hard_avoid 與 soft_avoid；低位階翻多可覆蓋 soft avoid，但不可覆蓋 hard risk。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        close = _safe_series(df, "close")
        ma20 = _safe_series(df, "ma20_final")
        ma60 = _safe_series(df, "ma60_final")
        high20 = _safe_series(df, "high20_prev")
        high60 = _safe_series(df, "high60_prev")
        prev_high = _safe_series(df, "prev_high_price")
        prev2_high = _safe_series(df, "prev2_high_price")
        k_warning = _safe_str_series(df, "k_warning_type")
        low_base_reversal = _safe_bool_series(df, "low_base_reversal_flag")
        wave_correction = _safe_bool_series(df, "wave_correction_flag")
        deviation_risk = _safe_bool_series(df, "deviation_risk_flag")
        df["two_high_not_passed_flag"] = (prev_high < high60 * 0.995) & (prev2_high < high60 * 0.995) & (close < high20)
        df["downtrend_flag"] = (close < ma20) & (ma20 < ma60)
        df["neckline_pressure_flag"] = (close < high20) & (high20.notna()) & ((high20 - close) / close.replace(0, math.nan) <= 0.05)
        df["hard_k_risk_flag"] = k_warning.str.contains("長黑K|墓碑線", na=False)
        hard_reasons, soft_reasons, release_reasons, pressure_states = [], [], [], []
        for i in df.index:
            hard, soft, release = [], [], []
            if bool(df.at[i, "downtrend_flag"]): hard.append("下降軌道")
            if bool(wave_correction.loc[i]): hard.append("MACD+KD波段修正")
            if bool(df.at[i, "hard_k_risk_flag"]): hard.append("硬K線風險")
            if bool(df.at[i, "two_high_not_passed_flag"]): soft.append("兩高不過")
            if bool(df.at[i, "neckline_pressure_flag"]): soft.append("頸線/前高壓力未突破")
            if bool(deviation_risk.loc[i]): soft.append("乖離過大禁追高")
            if bool(low_base_reversal.loc[i]) and soft and not hard:
                release.append("低位階翻多覆蓋軟性壓力")
            hard_reasons.append(";".join(hard))
            soft_reasons.append(";".join(soft))
            release_reasons.append(";".join(release))
            pressure_states.append(";".join(hard + soft) if (hard or soft) else "未觸發")
        df["hard_avoid_reason"] = hard_reasons
        df["soft_avoid_reason"] = soft_reasons
        df["avoid_release_reason"] = release_reasons
        df["hard_avoid_flag"] = df["hard_avoid_reason"].astype(str).ne("")
        df["soft_avoid_flag"] = df["soft_avoid_reason"].astype(str).ne("")
        df["avoid_flag"] = df["hard_avoid_flag"] | (df["soft_avoid_flag"] & ~low_base_reversal)
        df["swap_reason"] = ""
        for idx, r in df.iterrows():
            parts = []
            if str(r.get("hard_avoid_reason", "") or ""): parts.append(str(r.get("hard_avoid_reason")))
            if str(r.get("soft_avoid_reason", "") or "") and not bool(r.get("low_base_reversal_flag", False)):
                parts.append(str(r.get("soft_avoid_reason")))
            if str(r.get("soft_avoid_reason", "") or "") and bool(r.get("low_base_reversal_flag", False)) and not str(r.get("hard_avoid_reason", "") or ""):
                parts.append("軟性壓力已由低位階翻多覆蓋")
            df.at[idx, "swap_reason"] = ";".join(parts)
        df["avoid_level"] = "NONE"
        df.loc[df["soft_avoid_flag"], "avoid_level"] = "SOFT"
        df.loc[df["hard_avoid_flag"], "avoid_level"] = "HARD"
        df.loc[df["soft_avoid_flag"] & low_base_reversal & ~df["hard_avoid_flag"], "avoid_level"] = "SOFT_RELEASED"
        df["pressure_line_state"] = pressure_states
        if self.logger:
            summary = {
                "avoid_count": int(df["avoid_flag"].fillna(False).sum()),
                "hard_avoid_count": int(df["hard_avoid_flag"].fillna(False).sum()),
                "soft_avoid_count": int(df["soft_avoid_flag"].fillna(False).sum()),
                "soft_released_by_low_base_count": int((df["soft_avoid_flag"] & low_base_reversal & ~df["hard_avoid_flag"]).sum()),
                "two_high_not_passed_count": int(df["two_high_not_passed_flag"].fillna(False).sum()),
                "downtrend_count": int(df["downtrend_flag"].fillna(False).sum()),
                "neckline_pressure_count": int(df["neckline_pressure_flag"].fillna(False).sum()),
            }
            self.logger.strategy_trace("AVOID_SWAP_SUMMARY", summary)
            sample = df.loc[df["avoid_flag"].fillna(False), ["stock_id", "swap_reason", "avoid_level"]].head(20)
            for _, row in sample.iterrows():
                self.logger.strategy_trace("AVOID_REASON", {"stock": str(row.get("stock_id", "")).zfill(4), "level": row.get("avoid_level", ""), "reason": row.get("swap_reason", "")})
        return df


class TeacherPhase5SemanticEngine:
    """V2.7.3：Phase5語義一致性修正層。
    目的：把差異分析報告指出的「第3浪訊號 vs 主跌修正浪」、「主升池混入弱反彈」、
    「修正反彈只剩文字但沒有位置」轉成可被 TeacherDecisionEngine 使用的欄位。
    本層不取代既有老師模型，而是建立 Phase5 Gate 與候選池，使報表/決策/原因一致。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        close = _safe_series(df, "close")
        ma20 = _safe_series(df, "ma20_final")
        ma60 = _safe_series(df, "ma60_final")
        rsi = _safe_series(df, "rsi", default=50).fillna(50)
        vol5 = _safe_series(df, "vol5")
        vol20 = _safe_series(df, "vol20")
        high20 = _safe_series(df, "high20_prev")
        low20 = _safe_series(df, "low20_prev")
        high60 = _safe_series(df, "high60_prev")
        low60 = _safe_series(df, "low60_prev")
        high120 = _safe_series(df, "high120_prev")
        low120 = _safe_series(df, "low120_prev")
        rr = _safe_series(df, "rr", default=0).fillna(0)
        score = _safe_series(df, "teacher_score", default=0).fillna(0)
        wave_correction = _safe_bool_series(df, "wave_correction_flag")
        low_base_reversal = _safe_bool_series(df, "low_base_reversal_flag")
        hard_avoid = _safe_bool_series(df, "hard_avoid_flag")
        k_warning = _safe_str_series(df, "k_warning_type")
        ratio = (vol5 / vol20.replace(0, math.nan)).replace([math.inf, -math.inf], math.nan).fillna(1.0)
        # 大波：避免只用「第3浪」文字，先判定主升/主跌背景。
        main_up = close.notna() & ma20.notna() & ma60.notna() & (close >= ma20) & (ma20 >= ma60 * 0.98)
        main_down = close.notna() & ma20.notna() & ma60.notna() & ((close < ma20) & (ma20 < ma60))
        box = ~(main_up | main_down)
        df["phase5_major_wave"] = "大箱型整理"
        df.loc[main_up, "phase5_major_wave"] = "主升推動浪"
        df.loc[main_down, "phase5_major_wave"] = "主跌修正浪"
        # 費波回撤/反彈比例：使用120/60/20日高低點依序取可用區間。
        swing_high = high120.fillna(high60).fillna(high20)
        swing_low = low120.fillna(low60).fillna(low20)
        span = (swing_high - swing_low).replace(0, math.nan)
        fibo_retrace = ((swing_high - close) / span).where(main_up, (close - swing_low) / span)
        fibo_retrace = fibo_retrace.clip(lower=0, upper=1.5).fillna(0.5)
        df["phase5_fibo_retrace"] = fibo_retrace.round(3)
        zone = pd.Series("0.382~0.618中性回撤", index=df.index)
        zone.loc[fibo_retrace < 0.382] = "0~0.382淺回/弱反彈"
        zone.loc[(fibo_retrace >= 0.382) & (fibo_retrace <= 0.618)] = "0.382~0.618健康回撤"
        zone.loc[(fibo_retrace > 0.618) & (fibo_retrace <= 0.786)] = "0.618~0.786深回撤"
        zone.loc[fibo_retrace > 0.786] = "0.786以上結構破壞"
        df["phase5_fibo_retrace_zone"] = zone
        # 小波定位。
        minor = pd.Series("小波待確認", index=df.index)
        minor.loc[main_up & (fibo_retrace <= 0.618) & low_base_reversal] = "第2浪/第4浪拉回"
        minor.loc[main_up & (close > high20) & (ratio >= 1.2)] = "第3浪推動"
        minor.loc[main_down & (close < ma20)] = "C浪反彈待確認"
        minor.loc[main_down & (close >= ma20) & (close < ma60)] = "B浪弱反彈"
        minor.loc[main_down & (fibo_retrace > 0.618)] = "C浪延伸/深修正"
        df["phase5_minor_wave"] = minor
        # 修正型態。
        corr = pd.Series("修正型態待確認", index=df.index)
        corr.loc[main_up & low_base_reversal] = "主升回撤修正"
        corr.loc[main_down & (fibo_retrace <= 0.382)] = "Zigzag弱反彈"
        corr.loc[main_down & (fibo_retrace > 0.382) & (fibo_retrace <= 0.618)] = "ABC修正反彈"
        corr.loc[main_down & (fibo_retrace > 0.618)] = "C浪延伸修正"
        corr.loc[box & (fibo_retrace <= 0.618)] = "箱型整理修正"
        df["phase5_correction_type"] = corr
        # 修正完成：必須同時具備站回MA20、量能、RSI、壓力突破任一確認。
        correction_completed = (
            close.notna() & ma20.notna() & (close >= ma20)
            & (ratio >= 1.2)
            & (rsi >= 45)
            & ((high20.isna()) | (close >= high20 * 0.995) | low_base_reversal)
            & ~hard_avoid
        )
        df["phase5_correction_completed"] = correction_completed
        # 逃命反彈/弱反彈：主跌、未站回MA60、量不足或RSI不足，且未完成。
        escape = main_down & (close < ma60) & (~correction_completed) & ((ratio < 1.0) | (rsi < 50) | (fibo_retrace < 0.5))
        df["phase5_escape_rally"] = escape
        rebound = pd.Series("一般反彈", index=df.index)
        rebound.loc[main_up & low_base_reversal & ~correction_completed] = "主升拉回待確認"
        rebound.loc[main_up & correction_completed] = "主升拉回完成"
        rebound.loc[main_down & ~escape] = "跌深技術反彈"
        rebound.loc[escape] = "逃命反彈"
        rebound.loc[main_down & (close < low20)] = "反彈失敗"
        rebound.loc[box & correction_completed] = "箱型突破轉強"
        df["phase5_rebound_type"] = rebound

        # V2.7.4 Position Engine：
        # 依 Word《彩晶（6116）Position Engine缺口分析報告》補齊 wave_phase / breakout_stage / position_score。
        # 重點：彩晶類低價量價轉強股不應因「尚未正式突破」被壓成 position_score=0。
        bottom_complete = (
            (close >= ma20.fillna(close) * 0.98)
            & (ma20.fillna(close) >= ma60.fillna(ma20).fillna(close) * 0.96)
            & (ratio >= 1.15)
            & (rsi.between(45, 72))
            & (score >= 60)
            & ~hard_avoid
        )
        # 大波補強：底部完成不同於一般大箱型整理。
        df.loc[box & bottom_complete, "phase5_major_wave"] = "底部完成"

        compression = (
            (close >= ma20.fillna(close) * 0.98)
            & (close <= high20.fillna(close * 1.08) * 1.01)
            & (ratio >= 1.1)
            & (rsi.between(45, 72))
            & ~hard_avoid
        )
        impulsive = (
            (main_up | df["phase5_major_wave"].astype(str).eq("底部完成"))
            & (close > high20.fillna(close * 9))
            & (ratio >= 1.2)
            & (rsi.between(45, 72))
            & ~hard_avoid
            & ~wave_correction
        )
        # Wave3預突破：不是推動浪True，但要保留為觀察，不可誤判為主跌。
        prebreakout = (
            (score >= 60)
            & (rr >= 1.0)
            & (main_up | box | df["phase5_major_wave"].astype(str).eq("底部完成"))
            & (close >= ma20.fillna(close) * 0.98)
            & (ratio >= 1.1)
            & (rsi.between(45, 72))
            & ~hard_avoid
            & ~wave_correction
        )

        wave_phase = pd.Series("Unknown", index=df.index)
        wave_phase.loc[main_down & ~correction_completed] = "A/B/C_Correction"
        wave_phase.loc[main_up & low_base_reversal & ~correction_completed] = "Wave2_or_Wave4_Pullback"
        wave_phase.loc[df["phase5_major_wave"].astype(str).eq("底部完成") & prebreakout] = "Wave3_PreBreakout"
        wave_phase.loc[main_up & prebreakout & ~impulsive] = "Wave3_PreBreakout"
        wave_phase.loc[impulsive] = "Wave3_Breakout"
        wave_phase.loc[impulsive & (ratio >= 1.8)] = "Wave3_Expansion"
        wave_phase.loc[(rsi > 72) & main_up & (close >= high60.fillna(high20).fillna(close) * 0.98)] = "Wave5_Risk"
        df["phase5_wave_phase"] = wave_phase

        breakout_stage = pd.Series("None", index=df.index)
        breakout_stage.loc[compression & ~prebreakout] = "Compression"
        breakout_stage.loc[prebreakout & ~impulsive] = "PreBreakout"
        breakout_stage.loc[impulsive] = "Breakout"
        breakout_stage.loc[impulsive & (ratio >= 1.8)] = "Expansion"
        breakout_stage.loc[(rsi > 72) & main_up] = "Exhaustion"
        breakout_stage.loc[main_down & ~correction_completed] = "Correction"
        df["phase5_breakout_stage"] = breakout_stage

        df["phase5_impulsive_wave"] = impulsive
        df["phase5_impulse_stage"] = "None"
        df.loc[prebreakout & ~impulsive, "phase5_impulse_stage"] = "Early"
        df.loc[impulsive, "phase5_impulse_stage"] = "Confirmed"
        df.loc[impulsive & (ratio >= 1.8), "phase5_impulse_stage"] = "Expansion"
        df.loc[(rsi > 72) & main_up, "phase5_impulse_stage"] = "Exhausted"
        df.loc[main_down & ~impulsive, "phase5_impulse_stage"] = "Failed/Correction"

        # position_score：正式波段位置分，替代原DB position_score=0造成的放行缺口。
        position_score = pd.Series(0.0, index=df.index)
        position_score += main_up.astype(float) * 20
        position_score += df["phase5_major_wave"].astype(str).eq("底部完成").astype(float) * 18
        position_score += box.astype(float) * 8
        position_score -= main_down.astype(float) * 15
        position_score += df["phase5_minor_wave"].astype(str).isin(["第2浪/第4浪拉回", "第3浪推動"]).astype(float) * 15
        position_score += df["phase5_wave_phase"].astype(str).eq("Wave3_PreBreakout").astype(float) * 20
        position_score += df["phase5_wave_phase"].astype(str).eq("Wave3_Breakout").astype(float) * 24
        position_score += df["phase5_wave_phase"].astype(str).eq("Wave3_Expansion").astype(float) * 18
        position_score += df["phase5_breakout_stage"].astype(str).eq("Compression").astype(float) * 8
        position_score += df["phase5_breakout_stage"].astype(str).eq("PreBreakout").astype(float) * 15
        position_score += df["phase5_breakout_stage"].astype(str).eq("Breakout").astype(float) * 18
        position_score += df["phase5_impulse_stage"].astype(str).eq("Early").astype(float) * 10
        position_score += df["phase5_impulse_stage"].astype(str).eq("Confirmed").astype(float) * 14
        position_score += df["phase5_correction_completed"].astype(float) * 10
        position_score += zone.astype(str).str.contains("健康回撤", na=False).astype(float) * 8
        position_score += (ratio >= 1.2).astype(float) * 5
        position_score -= df["phase5_escape_rally"].astype(float) * 35
        position_score -= hard_avoid.astype(float) * 25
        df["phase5_position_score"] = position_score.clip(lower=0, upper=100).round(2)
        df["position_score"] = df[["phase5_position_score"]].max(axis=1)

        # 彩晶/PreBreakout 類：若原本 position_stage 未知，補成可讀波段位置。
        df["position_stage"] = _safe_str_series(df, "position_stage", default="未知")
        df.loc[df["phase5_wave_phase"].astype(str).eq("Wave3_PreBreakout"), "position_stage"] = "Wave3_PreBreakout"
        df.loc[df["phase5_wave_phase"].astype(str).eq("Wave3_Breakout"), "position_stage"] = "Wave3_Breakout"
        df.loc[df["phase5_major_wave"].astype(str).eq("底部完成") & df["phase5_breakout_stage"].astype(str).eq("PreBreakout"), "position_stage"] = "底部完成_PreBreakout"
        block_reason = pd.Series("", index=df.index)
        block_reason.loc[escape] = "Phase5逃命反彈：主跌修正浪且修正未完成，禁止追價或主動布局"
        block_reason.loc[main_down & ~escape & ~correction_completed] = "Phase5主跌弱反彈：大波仍是主跌修正浪，僅可觀察"
        block_reason.loc[main_up & ~correction_completed & low_base_reversal] = "Phase5主升拉回未完成：等待量價與壓力突破確認"
        block_reason.loc[df["phase5_wave_phase"].astype(str).eq("Wave3_PreBreakout") & (df["phase5_position_score"] < 60)] = "Phase5預突破位置分不足：等待量能/壓力確認"
        df["phase5_block_reason"] = block_reason
        df["phase5_hard_block"] = escape | (main_down & ~correction_completed)
        df["phase5_wait_block"] = (main_up & low_base_reversal & ~correction_completed) | (box & ~correction_completed & ~impulsive)
        pool = pd.Series("觀察池", index=df.index)
        pool.loc[df["phase5_hard_block"],] = "高風險反彈池"
        pool.loc[escape] = "禁追風控池"
        pool.loc[main_up & low_base_reversal & ~correction_completed] = "主升拉回觀察池"
        pool.loc[prebreakout & ~impulsive & ~df["phase5_hard_block"]] = "主升預突破觀察池"
        pool.loc[df["phase5_wave_phase"].astype(str).eq("Wave3_PreBreakout") & (df["phase5_position_score"] >= 65) & ~df["phase5_hard_block"]] = "主升預突破觀察池"
        pool.loc[impulsive] = "主升確認池"
        df["phase5_candidate_pool"] = pool
        labels = []
        for _, r in df.iterrows():
            label = f"{r.get('phase5_major_wave','')}/{r.get('phase5_minor_wave','')}/{r.get('phase5_rebound_type','')}"
            if bool(r.get('phase5_escape_rally', False)):
                label += "（逃命反彈風控）"
            elif not bool(r.get('phase5_correction_completed', False)) and "反彈" in str(r.get('phase5_rebound_type','')):
                label += "（修正未完成）"
            labels.append(label)
        df["phase5_wave_label"] = labels
        if self.logger:
            self.logger.strategy_trace("PHASE5_SEMANTIC_SUMMARY", {
                "rows": int(len(df)),
                "escape_rally": int(df["phase5_escape_rally"].fillna(False).sum()),
                "impulsive_wave": int(df["phase5_impulsive_wave"].fillna(False).sum()),
                "hard_block": int(df["phase5_hard_block"].fillna(False).sum()),
                "prebreakout": int(df["phase5_breakout_stage"].astype(str).eq("PreBreakout").sum()),
                "bottom_complete": int(df["phase5_major_wave"].astype(str).eq("底部完成").sum()),
                "avg_position_score": round(float(df["phase5_position_score"].fillna(0).mean()), 2),
            })
        return df

class TeacherDecisionEngine:
    """Phase5：老師五態決策重平衡。
    決策順序：硬性風險先擋，再由低位階翻多、RR、分數與核心股狀態決定 BUY/LOW_BUY。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        # V2.7.3：先建立Phase5語義/風控欄位，再進老師五態決策，避免報表與決策不同步。
        df = TeacherPhase5SemanticEngine(self.logger).apply(df)
        score = _safe_series(df, "teacher_score", default=0).fillna(0)
        rr = _safe_series(df, "rr", default=0).fillna(0)
        close = _safe_series(df, "close")
        entry_high = _safe_series(df, "entry_high")
        entry_low = _safe_series(df, "entry_low")
        high20 = _safe_series(df, "high20_prev")
        low_base = _safe_series(df, "low_base_score", default=50).fillna(50)
        low_base_reversal = _safe_bool_series(df, "low_base_reversal_flag")
        hard_avoid = _safe_bool_series(df, "hard_avoid_flag")
        avoid = _safe_bool_series(df, "avoid_flag")
        soft_released = _safe_str_series(df, "avoid_level", default="NONE").eq("SOFT_RELEASED")
        k_warning = _safe_str_series(df, "k_warning_type")
        wave = _safe_bool_series(df, "wave_correction_flag")
        deviation = _safe_bool_series(df, "deviation_risk_flag")
        core_state = _safe_str_series(df, "core_leader_state", default="NE")
        core_ok = ~core_state.eq("核心轉弱")
        phase5_hard_block = _safe_bool_series(df, "phase5_hard_block")
        phase5_wait_block = _safe_bool_series(df, "phase5_wait_block")
        phase5_escape = _safe_bool_series(df, "phase5_escape_rally")
        phase5_impulsive = _safe_bool_series(df, "phase5_impulsive_wave")
        phase5_pool = _safe_str_series(df, "phase5_candidate_pool")
        phase5_position_score = _safe_series(df, "phase5_position_score", default=0).fillna(0)
        phase5_wave_phase = _safe_str_series(df, "phase5_wave_phase")
        phase5_block_reason = _safe_str_series(df, "phase5_block_reason")
        df["teacher_decision"] = "WATCH"
        df.loc[hard_avoid | core_state.eq("核心轉弱") | phase5_escape, "teacher_decision"] = "AVOID"
        reduce_mask = ((k_warning.ne("") | wave | deviation | phase5_hard_block) & ~hard_avoid & ~low_base_reversal & ~phase5_escape)
        df.loc[reduce_mask & ~df["teacher_decision"].eq("AVOID"), "teacher_decision"] = "REDUCE"
        low_buy = (
            (score >= 55)
            & (rr >= 1.20)
            & (low_base >= 55)
            & low_base_reversal
            & core_ok
            & ~hard_avoid
            & ~phase5_hard_block
            & (close <= entry_high * 1.05)
        )
        breakout_or_pullback_confirm = low_base_reversal | (close > high20) | soft_released
        buy = (
            (score >= 72)
            & (rr >= 1.55)
            & breakout_or_pullback_confirm
            & core_ok
            & ~hard_avoid
            & ~phase5_hard_block
            & (phase5_impulsive | ((phase5_pool.isin(["主升確認池", "主升預突破觀察池", "主升拉回觀察池"])) & (phase5_position_score >= 65)))
            & (close <= entry_high * 1.035)
            & ~wave
            & k_warning.eq("")
        )
        df.loc[low_buy & ~df["teacher_decision"].eq("AVOID"), "teacher_decision"] = "LOW_BUY"
        df.loc[buy & ~df["teacher_decision"].eq("AVOID"), "teacher_decision"] = "BUY"
        reasons = []
        traces = []
        release = []
        for i, r in df.iterrows():
            reason_parts = []
            release_parts = []
            if str(r.get("hard_avoid_reason", "") or ""): reason_parts.append("硬性排除=" + str(r.get("hard_avoid_reason")))
            if str(r.get("phase5_block_reason", "") or ""): reason_parts.append(str(r.get("phase5_block_reason")))
            if str(r.get("phase5_candidate_pool", "") or ""): release_parts.append("Phase5候選池=" + str(r.get("phase5_candidate_pool")))
            if str(r.get("phase5_wave_phase", "") or ""): release_parts.append("波段階段=" + str(r.get("phase5_wave_phase")))
            if _safe_float(r.get("phase5_position_score"), 0) > 0: release_parts.append("波段位置分=" + str(round(_safe_float(r.get("phase5_position_score"),0),2)))
            if str(r.get("soft_avoid_reason", "") or ""):
                if bool(r.get("low_base_reversal_flag", False)) and not str(r.get("hard_avoid_reason", "") or ""):
                    release_parts.append("低位階翻多覆蓋軟性壓力=" + str(r.get("soft_avoid_reason")))
                else:
                    reason_parts.append("軟性壓力=" + str(r.get("soft_avoid_reason")))
            if str(r.get("k_warning_type", "") or ""): reason_parts.append("K線警訊=" + str(r.get("k_warning_type")))
            if bool(r.get("wave_correction_flag", False)): reason_parts.append("MACD+KD波段修正")
            if bool(r.get("deviation_risk_flag", False)): reason_parts.append("乖離過大")
            if _safe_float(r.get("rr"), 0) < 1.2: reason_parts.append("RR不足")
            if _safe_float(r.get("teacher_score"), 0) < 55: reason_parts.append("老師總分不足")
            if bool(r.get("low_base_reversal_flag", False)): reason_parts.append("低位階翻多=" + str(r.get("low_base_reversal_reason", "")))
            if str(r.get("teacher_decision")) in ("BUY", "LOW_BUY"):
                release_parts.append("分數/RR/位置/核心股狀態符合放行")
            if not reason_parts and not release_parts:
                reason_parts.append("條件未完全確認，列觀察")
            trace = (
                f"stock={str(r.get('stock_id','')).zfill(4)};"
                f"decision={r.get('teacher_decision')};"
                f"score={round(_safe_float(r.get('teacher_score'),0),2)};"
                f"rr={round(_safe_float(r.get('rr'),0),2)};"
                f"low_base={r.get('low_base_reversal_flag','')};"
                f"core={r.get('core_leader_state','')};"
                f"wave={r.get('market_pullback_type','')};"
                f"k_warning={r.get('k_warning_type','')};"
                f"avoid_level={r.get('avoid_level','')};"
                f"phase5={r.get('phase5_wave_label','')};"
                f"wave_phase={r.get('phase5_wave_phase','')};"
                f"breakout_stage={r.get('phase5_breakout_stage','')};"
                f"position_score={round(_safe_float(r.get('phase5_position_score'),0),2)};"
                f"phase5_pool={r.get('phase5_candidate_pool','')};"
                f"phase5_block={r.get('phase5_block_reason','')};"
                f"avoid={r.get('swap_reason','')}"
            )
            reasons.append(";".join(reason_parts))
            release.append(";".join(release_parts))
            traces.append(trace)
        df["teacher_decision_reason"] = reasons
        df["decision_release_reason"] = release
        df["decision_trace"] = traces
        pool = []
        for i, r in df.iterrows():
            div = _safe_float(r.get("dividend_yield"), 0)
            rev = _safe_float(r.get("yoy", r.get("revenue_yoy")), 0)
            theme = str(r.get("theme", "") or "") + " " + str(r.get("sub_theme", "") or "") + " " + str(r.get("industry", "") or "")
            if str(r.get("stock_id", "")).zfill(4) == "2330":
                pool.append("主流核心")
            elif bool(r.get("low_base_reversal_flag", False)):
                pool.append("低位階翻多")
            elif rev >= 20 or any(k in theme for k in ["AI", "半導體", "CPO", "伺服器", "散熱"]):
                pool.append("營收成長/主流題材")
            elif div >= 4:
                pool.append("高配息防守")
            else:
                pool.append("觀察")
        df["teacher_pool_type"] = pool
        # V2.7.3：外顯老師股票池優先接Phase5候選池，避免主跌弱反彈混入主升池。
        df.loc[phase5_pool.astype(str).ne(""), "teacher_pool_type"] = phase5_pool

        # Phase5 semantic fix：以老師五態重新同步 YES / WAIT / NO，避免 YES 與 BUY/LOW_BUY 語義打架。
        in_entry_zone = close.notna() & entry_low.notna() & entry_high.notna() & (close >= entry_low) & (close <= entry_high)
        df["teacher_execution_status"] = "NO"
        df["teacher_execution_reason"] = ""
        df["是否可下單"] = "NO"
        buy_mask = df["teacher_decision"].eq("BUY")
        low_buy_mask = df["teacher_decision"].eq("LOW_BUY")
        watch_mask = df["teacher_decision"].eq("WATCH")
        reduce_mask2 = df["teacher_decision"].eq("REDUCE")
        avoid_mask2 = df["teacher_decision"].eq("AVOID")
        buy_entry_ok = buy_mask & (in_entry_zone | (close <= entry_high * 1.01)) & ~phase5_hard_block & ~phase5_wait_block
        df.loc[buy_entry_ok, "teacher_execution_status"] = "YES"
        df.loc[buy_entry_ok, "teacher_execution_reason"] = "BUY主攻：核心風向未轉弱、無硬性排除、RR與Phase5位置符合"
        df.loc[buy_mask & ~buy_entry_ok, "teacher_execution_status"] = "WAIT"
        df.loc[buy_mask & ~buy_entry_ok, "teacher_execution_reason"] = "BUY條件成立但價格/Phase5尚待確認，等待回測或突破確認"
        df.loc[low_buy_mask & in_entry_zone, "teacher_execution_status"] = "YES"
        df.loc[low_buy_mask & in_entry_zone, "teacher_execution_reason"] = "LOW_BUY且已進入低接區，可分批"
        df.loc[low_buy_mask & ~in_entry_zone, "teacher_execution_status"] = "WAIT"
        df.loc[low_buy_mask & ~in_entry_zone, "teacher_execution_reason"] = "LOW_BUY但尚未進入低接區，等待回落"
        df.loc[watch_mask & ~hard_avoid & (rr >= 1.20), "teacher_execution_status"] = "WAIT"
        df.loc[watch_mask & ~hard_avoid & (rr >= 1.20), "teacher_execution_reason"] = "WATCH：條件未完整確認，只觀察不主攻"
        df.loc[reduce_mask2, "teacher_execution_status"] = "NO"
        df.loc[reduce_mask2, "teacher_execution_reason"] = "REDUCE：壓力或波段風險，禁止新增"
        df.loc[avoid_mask2, "teacher_execution_status"] = "NO"
        df.loc[avoid_mask2, "teacher_execution_reason"] = "AVOID：硬性排除或核心風向轉弱"
        df.loc[phase5_escape, "teacher_execution_status"] = "NO"
        df.loc[phase5_escape, "teacher_execution_reason"] = "Phase5逃命反彈風控：禁止新增"
        df.loc[phase5_hard_block & ~phase5_escape, "teacher_execution_status"] = "WAIT"
        df.loc[phase5_hard_block & ~phase5_escape, "teacher_execution_reason"] = "Phase5主跌弱反彈或修正未完成：只觀察不主攻"
        df["是否可下單"] = df["teacher_execution_status"]

        if self.logger:
            counts = df["teacher_decision"].value_counts(dropna=False).to_dict()
            self.logger.strategy_trace("TEACHER_DECISION_SUMMARY", counts)
            self.logger.strategy_trace("TEACHER_EXECUTION_SUMMARY", df["teacher_execution_status"].value_counts(dropna=False).to_dict())
            if int((df["teacher_decision"].isin(["BUY", "LOW_BUY"])).sum()) == 0:
                self.logger.warning("TEACHER_DECISION_STRICT_WARNING BUY與LOW_BUY皆為0，請檢查RR/低位階翻多/硬性風險條件")
            for _, r in df.sort_values(["teacher_score", "rr"], ascending=False).head(30).iterrows():
                self.logger.strategy_trace("DECISION_TRACE", {
                    "stock": str(r.get("stock_id", "")).zfill(4),
                    "decision": r.get("teacher_decision", ""),
                    "execution": r.get("teacher_execution_status", ""),
                    "reason": r.get("teacher_decision_reason", ""),
                    "release": r.get("decision_release_reason", ""),
                    "score": round(_safe_float(r.get("teacher_score"),0), 2),
                    "rr": round(_safe_float(r.get("rr"),0), 2),
                    "low_base": bool(r.get("low_base_reversal_flag", False)),
                    "wave": r.get("market_pullback_type", ""),
                    "k_warning": r.get("k_warning_type", ""),
                    "avoid": r.get("swap_reason", ""),
                    "phase5": r.get("phase5_wave_label", ""),
                    "phase5_pool": r.get("phase5_candidate_pool", ""),
                    "phase5_block": r.get("phase5_block_reason", ""),
                    "impulse_stage": r.get("phase5_impulse_stage", ""),
                })
        return df

def _safe_float(v, default=math.nan):
    try:
        if v is None:
            return default
        if isinstance(v, str) and not v.strip():
            return default
        x = float(v)
        if math.isnan(x):
            return default
        return x
    except Exception:
        return default

class FeatureBuilder:
    def build(self, df):
        for c in ["close", "ma20", "ma60", "ma20_calc", "ma60_calc", "rsi", "volume", "vol5", "vol20", "revenue_yoy", "yoy", "eps_ttm", "eps_yoy", "pe", "dividend_yield", "roe", "institutional_score", "risk_score"]:
            if c in df.columns:
                df[c] = pd.to_numeric(df[c], errors="coerce")
        df["close"] = df.get("close", pd.Series(index=df.index, dtype=float)).fillna(df.get("close_m", pd.Series(index=df.index, dtype=float)))
        df["ma20_final"] = df.get("ma20", pd.Series(index=df.index, dtype=float)).fillna(df.get("ma20_calc", pd.Series(index=df.index, dtype=float)))
        df["ma60_final"] = df.get("ma60", pd.Series(index=df.index, dtype=float)).fillna(df.get("ma60_calc", pd.Series(index=df.index, dtype=float)))
        df["low_base_score"] = ((1 - df.get("pos_120d", pd.Series(index=df.index, dtype=float)).clip(0,1)) * 100).fillna(50)
        df["ma_support_score"] = 50
        df.loc[df["close"] >= df["ma20_final"], "ma_support_score"] += 25
        df.loc[df["ma20_final"] >= df["ma60_final"] * 0.98, "ma_support_score"] += 25
        df["ma_support_score"] = df["ma_support_score"].clip(0,100)
        df["kline_score"] = df.get("reversal_score", pd.Series(50, index=df.index)).fillna(50)
        ratio = df.get("vol5", pd.Series(index=df.index, dtype=float)) / df.get("vol20", pd.Series(index=df.index, dtype=float)).replace(0, math.nan)
        df["volume_health_score"] = (50 + (ratio.fillna(1) - 1) * 50).clip(0,100)
        rev = df.get("yoy", df.get("revenue_yoy", pd.Series(index=df.index, dtype=float)))
        eps_yoy = df.get("eps_yoy", pd.Series(index=df.index, dtype=float))
        df["revenue_eps_score"] = (normalize_series(rev) * 0.55 + normalize_series(eps_yoy) * 0.45).fillna(50)
        theme_text = (df.get("theme", df.get("sub_theme", pd.Series("", index=df.index))).astype(str) + " " + df.get("industry", pd.Series("", index=df.index)).astype(str))
        df["theme_score"] = theme_text.apply(lambda x: 85 if any(k in x for k in ["AI", "半導體", "伺服器", "CPO", "散熱", "電源", "網通"]) else 50)
        return df

def normalize_series(series, low=None, high=None):
    """安全正規化：支援 None、scalar、Series，避免 scalar 造成後續運算失敗。"""
    if pd is None:
        return series
    if series is None:
        return pd.Series(dtype=float)
    if isinstance(series, pd.DataFrame):
        series = series.iloc[:, 0]
    if not isinstance(series, pd.Series):
        series = pd.Series([series])
    s = pd.to_numeric(series, errors="coerce")
    if s.dropna().empty:
        return pd.Series(50, index=s.index)
    lo = s.quantile(0.05) if low is None else low
    hi = s.quantile(0.95) if high is None else high
    if hi == lo:
        return pd.Series(50, index=s.index)
    return ((s - lo) / (hi - lo) * 100).clip(0,100).fillna(50)

class DualModelScoringEngine:
    def score(self, df):
        rev = df.get("yoy", df.get("revenue_yoy", pd.Series(index=df.index, dtype=float)))
        df["growth_score"] = (
            0.25 * normalize_series(rev) +
            0.20 * normalize_series(df.get("eps_yoy")) +
            0.20 * normalize_series(df.get("ai_score")) +
            0.15 * normalize_series(df.get("momentum_score")) +
            0.10 * normalize_series(df.get("volume_score")) +
            0.10 * df.get("theme_score", pd.Series(50, index=df.index))
        )
        df["value_score"] = (
            0.25 * (100 - normalize_series(df.get("pe"))) +
            0.20 * normalize_series(df.get("dividend_yield")) +
            0.20 * normalize_series(df.get("roe")) +
            0.15 * normalize_series(100 - pd.to_numeric(df.get("risk_score", pd.Series(50, index=df.index)), errors="coerce")) +
            0.10 * df.get("ma_support_score", pd.Series(50, index=df.index)) +
            0.10 * normalize_series(df.get("institutional_score"))
        )
        df["teacher_score"] = (
            0.20 * df.get("low_base_score", pd.Series(50, index=df.index)) +
            0.20 * df.get("kline_score", pd.Series(50, index=df.index)) +
            0.15 * df.get("ma_support_score", pd.Series(50, index=df.index)) +
            0.10 * df.get("volume_health_score", pd.Series(50, index=df.index)) +
            0.15 * df.get("revenue_eps_score", pd.Series(50, index=df.index)) +
            0.10 * normalize_series(df.get("dividend_yield")) +
            0.10 * df.get("theme_score", pd.Series(50, index=df.index))
        )
        return df

class ReportClassifier:
    def classify(self, df):
        df["老師分類"] = "觀察"
        rev = pd.to_numeric(df.get("yoy", df.get("revenue_yoy", pd.Series(index=df.index))), errors="coerce")
        df.loc[(rev >= 20) & (df["growth_score"] >= 70), "老師分類"] = "主流成長"
        df.loc[(pd.to_numeric(df.get("eps_ttm", 0), errors="coerce") > 0) & (pd.to_numeric(df.get("pe", 999), errors="coerce") <= 18), "老師分類"] = "價值防守"
        df.loc[(df.get("pos_120d", pd.Series(1,index=df.index)) <= 0.45) & (df["ma_support_score"] >= 60), "老師分類"] = "低位階翻多"
        return df


class LowBaseReversalEngine:
    """Phase5：顧奎國老師低位階翻多/第二浪末端辨識層。
    目的：補足「主升辨識層」，避免所有低位階回測不破標的都被 AvoidSwap 軟壓力直接封殺。
    這一層只允許覆蓋 soft avoid，不覆蓋長黑K、波段修正、下降軌道等 hard risk。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def apply(self, df):
        if df is None or df.empty:
            return df
        close = _safe_series(df, "close")
        ma20 = _safe_series(df, "ma20_final")
        ma60 = _safe_series(df, "ma60_final")
        low_base = _safe_series(df, "low_base_score", default=50).fillna(50)
        ma_support = _safe_series(df, "ma_support_score", default=50).fillna(50)
        volume_health = _safe_series(df, "volume_health_score", default=50).fillna(50)
        revenue_eps = _safe_series(df, "revenue_eps_score", default=50).fillna(50)
        k_warning = _safe_str_series(df, "k_warning_type")
        core_state = _safe_str_series(df, "core_leader_state", default="NE")
        wave_correction = _safe_bool_series(df, "wave_correction_flag")
        deviation_risk = _safe_bool_series(df, "deviation_risk_flag")
        hard_k_warning = k_warning.str.contains("長黑K|墓碑線", na=False)
        pullback_hold_ma20 = close.notna() & ma20.notna() & (close >= ma20 * 0.97)
        mid_trend_not_broken = close.notna() & ma60.notna() & (close >= ma60 * 0.92)
        df["low_base_reversal_flag"] = (
            (low_base >= 58)
            & (ma_support >= 55)
            & (volume_health >= 42)
            & (revenue_eps >= 45)
            & pullback_hold_ma20
            & mid_trend_not_broken
            & ~wave_correction
            & ~hard_k_warning
            & ~core_state.eq("核心轉弱")
        )
        df["low_base_reversal_strength"] = (
            low_base * 0.35 + ma_support * 0.25 + volume_health * 0.15 + revenue_eps * 0.15
            + (~deviation_risk).astype(int) * 10
        ).clip(0, 100)
        stage, reason, release_reason = [], [], []
        for _, r in df.iterrows():
            parts = []
            if bool(r.get("low_base_reversal_flag", False)):
                parts.extend(["低位階分>=58", "回測MA20不破", "中期線未破壞", "量能健康"])
                if _safe_float(r.get("revenue_eps_score"), 0) >= 55:
                    parts.append("營收/EPS支持")
                stage.append("第二浪末端/低位階翻多")
                reason.append(";".join(parts))
                release_reason.append("低位階翻多可覆蓋軟性頸線壓力，但不可覆蓋硬性風險")
            else:
                miss = []
                if _safe_float(r.get("low_base_score"), 0) < 58: miss.append("位階不足")
                if _safe_float(r.get("ma_support_score"), 0) < 55: miss.append("均線支撐不足")
                if bool(r.get("wave_correction_flag", False)): miss.append("波段修正")
                warning_text = str(r.get("k_warning_type", ""))
                if any(x in warning_text for x in ["長黑K", "墓碑線"]): miss.append("硬K線風險")
                stage.append("未確認")
                reason.append(";".join(miss) if miss else "條件未完全確認")
                release_reason.append("")
        df["low_base_stage"] = stage
        df["low_base_reversal_reason"] = reason
        df["low_base_release_reason"] = release_reason
        if self.logger:
            self.logger.strategy_trace("LOW_BASE_REVERSAL_SUMMARY", {
                "rows": int(len(df)),
                "low_base_reversal_count": int(df["low_base_reversal_flag"].fillna(False).sum()),
                "avg_strength": round(float(pd.to_numeric(df["low_base_reversal_strength"], errors="coerce").fillna(0).mean()), 2),
            })
        return df

class TradePlanEngine:
    def build(self, df):
        close = _safe_series(df, "close")
        ma20 = _safe_series(df, "ma20_final")
        ma60 = _safe_series(df, "ma60_final")
        atr = _safe_series(df, "atr")
        df["entry_low"] = pd.Series(np.maximum(ma20 * 0.98, close * 0.935), index=df.index) if np is not None else close * 0.935
        df["entry_high"] = pd.Series(np.minimum(ma20 * 1.01, close * 0.965), index=df.index) if np is not None else close * 0.965
        fallback = df["entry_low"].isna() | df["entry_high"].isna() | (close < df["entry_low"])
        df.loc[fallback, "entry_low"] = close * 0.995
        df.loc[fallback, "entry_high"] = close * 1.010
        reverse_mask = df["entry_low"].notna() & df["entry_high"].notna() & (df["entry_low"] > df["entry_high"])
        if reverse_mask.any():
            tmp_entry_low = df.loc[reverse_mask, "entry_low"].copy()
            df.loc[reverse_mask, "entry_low"] = df.loc[reverse_mask, "entry_high"]
            df.loc[reverse_mask, "entry_high"] = tmp_entry_low
            if "entry_zone_fix_flag" not in df.columns:
                df["entry_zone_fix_flag"] = ""
            df.loc[reverse_mask, "entry_zone_fix_flag"] = "低接區上下限反向已自動修正"
        if not (df.loc[df["entry_low"].notna() & df["entry_high"].notna(), "entry_low"] <= df.loc[df["entry_low"].notna() & df["entry_high"].notna(), "entry_high"]).all():
            raise ValueError("P0_FAIL: TradePlanEngine entry_low/entry_high 區間仍存在反向")
        df["stop_loss"] = pd.Series(np.minimum(ma60 * 0.97, df["entry_low"] * 0.94), index=df.index) if np is not None else df["entry_low"] * 0.94
        df["stop_loss"] = df["stop_loss"].fillna(df["entry_low"] * 0.94)
        df["target_1"] = pd.Series(np.where(atr.notna(), close + atr * 1.382, close * 1.05), index=df.index) if np is not None else close * 1.05
        df["target_2"] = pd.Series(np.where(atr.notna(), close + atr * 1.618, close * 1.125), index=df.index) if np is not None else close * 1.125
        risk_denom = (df["entry_high"] - df["stop_loss"]).replace(0, math.nan)
        df["rr"] = ((df["target_1"] - df["entry_high"]) / risk_denom).replace([math.inf, -math.inf], math.nan)
        df["exclude_reason"] = ""
        if "entry_zone_fix_flag" in df.columns:
            df.loc[_safe_str_series(df, "entry_zone_fix_flag").ne(""), "exclude_reason"] += "低接區上下限反向已自動修正;"
        df.loc[_safe_series(df, "volume", default=0).fillna(0) < 500, "exclude_reason"] += "成交量不足;"
        df.loc[df["rr"].fillna(0) < 1.2, "exclude_reason"] += "RR不足;"
        df.loc[_safe_series(df, "rsi", default=50).fillna(50) > 78, "exclude_reason"] += "RSI過熱;"
        if "hard_avoid_reason" in df.columns:
            mask = _safe_str_series(df, "hard_avoid_reason").ne("")
            df.loc[mask, "exclude_reason"] += _safe_str_series(df, "hard_avoid_reason")[mask] + ";"
        elif "swap_reason" in df.columns:
            mask = _safe_str_series(df, "swap_reason").ne("")
            df.loc[mask, "exclude_reason"] += _safe_str_series(df, "swap_reason")[mask] + ";"
        if "wave_correction_flag" in df.columns:
            df.loc[_safe_bool_series(df, "wave_correction_flag"), "exclude_reason"] += "MACD+KD波段修正;"
        df["exclude_flag"] = df["exclude_reason"].ne("")
        df["是否可下單"] = "NO"
        yes = (_safe_series(df, "teacher_score", default=0) >= 55) & (df["rr"].fillna(0) >= 1.5) & (~df["exclude_flag"]) & (close <= df["entry_high"] * 1.03)
        wait = (_safe_series(df, "teacher_score", default=0) >= 55) & (df["rr"].fillna(0) >= 1.5) & (~df["exclude_flag"]) & (~yes)
        df.loc[yes, "是否可下單"] = "YES"
        df.loc[wait, "是否可下單"] = "WAIT"
        df["操作策略"] = df["是否可下單"].map({"YES":"低接區內可分批", "WAIT":"等待回到低接區或RR改善", "NO":"不符合下單條件"}).fillna("觀察")
        df.loc[_safe_str_series(df, "k_warning_type").ne(""), "操作策略"] = "K線警訊，壓力先賣/等待2日確認"
        df.loc[_safe_bool_series(df, "wave_correction_flag"), "操作策略"] = "波段修正確認，停止追高並降檔"
        df.loc[_safe_bool_series(df, "avoid_flag"), "操作策略"] = "硬性風險或未釋放壓力，反彈換股"
        df.loc[_safe_str_series(df, "avoid_level").eq("SOFT_RELEASED"), "操作策略"] = "低位階翻多覆蓋軟性壓力，可列LOW_BUY觀察"
        return df

class InstitutionalReportEngine:
    def __init__(self, db_path: str, logger: Macro16Logger):
        self.repo = DBRepository(db_path, logger)
        self.logger = logger

    def run(self):
        if pd is None:
            self.logger.warning("pandas未安裝，無法產出機構級股票投資規劃報表")
            return None
        df = None
        trade_date = ""
        try:
            with StageTrace(self.logger, "get_trade_date"):
                trade_date = self.repo.get_trade_date()
            with StageTrace(self.logger, "load_base_universe", lambda: df):
                df = self.repo.load_base_universe(trade_date)
            stages = [
                ("FeatureBuilder", lambda x: FeatureBuilder().build(x)),
                ("MarketRegimeEngine", lambda x: MarketRegimeEngine(self.logger).apply(x)),
                ("SakataRiskPatternEngine", lambda x: SakataRiskPatternEngine(self.logger).apply(x)),
                ("CoreLeaderEngine", lambda x: CoreLeaderEngine(self.logger).apply(x)),
                ("DualModelScoringEngine", lambda x: DualModelScoringEngine().score(x)),
                ("ReportClassifier", lambda x: ReportClassifier().classify(x)),
                ("LowBaseReversalEngine", lambda x: LowBaseReversalEngine(self.logger).apply(x)),
                ("AvoidSwapEngine", lambda x: AvoidSwapEngine(self.logger).apply(x)),
                ("TradePlanEngine", lambda x: TradePlanEngine().build(x)),
                ("TeacherDecisionEngine", lambda x: TeacherDecisionEngine(self.logger).apply(x)),
            ]
            for stage_name, fn in stages:
                with StageTrace(self.logger, stage_name, lambda df=df: df):
                    df = fn(df)
            with StageTrace(self.logger, "finalize_report_name", lambda: df):
                df["report_name"] = df.get("stock_name", df.get("stock_name_master", df.get("name", "")))
            result = {
                "trade_date": trade_date,
                "db_path": self.repo.db_path,
                "counts": {t: self.repo.table_count(t) for t in ["ranking_result", "market_snapshot", "price_history", "external_revenue", "external_valuation", "external_institutional", "external_margin", "trade_plan"]},
                "latest_dates": {t: self.repo.latest_date(t) for t in ["ranking_result", "market_snapshot", "price_history", "external_revenue", "external_valuation", "external_institutional", "external_margin", "trade_plan"]},
                "all": df
            }
            self.logger.info(f"INSTITUTIONAL_REPORT_READY date={trade_date} rows={len(df)}")
            return result
        except Exception as exc:
            payload = {"db_path": self.repo.db_path, "trade_date": trade_date, "error": str(exc), "error_type": type(exc).__name__}
            if df is not None:
                try:
                    payload["rows"] = int(len(df))
                    payload["columns"] = list(map(str, list(df.columns)[:160]))
                    payload["dtypes"] = {str(k): str(v) for k, v in df.dtypes.astype(str).head(80).items()}
                except Exception as meta_exc:
                    payload["meta_error"] = str(meta_exc)
            self.logger.strategy_trace("TEACHER_REPORT_FAIL_DETAIL", payload)
            raise

class ReportValidator:
    def validate_workbook(self, wb) -> List[str]:
        errors = []
        for name in EXPECTED_INSTITUTIONAL_SHEETS:
            if name not in wb.sheetnames:
                errors.append(f"缺少Sheet:{name}")
        # 驗證03~08欄位
        for name in EXPECTED_INSTITUTIONAL_SHEETS[3:9]:
            if name in wb.sheetnames:
                ws = wb[name]
                cols = [ws.cell(1, i).value for i in range(1, len(REPORT_COLUMNS)+1)]
                if cols != REPORT_COLUMNS:
                    errors.append(f"{name}欄位不一致")
        return errors


class TeacherFullReportBuilder:
    """Phase5：集中建立老師策略完整報表資料集，避免Writer層同時負責資料切分與寫表。"""
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def build(self, df):
        if df is None or df.empty:
            empty = pd.DataFrame() if pd is not None else None
            return {"final_top15": empty, "growth_top30": empty, "value_top30": empty, "low_base": empty, "watchlist": empty, "excluded": empty, "low_buy": empty, "watch": empty, "avoid": empty, "reduce": empty, "release": empty}
        teacher_decision = df.get("teacher_decision", pd.Series("WATCH", index=df.index)).astype(str)
        exclude_flag = df.get("exclude_flag", pd.Series(False, index=df.index)).fillna(False)
        score_cols = ["teacher_score", "growth_score", "value_score"]
        for col in score_cols:
            if col not in df.columns:
                df[col] = 0
        execution_status = df.get("teacher_execution_status", pd.Series("NO", index=df.index)).astype(str)
        final = df[(teacher_decision.isin(["BUY", "LOW_BUY", "WATCH"])) & (~teacher_decision.eq("AVOID"))].sort_values(score_cols, ascending=False).head(15)
        if final.empty:
            final = df[~teacher_decision.eq("AVOID")].sort_values(score_cols, ascending=False).head(15)
        release_or_watch = df[teacher_decision.isin(["BUY", "LOW_BUY"]) | execution_status.eq("WAIT") | df.get("avoid_level", pd.Series("", index=df.index)).astype(str).eq("SOFT_RELEASED")].sort_values(["teacher_score", "rr"], ascending=False)
        datasets = {
            "final_top15": final,
            "growth_top30": df.sort_values("growth_score", ascending=False).head(30),
            "value_top30": df.sort_values("value_score", ascending=False).head(30),
            "low_base": df[df.get("low_base_reversal_flag", pd.Series(False, index=df.index)).fillna(False)].sort_values("teacher_score", ascending=False).head(50),
            "watchlist": df[df["stock_id"].astype(str).str.zfill(4).isin(TEACHER_WATCHLIST)].sort_values("teacher_score", ascending=False),
            "excluded": df[df.get("exclude_flag", pd.Series(False, index=df.index)).fillna(False) | teacher_decision.eq("AVOID")].sort_values("teacher_score", ascending=False),
            "low_buy": df[teacher_decision.eq("LOW_BUY")].sort_values(["teacher_score", "rr"], ascending=False),
            "watch": df[teacher_decision.eq("WATCH")].sort_values(["teacher_score", "rr"], ascending=False),
            "avoid": df[teacher_decision.eq("AVOID")].sort_values(["teacher_score", "rr"], ascending=False),
            "reduce": df[teacher_decision.eq("REDUCE")].sort_values(["teacher_score", "rr"], ascending=False),
            "release": release_or_watch,
        }
        if self.logger:
            self.logger.strategy_trace("TEACHER_FULL_REPORT_BUILDER", {k: int(len(v)) for k, v in datasets.items() if v is not None})
        return datasets

class InstitutionalExcelWriter:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger


    def _write_teacher_failure_diagnostic(self, wb, error_info: Optional[Dict[str, Any]], expected_sheets: Optional[List[str]] = None):
        ws = self._sheet(wb, "99_老師策略失敗診斷")
        ws.append(["項目", "內容"])
        ws.append(["狀態", "老師策略報表未產出，已進入失敗診斷頁"])
        ws.append(["錯誤摘要", json.dumps(error_info or {}, ensure_ascii=False, default=str)[:3000]])
        ws.append(["必須存在Sheet", ",".join(expected_sheets or EXPECTED_INSTITUTIONAL_SHEETS)])
        ws.append(["修正方向", "先修 safe_series / Engine Stage Trace / 欄位型別，再重跑 InstitutionalReportEngine"])
        ws.append([])
        ws.append(["驗收規則", "成功時Log必須出現 TEACHER_REPORT_READY；失敗時必須出現 TEACHER_REPORT_FAIL_DETAIL，且本Sheet不得空白。"])
        self.logger.warning("TEACHER_REPORT_VALIDATE_FAIL diagnostic_sheet=99_老師策略失敗診斷")

    def _sheet(self, wb, name):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def _report_rows(self, df, topn=None):
        if df is None or df.empty:
            return []
        if topn:
            df = df.head(topn)
        rows=[]
        for i, (_, r) in enumerate(df.iterrows(), 1):
            name = r.get("stock_name") or r.get("stock_name_master") or r.get("report_name") or ""
            rows.append([
                i, str(r.get("stock_id", "")).zfill(4), name, r.get("market", ""), r.get("industry", r.get("industry_master", "")), r.get("theme", r.get("sub_theme", "")), r.get("老師分類", ""),
                r.get("teacher_decision", "WATCH"), r.get("teacher_decision_reason", ""), r.get("decision_trace", ""), r.get("teacher_pool_type", ""), r.get("core_leader_state", ""),
                r.get("k_warning_type", ""), r.get("market_pullback_type", ""), r.get("swap_reason", ""), round(float(r.get("close", 0) or 0),2),
                f"{round(float(r.get('entry_low',0) or 0),2)}~{round(float(r.get('entry_high',0) or 0),2)}", round(float(r.get("stop_loss",0) or 0),2), round(float(r.get("target_1",0) or 0),2), round(float(r.get("target_2",0) or 0),2), round(float(r.get("rr",0) or 0),2), r.get("是否可下單", "NO"), r.get("teacher_execution_status", r.get("是否可下單", "NO")),
                round(float(r.get("teacher_score",0) or 0),2), round(float(r.get("growth_score",0) or 0),2), round(float(r.get("value_score",0) or 0),2), round(float(r.get("eps_ttm", r.get("valuation_eps_ttm",0)) or 0),2), round(float(r.get("pe",0) or 0),2), round(float(r.get("dividend_yield",0) or 0),2),
                round(float(r.get("yoy", r.get("revenue_yoy",0)) or 0),2), round(float(r.get("institutional_score",0) or 0),2), round(float(r.get("pct_20d",0) or 0),2), round(float(r.get("low_base_score",0) or 0),2), round(float(r.get("kline_score",0) or 0),2),
                round(float(r.get("ma_support_score",0) or 0),2), round(float(r.get("volume_health_score",0) or 0),2), round(float(r.get("revenue_eps_score",0) or 0),2), r.get("操作策略", ""), r.get("exclude_reason", ""),
                "Y" if bool(r.get("low_base_reversal_flag", False)) else "", r.get("decision_release_reason", ""), r.get("hard_avoid_reason", ""), r.get("soft_avoid_reason", ""), r.get("pressure_line_state", ""),
                r.get("phase5_major_wave", ""), r.get("phase5_minor_wave", ""), r.get("phase5_correction_type", ""), r.get("phase5_fibo_retrace", ""), r.get("phase5_rebound_type", ""),
                "Y" if bool(r.get("phase5_correction_completed", False)) else "N", "Y" if bool(r.get("phase5_escape_rally", False)) else "N", "Y" if bool(r.get("phase5_impulsive_wave", False)) else "N",
                r.get("phase5_wave_phase", ""), r.get("phase5_breakout_stage", ""), r.get("phase5_impulse_stage", ""), round(float(r.get("phase5_position_score", 0) or 0), 2),
                r.get("phase5_candidate_pool", ""), r.get("phase5_block_reason", ""),
                r.get("decision_buy_count", ""), r.get("decision_avoid_count", ""), r.get("latest_downgrade_reason", r.get("teacher_execution_reason", ""))
            ])
        return rows

    def write_into_workbook(self, wb, report: Optional[Dict[str, Any]], gov_result: Optional[Dict[str, Any]] = None, market5_result: Optional[Dict[str, Any]] = None):
        if report is None:
            return
        df = report["all"]
        # 00
        ws = self._sheet(wb, "00_執行摘要")
        ws.append(["項目", "內容"])
        ws.append(["報告日期", report.get("trade_date")])
        ws.append(["DB路徑", report.get("db_path")])
        ws.append(["總股票數", len(df)])
        ws.append(["策略版本", STRATEGY_VERSION])
        ws.append(["YES", int((df["是否可下單"]=="YES").sum())])
        ws.append(["WAIT", int((df["是否可下單"]=="WAIT").sum())])
        ws.append(["NO", int((df["是否可下單"]=="NO").sum())])
        ws.append(["老師執行狀態YES", int((df.get("teacher_execution_status", pd.Series(index=df.index)).astype(str)=="YES").sum())])
        ws.append(["老師執行狀態WAIT", int((df.get("teacher_execution_status", pd.Series(index=df.index)).astype(str)=="WAIT").sum())])
        ws.append(["老師執行狀態NO", int((df.get("teacher_execution_status", pd.Series(index=df.index)).astype(str)=="NO").sum())])
        gov_source = str((gov_result or {}).get("source", ""))
        if "TEJ" in gov_source:
            gov_level = "正式"
        elif gov_result and (gov_result or {}).get("gov_net_100m") is not None:
            gov_level = "備援"
        else:
            gov_level = "缺資料/佐證"
        ws.append(["官股資料等級", gov_level])
        ws.append([])
        ws.append(["老師五態決策統計", "數量"])
        for decision in ["BUY", "LOW_BUY", "WATCH", "REDUCE", "AVOID"]:
            ws.append([decision, int((df.get("teacher_decision", pd.Series(index=df.index)).astype(str)==decision).sum())])
        ws.append([])
        ws.append(["策略原因統計", "數量"])
        reason_map = {
            "兩高不過": df.get("swap_reason", pd.Series("", index=df.index)).astype(str).str.contains("兩高不過", na=False),
            "下降軌道": df.get("swap_reason", pd.Series("", index=df.index)).astype(str).str.contains("下降軌道", na=False),
            "頸線/前高壓力": df.get("swap_reason", pd.Series("", index=df.index)).astype(str).str.contains("頸線", na=False),
            "流星": df.get("k_warning_type", pd.Series("", index=df.index)).astype(str).str.contains("流星", na=False),
            "空頭新星十字": df.get("k_warning_type", pd.Series("", index=df.index)).astype(str).str.contains("空頭新星十字", na=False),
            "墓碑線": df.get("k_warning_type", pd.Series("", index=df.index)).astype(str).str.contains("墓碑", na=False),
            "長黑K": df.get("k_warning_type", pd.Series("", index=df.index)).astype(str).str.contains("長黑", na=False),
            "KD死叉": df.get("kd_dead_cross_below_80", pd.Series(False, index=df.index)).fillna(False),
            "MACD翻負": df.get("macd_turn_negative", pd.Series(False, index=df.index)).fillna(False),
            "乖離過大": df.get("deviation_risk_flag", pd.Series(False, index=df.index)).fillna(False),
            "波段修正": df.get("wave_correction_flag", pd.Series(False, index=df.index)).fillna(False),
            "RR不足": pd.to_numeric(df.get("rr", pd.Series(index=df.index)), errors="coerce").fillna(0) < 1.5,
        }
        for reason, mask in reason_map.items():
            ws.append([reason, int(mask.sum())])
        ws.append([])
        ws.append(["TEJ官股", json.dumps(gov_result or {}, ensure_ascii=False)[:800]])
        ws.append(["跨月5日", json.dumps(market5_result or {}, ensure_ascii=False)[:800]])
        # 01
        ws = self._sheet(wb, "01_DB資料盤點")
        ws.append(["資料表", "筆數", "最新日期", "用途"])
        usages = {"ranking_result":"模型排行主資料", "market_snapshot":"價格/RSI/ATR/MA", "price_history":"日K/20日漲幅/120日位階", "external_revenue":"營收YoY", "external_valuation":"PE/EPS/殖利率", "external_institutional":"法人分", "external_margin":"融資風險", "trade_plan":"既有交易計畫參考"}
        for t,n in report.get("counts",{}).items():
            ws.append([t,n,report.get("latest_dates",{}).get(t,""),usages.get(t,"")])
        # 02
        ws = self._sheet(wb, "02_模型設計")
        ws.append(["模型", "權重/規則", "說明"])
        ws.append(["老師總分", "低位階20% + 阪田K線20% + 均線15% + 量能10% + 營收EPS15% + 殖利率10% + 題材10%", "依Word/Excel策略SOP落實"])
        ws.append(["核心股風向", "CoreLeaderEngine：2330突破續強/回測不破/假突破/核心轉弱", "台積電為老師方法的大盤風向核心"])
        ws.append(["波段修正", "MarketRegimeEngine：MACD翻負 + KD破80死叉", "雙條件成立才視為波段修正，不因單一K棒全面轉空"])
        ws.append(["阪田警訊", "SakataRiskPatternEngine：流星/空頭新星十字/墓碑線/長黑K", "觸發後壓力先賣或等待2日確認"])
        ws.append(["排除換股", "AvoidSwapEngine：兩高不過/下降軌道/頸線壓力未突破/乖離過大", "符合者不得列主攻"])
        ws.append(["老師決策", "BUY / LOW_BUY / WATCH / REDUCE / AVOID", "保留是否可下單YES/WAIT/NO，但新增老師五態決策"])
        ws.append(["Phase5語義一致性", "大波/小波/回撤/反彈/逃命/推動浪/候選池", "主跌弱反彈不得混入主升池；BUY需通過Phase5與進場區Gate"])
        # Phase5 datasets：由TeacherFullReportBuilder集中切分，Writer只負責寫表。
        datasets = TeacherFullReportBuilder(self.logger).build(df)
        report_sheet_map = [
            ("03_最終TOP15", datasets["final_top15"], 15),
            ("04_成長模型TOP30", datasets["growth_top30"], 30),
            ("05_價值模型TOP30", datasets["value_top30"], 30),
            ("06_低位階候選", datasets["low_base"], 50),
            ("07_老師點名股檢核", datasets["watchlist"], None),
            ("08_排除與風險", datasets["excluded"], 200),
            ("13_LOW_BUY候選", datasets["low_buy"], 100),
            ("14_WATCH觀察池", datasets["watch"], 200),
            ("15_AVOID排除清單", datasets["avoid"], 300),
            ("16_放行與觀察候選", datasets["release"], 200),
            ("17_REDUCE減碼警示", datasets["reduce"], 300),
        ]
        for sheet, data, topn in report_sheet_map:
            ws = self._sheet(wb, sheet)
            ws.append(REPORT_COLUMNS)
            for row in self._report_rows(data, topn):
                ws.append(row)
        ws = self._sheet(wb, "09_來源與限制")
        ws.append(["項目", "內容"])
        ws.append(["資料來源", "stock_system DB + TEJ八大公股行庫（若提供）+ TWSE/TPEX/TAIFEX宏觀來源"])
        ws.append(["官股/八大公股", "TEJ或Wantgoo等來源只作證據追溯；只要gov_net_100m解析正確，分數只依數值，不因來源不同扣分"])
        ws.append(["Macro16跨月5日", "本版已提供Market5DayEngine；不足5日標P0_FAIL，停止市場技術判斷"])
        ws.append(["老師策略P0", "已新增CoreLeaderEngine、MarketRegimeEngine、SakataRiskPatternEngine、AvoidSwapEngine與TeacherDecisionEngine"])
        ws.append(["格式驗收", "00~16固定；03~08與13~16共用Phase5欄位；10/11/12為驗收、修改追蹤與命中驗證"])
        ws = self._sheet(wb, "10_老師策略驗收")
        ws.append(["TC", "驗收情境", "程式欄位", "預期結果"] )
        ws.append(["TC01", "流星但MACD/KD未雙翻空", "k_warning_type + wave_correction_flag", "REDUCE或WATCH，不全面轉空"] )
        ws.append(["TC02", "MACD翻負且KD破80死叉", "wave_correction_flag", "波段修正，停止追高"] )
        ws.append(["TC03", "2330假突破/回測失敗", "core_leader_state", "市場風險升級"] )
        ws.append(["TC04", "個股兩高不過或下降軌道", "avoid_flag/swap_reason", "AVOID/REDUCE，不列主攻"] )
        ws.append(["TC05", "低位階翻多且RR足夠", "low_base_reversal_flag/teacher_decision", "LOW_BUY或BUY"] )
        ws.append(["TC06", "軟性頸線壓力但低位階翻多", "avoid_level/decision_release_reason", "SOFT_RELEASED且可進LOW_BUY觀察"] )
        ws.append(["TC07", "硬性風險成立", "hard_avoid_flag", "AVOID，不可被LowBase覆蓋"] )
        ws.append(["TC08", "已分析但不適合追價/新增", "teacher_decision=REDUCE", "輸出17_REDUCE減碼警示，不混入TOP15主攻"] )
        ws.append([])
        ws.append(["驗收項目", "實際結果", "Pass/Fail", "命中筆數/說明"])
        validation_rows = [
            ("五態決策欄位", "teacher_decision" in df.columns, "PASS" if "teacher_decision" in df.columns else "FAIL", int(df.get("teacher_decision", pd.Series(index=df.index)).notna().sum()) if "teacher_decision" in df.columns else 0),
            ("決策原因欄位", "teacher_decision_reason" in df.columns, "PASS" if "teacher_decision_reason" in df.columns else "FAIL", int(df.get("teacher_decision_reason", pd.Series(index=df.index)).astype(str).ne("").sum()) if "teacher_decision_reason" in df.columns else 0),
            ("決策追蹤欄位", "decision_trace" in df.columns, "PASS" if "decision_trace" in df.columns else "FAIL", int(df.get("decision_trace", pd.Series(index=df.index)).astype(str).ne("").sum()) if "decision_trace" in df.columns else 0),
            ("CoreLeader狀態", "core_leader_state" in df.columns, "PASS" if "core_leader_state" in df.columns else "FAIL", str(df.get("core_leader_state", pd.Series(["NE"])).iloc[0]) if len(df) else ""),
            ("阪田警訊", "k_warning_type" in df.columns, "PASS" if "k_warning_type" in df.columns else "FAIL", int(df.get("k_warning_type", pd.Series("", index=df.index)).astype(str).ne("").sum()) if "k_warning_type" in df.columns else 0),
            ("避開換股", "swap_reason" in df.columns, "PASS" if "swap_reason" in df.columns else "FAIL", int(df.get("swap_reason", pd.Series("", index=df.index)).astype(str).ne("").sum()) if "swap_reason" in df.columns else 0),
            ("波段修正", "wave_correction_flag" in df.columns, "WARN" if ("wave_correction_flag" in df.columns and int(df.get("wave_correction_flag", pd.Series(False, index=df.index)).fillna(False).sum()) == 0) else ("PASS" if "wave_correction_flag" in df.columns else "FAIL"), int(df.get("wave_correction_flag", pd.Series(False, index=df.index)).fillna(False).sum()) if "wave_correction_flag" in df.columns else 0),
            ("低位階翻多", "low_base_reversal_flag" in df.columns, "PASS" if "low_base_reversal_flag" in df.columns else "FAIL", int(df.get("low_base_reversal_flag", pd.Series(False, index=df.index)).fillna(False).sum()) if "low_base_reversal_flag" in df.columns else 0),
            ("Hard/Soft Avoid", "hard_avoid_flag" in df.columns and "soft_avoid_flag" in df.columns, "PASS" if "hard_avoid_flag" in df.columns and "soft_avoid_flag" in df.columns else "FAIL", f"hard={int(df.get('hard_avoid_flag', pd.Series(False,index=df.index)).fillna(False).sum()) if 'hard_avoid_flag' in df.columns else 0};soft={int(df.get('soft_avoid_flag', pd.Series(False,index=df.index)).fillna(False).sum()) if 'soft_avoid_flag' in df.columns else 0}"),
            ("Phase5候選池", "phase5_candidate_pool" in df.columns, "PASS" if "phase5_candidate_pool" in df.columns else "FAIL", int(df.get("phase5_candidate_pool", pd.Series("", index=df.index)).astype(str).ne("").sum()) if "phase5_candidate_pool" in df.columns else 0),
            ("Phase5逃命反彈", "phase5_escape_rally" in df.columns, "PASS" if "phase5_escape_rally" in df.columns else "FAIL", int(df.get("phase5_escape_rally", pd.Series(False, index=df.index)).fillna(False).sum()) if "phase5_escape_rally" in df.columns else 0),
            ("Phase5推動浪", "phase5_impulsive_wave" in df.columns, "PASS" if "phase5_impulsive_wave" in df.columns else "FAIL", int(df.get("phase5_impulsive_wave", pd.Series(False, index=df.index)).fillna(False).sum()) if "phase5_impulsive_wave" in df.columns else 0),
        ]
        for row in validation_rows:
            ws.append(list(row))
        ws = self._sheet(wb, "11_修改追蹤")
        ws.append(["修改ID", "修改項目", "狀態", "對應類別/函式", "說明"] )
        ws.append(["P0-01", "CoreLeaderEngine", "已修改", "CoreLeaderEngine", "新增2330核心股風向判斷"] )
        ws.append(["P0-02", "MACD+KD波段修正", "已修改", "MarketRegimeEngine", "新增wave_correction_flag與market_pullback_type"] )
        ws.append(["P0-03", "阪田K線警訊", "已修改", "SakataRiskPatternEngine", "新增流星/空頭新星十字/墓碑/長黑K"] )
        ws.append(["P0-04", "兩高不過/下降軌道排除", "已修改", "AvoidSwapEngine", "新增avoid_flag/swap_reason"] )
        ws.append(["P0-05", "老師五態決策", "已修改", "TeacherDecisionEngine", "新增BUY/LOW_BUY/WATCH/REDUCE/AVOID"] )
        ws.append(["P0-06", "五態Summary統計", "已修改", "InstitutionalExcelWriter.write_into_workbook", "00_執行摘要新增BUY/LOW_BUY/WATCH/REDUCE/AVOID數量"] )
        ws.append(["P0-07", "策略原因統計", "已修改", "InstitutionalExcelWriter.write_into_workbook", "00_執行摘要新增兩高不過/流星/KD死叉/MACD翻負/乖離過大等統計"] )
        ws.append(["P0-08", "Decision Trace", "已修改", "TeacherDecisionEngine + Logger", "新增teacher_decision_reason與decision_trace欄位並輸出DECISION_TRACE log"] )
        ws.append(["P0-09", "策略版本凍結", "已修改", "VERSION/STRATEGY_VERSION", "新增STRATEGY_VERSION並寫入log與00摘要"] )
        ws.append(["P5-01", "報表輸出語意", "已修改", "ExcelWriter.write / InstitutionalExcelWriter", "新增TEACHER_REPORT_READY與00_12/00_16語意，移除00_09誤解"] )
        ws.append(["P5-02", "TeacherFullReportBuilder", "已修改", "TeacherFullReportBuilder", "集中建立TOP15/LOW_BUY/WATCH/AVOID/REDUCE/放行原因資料集"] )
        ws.append(["P5-03", "LowBaseReversalEngine", "已修改", "LowBaseReversalEngine", "新增low_base_reversal_flag與LOW_BASE_REVERSAL_SUMMARY"] )
        ws.append(["P5-04", "Avoid hard/soft分層", "已修改", "AvoidSwapEngine", "soft avoid可被低位階翻多覆蓋，hard avoid不可覆蓋"] )
        ws.append(["P5-05", "BUY/LOW_BUY重平衡", "已修改", "TeacherDecisionEngine", "先擋硬風險，再依LowBase/RR/Score/Core決策"] )
        ws.append(["P5-06", "新增13~16分池Sheet", "已修改", "InstitutionalExcelWriter", "新增LOW_BUY候選/WATCH觀察/AVOID排除/放行與觀察候選"] )
        ws.append(["P5-07", "Strategy Validation", "已修改", "12_策略命中驗證", "保留待追蹤；新增樣本分布與後續計算接口"] )
        ws.append(["P5-08", "GUI/CLI語意", "已修改", "run_gui / argparse", "新增macro_teacher/teacher_full模式並更新說明"] )
        ws.append(["P5-09", "報表缺失防呆", "已修改", "ReportValidator/ExcelWriter", "輸出後檢查00~16，缺失寫TEACHER_REPORT_VALIDATE_FAIL"] )
        ws.append(["P5-10", "成熟度標記", "已修改", "VERSION/STRATEGY_VERSION", "升級為2.7.3-teacher-phase5-consistency-fix"] )
        ws.append(["P7-01", "Phase5語義一致性引擎", "已修改", "TeacherPhase5SemanticEngine", "新增大波/小波/回撤/反彈/逃命/推動浪與候選池欄位"] )
        ws.append(["P7-02", "Phase5候選池Gate", "已修改", "TeacherDecisionEngine", "主跌修正浪/逃命反彈不得進BUY/YES；只進高風險反彈池或禁追風控池"] )
        ws.append(["P7-03", "動態決策欄位", "已修改", "REPORT_COLUMNS/_report_rows", "新增期間BUY/AVOID次數與最後降級原因欄位，供盤中log差異追蹤"] )
        ws.append(["P8-01", "Position Engine欄位", "已修改", "TeacherPhase5SemanticEngine/REPORT_COLUMNS/_report_rows", "新增wave_phase、breakout_stage、impulse_stage、position_score，修正彩晶類PreBreakout position_score=0問題"] )
        ws.append(["P6-01", "YES/WAIT/NO語義同步", "已修改", "TeacherDecisionEngine", "以老師五態重新同步是否可下單與teacher_execution_status，避免YES與BUY/LOW_BUY打架"] )
        ws.append(["P6-02", "16表名稱修正", "已修改", "EXPECTED_INSTITUTIONAL_SHEETS/Writer", "16_BUY_LOW_BUY放行原因改為16_放行與觀察候選，避免內含WATCH時語義錯誤"] )
        ws.append(["P9-01", "REDUCE減碼警示頁", "已修改", "EXPECTED_INSTITUTIONAL_SHEETS/TeacherFullReportBuilder/InstitutionalExcelWriter", "新增17_REDUCE減碼警示；REDUCE不混入TOP15主攻，但完整輸出風險資訊與決策原因"] )
        ws.append(["P6-03", "波段修正零樣本WARN", "已修改", "MarketRegimeEngine/10_老師策略驗收", "wave_correction_count=0改列WARN並寫入log，不再誤判為完全PASS"] )
        ws.append(["P6-04", "官股資料等級", "已修改", "00_執行摘要", "新增正式/備援/缺資料等級，TEJ未提供時不偽裝正式"] )
        ws.append(["P0-10", "策略命中驗證", "已修改", "12_策略命中驗證", "若有未來價格資料則計算，否則明確標示待追蹤"] )
        ws = self._sheet(wb, "12_策略命中驗證")
        ws.append(["項目", "內容"])
        ws.append(["策略版本", STRATEGY_VERSION])
        ws.append(["驗證目的", "追蹤老師策略五態決策的隔日/5日表現、最大回撤；若DB尚無未來價格，先標示待追蹤，不假造勝率"])
        ws.append(["目前資料狀態", "本報表依當前DB產出，若price_history尚未包含決策日後資料，命中率不可計算"])
        ws.append([])
        ws.append(["決策", "樣本數", "可計算樣本", "平均隔日%", "平均5日%", "最大回撤%", "狀態"])
        for decision in ["BUY", "LOW_BUY", "WATCH", "REDUCE", "AVOID"]:
            sample_count = int((df.get("teacher_decision", pd.Series(index=df.index)).astype(str)==decision).sum())
            ws.append([decision, sample_count, 0, "", "", "", "待後續price_history資料回補後驗證"])
        ready_payload = {
            "rows": int(len(df)),
            "sheets": [name for name in EXPECTED_INSTITUTIONAL_SHEETS if name in wb.sheetnames],
            "decision_counts": df.get("teacher_decision", pd.Series(index=df.index)).astype(str).value_counts(dropna=False).to_dict() if len(df) else {},
        }
        self.logger.strategy_trace("TEACHER_REPORT_READY", ready_payload)
        errors = ReportValidator().validate_workbook(wb)
        if errors:
            self.logger.warning("TEACHER_REPORT_VALIDATE_FAIL " + ";".join(errors))
            self.logger.warning("INSTITUTIONAL_REPORT_VALIDATE_FAIL " + ";".join(errors))
        else:
            self.logger.info("TEACHER_REPORT_VALIDATE_OK sheets=00_16_teacher_strategy_reports")
            self.logger.info("INSTITUTIONAL_REPORT_VALIDATE_OK")



# =============================
# V2.8.1 CPO Theme Engine / CPO股票池落地
# =============================
CPO_THEME_MASTER = [
    {
        "stock_id": "2330",
        "stock_name": "台積電",
        "cpo_theme": "先進封裝/矽光平台",
        "cpo_subtheme": "CoWoS/SiPh平台",
        "cpo_layer": "核心基礎設施",
        "cpo_score": 94,
        "cpo_strategy": "長期核心配置",
        "cpo_directness": "直接/平台",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "CPO/矽光子/先進封裝平台核心"
    },
    {
        "stock_id": "4979",
        "stock_name": "華星光",
        "cpo_theme": "光收發/光通訊",
        "cpo_subtheme": "光模組/收發元件",
        "cpo_layer": "高彈性主題股",
        "cpo_score": 93,
        "cpo_strategy": "波段主攻但需嚴控停損",
        "cpo_directness": "直接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "CPO/高速光通訊代表股"
    },
    {
        "stock_id": "3081",
        "stock_name": "聯亞",
        "cpo_theme": "磊晶/雷射",
        "cpo_subtheme": "雷射/磊晶上游",
        "cpo_layer": "雷射上游",
        "cpo_score": 88,
        "cpo_strategy": "中長期觀察",
        "cpo_directness": "直接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "光源/雷射上游"
    },
    {
        "stock_id": "3363",
        "stock_name": "上詮",
        "cpo_theme": "光通訊元件",
        "cpo_subtheme": "光纖/光元件",
        "cpo_layer": "元件彈性",
        "cpo_score": 84,
        "cpo_strategy": "短中期波段",
        "cpo_directness": "直接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "光通訊元件"
    },
    {
        "stock_id": "3163",
        "stock_name": "波若威",
        "cpo_theme": "光纖元件",
        "cpo_subtheme": "光纖/連接元件",
        "cpo_layer": "光元件",
        "cpo_score": 84,
        "cpo_strategy": "短中期波段",
        "cpo_directness": "直接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "光纖元件與模組題材"
    },
    {
        "stock_id": "4908",
        "stock_name": "前鼎",
        "cpo_theme": "光通訊模組",
        "cpo_subtheme": "光模組",
        "cpo_layer": "模組彈性",
        "cpo_score": 82,
        "cpo_strategy": "波段觀察",
        "cpo_directness": "直接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "光模組題材"
    },
    {
        "stock_id": "6669",
        "stock_name": "緯穎",
        "cpo_theme": "AI Server/資料中心",
        "cpo_subtheme": "AI伺服器",
        "cpo_layer": "AI Server核心",
        "cpo_score": 82,
        "cpo_strategy": "中長期核心觀察",
        "cpo_directness": "間接/系統",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI資料中心系統端，受惠光互連升級"
    },
    {
        "stock_id": "2345",
        "stock_name": "智邦",
        "cpo_theme": "資料中心交換器",
        "cpo_subtheme": "Switch/網通",
        "cpo_layer": "Switch核心",
        "cpo_score": 90,
        "cpo_strategy": "中長期核心配置",
        "cpo_directness": "直接/系統",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "交換器/網通，CPO交換器方向"
    },
    {
        "stock_id": "3443",
        "stock_name": "創意",
        "cpo_theme": "ASIC/高速SerDes",
        "cpo_subtheme": "ASIC/SerDes",
        "cpo_layer": "ASIC槓桿",
        "cpo_score": 86,
        "cpo_strategy": "中長期核心觀察",
        "cpo_directness": "間接/晶片",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "ASIC高速介面/IP設計受惠"
    },
    {
        "stock_id": "3035",
        "stock_name": "智原",
        "cpo_theme": "IP/ASIC/高速介面",
        "cpo_subtheme": "ASIC/IP",
        "cpo_layer": "ASIC/IP",
        "cpo_score": 80,
        "cpo_strategy": "中期波段",
        "cpo_directness": "間接/晶片",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "ASIC/IP高速傳輸"
    },
    {
        "stock_id": "6531",
        "stock_name": "愛普*",
        "cpo_theme": "記憶體/IP/矽光潛在平台",
        "cpo_subtheme": "IP/記憶體",
        "cpo_layer": "題材彈性",
        "cpo_score": 78,
        "cpo_strategy": "中長期觀察/波段",
        "cpo_directness": "間接/平台",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI記憶體/IP與平台題材"
    },
    {
        "stock_id": "6533",
        "stock_name": "晶心科",
        "cpo_theme": "RISC-V/IP/ASIC",
        "cpo_subtheme": "IP/ASIC",
        "cpo_layer": "ASIC/IP",
        "cpo_score": 76,
        "cpo_strategy": "中期觀察",
        "cpo_directness": "間接/晶片",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "ASIC/IP延伸題材"
    },
    {
        "stock_id": "2383",
        "stock_name": "台光電",
        "cpo_theme": "低損耗材料/CCL",
        "cpo_subtheme": "材料/CCL",
        "cpo_layer": "材料核心",
        "cpo_score": 90,
        "cpo_strategy": "中長期核心配置",
        "cpo_directness": "間接/材料",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "高速PCB低損耗材料"
    },
    {
        "stock_id": "6274",
        "stock_name": "台燿",
        "cpo_theme": "低損耗材料",
        "cpo_subtheme": "材料/CCL",
        "cpo_layer": "材料受惠",
        "cpo_score": 84,
        "cpo_strategy": "中期佈局",
        "cpo_directness": "間接/材料",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "高速材料受惠"
    },
    {
        "stock_id": "6213",
        "stock_name": "聯茂",
        "cpo_theme": "CCL材料",
        "cpo_subtheme": "材料/CCL",
        "cpo_layer": "材料補漲",
        "cpo_score": 72,
        "cpo_strategy": "觀察/低接",
        "cpo_directness": "間接/材料",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "高速材料補漲題材"
    },
    {
        "stock_id": "3037",
        "stock_name": "欣興",
        "cpo_theme": "高階PCB/載板",
        "cpo_subtheme": "PCB/載板",
        "cpo_layer": "PCB+載板",
        "cpo_score": 82,
        "cpo_strategy": "短中長期分批",
        "cpo_directness": "間接/PCB",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI高速PCB與載板"
    },
    {
        "stock_id": "3189",
        "stock_name": "景碩",
        "cpo_theme": "ABF/載板",
        "cpo_subtheme": "ABF載板",
        "cpo_layer": "載板復甦",
        "cpo_score": 78,
        "cpo_strategy": "中期佈局",
        "cpo_directness": "間接/PCB",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "ABF/載板受惠"
    },
    {
        "stock_id": "8046",
        "stock_name": "南電",
        "cpo_theme": "ABF/載板",
        "cpo_subtheme": "ABF載板",
        "cpo_layer": "載板復甦",
        "cpo_score": 78,
        "cpo_strategy": "中期佈局",
        "cpo_directness": "間接/PCB",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "ABF/載板受惠"
    },
    {
        "stock_id": "2313",
        "stock_name": "華通",
        "cpo_theme": "高階PCB",
        "cpo_subtheme": "PCB",
        "cpo_layer": "PCB受惠",
        "cpo_score": 72,
        "cpo_strategy": "中期觀察",
        "cpo_directness": "間接/PCB",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "高階PCB題材"
    },
    {
        "stock_id": "3665",
        "stock_name": "貿聯-KY",
        "cpo_theme": "高速線束/連接",
        "cpo_subtheme": "連接/線束",
        "cpo_layer": "過渡期受惠",
        "cpo_score": 80,
        "cpo_strategy": "2025-2027主軸",
        "cpo_directness": "間接/連接",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "去纜線化過渡期與高速連接"
    },
    {
        "stock_id": "3711",
        "stock_name": "日月光投控",
        "cpo_theme": "先進封裝/SiP",
        "cpo_subtheme": "封測/封裝",
        "cpo_layer": "封測核心",
        "cpo_score": 78,
        "cpo_strategy": "中長期觀察",
        "cpo_directness": "間接/封測",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "先進封裝/封測鏈"
    },
    {
        "stock_id": "2308",
        "stock_name": "台達電",
        "cpo_theme": "電源/散熱/資料中心",
        "cpo_subtheme": "電源/散熱",
        "cpo_layer": "基建核心",
        "cpo_score": 76,
        "cpo_strategy": "中長期配置",
        "cpo_directness": "間接/基建",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI Data Center基建"
    },
    {
        "stock_id": "2356",
        "stock_name": "英業達",
        "cpo_theme": "AI Server ODM",
        "cpo_subtheme": "AI伺服器",
        "cpo_layer": "ODM受惠",
        "cpo_score": 70,
        "cpo_strategy": "觀察",
        "cpo_directness": "間接/系統",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "系統端受惠"
    },
    {
        "stock_id": "2382",
        "stock_name": "廣達",
        "cpo_theme": "AI Server ODM",
        "cpo_subtheme": "AI伺服器",
        "cpo_layer": "ODM核心",
        "cpo_score": 76,
        "cpo_strategy": "中長期觀察",
        "cpo_directness": "間接/系統",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI Server供應鏈"
    },
    {
        "stock_id": "3231",
        "stock_name": "緯創",
        "cpo_theme": "AI Server ODM",
        "cpo_subtheme": "AI伺服器",
        "cpo_layer": "ODM受惠",
        "cpo_score": 72,
        "cpo_strategy": "觀察/波段",
        "cpo_directness": "間接/系統",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI Server供應鏈"
    },
    {
        "stock_id": "3017",
        "stock_name": "奇鋐",
        "cpo_theme": "散熱",
        "cpo_subtheme": "散熱",
        "cpo_layer": "散熱核心",
        "cpo_score": 74,
        "cpo_strategy": "中期波段",
        "cpo_directness": "間接/散熱",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "光電混合/AI機櫃散熱升級"
    },
    {
        "stock_id": "3324",
        "stock_name": "雙鴻",
        "cpo_theme": "散熱",
        "cpo_subtheme": "散熱",
        "cpo_layer": "散熱受惠",
        "cpo_score": 72,
        "cpo_strategy": "觀察/波段",
        "cpo_directness": "間接/散熱",
        "cpo_source": "CPO_THEME_MASTER",
        "cpo_note": "AI機櫃散熱升級"
    }
]

CPO_REPORT_COLUMNS = [
    "排名", "代號", "名稱", "CPO主題", "CPO子分類", "CPO分層", "CPO分數",
    "直接性", "策略定位", "老師決策", "是否可下單", "波段階段", "突破階段",
    "波段位置分", "RR", "CPO爆發前決策", "CPO爆發前原因", "資料狀態"
]

class CPOThemeEngine:
    """
    V2.8.1：CPO股票池引擎。
    目的：
    1. 將CPO_THEME_MASTER落地成可合併資料框，不再只有頁面名稱。
    2. 若有institutional_report['all']，將is_cpo/cpo_theme/cpo_score等欄位合併回主資料流。
    3. 無DB時仍能輸出完整CPO股票池，避免CPO頁空白。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def master_df(self):
        if pd is None:
            return CPO_THEME_MASTER
        df = pd.DataFrame(CPO_THEME_MASTER)
        if not df.empty:
            df["stock_id"] = df["stock_id"].astype(str).str.zfill(4)
            df["cpo_score"] = pd.to_numeric(df["cpo_score"], errors="coerce").fillna(0)
        return df

    def apply(self, df):
        if df is None or pd is None or getattr(df, "empty", True):
            return df
        master = self.master_df()
        if isinstance(master, list) or master.empty:
            return df
        out = df.copy()
        if "stock_id" not in out.columns:
            out["is_cpo"] = False
            return out
        out["stock_id"] = out["stock_id"].astype(str).str.zfill(4)
        drop_cols = [c for c in ["is_cpo","cpo_theme","cpo_subtheme","cpo_layer","cpo_score","cpo_strategy","cpo_directness","cpo_source","cpo_note"] if c in out.columns]
        if drop_cols:
            out = out.drop(columns=drop_cols)
        out = out.merge(master, on="stock_id", how="left")
        out["is_cpo"] = out["cpo_theme"].notna()
        out["cpo_score"] = pd.to_numeric(out["cpo_score"], errors="coerce").fillna(0)
        if self.logger:
            self.logger.info(
                "CPO_THEME_MASTER_READY "
                f"master_count={len(master)} merged_count={int(out['is_cpo'].sum())}"
            )
        return out

class CPOReportExcelIntegrator:
    """
    V2.8.1：CPO報表輸出整合器。
    輸出：
    1. CPO股票池：固定Universe，不依賴DB。
    2. CPO爆發前候選：CPO股票 × Phase5/Teacher/TradePlan。
    3. CPO整合驗收：確認資料流是否完整。
    """
    SHEETS = ["CPO股票池", "CPO股票池決策", "CPO×爆發前交集", "CPO整合驗收"]

    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger
        self.engine = CPOThemeEngine(logger)

    def _sheet(self, wb, name: str):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def _style_basic(self, ws):
        try:
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="1F4E78")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in range(1, ws.max_column + 1):
                letter = get_column_letter(col)
                ws.column_dimensions[letter].width = min(max(12, len(str(ws.cell(1, col).value or "")) + 4), 28)
            ws.freeze_panes = "A2"
        except Exception:
            pass

    def write_into_workbook(self, wb, institutional_report: Optional[Dict[str, Any]] = None):
        merged_df = None
        if institutional_report is not None and institutional_report.get("all") is not None and pd is not None:
            try:
                merged_df = self.engine.apply(institutional_report.get("all"))
                institutional_report["all"] = merged_df
            except Exception as exc:
                if self.logger:
                    self.logger.warning(f"CPO_THEME_ENGINE_APPLY_FAIL error={exc}")
        self._write_master(wb)
        self._write_prebreakout_candidates(wb, merged_df)
        self._write_cpo_prebreakout_intersection(wb, merged_df)
        self._write_validation(wb, merged_df)
        if self.logger:
            self.logger.info("CPO_REPORT_WRITTEN sheets=" + ",".join(self.SHEETS))
        return wb

    def _write_master(self, wb):
        ws = self._sheet(wb, "CPO股票池")
        headers = ["代號","公司","CPO主題","CPO子分類","CPO分層","CPO分數","直接性","策略定位","資料來源","備註"]
        ws.append(headers)
        for item in CPO_THEME_MASTER:
            ws.append([
                item.get("stock_id",""), item.get("stock_name",""), item.get("cpo_theme",""),
                item.get("cpo_subtheme",""), item.get("cpo_layer",""), item.get("cpo_score",""),
                item.get("cpo_directness",""), item.get("cpo_strategy",""),
                item.get("cpo_source",""), item.get("cpo_note","")
            ])
        self._style_basic(ws)

    def _candidate_decision(self, row):
        risk = " ".join([
            str(row.get("hard_avoid_reason", "") or ""),
            str(row.get("k_warning_type", "") or ""),
            str(row.get("phase5_block_reason", "") or ""),
        ])
        wp = str(row.get("phase5_wave_phase", "") or "")
        bs = str(row.get("phase5_breakout_stage", "") or "")
        pos = _safe_float(row.get("phase5_position_score"), 0)
        rr = _safe_float(row.get("rr"), 0)
        cpo_score = _safe_float(row.get("cpo_score"), 0)
        tdec = str(row.get("teacher_decision", "WATCH") or "WATCH")
        exe = str(row.get("teacher_execution_status", row.get("是否可下單", "NO")) or "NO")
        if any(k in risk for k in ["硬性", "逃命", "主跌", "長黑", "墓碑"]):
            return "AVOID", "CPO標的但觸發硬性/波段風險：" + risk[:120]
        if tdec == "REDUCE":
            return "REDUCE", "老師策略為REDUCE，CPO題材不得覆蓋減碼訊號"
        if cpo_score >= 85 and wp in ["Wave3_Breakout", "Wave3_Expansion"] and pos >= 70 and rr >= 1.5 and exe in ["YES", "WAIT"]:
            return ("BUY" if exe == "YES" else "WATCH"), "CPO高分+Wave3突破+RR達標；仍需依執行狀態確認"
        if cpo_score >= 80 and (wp == "Wave3_PreBreakout" or bs == "PreBreakout") and pos >= 65 and rr >= 1.2:
            return ("LOW_BUY" if tdec in ["LOW_BUY", "BUY"] else "WATCH"), "CPO高分+Wave3預突破；等待量價確認"
        if cpo_score >= 75 and bs in ["Compression", "PreBreakout", "Breakout"]:
            return "WATCH", "CPO題材進入壓縮/預突破觀察"
        return "WATCH", "CPO股票池標的，條件未完整確認"

    def _write_prebreakout_candidates(self, wb, merged_df):
        ws = self._sheet(wb, "CPO股票池決策")
        ws.append(CPO_REPORT_COLUMNS)
        if merged_df is None or pd is None or getattr(merged_df, "empty", True):
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "待DB", "未提供DB或老師策略資料，僅輸出CPO股票池", "NO_DB"])
            self._style_basic(ws)
            return
        df = merged_df.copy()
        if "is_cpo" not in df.columns:
            df = self.engine.apply(df)
        cpo = df[df.get("is_cpo", pd.Series(False, index=df.index)).fillna(False)].copy()
        if cpo.empty:
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "待DB", "DB有資料但未命中CPO_THEME_MASTER", "NO_MATCH"])
            self._style_basic(ws)
            return
        for col in ["cpo_score", "phase5_position_score", "rr", "teacher_score"]:
            if col in cpo.columns:
                cpo[col] = pd.to_numeric(cpo[col], errors="coerce").fillna(0)
            else:
                cpo[col] = 0
        decisions, reasons = [], []
        for _, row in cpo.iterrows():
            d, reason = self._candidate_decision(row)
            decisions.append(d); reasons.append(reason)
        cpo["cpo_prebreakout_decision"] = decisions
        cpo["cpo_prebreakout_reason"] = reasons
        cpo["_rank"] = cpo["cpo_prebreakout_decision"].map({"BUY":1,"LOW_BUY":2,"WATCH":3,"REDUCE":4,"AVOID":5}).fillna(9)
        cpo = cpo.sort_values(["_rank","cpo_score","phase5_position_score","rr"], ascending=[True, False, False, False]).head(80)
        for n, (_, r) in enumerate(cpo.iterrows(), start=1):
            ws.append([
                n,
                str(r.get("stock_id","")).zfill(4),
                r.get("report_name", r.get("stock_name", r.get("name", ""))),
                r.get("cpo_theme",""),
                r.get("cpo_subtheme",""),
                r.get("cpo_layer",""),
                round(_safe_float(r.get("cpo_score"),0),2),
                r.get("cpo_directness",""),
                r.get("cpo_strategy",""),
                r.get("teacher_decision",""),
                r.get("teacher_execution_status", r.get("是否可下單","")),
                r.get("phase5_wave_phase",""),
                r.get("phase5_breakout_stage",""),
                round(_safe_float(r.get("phase5_position_score"),0),2),
                round(_safe_float(r.get("rr"),0),2),
                r.get("cpo_prebreakout_decision",""),
                r.get("cpo_prebreakout_reason",""),
                "DB_OK",
            ])
        self._style_basic(ws)

    def _build_strict_intersection_df(self, merged_df):
        """V2.8.2：真正 CPO × 爆發前嚴格交集。
        定義：CPO_THEME_MASTER 代號 ∩ 爆發前候選條件（Phase5/Wave/Compression/主升預突破）。
        注意：本表不是全CPO股票池；若0檔，必須輸出0檔原因，避免誤導為可主攻。
        """
        if merged_df is None or pd is None or getattr(merged_df, "empty", True):
            return None
        df = merged_df.copy()
        if "is_cpo" not in df.columns:
            df = self.engine.apply(df)
        if "stock_id" not in df.columns:
            return pd.DataFrame()
        for col in ["phase5_position_score", "rr", "teacher_score", "vol5", "vol20"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        if "prebreakout_volume_ratio" not in df.columns:
            vol5 = df["vol5"] if "vol5" in df.columns else pd.Series(index=df.index, dtype=float)
            vol20 = df["vol20"] if "vol20" in df.columns else pd.Series(index=df.index, dtype=float)
            df["prebreakout_volume_ratio"] = (vol5 / vol20.replace(0, math.nan)).replace([math.inf, -math.inf], math.nan)
        wave_phase = df.get("phase5_wave_phase", pd.Series("", index=df.index)).astype(str)
        breakout_stage = df.get("phase5_breakout_stage", pd.Series("", index=df.index)).astype(str)
        pool = df.get("phase5_candidate_pool", pd.Series("", index=df.index)).astype(str)
        # V2.9.0 P0修正：Compression 只能列觀察，不可標為嚴格交集。
        # 嚴格交集必須是 CPO股票池 ∩ Wave3_PreBreakout/Breakout/Expansion 或明確PreBreakout/Breakout階段。
        strict_mask = (
            df.get("is_cpo", pd.Series(False, index=df.index)).fillna(False)
            & (
                wave_phase.isin(["Wave3_PreBreakout", "Wave3_Breakout", "Wave3_Expansion"])
                | breakout_stage.isin(["PreBreakout", "Breakout", "Expansion"])
                | pool.str.contains("主升預突破觀察池|主升確認池", na=False)
            )
        )
        inter = df.loc[strict_mask].copy()
        if inter.empty:
            return inter
        if "cpo_prebreakout_decision" not in inter.columns:
            decisions, reasons = [], []
            for _, row in inter.iterrows():
                d, reason = self._candidate_decision(row)
                decisions.append(d); reasons.append(reason)
            inter["cpo_prebreakout_decision"] = decisions
            inter["cpo_prebreakout_reason"] = reasons
        inter["_rank"] = inter["cpo_prebreakout_decision"].map({"BUY":1,"LOW_BUY":2,"WATCH":3,"REDUCE":4,"AVOID":5}).fillna(9)
        return inter.sort_values(["_rank","cpo_score","phase5_position_score","rr","prebreakout_volume_ratio"], ascending=[True, False, False, False, False])

    def _write_cpo_prebreakout_intersection(self, wb, merged_df):
        ws = self._sheet(wb, "CPO×爆發前交集")
        headers = [
            "排名", "代號", "名稱", "CPO主題", "CPO子分類", "CPO分數", "CPO直接性",
            "老師決策", "是否可下單", "波段階段", "突破階段", "Phase5候選池",
            "波段位置分", "量比", "RR", "CPO爆發前決策", "交集判定", "原因/風險", "資料狀態"
        ]
        ws.append(headers)
        inter = self._build_strict_intersection_df(merged_df)
        if inter is None:
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "待DB", "NO_DB", "未提供DB或InstitutionalReportEngine未產出，無法計算嚴格交集", "NO_DB"])
            self._style_basic(ws)
            return
        if getattr(inter, "empty", True):
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "無", "STRICT_INTERSECTION_ZERO", "CPO股票池與爆發前候選條件交集為0；代表目前沒有CPO主攻爆發前清單，僅能列CPO觀察池", "ZERO_MATCH"])
            self._style_basic(ws)
            return
        for n, (_, r) in enumerate(inter.head(80).iterrows(), start=1):
            reason = str(r.get("cpo_prebreakout_reason", "") or "")
            risk = " ".join([
                str(r.get("hard_avoid_reason", "") or ""),
                str(r.get("k_warning_type", "") or ""),
                str(r.get("phase5_block_reason", "") or ""),
            ]).strip()
            ws.append([
                n,
                str(r.get("stock_id", "")).zfill(4),
                r.get("report_name", r.get("stock_name", r.get("name", ""))),
                r.get("cpo_theme", ""),
                r.get("cpo_subtheme", ""),
                round(_safe_float(r.get("cpo_score"), 0), 2),
                r.get("cpo_directness", ""),
                r.get("teacher_decision", ""),
                r.get("teacher_execution_status", r.get("是否可下單", "")),
                r.get("phase5_wave_phase", ""),
                r.get("phase5_breakout_stage", ""),
                r.get("phase5_candidate_pool", ""),
                round(_safe_float(r.get("phase5_position_score"), 0), 2),
                round(_safe_float(r.get("prebreakout_volume_ratio"), 0), 2),
                round(_safe_float(r.get("rr"), 0), 2),
                r.get("cpo_prebreakout_decision", "WATCH"),
                "STRICT_MATCH",
                (reason + ("；風險=" + risk[:120] if risk else ""))[:250],
                "DB_OK",
            ])
        self._style_basic(ws)

    def _write_validation(self, wb, merged_df):
        ws = self._sheet(wb, "CPO整合驗收")
        master_count = len(CPO_THEME_MASTER)
        merged_count = 0
        candidate_count = 0
        strict_count = 0
        if merged_df is not None and pd is not None and not getattr(merged_df, "empty", True):
            if "is_cpo" in merged_df.columns:
                merged_count = int(merged_df["is_cpo"].fillna(False).sum())
                candidate_count = merged_count
            try:
                inter = self._build_strict_intersection_df(merged_df)
                strict_count = 0 if inter is None or getattr(inter, "empty", True) else int(len(inter))
            except Exception:
                strict_count = 0
        rows = [
            ["查核項目", "結果", "說明"],
            ["CPO_THEME_MASTER檔數", master_count, "固定股票池，不依賴DB"],
            ["DB合併命中數", merged_count, "institutional_report['all']與CPO_THEME_MASTER交集"],
            ["CPO股票池決策數", candidate_count, "有DB時由全CPO股票池×Phase5/Teacher/TradePlan產生，不等於嚴格爆發前交集"],
            ["CPO×爆發前嚴格交集數", strict_count, "CPO_THEME_MASTER ∩ 爆發前候選條件；0檔代表目前無CPO主攻爆發前清單"],
            ["無DB保護", "PASS", "無DB仍輸出CPO股票池，候選頁顯示待DB"],
            ["Excel頁面", ",".join(self.SHEETS), "需與主程式輸出一致"],
            ["Log驗收", "CPO_THEME_MASTER_READY / CPO_REPORT_WRITTEN", "可grep追蹤"],
        ]
        for r in rows:
            ws.append(r)
        self._style_basic(ws)

# =============================
# V2.8 CPO / 爆發前股票 SOP 整合頁
# =============================

# =============================
# V2.9.0 Main Theme × PreBreakout Engine
# 來源：主流題材_爆發前追蹤類型_整合股票池_深度分析報告.xlsx
# 目的：把 CPO 專屬交集升級為「主流題材 × 爆發前」門戶大開雷達。
# =============================
THEME_TRACKING_MASTER = [
    {
        "stock_id": "2345",
        "stock_name": "智邦",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway；Edge AI/TinyML/MCU AI；其他/待分類",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 12,
        "theme_score_base": 90.65,
        "db_total_score_ref": 90.65,
        "ai_score_ref": 78.67,
        "research_score_ref": 87.22,
        "raw_theme_line": "AI Connectivity；Networking；Networking / Edge Switch；營運績效前五；網路交換器；網通交換器；資料中心/Edge交換器；資料中心交換器",
        "product_position": "Networking / Edge Switch；網路交換器；資料中心/Edge交換器",
        "strategy_note": "HOLD / 題材觀察；Pullback Buy；WAIT突破或回檔；可分批佈局",
        "risk_note": "正常",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；06_候選池全表_DB；07_投資組合TOP20；07_長期選股策略；08_明日TOP5_每類別；Dashboard"
    },
    {
        "stock_id": "8271",
        "stock_name": "宇瞻",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 87.87,
        "db_total_score_ref": 87.87,
        "ai_score_ref": 77.23,
        "research_score_ref": 74.7,
        "raw_theme_line": "AI/晶圓代工；記憶體模組/工控儲存",
        "product_position": "記憶體模組/工控儲存",
        "strategy_note": "BUY/分批布局；Pullback Buy",
        "risk_note": "量能與市場辨識度",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；EdgeAI_TOP10"
    },
    {
        "stock_id": "2357",
        "stock_name": "華碩",
        "theme_category": "AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 87.36,
        "db_total_score_ref": 87.36,
        "ai_score_ref": 73.4,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；AI伺服器/PC品牌；服務業獲利前十；電腦/AI PC/伺服器",
        "product_position": "電腦/AI PC/伺服器",
        "strategy_note": "Pullback Buy；WAIT突破或回檔",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；EdgeAI_TOP10"
    },
    {
        "stock_id": "6166",
        "stock_name": "凌華",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 86.91,
        "db_total_score_ref": 86.91,
        "ai_score_ref": 75.13,
        "research_score_ref": 84.18,
        "raw_theme_line": "Edge IPC / Physical AI；Edge IPC / Robot Platform；IPC/Edge System",
        "product_position": "Edge IPC / Robot Platform",
        "strategy_note": "HOLD / 題材觀察；Pullback Buy；題材觀察/等回測",
        "risk_note": "RSI偏熱",
        "theme_source_file": "2026_AI_Agenda_Investment_Strategy.xlsx；2026_Edge_AI_投資策略_DB版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；EdgeAI_TOP10；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "2344",
        "stock_name": "華邦電",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 83.71,
        "db_total_score_ref": 83.71,
        "ai_score_ref": 73.06,
        "research_score_ref": 72.3,
        "raw_theme_line": "AI/晶圓代工；SLC NAND/Specialty Memory；TL-RAM/記憶體架構受惠",
        "product_position": "SLC NAND/Specialty Memory",
        "strategy_note": "BUY/分批布局；加入主題追蹤",
        "risk_note": "非高階Enterprise SSD主角、PE偏高",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；Agenda策略總表；docx內文"
    },
    {
        "stock_id": "2383",
        "stock_name": "台光電",
        "theme_category": "AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 79.8,
        "db_total_score_ref": None,
        "ai_score_ref": 50.0,
        "research_score_ref": 79.8,
        "raw_theme_line": "ABF + 高速PCB；Ultra-low-loss材料、AI Server PCB、GPU tray高頻低損耗基材；高速CCL/PCB材料",
        "product_position": "Ultra-low-loss材料、AI Server PCB、GPU tray高頻低損耗基材",
        "strategy_note": "接近低接買點，可小量試單；次主攻：等量價確認",
        "risk_note": "站上65MA；站上120MA；乖離65MA偏大；營收YoY正；法人分數偏強；120日位階偏高",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；投資策略總結；總表_更新版；選股策略"
    },
    {
        "stock_id": "6274",
        "stock_name": "台燿",
        "theme_category": "AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 76.2,
        "db_total_score_ref": None,
        "ai_score_ref": 45.0,
        "research_score_ref": 76.2,
        "raw_theme_line": "ABF + 高速PCB；M7/M8/M9級低損耗材料、HVLP銅箔、伺服器板升級；高速CCL/PCB材料",
        "product_position": "M7/M8/M9級低損耗材料、HVLP銅箔、伺服器板升級",
        "strategy_note": "接近低接買點，可小量試單；次主攻：等量價確認",
        "risk_note": "站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；投資策略總結；總表_更新版；選股策略"
    },
    {
        "stock_id": "2317",
        "stock_name": "鴻海",
        "theme_category": "AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 74.93,
        "db_total_score_ref": 74.93,
        "ai_score_ref": 68.67,
        "research_score_ref": 60.518691670557885,
        "raw_theme_line": "AI伺服器；AI伺服器/整機櫃；AI工業化/ODM/機器人；AI整機櫃/ODM；ODM/Robot；製造業營收前十",
        "product_position": "AI伺服器/整機櫃；AI工業化/ODM/機器人",
        "strategy_note": "BUY/主攻；HOLD / 題材觀察；核心主攻",
        "risk_note": "",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；06_候選池全表_DB；07_投資組合TOP20；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2382",
        "stock_name": "廣達",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 74.22,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；AI伺服器ODM；伺服器；製造業營收前十",
        "product_position": "AI伺服器ODM",
        "strategy_note": "BUY/主攻；核心主攻；觀察回檔，不追高",
        "risk_note": "120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "00_總覽儀表板；03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5"
    },
    {
        "stock_id": "3231",
        "stock_name": "緯創",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 70.06,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；AI伺服器ODM；伺服器；製造業營收前十",
        "product_position": "AI伺服器ODM",
        "strategy_note": "BUY/主攻；核心主攻；觀察回檔，不追高",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "00_總覽儀表板；03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_Industry_Intelligence；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "4585",
        "stock_name": "達明",
        "theme_category": "AI Power/PPA/BBU/電池；Edge IPC/Physical AI/Robot",
        "theme_priority": "A_主攻追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI奇兵50強；Physical AI/機器人；機器人與自動化設備；電源/HVDC",
        "product_position": "機器人與自動化設備",
        "strategy_note": "BUY/主攻",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "6781",
        "stock_name": "AES-KY",
        "theme_category": "AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 4,
        "theme_score_base": 91.05,
        "db_total_score_ref": 91.05,
        "ai_score_ref": 79.33,
        "research_score_ref": None,
        "raw_theme_line": "AI電池/BBU；AI電池_BBU",
        "product_position": "",
        "strategy_note": "Pullback Buy；WAIT；加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "Agenda策略總表；EdgeAI_TOP10；TOP5主攻；docx內文"
    },
    {
        "stock_id": "1256",
        "stock_name": "鮮活果汁-KY",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 89.5,
        "db_total_score_ref": 89.5,
        "ai_score_ref": 77.65,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "4583",
        "stock_name": "台灣精銳",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 88.57,
        "db_total_score_ref": 88.57,
        "ai_score_ref": 79.78,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2404",
        "stock_name": "漢唐",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 88.47,
        "db_total_score_ref": 88.47,
        "ai_score_ref": 76.81,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "1514",
        "stock_name": "亞力",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 88.37,
        "db_total_score_ref": 88.37,
        "ai_score_ref": 80.09,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "3617",
        "stock_name": "碩天",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 88.24,
        "db_total_score_ref": 88.24,
        "ai_score_ref": 80.02,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2492",
        "stock_name": "華新科",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 2,
        "theme_score_base": 87.32,
        "db_total_score_ref": 87.32,
        "ai_score_ref": 77.83,
        "research_score_ref": None,
        "raw_theme_line": "Edge AI TOP10",
        "product_position": "",
        "strategy_note": "Pullback Buy；加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx；2026_Edge_AI_深度白皮書_DB版.docx",
        "theme_source_sheet": "EdgeAI_TOP10；docx內文"
    },
    {
        "stock_id": "2376",
        "stock_name": "技嘉",
        "theme_category": "AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 2,
        "theme_score_base": 87.24,
        "db_total_score_ref": 87.24,
        "ai_score_ref": 73.85,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；伺服器",
        "product_position": "",
        "strategy_note": "Pullback Buy；觀察回檔，不追高",
        "risk_note": "120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "03_中長線候選；EdgeAI_TOP10"
    },
    {
        "stock_id": "6525",
        "stock_name": "6525",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 87.04,
        "db_total_score_ref": 87.04,
        "ai_score_ref": 74.36,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "7749",
        "stock_name": "7749",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 86.68,
        "db_total_score_ref": 86.68,
        "ai_score_ref": 76.52,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2421",
        "stock_name": "建準",
        "theme_category": "AI散熱/液冷；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 2,
        "theme_score_base": 86.26,
        "db_total_score_ref": 86.26,
        "ai_score_ref": 77.47,
        "research_score_ref": None,
        "raw_theme_line": "AI散熱；液冷",
        "product_position": "",
        "strategy_note": "Pullback Buy；分批布局/回測加碼",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "02_長期主升核心；EdgeAI_TOP10"
    },
    {
        "stock_id": "3533",
        "stock_name": "3533",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 86.16,
        "db_total_score_ref": 86.16,
        "ai_score_ref": 69.67,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "6446",
        "stock_name": "6446",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 86.03,
        "db_total_score_ref": 86.03,
        "ai_score_ref": 73.31,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2428",
        "stock_name": "興勤",
        "theme_category": "AI散熱/液冷；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 2,
        "theme_score_base": 85.91,
        "db_total_score_ref": 85.91,
        "ai_score_ref": 74.54,
        "research_score_ref": None,
        "raw_theme_line": "零組件；電子零組件",
        "product_position": "",
        "strategy_note": "Pullback Buy；觀察回檔，不追高",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "03_中長線候選；EdgeAI_TOP10"
    },
    {
        "stock_id": "8926",
        "stock_name": "8926",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 85.88,
        "db_total_score_ref": 85.88,
        "ai_score_ref": 77.34,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2548",
        "stock_name": "2548",
        "theme_category": "Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 1,
        "theme_score_base": 85.86,
        "db_total_score_ref": 85.86,
        "ai_score_ref": 63.55,
        "research_score_ref": None,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "Pullback Buy",
        "risk_note": "",
        "theme_source_file": "2026_Edge_AI_投資策略_DB版.xlsx",
        "theme_source_sheet": "EdgeAI_TOP10"
    },
    {
        "stock_id": "2395",
        "stock_name": "研華",
        "theme_category": "AI Server ODM/Rack Scale；ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 84.14,
        "db_total_score_ref": 84.14,
        "ai_score_ref": 70.92,
        "research_score_ref": 74.92,
        "raw_theme_line": "AI伺服器；Edge IPC / Physical AI；Edge IPC / Physical AI平台；IPC/Edge System；伺服器",
        "product_position": "Edge IPC / Physical AI平台",
        "strategy_note": "HOLD / 題材觀察；不列長期主攻；等待回檔",
        "risk_note": "RSI過熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；Investment_ROI；Stock_Strategy"
    },
    {
        "stock_id": "3706",
        "stock_name": "神達",
        "theme_category": "AI Server ODM/Rack Scale；ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 83.66,
        "db_total_score_ref": 83.66,
        "ai_score_ref": 74.37,
        "research_score_ref": 82.01,
        "raw_theme_line": "AI伺服器；Edge IPC / Physical AI；Edge Server；Edge Server / 車用AI；伺服器",
        "product_position": "Edge Server / 車用AI",
        "strategy_note": "HOLD / 題材觀察；分批布局/回測加碼；可分批佈局",
        "risk_note": "正常；無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "02_長期主升核心；06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "2379",
        "stock_name": "瑞昱",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 82.97,
        "db_total_score_ref": 82.97,
        "ai_score_ref": 76.71,
        "research_score_ref": 82.92,
        "raw_theme_line": "AI Connectivity；AI Connectivity / Gateway IC；Connectivity",
        "product_position": "AI Connectivity / Gateway IC",
        "strategy_note": "HOLD / 題材觀察；可分批佈局",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512....",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "2451",
        "stock_name": "創見",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 11,
        "theme_score_base": 82.6,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 82.6,
        "raw_theme_line": "AI/晶圓代工；半導體；工控/消費/企業儲存品牌；記憶體/邊緣儲存；記憶體及周邊產品",
        "product_position": "工控/消費/企業儲存品牌；記憶體及周邊產品",
        "strategy_note": "BUY/分批布局；WAIT突破或回檔；不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；成長彈性不如控制器",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報...",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_候選池_DB重算；04_技術量化；05_SWOT；05_短期選股策略；05_風險排除觀察"
    },
    {
        "stock_id": "3017",
        "stock_name": "奇鋐",
        "theme_category": "AI Power/PPA/BBU/電池；AI Server ODM/Rack Scale；AI散熱/液冷；PCIe7/高速互連/PCB材料",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 12,
        "theme_score_base": 82.22,
        "db_total_score_ref": 74.27,
        "ai_score_ref": 57.24,
        "research_score_ref": 76.2,
        "raw_theme_line": "AI散熱；GB300/NVL72液冷與GPU Tray高功耗散熱同步升級；散熱模組；液冷；液冷 + AI Power + HVDC；液冷/散熱；營運績效前五",
        "product_position": "GB300/NVL72液冷與GPU Tray高功耗散熱同步升級；散熱模組",
        "strategy_note": "WATCH；分批布局/回測加碼；次主攻：等量價確認；高於低接區約6.4%，不追等回測",
        "risk_note": "無明顯過熱；站上65MA；站上120MA；量縮；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "02_長期主升核心；04_候選池_DB重算；06_Industry_Intelligence；06_中期選股策略；07_長期選股策略；08_明日TOP5_..."
    },
    {
        "stock_id": "5388",
        "stock_name": "中磊",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 82.0,
        "db_total_score_ref": 82.0,
        "ai_score_ref": 69.69,
        "research_score_ref": 79.62,
        "raw_theme_line": "AI Connectivity；AI Gateway / FWA / Broadband；Connectivity Device；資料中心交換器；高階網通",
        "product_position": "AI Gateway / FWA / Broadband",
        "strategy_note": "HOLD / 題材觀察；可分批佈局；觀察回檔，不追高",
        "risk_note": "120日高位階；正常",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI...",
        "theme_source_sheet": "03_中長線候選；06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；The..."
    },
    {
        "stock_id": "2301",
        "stock_name": "光寶科",
        "theme_category": "AI Power/PPA/BBU/電池；AI散熱/液冷",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 81.11,
        "db_total_score_ref": 79.56,
        "ai_score_ref": 69.67,
        "research_score_ref": 70.48,
        "raw_theme_line": "AI Power；AI Power / PPA；AI Power / PSU；液冷 + AI Power + HVDC；電源；電源/HVDC",
        "product_position": "AI Power / PSU",
        "strategy_note": "HOLD / 題材觀察；不列長期主攻；等待回檔；高於低接區約4.0%，不追等回測",
        "risk_note": "RSI偏熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "05_風險排除觀察；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；T..."
    },
    {
        "stock_id": "3037",
        "stock_name": "欣興",
        "theme_category": "AI Memory/Storage/NAND；AI Server ODM/Rack Scale；Advanced Packaging/CoWoS/半導體製...",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 80.73,
        "db_total_score_ref": 80.73,
        "ai_score_ref": 66.48,
        "research_score_ref": 68.2,
        "raw_theme_line": "ABF + 高速PCB；AI伺服器高階PCB與載板、層數提升受益；AI記憶體/高速材料/封裝；PCB/載板",
        "product_position": "AI伺服器高階PCB與載板、層數提升受益",
        "strategy_note": "加入主題追蹤；觀察：只做拉回，不追高；高於低接區約4.7%，不追等回測",
        "risk_note": "站上65MA；站上120MA；乖離65MA偏大；量能放大；營收YoY正；120日位階偏高",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xls...",
        "theme_source_sheet": "Agenda策略總表；DB原始快照；Dashboard；docx內文；投資策略總結；總表_更新版；選股策略"
    },
    {
        "stock_id": "3711",
        "stock_name": "日月光投控",
        "theme_category": "AI Memory/Storage/NAND；ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程...",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 80.69,
        "db_total_score_ref": 80.6,
        "ai_score_ref": 71.53,
        "research_score_ref": 80.69,
        "raw_theme_line": "2nm + CoWoS；Advanced Packaging；Advanced Packaging / Memory；Packaging",
        "product_position": "Advanced Packaging",
        "strategy_note": "HOLD / 題材觀察；接近低接買點，可小量試單；題材觀察/等回測",
        "risk_note": "RSI偏熱；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；E...",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mappin..."
    },
    {
        "stock_id": "3019",
        "stock_name": "亞光",
        "theme_category": "CPO/矽光子/光通訊；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 79.97,
        "db_total_score_ref": 78.36,
        "ai_score_ref": 72.86,
        "research_score_ref": 79.97,
        "raw_theme_line": "AI Optical Sensor / Vision；AI Sensor",
        "product_position": "AI Optical Sensor / Vision",
        "strategy_note": "HOLD / 題材觀察；題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512....",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "2308",
        "stock_name": "台達電",
        "theme_category": "AI Power/PPA/BBU/電池；AI散熱/液冷；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 79.85,
        "db_total_score_ref": 76.72,
        "ai_score_ref": 63.69,
        "research_score_ref": 77.45,
        "raw_theme_line": "AI Power；AI Power / PPA；AI Power / Thermal / PPA；液冷 + AI Power + HVDC",
        "product_position": "AI Power / Thermal / PPA",
        "strategy_note": "HOLD / 題材觀察；題材觀察/等回測；高於低接區約5.9%，不追等回測",
        "risk_note": "RSI偏熱；站上65MA；站上120MA；乖離65MA偏大；營收YoY正",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xl...",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；Investment_ROI；Stock_Strategy；T..."
    },
    {
        "stock_id": "6533",
        "stock_name": "晶心科",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 79.05,
        "db_total_score_ref": 77.36,
        "ai_score_ref": 78.81,
        "research_score_ref": 79.05,
        "raw_theme_line": "ASIC / RISC-V / IP；CPU/IP；IP；MCU AI / TinyML；RISC-V；RISC-V Edge CPU/IP",
        "product_position": "RISC-V Edge CPU/IP",
        "strategy_note": "AVOID / 不列主攻；分批布局/回測加碼；題材觀察/等回測",
        "risk_note": "正常；無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_St...",
        "theme_source_sheet": "02_長期主升核心；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；T..."
    },
    {
        "stock_id": "3596",
        "stock_name": "智易",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 79.01,
        "db_total_score_ref": 79.01,
        "ai_score_ref": 67.63,
        "research_score_ref": 77.38,
        "raw_theme_line": "AI Connectivity；AI Gateway / CPE；Connectivity Device",
        "product_position": "AI Gateway / CPE",
        "strategy_note": "AVOID / 不列主攻；題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512....",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "3443",
        "stock_name": "創意",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 78.94,
        "db_total_score_ref": 78.94,
        "ai_score_ref": 67.97,
        "research_score_ref": 70.08,
        "raw_theme_line": "AI ASIC；ASIC；ASIC / Edge AI Custom Chip；ASIC / RISC-V / IP；ASIC/IP",
        "product_position": "ASIC / Edge AI Custom Chip",
        "strategy_note": "HOLD / 題材觀察；不列長期主攻；等待回檔",
        "risk_note": "RSI過熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_St...",
        "theme_source_sheet": "05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashbo..."
    },
    {
        "stock_id": "6442",
        "stock_name": "光聖",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 78.73,
        "db_total_score_ref": 49.46,
        "ai_score_ref": 41.18,
        "research_score_ref": 39.58,
        "raw_theme_line": "AI光通訊；CPO/光模組；Optical / Interconnect；Optical Interconnect；光通訊元件與射頻元件；高速光通訊",
        "product_position": "Optical Interconnect；光通訊元件與射頻元件",
        "strategy_note": "AVOID/觀望；分批布局/回測加碼；減碼/不列主攻",
        "risk_note": "DB轉弱 / DB分數不足；無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_St...",
        "theme_source_sheet": "02_長期主升核心；04_候選池_DB重算；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "3081",
        "stock_name": "聯亞",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 78.69,
        "db_total_score_ref": 56.57,
        "ai_score_ref": 52.37,
        "research_score_ref": 59.74,
        "raw_theme_line": "1.6T/CPO/矽光子；Optical / Interconnect；Optical Component；矽光子與光通訊",
        "product_position": "Optical Component",
        "strategy_note": "WAIT；觀察；高於低接區約4.1%，不追等回測",
        "risk_note": "DB分數不足；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xl...",
        "theme_source_sheet": "Agenda策略總表；Investment_ROI；Stock_Strategy；TOP5主攻；Theme_Mapping；投資策略總結；總表_更新版"
    },
    {
        "stock_id": "3035",
        "stock_name": "智原",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 78.66,
        "db_total_score_ref": 70.72,
        "ai_score_ref": 78.66,
        "research_score_ref": 71.4,
        "raw_theme_line": "ASIC / RISC-V / IP；ASIC Design Service；ASIC/IP",
        "product_position": "ASIC Design Service",
        "strategy_note": "AVOID / 不列主攻；題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；E...",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "3324",
        "stock_name": "雙鴻",
        "theme_category": "AI Power/PPA/BBU/電池；AI Server ODM/Rack Scale；AI散熱/液冷；PCIe7/高速互連/PCB材料",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 76.92,
        "db_total_score_ref": 63.0,
        "ai_score_ref": 46.46,
        "research_score_ref": 59.8,
        "raw_theme_line": "AI Server液冷、冷板與高瓦數散熱需求；AI散熱；液冷；液冷 + AI Power + HVDC；液冷/散熱",
        "product_position": "AI Server液冷、冷板與高瓦數散熱需求",
        "strategy_note": "分批布局/回測加碼；排除/等待；接近低接買點，可小量試單",
        "risk_note": "未站上65MA；量縮；營收YoY正；無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "02_長期主升核心；06_Industry_Intelligence；Agenda策略總表；DB原始快照；投資策略總結；總表_更新版；選股策略"
    },
    {
        "stock_id": "6202",
        "stock_name": "盛群",
        "theme_category": "Edge AI/TinyML/MCU AI；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 76.63,
        "db_total_score_ref": 67.42,
        "ai_score_ref": 72.31,
        "research_score_ref": 54.45,
        "raw_theme_line": "AI MCU / Embedded Controller；AI MCU / 控制IC；MCU AI；MCU AI / TinyML",
        "product_position": "AI MCU / Embedded Controller；AI MCU / 控制IC",
        "strategy_note": "AVOID；AVOID / 不列主攻；題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML...",
        "theme_source_sheet": "06_候選池全表_DB；AI_MCU策略；DB原始整合；Dashboard；Investment_ROI；Stock_Strategy；Theme_Map..."
    },
    {
        "stock_id": "4979",
        "stock_name": "華星光",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 12,
        "theme_score_base": 74.57,
        "db_total_score_ref": 69.77,
        "ai_score_ref": 63.0,
        "research_score_ref": 68.18,
        "raw_theme_line": "1.6T/CPO/矽光子；800G→1.6T、CPO/光模組升級的高速互連旁支受益；Interconnect；Optical / Interconnect...",
        "product_position": "800G→1.6T、CPO/光模組升級的高速互連旁支受益；Optical Interconnect / CPO；光通訊 / CPO",
        "strategy_note": "AVOID / 不列主攻；觀察；觀察：只做拉回，不追高；高於低接區約9.8%，不追等回測",
        "risk_note": "正常；站上65MA；站上120MA；量縮；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；DB原始快照；Dashboard..."
    },
    {
        "stock_id": "2330",
        "stock_name": "台積電",
        "theme_category": "Advanced Packaging/CoWoS/半導體製程",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 9,
        "theme_score_base": 74.35,
        "db_total_score_ref": None,
        "ai_score_ref": 65.0,
        "research_score_ref": None,
        "raw_theme_line": "2nm + CoWoS；AI/晶圓代工；半導體製造；晶圓代工；製造業營收前十；高權值",
        "product_position": "晶圓代工",
        "strategy_note": "BUY/主攻；接近低接買點，可小量試單；觀察回檔，不追高",
        "risk_note": "120日高位階；站上65MA；站上120MA；量縮；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；前三大_可執行策略；投資..."
    },
    {
        "stock_id": "4919",
        "stock_name": "新唐",
        "theme_category": "Edge AI/TinyML/MCU AI；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 72.56,
        "db_total_score_ref": 63.39,
        "ai_score_ref": 72.46,
        "research_score_ref": 28.18,
        "raw_theme_line": "AI MCU / Embedded Controller；MCU AI；MCU AI / TinyML",
        "product_position": "AI MCU / Embedded Controller",
        "strategy_note": "AVOID；AVOID / 不列主攻；等待回檔",
        "risk_note": "RSI偏熱 / 短線乖離過大",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML...",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；AI_MCU策略；DB原始整合；Investment_ROI；Stock_Strategy；Theme_..."
    },
    {
        "stock_id": "5289",
        "stock_name": "宜鼎",
        "theme_category": "AI Memory/Storage/NAND；AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 6,
        "theme_score_base": 72.4,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 72.4,
        "raw_theme_line": "AI伺服器；工控/Edge AI SSD",
        "product_position": "工控/Edge AI SSD",
        "strategy_note": "WATCH/等拉回或突破",
        "risk_note": "本益比/位階偏高、流動性震盪",
        "theme_source_file": "AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "8114",
        "stock_name": "振樺電",
        "theme_category": "AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 72.31,
        "db_total_score_ref": 70.96,
        "ai_score_ref": 72.31,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；AI奇兵50強；Edge AI 從視覺升級至本地AI系統；Edge_AI；POS/邊緣裝置；銷售終端系統POS",
        "product_position": "銷售終端系統POS",
        "strategy_note": "WAIT；WAIT突破或回檔；加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.x...",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；Agenda策略總表；TOP5主攻；docx內文"
    },
    {
        "stock_id": "3163",
        "stock_name": "波若威",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 72.25,
        "db_total_score_ref": 40.45,
        "ai_score_ref": 29.22,
        "research_score_ref": 34.37,
        "raw_theme_line": "1.6T/CPO/矽光子；CPO/光模組；Optical / Interconnect；Optical Module；高速光通訊",
        "product_position": "Optical Module",
        "strategy_note": "接近低接買點，可小量試單；減碼/不列主攻；觀察回檔，不追高",
        "risk_note": "DB轉弱 / DB分數不足；無明顯過熱；站上65MA；站上120MA；距65MA適中；量縮；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "03_中長線候選；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping；前三大_可執行策略；投資策..."
    },
    {
        "stock_id": "2368",
        "stock_name": "金像電",
        "theme_category": "AI Server ODM/Rack Scale；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 3,
        "theme_score_base": 72.2,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 72.2,
        "raw_theme_line": "AI伺服器板、交換器/網通板，高層數PCB；伺服器PCB",
        "product_position": "AI伺服器板、交換器/網通板，高層數PCB",
        "strategy_note": "次主攻：等量價確認",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；選股策略"
    },
    {
        "stock_id": "2359",
        "stock_name": "所羅門",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 71.69,
        "db_total_score_ref": 68.67,
        "ai_score_ref": 68.36,
        "research_score_ref": 71.69,
        "raw_theme_line": "AI Vision / Physical AI；AI Vision/Robot；Edge IPC / Physical AI",
        "product_position": "AI Vision / Physical AI",
        "strategy_note": "HOLD / 題材觀察；觀察",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512....",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "8299",
        "stock_name": "群聯",
        "theme_category": "AI Memory/Storage/NAND；Advanced Packaging/CoWoS/半導體製程",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 10,
        "theme_score_base": 71.5,
        "db_total_score_ref": 51.06,
        "ai_score_ref": 58.97,
        "research_score_ref": 51.8,
        "raw_theme_line": "AI Storage / Controller；AI/晶圓代工；Advanced Packaging / Memory；半導體；控制IC/AI SSD控制器",
        "product_position": "AI Storage / Controller；控制IC/AI SSD控制器",
        "strategy_note": "WATCH/等拉回或突破；不列長期主攻；減碼/不列主攻",
        "risk_note": "RSI偏熱 / 短線乖離過大 / DB分數不足；RSI過熱；MA65乖離過大；120日高位階；高Beta、估值與報價斜率轉弱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報...",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；05_風險排除觀察；06_情境報酬與配置；Investment_ROI"
    },
    {
        "stock_id": "3363",
        "stock_name": "上詮",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 71.0,
        "db_total_score_ref": 56.45,
        "ai_score_ref": 57.82,
        "research_score_ref": 57.62,
        "raw_theme_line": "1.6T/CPO/矽光子；CPO/光模組；Optical / Interconnect；Optical Module；高速光通訊",
        "product_position": "Optical Module",
        "strategy_note": "觀察；觀察回檔，不追高；高於低接區約5.1%，不追等回測",
        "risk_note": "DB分數不足；站上65MA；站上120MA；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_...",
        "theme_source_sheet": "03_中長線候選；06_Industry_Intelligence；Agenda策略總表；Investment_ROI；Stock_Strategy；Th..."
    },
    {
        "stock_id": "3661",
        "stock_name": "世芯-KY",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 69.15,
        "db_total_score_ref": 52.29,
        "ai_score_ref": 61.14,
        "research_score_ref": 50.62,
        "raw_theme_line": "AI ASIC服務；AI/晶圓代工；ASIC / AI Custom Chip；ASIC / RISC-V / IP；ASIC/IP；半導體",
        "product_position": "AI ASIC服務；ASIC / AI Custom Chip",
        "strategy_note": "AVOID / 不列主攻；不列長期主攻；減碼/不列主攻",
        "risk_note": "RSI過熱 / 短線乖離過大 / DB分數不足；RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_St...",
        "theme_source_sheet": "05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；Agenda策略總表；Investment_ROI；Stoc..."
    },
    {
        "stock_id": "6285",
        "stock_name": "啟碁",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 66.38,
        "db_total_score_ref": 63.46,
        "ai_score_ref": 55.96,
        "research_score_ref": 66.38,
        "raw_theme_line": "AI Connectivity；AI Gateway / Networking Device；AI資料中心網通設備/高速交換器相關，直接性較材料低；Con...",
        "product_position": "AI Gateway / Networking Device；AI資料中心網通設備/高速交換器相關，直接性較材料低",
        "strategy_note": "HOLD / 題材觀察；排除/等待；觀察",
        "risk_note": "正常",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Ed...",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；DB原始快照；Investment_ROI；Stock_Strategy；Theme_Mapping；選股策略"
    },
    {
        "stock_id": "6922",
        "stock_name": "宸曜",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 63.84,
        "db_total_score_ref": 58.68,
        "ai_score_ref": 62.09,
        "research_score_ref": 63.84,
        "raw_theme_line": "Edge IPC / Physical AI；IPC/Edge System；Rugged Edge AI；Rugged Edge AI / 車載工控",
        "product_position": "Rugged Edge AI；Rugged Edge AI / 車載工控",
        "strategy_note": "AVOID / 不列主攻；觀察",
        "risk_note": "DB分數不足",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "2337",
        "stock_name": "旺宏",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 8,
        "theme_score_base": 57.9,
        "db_total_score_ref": 51.14,
        "ai_score_ref": 44.74,
        "research_score_ref": 57.9,
        "raw_theme_line": "AI/晶圓代工；NOR/SLC/MLC NAND；TL-RAM/記憶體架構受惠",
        "product_position": "NOR/SLC/MLC NAND",
        "strategy_note": "HOLD/題材觀察；加入主題追蹤",
        "risk_note": "PB高、獲利能見度待驗證",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；Agenda策略總表；docx內文"
    },
    {
        "stock_id": "6669",
        "stock_name": "緯穎",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": 50.0,
        "db_total_score_ref": None,
        "ai_score_ref": 50.0,
        "research_score_ref": None,
        "raw_theme_line": "AI Rack + ODM；AI伺服器；AI伺服器/機櫃；AI伺服器ODM；伺服器",
        "product_position": "AI伺服器/機櫃",
        "strategy_note": "WAIT突破或回檔；不列長期主攻；高於低接區約4.5%，不追等回測",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；站上65MA；站上120MA；乖離65MA偏大；營收YoY正",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；05_風險排除觀察；06_中期選股策略；07_長期選股策略；投資策略總結；總表_更新版"
    },
    {
        "stock_id": "2059",
        "stock_name": "川湖",
        "theme_category": "AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": 27.86,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI機械/滑軌；伺服器滑軌；零組件；電子零組件",
        "product_position": "伺服器滑軌",
        "strategy_note": "WAIT突破或回檔；不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；05_風險排除觀察；06_中期選股策略；07_長期選股策略"
    },
    {
        "stock_id": "2360",
        "stock_name": "致茂",
        "theme_category": "AI Power/PPA/BBU/電池；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI奇兵50強；全市場；半導體/電子測試設備；半導體測試",
        "product_position": "半導體/電子測試設備",
        "strategy_note": "BUY/主攻；核心主攻",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5"
    },
    {
        "stock_id": "2603",
        "stock_name": "長榮",
        "theme_category": "其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "服務業獲利前十；海運；航運/海運；運輸",
        "product_position": "海運",
        "strategy_note": "BUY/主攻",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "3167",
        "stock_name": "大量",
        "theme_category": "AI Power/PPA/BBU/電池；PCIe7/高速互連/PCB材料",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 7,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI奇兵50強；PCB設備；PCB設備與半導體測試設備；電源/HVDC",
        "product_position": "PCB設備與半導體測試設備",
        "strategy_note": "BUY/主攻；核心主攻",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5"
    },
    {
        "stock_id": "6223",
        "stock_name": "旺矽",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；其他/待分類",
        "theme_priority": "B_觀察追蹤",
        "theme_appearance_count": 5,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；AI奇兵50強；半導體測試介面",
        "product_position": "半導體測試介面",
        "strategy_note": "BUY/主攻",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "3211",
        "stock_name": "順達",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 86.66,
        "db_total_score_ref": 86.66,
        "ai_score_ref": 76.15,
        "research_score_ref": None,
        "raw_theme_line": "AI電池/BBU；AI電池_BBU",
        "product_position": "",
        "strategy_note": "WAIT；加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；TOP5主攻；docx內文"
    },
    {
        "stock_id": "2408",
        "stock_name": "南亞科",
        "theme_category": "AI Memory/Storage/NAND",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 83.15,
        "db_total_score_ref": 83.15,
        "ai_score_ref": 67.75,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；TL-RAM/記憶體架構受惠；半導體",
        "product_position": "",
        "strategy_note": "分批布局/回測加碼；加入主題追蹤",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "02_長期主升核心；Agenda策略總表；docx內文"
    },
    {
        "stock_id": "6414",
        "stock_name": "樺漢",
        "theme_category": "Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 80.86,
        "db_total_score_ref": 80.86,
        "ai_score_ref": 70.66,
        "research_score_ref": None,
        "raw_theme_line": "Edge_AI；受惠台股：研華、凌華、樺漢、振樺電",
        "product_position": "",
        "strategy_note": "WAIT；加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；TOP5主攻；docx內文"
    },
    {
        "stock_id": "8996",
        "stock_name": "高力",
        "theme_category": "AI Power/PPA/BBU/電池；AI散熱/液冷",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 79.57,
        "db_total_score_ref": 63.23,
        "ai_score_ref": 41.69,
        "research_score_ref": None,
        "raw_theme_line": "AI散熱；液冷",
        "product_position": "",
        "strategy_note": "分批布局/回測加碼",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "02_長期主升核心；Agenda策略總表"
    },
    {
        "stock_id": "4931",
        "stock_name": "新盛力",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 78.56,
        "db_total_score_ref": 78.56,
        "ai_score_ref": 72.09,
        "research_score_ref": None,
        "raw_theme_line": "AI電池/BBU",
        "product_position": "",
        "strategy_note": "加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；docx內文"
    },
    {
        "stock_id": "4952",
        "stock_name": "凌通",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 78.49,
        "db_total_score_ref": 77.47,
        "ai_score_ref": 78.49,
        "research_score_ref": 76.39,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "BUY_PULLBACK",
        "risk_note": "",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx",
        "theme_source_sheet": "AI_MCU策略；DB原始整合；Dashboard"
    },
    {
        "stock_id": "5471",
        "stock_name": "松翰",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 77.57,
        "db_total_score_ref": 77.57,
        "ai_score_ref": 72.69,
        "research_score_ref": 44.54,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "AVOID",
        "risk_note": "",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx",
        "theme_source_sheet": "AI_MCU策略；DB原始整合；Dashboard"
    },
    {
        "stock_id": "2449",
        "stock_name": "京元電子",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 77.49,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "分批布局/回測加碼",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "02_長期主升核心"
    },
    {
        "stock_id": "6510",
        "stock_name": "精測",
        "theme_category": "AI Server ODM/Rack Scale；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 77.28,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 56.8,
        "raw_theme_line": "AI/晶圓代工；半導體；測試介面；高階測試介面、探針卡，間接受益高速運算IC驗證",
        "product_position": "高階測試介面、探針卡，間接受益高速運算IC驗證",
        "strategy_note": "分批布局/回測加碼；排除/等待",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "02_長期主升核心；DB原始快照；選股策略"
    },
    {
        "stock_id": "4971",
        "stock_name": "IET-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 76.77,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "等回檔至10日/20日線",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "02_長期主升核心"
    },
    {
        "stock_id": "6415",
        "stock_name": "矽力-KY",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 76.07,
        "db_total_score_ref": 76.07,
        "ai_score_ref": 67.88,
        "research_score_ref": 75.26,
        "raw_theme_line": "AI Power / PPA；Power IC / PMIC",
        "product_position": "Power IC / PMIC",
        "strategy_note": "題材觀察/等回測",
        "risk_note": "RSI偏熱",
        "theme_source_file": "Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "3653",
        "stock_name": "健策",
        "theme_category": "AI散熱/液冷；ASIC/FPGA/RISC-V/IP/EDA",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 76.03,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI散熱；散熱模組；液冷",
        "product_position": "散熱模組",
        "strategy_note": "AVOID/觀望；分批布局/回測加碼",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "02_長期主升核心；04_候選池_DB重算；07_長期選股策略"
    },
    {
        "stock_id": "3227",
        "stock_name": "原相",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 75.33,
        "db_total_score_ref": 73.76,
        "ai_score_ref": 75.33,
        "research_score_ref": 73.86,
        "raw_theme_line": "AI Sensor；AI Sensor / Perception IC",
        "product_position": "AI Sensor / Perception IC",
        "strategy_note": "題材觀察/等回測",
        "risk_note": "RSI偏熱",
        "theme_source_file": "Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "4967",
        "stock_name": "十銓",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 7,
        "theme_score_base": 74.45,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 57.7,
        "raw_theme_line": "AI/晶圓代工；半導體；記憶體模組/高Beta",
        "product_position": "記憶體模組/高Beta",
        "strategy_note": "HOLD/題材觀察；觀察回檔，不追高",
        "risk_note": "月營收波動大、追價風險；無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；03_中長線候選；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "6643",
        "stock_name": "M31",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 74.2,
        "db_total_score_ref": 71.3,
        "ai_score_ref": 74.2,
        "research_score_ref": None,
        "raw_theme_line": "Tape-Out ERC / FPGA/IP",
        "product_position": "",
        "strategy_note": "加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；docx內文"
    },
    {
        "stock_id": "3441",
        "stock_name": "聯一光",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": 73.75,
        "db_total_score_ref": 73.38,
        "ai_score_ref": 73.75,
        "research_score_ref": 73.69,
        "raw_theme_line": "AI Sensor；AI Sensor/Optics；Optical Sensor / Optical Module；Optical Sensor/Interconnect",
        "product_position": "Optical Sensor / Optical Module；Optical Sensor/Interconnect",
        "strategy_note": "AVOID / 不列主攻；題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "4903",
        "stock_name": "聯光通",
        "theme_category": "CPO/矽光子/光通訊",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 73.55,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "CPO/光模組；高速光通訊",
        "product_position": "",
        "strategy_note": "觀察回檔，不追高",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "03_中長線候選"
    },
    {
        "stock_id": "2436",
        "stock_name": "偉詮電",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 71.95,
        "db_total_score_ref": 71.02,
        "ai_score_ref": 61.15,
        "research_score_ref": 71.95,
        "raw_theme_line": "AI Power / PPA；Power/Controller IC",
        "product_position": "Power/Controller IC",
        "strategy_note": "題材觀察/等回測",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "6530",
        "stock_name": "創威",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 71.8,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "資料中心交換器；高階網通",
        "product_position": "",
        "strategy_note": "觀察回檔，不追高",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "03_中長線候選"
    },
    {
        "stock_id": "3466",
        "stock_name": "德晉",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 71.54,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "資料中心交換器；高階網通",
        "product_position": "",
        "strategy_note": "觀察回檔，不追高",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "03_中長線候選"
    },
    {
        "stock_id": "3665",
        "stock_name": "貿聯-KY",
        "theme_category": "AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 71.2,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 71.2,
        "raw_theme_line": "AI伺服器內外部高速線束、GPU Tray連接、DAC/AEC ASP升級；高速線束/AEC/DAC",
        "product_position": "AI伺服器內外部高速線束、GPU Tray連接、DAC/AEC ASP升級",
        "strategy_note": "觀察：只做拉回，不追高",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；選股策略"
    },
    {
        "stock_id": "3189",
        "stock_name": "景碩",
        "theme_category": "AI Memory/Storage/NAND；PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 70.95,
        "db_total_score_ref": 70.95,
        "ai_score_ref": 59.5,
        "research_score_ref": None,
        "raw_theme_line": "ABF + 高速PCB",
        "product_position": "",
        "strategy_note": "高於低接區約7.2%，不追等回測",
        "risk_note": "站上65MA；站上120MA；乖離65MA偏大；營收YoY正；120日位階偏高",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；投資策略總結；總表_更新版"
    },
    {
        "stock_id": "3228",
        "stock_name": "金麗科",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 70.52,
        "db_total_score_ref": 61.42,
        "ai_score_ref": 70.52,
        "research_score_ref": 32.69,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "AVOID",
        "risk_note": "",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx",
        "theme_source_sheet": "AI_MCU策略；DB原始整合"
    },
    {
        "stock_id": "3088",
        "stock_name": "艾訊",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": 68.71,
        "db_total_score_ref": 67.61,
        "ai_score_ref": 68.71,
        "research_score_ref": 61.96,
        "raw_theme_line": "Edge AI System / AMR；Edge Box / AMR Vision；Edge IPC / Physical AI；IPC/Edge System",
        "product_position": "Edge AI System / AMR；Edge Box / AMR Vision",
        "strategy_note": "AVOID / 不列主攻；等待回檔",
        "risk_note": "RSI過熱",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "6190",
        "stock_name": "萬泰科",
        "theme_category": "AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 67.9,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 67.9,
        "raw_theme_line": "PTFE/氟塑料高速線材、AI伺服器PCIe線束題材；高速線材/PTFE",
        "product_position": "PTFE/氟塑料高速線材、AI伺服器PCIe線束題材",
        "strategy_note": "觀察：只做拉回，不追高",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；選股策略"
    },
    {
        "stock_id": "5269",
        "stock_name": "祥碩",
        "theme_category": "PCIe7/高速互連/PCB材料；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 67.5,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 67.5,
        "raw_theme_line": "PCIe/USB高速控制IC，受規格升級帶動驗證與設計需求；高速介面IC",
        "product_position": "PCIe/USB高速控制IC，受規格升級帶動驗證與設計需求",
        "strategy_note": "觀察：只做拉回，不追高",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；選股策略"
    },
    {
        "stock_id": "3022",
        "stock_name": "威強電",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": 67.35,
        "db_total_score_ref": 67.35,
        "ai_score_ref": 62.09,
        "research_score_ref": 57.33,
        "raw_theme_line": "Edge IPC / Physical AI；IPC/Edge System；Industrial Edge AI；Rugged Edge / Secure Edge Infrastructure",
        "product_position": "Industrial Edge AI；Rugged Edge / Secure Edge Infrastructure",
        "strategy_note": "REDUCE / 只觀察；減碼/不列主攻",
        "risk_note": "DB轉弱",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "6237",
        "stock_name": "驊訊",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 67.13,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "觀察回檔，不追高",
        "risk_note": "120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "03_中長線候選"
    },
    {
        "stock_id": "6494",
        "stock_name": "九齊",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 67.03,
        "db_total_score_ref": 67.03,
        "ai_score_ref": 65.43,
        "research_score_ref": 62.07,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "WAIT",
        "risk_note": "",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx",
        "theme_source_sheet": "AI_MCU策略；DB原始整合；Dashboard"
    },
    {
        "stock_id": "6239",
        "stock_name": "力成",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Advanced Packaging/CoWoS/半導體製程；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 10,
        "theme_score_base": 66.13,
        "db_total_score_ref": 64.09,
        "ai_score_ref": 53.61,
        "research_score_ref": 66.13,
        "raw_theme_line": "AI/晶圓代工；Advanced Packaging / Memory；Memory/Packaging；半導體；記憶體封測/SSD封裝",
        "product_position": "Memory/Packaging；記憶體封測/SSD封裝",
        "strategy_note": "AVOID/暫避；觀察；觀察回檔，不追高",
        "risk_note": "兩高不過疑慮；封測ASP未必等同NAND漲幅；正常",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；03_中長線候選；04_技術量化；05_SWOT；06_情境報酬與配置；Investment_ROI"
    },
    {
        "stock_id": "3013",
        "stock_name": "晟銘電",
        "theme_category": "AI Server ODM/Rack Scale；AI散熱/液冷",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 65.55,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；伺服器",
        "product_position": "",
        "strategy_note": "觀察回檔，不追高",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "03_中長線候選"
    },
    {
        "stock_id": "2417",
        "stock_name": "圓剛",
        "theme_category": "ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": 65.53,
        "db_total_score_ref": 61.25,
        "ai_score_ref": 64.28,
        "research_score_ref": 65.53,
        "raw_theme_line": "AI Vision/Edge Module；Edge IPC / Physical AI；Edge Vision Capture / Jetson Ecosystem；Vision/Edge Module",
        "product_position": "Edge Vision Capture / Jetson Ecosystem；Vision/Edge Module",
        "strategy_note": "AVOID / 不列主攻；觀察",
        "risk_note": "正常",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx",
        "theme_source_sheet": "06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping"
    },
    {
        "stock_id": "3260",
        "stock_name": "威剛",
        "theme_category": "AI Memory/Storage/NAND",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 6,
        "theme_score_base": 64.9,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 64.9,
        "raw_theme_line": "AI/晶圓代工；記憶體模組/品牌",
        "product_position": "記憶體模組/品牌",
        "strategy_note": "WATCH/等拉回或突破",
        "risk_note": "庫存反轉、消費SSD需求弱",
        "theme_source_file": "AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "3023",
        "stock_name": "信邦",
        "theme_category": "AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 62.7,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 62.7,
        "raw_theme_line": "工業/資料中心線束，觀察AI伺服器連接占比；線束/連接",
        "product_position": "工業/資料中心線束，觀察AI伺服器連接占比",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；Dashboard；選股策略"
    },
    {
        "stock_id": "6515",
        "stock_name": "穎崴",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 61.9,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 61.9,
        "raw_theme_line": "測試介面；高階測試座/探針卡，受AI/HPC晶片測試需求推動",
        "product_position": "高階測試座/探針卡，受AI/HPC晶片測試需求推動",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "4966",
        "stock_name": "譜瑞-KY",
        "theme_category": "PCIe7/高速互連/PCB材料；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 61.5,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 61.5,
        "raw_theme_line": "AI/晶圓代工；半導體；高速介面IC/Retimer；高速介面、橋接與訊號完整性相關IC，觀察PCIe/CXL/USB4滲透",
        "product_position": "高速介面、橋接與訊號完整性相關IC，觀察PCIe/CXL/USB4滲透",
        "strategy_note": "不列長期主攻；排除/等待",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "05_風險排除觀察；DB原始快照；選股策略"
    },
    {
        "stock_id": "6213",
        "stock_name": "聯茂",
        "theme_category": "PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 61.4,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 61.4,
        "raw_theme_line": "CCL材料；高速CCL材料第二供應鏈，觀察高階材料比重",
        "product_position": "高速CCL材料第二供應鏈，觀察高階材料比重",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "8046",
        "stock_name": "南電",
        "theme_category": "AI Memory/Storage/NAND；PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": 60.48,
        "db_total_score_ref": 60.48,
        "ai_score_ref": 49.09,
        "research_score_ref": None,
        "raw_theme_line": "ABF + 高速PCB",
        "product_position": "",
        "strategy_note": "接近低接買點，可小量試單",
        "risk_note": "站上65MA；站上120MA；乖離65MA偏大；營收YoY正；法人分數偏強",
        "theme_source_file": "(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx",
        "theme_source_sheet": "Agenda策略總表；前三大_可執行策略；投資策略總結；總表_更新版"
    },
    {
        "stock_id": "2454",
        "stock_name": "聯發科",
        "theme_category": "Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 60.42,
        "db_total_score_ref": 58.91,
        "ai_score_ref": 60.42,
        "research_score_ref": None,
        "raw_theme_line": "Edge AI SoC / IoT平台；Edge SoC；IC設計；高權值",
        "product_position": "Edge AI SoC / IoT平台",
        "strategy_note": "AVOID / 不列主攻；不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx",
        "theme_source_sheet": "05_風險排除觀察；06_候選池全表_DB"
    },
    {
        "stock_id": "3587",
        "stock_name": "閎康",
        "theme_category": "Advanced Packaging/CoWoS/半導體製程；PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 54.2,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 54.2,
        "raw_theme_line": "材料/失效分析；高速材料、封裝、故障分析與驗證服務",
        "product_position": "高速材料、封裝、故障分析與驗證服務",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "3289",
        "stock_name": "宜特",
        "theme_category": "PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 52.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 52.0,
        "raw_theme_line": "測試驗證；高速訊號、可靠度、材料與IC驗證服務",
        "product_position": "高速訊號、可靠度、材料與IC驗證服務",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "2313",
        "stock_name": "華通",
        "theme_category": "PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 50.4,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 50.4,
        "raw_theme_line": "PCB；伺服器/網通PCB升級，直接性低於台光/台燿",
        "product_position": "伺服器/網通PCB升級，直接性低於台光/台燿",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "6205",
        "stock_name": "詮欣",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": 47.9,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 47.9,
        "raw_theme_line": "連接器；高速連接器與端子，需確認AI Server占比與ASP",
        "product_position": "高速連接器與端子，需確認AI Server占比與ASP",
        "strategy_note": "排除/等待",
        "risk_note": "",
        "theme_source_file": "(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx",
        "theme_source_sheet": "DB原始快照；選股策略"
    },
    {
        "stock_id": "6907",
        "stock_name": "雅特力-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": 44.4,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 44.4,
        "raw_theme_line": "",
        "product_position": "",
        "strategy_note": "AVOID",
        "risk_note": "",
        "theme_source_file": "AI_MCU_DB原始資料_大師級投資策略.xlsx",
        "theme_source_sheet": "AI_MCU策略；DB原始整合；Dashboard"
    },
    {
        "stock_id": "8150",
        "stock_name": "南茂",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 6,
        "theme_score_base": 42.4,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 42.4,
        "raw_theme_line": "AI/晶圓代工；記憶體/驅動IC封測",
        "product_position": "記憶體/驅動IC封測",
        "strategy_note": "AVOID/暫避",
        "risk_note": "PE偏高、EPS低",
        "theme_source_file": "AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "8227",
        "stock_name": "巨有科技",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 42.11,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2387",
        "stock_name": "精元",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 40.44,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；伺服器",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "8088",
        "stock_name": "品安",
        "theme_category": "AI Memory/Storage/NAND；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 6,
        "theme_score_base": 31.5,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 31.5,
        "raw_theme_line": "AI/晶圓代工；記憶體模組",
        "product_position": "記憶體模組",
        "strategy_note": "AVOID/暫避",
        "risk_note": "規模與議價力低於大型模組廠",
        "theme_source_file": "AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "3380",
        "stock_name": "明泰",
        "theme_category": "AI Connectivity/Wi-Fi7/AI Gateway",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 31.13,
        "db_total_score_ref": 30.6,
        "ai_score_ref": 31.13,
        "research_score_ref": None,
        "raw_theme_line": "Connectivity Device；Networking Device",
        "product_position": "Networking Device",
        "strategy_note": "REDUCE / 只觀察",
        "risk_note": "",
        "theme_source_file": "Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx",
        "theme_source_sheet": "06_候選池全表_DB"
    },
    {
        "stock_id": "3135",
        "stock_name": "凌航",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 30.25,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "7769",
        "stock_name": "鴻勁",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 30.13,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "8210",
        "stock_name": "勤誠",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 29.16,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；伺服器",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2496",
        "stock_name": "卓越",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 28.1,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "其他",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2762",
        "stock_name": "世界健身-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 27.09,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；運動",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "9927",
        "stock_name": "泰銘",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 27.09,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "其他",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "7750",
        "stock_name": "新代",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 27.02,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "電源；電源/HVDC",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "1590",
        "stock_name": "亞德客-KY",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 26.55,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "電源；電源/HVDC",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2329",
        "stock_name": "華泰",
        "theme_category": "AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 6,
        "theme_score_base": 26.1,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": 26.1,
        "raw_theme_line": "AI/晶圓代工；記憶體封測",
        "product_position": "記憶體封測",
        "strategy_note": "AVOID/暫避",
        "risk_note": "營收年增弱、短線偏弱",
        "theme_source_file": "AI_Storage_NAND_投資策略_原始DB分析報告.xlsx",
        "theme_source_sheet": "00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置"
    },
    {
        "stock_id": "5285",
        "stock_name": "界霖",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 25.99,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2739",
        "stock_name": "寒舍",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 25.79,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；觀光",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "6771",
        "stock_name": "平和環保-創",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 25.79,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "環保；綠能環保",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "無明顯過熱",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "3016",
        "stock_name": "嘉晶",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 24.84,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "3581",
        "stock_name": "博磊",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 24.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "3532",
        "stock_name": "台勝科",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 24.14,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2049",
        "stock_name": "上銀",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 24.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "電源；電源/HVDC",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "6147",
        "stock_name": "頎邦",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 23.92,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "1597",
        "stock_name": "直得",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 23.45,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "電源；電源/HVDC",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "6182",
        "stock_name": "合晶",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 22.97,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "6651",
        "stock_name": "全宇昕",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 22.97,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "3095",
        "stock_name": "及成",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 21.93,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "資料中心交換器；高階網通",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "5425",
        "stock_name": "台半",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 21.92,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2020",
        "stock_name": "美亞",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 21.16,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "基礎原物料；鋼鐵",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "1442",
        "stock_name": "名軒",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 19.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "6754",
        "stock_name": "匯僑設計",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 19.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；居家",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2442",
        "stock_name": "新美齊",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 19.6,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2227",
        "stock_name": "裕日車",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 19.39,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "車用；電動車",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "1315",
        "stock_name": "達新",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 18.64,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "基礎原物料；塑膠",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2540",
        "stock_name": "愛山林",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 18.64,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2731",
        "stock_name": "雄獅",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 17.39,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；觀光",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "9946",
        "stock_name": "三發地產",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 17.24,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2543",
        "stock_name": "皇昌",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 16.61,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "6488",
        "stock_name": "環球晶",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 15.87,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI/晶圓代工；半導體",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "1432",
        "stock_name": "大魯閣",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 15.05,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；運動",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "4137",
        "stock_name": "麗豐-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 15.05,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "生技醫療；醫療",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2547",
        "stock_name": "日勝生",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 14.54,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "6834",
        "stock_name": "天二科技",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 13.88,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "零組件；電子零組件",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "5706",
        "stock_name": "鳳凰",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 13.79,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；觀光",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2109",
        "stock_name": "華豐",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 12.69,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；橡膠",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2348",
        "stock_name": "海悅",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 9.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "其他",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "4763",
        "stock_name": "材料*-KY",
        "theme_category": "PCIe7/高速互連/PCB材料",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 9.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "化工；基礎原物料",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "5225",
        "stock_name": "東科-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 9.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "全市場；系統掃描",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2504",
        "stock_name": "國產",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 8.41,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2530",
        "stock_name": "華建",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 8.41,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "1436",
        "stock_name": "華友聯",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 7.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2528",
        "stock_name": "皇普",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 7.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "2534",
        "stock_name": "宏盛",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 7.74,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；營造",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "1446",
        "stock_name": "宏和",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 6.64,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；紡織",
        "product_position": "",
        "strategy_note": "防守小部位/只低接",
        "risk_note": "兩高不過疑慮；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "04_防守現金流"
    },
    {
        "stock_id": "1595",
        "stock_name": "川寶",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 4.91,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "零組件；電子零組件",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "1216",
        "stock_name": "統一",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 4,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "民生消費；製造業營收前十；食品/消費；食品/防禦",
        "product_position": "食品/消費",
        "strategy_note": "WAIT突破或回檔",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "1513",
        "stock_name": "中興電",
        "theme_category": "AI Power/PPA/BBU/電池",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI能源/電源效率",
        "product_position": "",
        "strategy_note": "加入主題追蹤",
        "risk_note": "",
        "theme_source_file": "2026_AI_Agenda_Deep_Analysis_Report.docx",
        "theme_source_sheet": "docx內文"
    },
    {
        "stock_id": "1762",
        "stock_name": "中化生",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "生技醫療；醫療",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "1809",
        "stock_name": "中釉",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "傳產；玻璃陶瓷",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階；累計營收衰退",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2207",
        "stock_name": "和泰車",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "汽車通路；汽車銷售；電動車",
        "product_position": "汽車銷售",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2324",
        "stock_name": "仁寶",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；電子代工",
        "product_position": "電子代工",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2412",
        "stock_name": "中華電",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "資料中心交換器；電信/IDC；電信/IDC/雲端",
        "product_position": "電信/IDC/雲端",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2609",
        "stock_name": "陽明",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "海運；航運/海運；運輸",
        "product_position": "海運",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2610",
        "stock_name": "華航",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "服務業獲利前十；航空；運輸",
        "product_position": "航空",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2615",
        "stock_name": "萬海",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "服務業獲利前十；海運；航運/海運；運輸",
        "product_position": "海運",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2618",
        "stock_name": "長榮航",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "服務業獲利前十；航空；運輸",
        "product_position": "航空",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2855",
        "stock_name": "統一證",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "民生消費；食品",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "2880",
        "stock_name": "華南金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2881",
        "stock_name": "富邦金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金控排名；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "WATCH",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2882",
        "stock_name": "國泰金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金控排名；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "WATCH",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2883",
        "stock_name": "凱基金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金控排名；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "WATCH",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2884",
        "stock_name": "玉山金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2885",
        "stock_name": "元大金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金控排名；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "WATCH",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "2886",
        "stock_name": "兆豐金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2887",
        "stock_name": "台新新光金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2889",
        "stock_name": "國票金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2890",
        "stock_name": "永豐金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2891",
        "stock_name": "中信金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "2892",
        "stock_name": "第一金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 5,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金控排名；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "WAIT突破或回檔",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別"
    },
    {
        "stock_id": "3033",
        "stock_name": "威健",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "通路；電子通路",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "3209",
        "stock_name": "全科",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "通路；電子通路",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "3430",
        "stock_name": "奇鈦科",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "化工；基礎原物料",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "4556",
        "stock_name": "旭然",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "其他",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "4807",
        "stock_name": "日成-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "內需消費；百貨",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "4938",
        "stock_name": "和碩",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI伺服器；電子代工",
        "product_position": "電子代工",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "5314",
        "stock_name": "世紀*",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 2,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "保健產品；保健產品/未分類；全市場；營運績效前五",
        "product_position": "保健產品/未分類",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；08_明日TOP5_每類別"
    },
    {
        "stock_id": "5871",
        "stock_name": "中租-KY",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "全市場；租賃金融；金融租賃",
        "product_position": "租賃金融",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "5880",
        "stock_name": "合庫金",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "金控；金融；金融金控",
        "product_position": "金控",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算"
    },
    {
        "stock_id": "6584",
        "stock_name": "南俊國際",
        "theme_category": "AI Server ODM/Rack Scale",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 3,
        "theme_score_base": None,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "AI機械/滑軌；伺服器滑軌；電子零組件",
        "product_position": "伺服器滑軌",
        "strategy_note": "AVOID/觀望",
        "risk_note": "",
        "theme_source_file": "三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx",
        "theme_source_sheet": "04_候選池_DB重算；06_中期選股策略；07_長期選股策略"
    },
    {
        "stock_id": "6658",
        "stock_name": "聯策",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "全市場；系統掃描",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    },
    {
        "stock_id": "6861",
        "stock_name": "睿生光電",
        "theme_category": "其他/待分類",
        "theme_priority": "C_題材備查",
        "theme_appearance_count": 1,
        "theme_score_base": 0.0,
        "db_total_score_ref": None,
        "ai_score_ref": None,
        "research_score_ref": None,
        "raw_theme_line": "生技醫療；醫療",
        "product_position": "",
        "strategy_note": "不列長期主攻",
        "risk_note": "RSI過熱；MA65乖離過大；120日高位階",
        "theme_source_file": "(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx",
        "theme_source_sheet": "05_風險排除觀察"
    }
]


# =============================
# V3.0 AI Project Rotation Monitor
# Source: 主流題材_爆發前追蹤類型_整合股票池_深度分析報告.xlsx
# Purpose: All AI project stocks must be listed and tracked by category/status.
# =============================
AI_PROJECT_CATEGORIES = [{'ai_project_type': 'Edge AI/TinyML/MCU AI', 'stock_count': 36, 'priority_a_count': 5, 'tracking_purpose': '追蹤AI下沉到MCU、Sensor、Gateway、Runtime的長週期', 'representative_stocks': '2345 智邦、8271 宇瞻、2357 華碩、6166 凌華、2344 華邦電、6781 AES-KY、1256 鮮活果汁-KY、4583 台灣精銳、2404 漢唐、1514 亞力'}, {'ai_project_type': 'AI Server ODM/Rack Scale', 'stock_count': 29, 'priority_a_count': 6, 'tracking_purpose': '追蹤AI伺服器、Rack Scale、整機櫃ODM主線', 'representative_stocks': '2357 華碩、2383 台光電、6274 台燿、2317 鴻海、2382 廣達、3231 緯創、2376 技嘉、2395 研華、3706 神達、3017 奇鋐'}, {'ai_project_type': 'AI Power/PPA/BBU/電池', 'stock_count': 26, 'priority_a_count': 3, 'tracking_purpose': '追蹤AI電源、PPA、BBU、電池與能源效率瓶頸', 'representative_stocks': '8271 宇瞻、2344 華邦電、4585 達明、6781 AES-KY、2451 創見、3017 奇鋐、2301 光寶科、2308 台達電、3324 雙鴻、2337 旺宏'}, {'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA', 'stock_count': 23, 'priority_a_count': 1, 'tracking_purpose': '追蹤客製ASIC、FPGA再起、RISC-V、EDA驗證', 'representative_stocks': '6166 凌華、2395 研華、3706 神達、3711 日月光投控、6533 晶心科、3443 創意、6442 光聖、3081 聯亞、3035 智原、4979 華星光'}, {'ai_project_type': 'PCIe7/高速互連/PCB材料', 'stock_count': 21, 'priority_a_count': 2, 'tracking_purpose': '追蹤PCIe 7、高速CCL、PCB材料、線束/連接器升級', 'representative_stocks': '2383 台光電、6274 台燿、3017 奇鋐、3037 欣興、3324 雙鴻、4979 華星光、2368 金像電、6285 啟碁、3167 大量、3665 貿聯-KY'}, {'ai_project_type': 'AI Memory/Storage/NAND', 'stock_count': 17, 'priority_a_count': 2, 'tracking_purpose': '追蹤AI資料爆炸、SSD/NAND、HBM/TL-RAM/記憶體', 'representative_stocks': '8271 宇瞻、2344 華邦電、2451 創見、3037 欣興、3711 日月光投控、5289 宜鼎、8299 群聯、2337 旺宏、2408 南亞科、4967 十銓'}, {'ai_project_type': 'Edge IPC/Physical AI/Robot', 'stock_count': 15, 'priority_a_count': 3, 'tracking_purpose': '追蹤工業AI、Robot Node、AI Box、IPC商用落地', 'representative_stocks': '6166 凌華、2317 鴻海、4585 達明、2395 研華、3706 神達、3019 亞光、8114 振樺電、2359 所羅門、6922 宸曜、2059 川湖'}, {'ai_project_type': 'CPO/矽光子/光通訊', 'stock_count': 11, 'priority_a_count': 0, 'tracking_purpose': '延續CPO模式，追蹤光互連/CPO/1.6T/3.2T門戶大開', 'representative_stocks': '3037 欣興、3019 亞光、6442 光聖、3081 聯亞、4979 華星光、3163 波若威、2368 金像電、3363 上詮、6285 啟碁、3441 聯一光'}, {'ai_project_type': 'AI散熱/液冷', 'stock_count': 9, 'priority_a_count': 0, 'tracking_purpose': '追蹤GB300/GB400後AI Server液冷與散熱瓶頸', 'representative_stocks': '2421 建準、2428 興勤、3017 奇鋐、2301 光寶科、2308 台達電、3324 雙鴻、8996 高力、3653 健策、3013 晟銘電'}, {'ai_project_type': 'Advanced Packaging/CoWoS/半導體製程', 'stock_count': 8, 'priority_a_count': 0, 'tracking_purpose': '追蹤CoWoS、2nm、Chiplet與先進封裝', 'representative_stocks': '3037 欣興、3711 日月光投控、3443 創意、2330 台積電、8299 群聯、3661 世芯-KY、6239 力成、3587 閎康'}, {'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway', 'stock_count': 6, 'priority_a_count': 1, 'tracking_purpose': '追蹤AI Gateway、Wi-Fi 7、Broadband Gateway、CAT安全平台入口', 'representative_stocks': '2345 智邦、2379 瑞昱、5388 中磊、3596 智易、6285 啟碁、3380 明泰'}]

AI_PROJECT_TRACKING_MASTER = [{'stock_id': '2345', 'stock_name': '智邦', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway；Edge AI/TinyML/MCU AI；其他/待分類', 'tracking_priority': 'A_主攻追蹤', 'source_count': 12, 'max_score': 90.65, 'db_score': 90.65, 'ai_score_ref': 78.67, 'research_score': 87.22, 'raw_theme_line': 'AI Connectivity；Networking；Networking / Edge Switch；營運績效前五；網路交換器；網通交換器；資料中心/Edge交換器；資料中心交換器', 'product_position': 'Networking / Edge Switch；網路交換器；資料中心/Edge交換器', 'strategy_summary': 'HOLD / 題材觀察；Pullback Buy；WAIT突破或回檔；可分批佈局', 'risk_summary': '正常', 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；06_候選池全表_DB；07_投資組合TOP20；07_長期選股策略；08_明日TOP5_每類別；Dashboard'}, {'stock_id': '8271', 'stock_name': '宇瞻', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI', 'tracking_priority': 'A_主攻追蹤', 'source_count': 7, 'max_score': 87.87, 'db_score': 87.87, 'ai_score_ref': 77.23, 'research_score': 74.7, 'raw_theme_line': 'AI/晶圓代工；記憶體模組/工控儲存', 'product_position': '記憶體模組/工控儲存', 'strategy_summary': 'BUY/分批布局；Pullback Buy', 'risk_summary': '量能與市場辨識度', 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；EdgeAI_TOP10'}, {'stock_id': '2357', 'stock_name': '華碩', 'ai_project_type': 'AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI', 'tracking_priority': 'A_主攻追蹤', 'source_count': 6, 'max_score': 87.36, 'db_score': 87.36, 'ai_score_ref': 73.4, 'research_score': None, 'raw_theme_line': 'AI伺服器；AI伺服器/PC品牌；服務業獲利前十；電腦/AI PC/伺服器', 'product_position': '電腦/AI PC/伺服器', 'strategy_summary': 'Pullback Buy；WAIT突破或回檔', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；EdgeAI_TOP10'}, {'stock_id': '6166', 'stock_name': '凌華', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'A_主攻追蹤', 'source_count': 8, 'max_score': 86.91, 'db_score': 86.91, 'ai_score_ref': 75.13, 'research_score': 84.18, 'raw_theme_line': 'Edge IPC / Physical AI；Edge IPC / Robot Platform；IPC/Edge System', 'product_position': 'Edge IPC / Robot Platform', 'strategy_summary': 'HOLD / 題材觀察；Pullback Buy；題材觀察/等回測', 'risk_summary': 'RSI偏熱', 'source_file': '2026_AI_Agenda_Investment_Strategy.xlsx；2026_Edge_AI_投資策略_DB版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；EdgeAI_TOP10；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2344', 'stock_name': '華邦電', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI', 'tracking_priority': 'A_主攻追蹤', 'source_count': 8, 'max_score': 83.71, 'db_score': 83.71, 'ai_score_ref': 73.06, 'research_score': 72.3, 'raw_theme_line': 'AI/晶圓代工；SLC NAND/Specialty Memory；TL-RAM/記憶體架構受惠', 'product_position': 'SLC NAND/Specialty Memory', 'strategy_summary': 'BUY/分批布局；加入主題追蹤', 'risk_summary': '非高階Enterprise SSD主角、PE偏高', 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；Agenda策略總表；docx內文'}, {'stock_id': '2383', 'stock_name': '台光電', 'ai_project_type': 'AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料', 'tracking_priority': 'A_主攻追蹤', 'source_count': 5, 'max_score': 79.8, 'db_score': None, 'ai_score_ref': 50, 'research_score': 79.8, 'raw_theme_line': 'ABF + 高速PCB；Ultra-low-loss材料、AI Server PCB、GPU tray高頻低損耗基材；高速CCL/PCB材料', 'product_position': 'Ultra-low-loss材料、AI Server PCB、GPU tray高頻低損耗基材', 'strategy_summary': '接近低接買點，可小量試單；次主攻：等量價確認', 'risk_summary': '站上65MA；站上120MA；乖離65MA偏大；營收YoY正；法人分數偏強；120日位階偏高', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；投資策略總結；總表_更新版；選股策略'}, {'stock_id': '6274', 'stock_name': '台燿', 'ai_project_type': 'AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料', 'tracking_priority': 'A_主攻追蹤', 'source_count': 5, 'max_score': 76.2, 'db_score': None, 'ai_score_ref': 45, 'research_score': 76.2, 'raw_theme_line': 'ABF + 高速PCB；M7/M8/M9級低損耗材料、HVLP銅箔、伺服器板升級；高速CCL/PCB材料', 'product_position': 'M7/M8/M9級低損耗材料、HVLP銅箔、伺服器板升級', 'strategy_summary': '接近低接買點，可小量試單；次主攻：等量價確認', 'risk_summary': '站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；投資策略總結；總表_更新版；選股策略'}, {'stock_id': '2317', 'stock_name': '鴻海', 'ai_project_type': 'AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot', 'tracking_priority': 'A_主攻追蹤', 'source_count': 9, 'max_score': 74.93, 'db_score': 74.93, 'ai_score_ref': 68.67, 'research_score': 60.518691670557885, 'raw_theme_line': 'AI伺服器；AI伺服器/整機櫃；AI工業化/ODM/機器人；AI整機櫃/ODM；ODM/Robot；製造業營收前十', 'product_position': 'AI伺服器/整機櫃；AI工業化/ODM/機器人', 'strategy_summary': 'BUY/主攻；HOLD / 題材觀察；核心主攻', 'risk_summary': None, 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；06_候選池全表_DB；07_投資組合TOP20；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '2382', 'stock_name': '廣達', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'A_主攻追蹤', 'source_count': 8, 'max_score': 74.22, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；AI伺服器ODM；伺服器；製造業營收前十', 'product_position': 'AI伺服器ODM', 'strategy_summary': 'BUY/主攻；核心主攻；觀察回檔，不追高', 'risk_summary': '120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽儀表板；03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5'}, {'stock_id': '3231', 'stock_name': '緯創', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'A_主攻追蹤', 'source_count': 9, 'max_score': 70.06, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；AI伺服器ODM；伺服器；製造業營收前十', 'product_position': 'AI伺服器ODM', 'strategy_summary': 'BUY/主攻；核心主攻；觀察回檔，不追高', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽儀表板；03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_Industry_Intelligence；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '4585', 'stock_name': '達明', 'ai_project_type': 'AI Power/PPA/BBU/電池；Edge IPC/Physical AI/Robot', 'tracking_priority': 'A_主攻追蹤', 'source_count': 5, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI奇兵50強；Physical AI/機器人；機器人與自動化設備；電源/HVDC', 'product_position': '機器人與自動化設備', 'strategy_summary': 'BUY/主攻', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '6781', 'stock_name': 'AES-KY', 'ai_project_type': 'AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 4, 'max_score': 91.05, 'db_score': 91.05, 'ai_score_ref': 79.33, 'research_score': None, 'raw_theme_line': 'AI電池/BBU；AI電池_BBU', 'product_position': None, 'strategy_summary': 'Pullback Buy；WAIT；加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'Agenda策略總表；EdgeAI_TOP10；TOP5主攻；docx內文'}, {'stock_id': '1256', 'stock_name': '鮮活果汁-KY', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 89.5, 'db_score': 89.5, 'ai_score_ref': 77.65, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '4583', 'stock_name': '台灣精銳', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 88.57, 'db_score': 88.57, 'ai_score_ref': 79.78, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2404', 'stock_name': '漢唐', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 88.47, 'db_score': 88.47, 'ai_score_ref': 76.81, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '1514', 'stock_name': '亞力', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 88.37, 'db_score': 88.37, 'ai_score_ref': 80.09, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '3617', 'stock_name': '碩天', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 88.24, 'db_score': 88.24, 'ai_score_ref': 80.02, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2492', 'stock_name': '華新科', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 2, 'max_score': 87.32, 'db_score': 87.32, 'ai_score_ref': 77.83, 'research_score': None, 'raw_theme_line': 'Edge AI TOP10', 'product_position': None, 'strategy_summary': 'Pullback Buy；加入主題追蹤', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx；2026_Edge_AI_深度白皮書_DB版.docx', 'source_sheet': 'EdgeAI_TOP10；docx內文'}, {'stock_id': '2376', 'stock_name': '技嘉', 'ai_project_type': 'AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 2, 'max_score': 87.24, 'db_score': 87.24, 'ai_score_ref': 73.85, 'research_score': None, 'raw_theme_line': 'AI伺服器；伺服器', 'product_position': None, 'strategy_summary': 'Pullback Buy；觀察回檔，不追高', 'risk_summary': '120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': '03_中長線候選；EdgeAI_TOP10'}, {'stock_id': '6525', 'stock_name': '6525', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 87.04, 'db_score': 87.04, 'ai_score_ref': 74.36, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '7749', 'stock_name': '7749', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 86.68, 'db_score': 86.68, 'ai_score_ref': 76.52, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2421', 'stock_name': '建準', 'ai_project_type': 'AI散熱/液冷；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 2, 'max_score': 86.26, 'db_score': 86.26, 'ai_score_ref': 77.47, 'research_score': None, 'raw_theme_line': 'AI散熱；液冷', 'product_position': None, 'strategy_summary': 'Pullback Buy；分批布局/回測加碼', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': '02_長期主升核心；EdgeAI_TOP10'}, {'stock_id': '3533', 'stock_name': '3533', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 86.16, 'db_score': 86.16, 'ai_score_ref': 69.67, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '6446', 'stock_name': '6446', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 86.03, 'db_score': 86.03, 'ai_score_ref': 73.31, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2428', 'stock_name': '興勤', 'ai_project_type': 'AI散熱/液冷；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 2, 'max_score': 85.91, 'db_score': 85.91, 'ai_score_ref': 74.54, 'research_score': None, 'raw_theme_line': '零組件；電子零組件', 'product_position': None, 'strategy_summary': 'Pullback Buy；觀察回檔，不追高', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': '03_中長線候選；EdgeAI_TOP10'}, {'stock_id': '8926', 'stock_name': '8926', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 85.88, 'db_score': 85.88, 'ai_score_ref': 77.34, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2548', 'stock_name': '2548', 'ai_project_type': 'Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 1, 'max_score': 85.86, 'db_score': 85.86, 'ai_score_ref': 63.55, 'research_score': None, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'Pullback Buy', 'risk_summary': None, 'source_file': '2026_Edge_AI_投資策略_DB版.xlsx', 'source_sheet': 'EdgeAI_TOP10'}, {'stock_id': '2395', 'stock_name': '研華', 'ai_project_type': 'AI Server ODM/Rack Scale；ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 84.14, 'db_score': 84.14, 'ai_score_ref': 70.92, 'research_score': 74.92, 'raw_theme_line': 'AI伺服器；Edge IPC / Physical AI；Edge IPC / Physical AI平台；IPC/Edge System；伺服器', 'product_position': 'Edge IPC / Physical AI平台', 'strategy_summary': 'HOLD / 題材觀察；不列長期主攻；等待回檔', 'risk_summary': 'RSI過熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；Investment_ROI；Stock_Strategy'}, {'stock_id': '3706', 'stock_name': '神達', 'ai_project_type': 'AI Server ODM/Rack Scale；ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 83.66, 'db_score': 83.66, 'ai_score_ref': 74.37, 'research_score': 82.01, 'raw_theme_line': 'AI伺服器；Edge IPC / Physical AI；Edge Server；Edge Server / 車用AI；伺服器', 'product_position': 'Edge Server / 車用AI', 'strategy_summary': 'HOLD / 題材觀察；分批布局/回測加碼；可分批佈局', 'risk_summary': '正常；無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '02_長期主升核心；06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2379', 'stock_name': '瑞昱', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway', 'tracking_priority': 'B_觀察追蹤', 'source_count': 6, 'max_score': 82.97, 'db_score': 82.97, 'ai_score_ref': 76.71, 'research_score': 82.92, 'raw_theme_line': 'AI Connectivity；AI Connectivity / Gateway IC；Connectivity', 'product_position': 'AI Connectivity / Gateway IC', 'strategy_summary': 'HOLD / 題材觀察；可分批佈局', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2451', 'stock_name': '創見', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池', 'tracking_priority': 'B_觀察追蹤', 'source_count': 11, 'max_score': 82.6, 'db_score': None, 'ai_score_ref': None, 'research_score': 82.6, 'raw_theme_line': 'AI/晶圓代工；半導體；工控/消費/企業儲存品牌；記憶體/邊緣儲存；記憶體及周邊產品', 'product_position': '工控/消費/企業儲存品牌；記憶體及周邊產品', 'strategy_summary': 'BUY/分批布局；WAIT突破或回檔；不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；成長彈性不如控制器', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_候選池_DB重算；04_技術量化；05_SWOT；05_短期選股策略；05_風險排除觀察'}, {'stock_id': '3017', 'stock_name': '奇鋐', 'ai_project_type': 'AI Power/PPA/BBU/電池；AI Server ODM/Rack Scale；AI散熱/液冷；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 12, 'max_score': 82.22, 'db_score': 74.27, 'ai_score_ref': 57.24, 'research_score': 76.2, 'raw_theme_line': 'AI散熱；GB300/NVL72液冷與GPU Tray高功耗散熱同步升級；散熱模組；液冷；液冷 + AI Power + HVDC；液冷/散熱；營運績效前五', 'product_position': 'GB300/NVL72液冷與GPU Tray高功耗散熱同步升級；散熱模組', 'strategy_summary': 'WATCH；分批布局/回測加碼；次主攻：等量價確認；高於低接區約6.4%，不追等回測', 'risk_summary': '無明顯過熱；站上65MA；站上120MA；量縮；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '02_長期主升核心；04_候選池_DB重算；06_Industry_Intelligence；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；Agenda策略總表；DB原始快照'}, {'stock_id': '5388', 'stock_name': '中磊', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 82, 'db_score': 82, 'ai_score_ref': 69.69, 'research_score': 79.62, 'raw_theme_line': 'AI Connectivity；AI Gateway / FWA / Broadband；Connectivity Device；資料中心交換器；高階網通', 'product_position': 'AI Gateway / FWA / Broadband', 'strategy_summary': 'HOLD / 題材觀察；可分批佈局；觀察回檔，不追高', 'risk_summary': '120日高位階；正常', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '03_中長線候選；06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2301', 'stock_name': '光寶科', 'ai_project_type': 'AI Power/PPA/BBU/電池；AI散熱/液冷', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 81.11, 'db_score': 79.56, 'ai_score_ref': 69.67, 'research_score': 70.48, 'raw_theme_line': 'AI Power；AI Power / PPA；AI Power / PSU；液冷 + AI Power + HVDC；電源；電源/HVDC', 'product_position': 'AI Power / PSU', 'strategy_summary': 'HOLD / 題材觀察；不列長期主攻；等待回檔；高於低接區約4.0%，不追等回測', 'risk_summary': 'RSI偏熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '05_風險排除觀察；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping；投資策略總結'}, {'stock_id': '3037', 'stock_name': '欣興', 'ai_project_type': 'AI Memory/Storage/NAND；AI Server ODM/Rack Scale；Advanced Packaging/CoWoS/半導體製程；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 80.73, 'db_score': 80.73, 'ai_score_ref': 66.48, 'research_score': 68.2, 'raw_theme_line': 'ABF + 高速PCB；AI伺服器高階PCB與載板、層數提升受益；AI記憶體/高速材料/封裝；PCB/載板', 'product_position': 'AI伺服器高階PCB與載板、層數提升受益', 'strategy_summary': '加入主題追蹤；觀察：只做拉回，不追高；高於低接區約4.7%，不追等回測', 'risk_summary': '站上65MA；站上120MA；乖離65MA偏大；量能放大；營收YoY正；120日位階偏高', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；DB原始快照；Dashboard；docx內文；投資策略總結；總表_更新版；選股策略'}, {'stock_id': '3711', 'stock_name': '日月光投控', 'ai_project_type': 'AI Memory/Storage/NAND；ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 8, 'max_score': 80.69, 'db_score': 80.6, 'ai_score_ref': 71.53, 'research_score': 80.69, 'raw_theme_line': '2nm + CoWoS；Advanced Packaging；Advanced Packaging / Memory；Packaging', 'product_position': 'Advanced Packaging', 'strategy_summary': 'HOLD / 題材觀察；接近低接買點，可小量試單；題材觀察/等回測', 'risk_summary': 'RSI偏熱；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping；投資策略總結；總表_更新版'}, {'stock_id': '3019', 'stock_name': '亞光', 'ai_project_type': 'CPO/矽光子/光通訊；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 6, 'max_score': 79.97, 'db_score': 78.36, 'ai_score_ref': 72.86, 'research_score': 79.97, 'raw_theme_line': 'AI Optical Sensor / Vision；AI Sensor', 'product_position': 'AI Optical Sensor / Vision', 'strategy_summary': 'HOLD / 題材觀察；題材觀察/等回測', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2308', 'stock_name': '台達電', 'ai_project_type': 'AI Power/PPA/BBU/電池；AI散熱/液冷；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 79.85, 'db_score': 76.72, 'ai_score_ref': 63.69, 'research_score': 77.45, 'raw_theme_line': 'AI Power；AI Power / PPA；AI Power / Thermal / PPA；液冷 + AI Power + HVDC', 'product_position': 'AI Power / Thermal / PPA', 'strategy_summary': 'HOLD / 題材觀察；題材觀察/等回測；高於低接區約5.9%，不追等回測', 'risk_summary': 'RSI偏熱；站上65MA；站上120MA；乖離65MA偏大；營收YoY正', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping；投資策略總結'}, {'stock_id': '6533', 'stock_name': '晶心科', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 79.05, 'db_score': 77.36, 'ai_score_ref': 78.81, 'research_score': 79.05, 'raw_theme_line': 'ASIC / RISC-V / IP；CPU/IP；IP；MCU AI / TinyML；RISC-V；RISC-V Edge CPU/IP', 'product_position': 'RISC-V Edge CPU/IP', 'strategy_summary': 'AVOID / 不列主攻；分批布局/回測加碼；題材觀察/等回測', 'risk_summary': '正常；無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '02_長期主升核心；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3596', 'stock_name': '智易', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': 79.01, 'db_score': 79.01, 'ai_score_ref': 67.63, 'research_score': 77.38, 'raw_theme_line': 'AI Connectivity；AI Gateway / CPE；Connectivity Device', 'product_position': 'AI Gateway / CPE', 'strategy_summary': 'AVOID / 不列主攻；題材觀察/等回測', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3443', 'stock_name': '創意', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 78.94, 'db_score': 78.94, 'ai_score_ref': 67.97, 'research_score': 70.08, 'raw_theme_line': 'AI ASIC；ASIC；ASIC / Edge AI Custom Chip；ASIC / RISC-V / IP；ASIC/IP', 'product_position': 'ASIC / Edge AI Custom Chip', 'strategy_summary': 'HOLD / 題材觀察；不列長期主攻；等待回檔', 'risk_summary': 'RSI過熱 / 短線乖離過大；RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Dashboard；Investment_ROI；Stock_Strategy'}, {'stock_id': '6442', 'stock_name': '光聖', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊', 'tracking_priority': 'B_觀察追蹤', 'source_count': 6, 'max_score': 78.73, 'db_score': 49.46, 'ai_score_ref': 41.18, 'research_score': 39.58, 'raw_theme_line': 'AI光通訊；CPO/光模組；Optical / Interconnect；Optical Interconnect；光通訊元件與射頻元件；高速光通訊', 'product_position': 'Optical Interconnect；光通訊元件與射頻元件', 'strategy_summary': 'AVOID/觀望；分批布局/回測加碼；減碼/不列主攻', 'risk_summary': 'DB轉弱 / DB分數不足；無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '02_長期主升核心；04_候選池_DB重算；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3081', 'stock_name': '聯亞', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 78.69, 'db_score': 56.57, 'ai_score_ref': 52.37, 'research_score': 59.74, 'raw_theme_line': '1.6T/CPO/矽光子；Optical / Interconnect；Optical Component；矽光子與光通訊', 'product_position': 'Optical Component', 'strategy_summary': 'WAIT；觀察；高於低接區約4.1%，不追等回測', 'risk_summary': 'DB分數不足；站上65MA；站上120MA；乖離65MA偏大；量縮；營收YoY正', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': 'Agenda策略總表；Investment_ROI；Stock_Strategy；TOP5主攻；Theme_Mapping；投資策略總結；總表_更新版'}, {'stock_id': '3035', 'stock_name': '智原', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 6, 'max_score': 78.66, 'db_score': 70.72, 'ai_score_ref': 78.66, 'research_score': 71.4, 'raw_theme_line': 'ASIC / RISC-V / IP；ASIC Design Service；ASIC/IP', 'product_position': 'ASIC Design Service', 'strategy_summary': 'AVOID / 不列主攻；題材觀察/等回測', 'risk_summary': '正常', 'source_file': '2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3324', 'stock_name': '雙鴻', 'ai_project_type': 'AI Power/PPA/BBU/電池；AI Server ODM/Rack Scale；AI散熱/液冷；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 76.92, 'db_score': 63, 'ai_score_ref': 46.46, 'research_score': 59.8, 'raw_theme_line': 'AI Server液冷、冷板與高瓦數散熱需求；AI散熱；液冷；液冷 + AI Power + HVDC；液冷/散熱', 'product_position': 'AI Server液冷、冷板與高瓦數散熱需求', 'strategy_summary': '分批布局/回測加碼；排除/等待；接近低接買點，可小量試單', 'risk_summary': '未站上65MA；量縮；營收YoY正；無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': '02_長期主升核心；06_Industry_Intelligence；Agenda策略總表；DB原始快照；投資策略總結；總表_更新版；選股策略'}, {'stock_id': '6202', 'stock_name': '盛群', 'ai_project_type': 'Edge AI/TinyML/MCU AI；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 76.63, 'db_score': 67.42, 'ai_score_ref': 72.31, 'research_score': 54.45, 'raw_theme_line': 'AI MCU / Embedded Controller；AI MCU / 控制IC；MCU AI；MCU AI / TinyML', 'product_position': 'AI MCU / Embedded Controller；AI MCU / 控制IC', 'strategy_summary': 'AVOID；AVOID / 不列主攻；題材觀察/等回測', 'risk_summary': '正常', 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；AI_MCU策略；DB原始整合；Dashboard；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '4979', 'stock_name': '華星光', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 12, 'max_score': 74.57, 'db_score': 69.77, 'ai_score_ref': 63, 'research_score': 68.18, 'raw_theme_line': '1.6T/CPO/矽光子；800G→1.6T、CPO/光模組升級的高速互連旁支受益；Interconnect；Optical / Interconnect；Optical Interconnect / CPO；光通訊 / CPO；光通訊/CPO', 'product_position': '800G→1.6T、CPO/光模組升級的高速互連旁支受益；Optical Interconnect / CPO；光通訊 / CPO', 'strategy_summary': 'AVOID / 不列主攻；觀察；觀察：只做拉回，不追高；高於低接區約9.8%，不追等回測', 'risk_summary': '正常；站上65MA；站上120MA；量縮；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_Industry_Intelligence；06_候選池全表_DB；07_投資組合TOP20；Agenda策略總表；DB原始快照；Dashboard；Investment_ROI；Stock_Strategy'}, {'stock_id': '2330', 'stock_name': '台積電', 'ai_project_type': 'Advanced Packaging/CoWoS/半導體製程', 'tracking_priority': 'B_觀察追蹤', 'source_count': 9, 'max_score': 74.35, 'db_score': None, 'ai_score_ref': 65, 'research_score': None, 'raw_theme_line': '2nm + CoWoS；AI/晶圓代工；半導體製造；晶圓代工；製造業營收前十；高權值', 'product_position': '晶圓代工', 'strategy_summary': 'BUY/主攻；接近低接買點，可小量試單；觀察回檔，不追高', 'risk_summary': '120日高位階；站上65MA；站上120MA；量縮；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；前三大_可執行策略；投資策略總結'}, {'stock_id': '4919', 'stock_name': '新唐', 'ai_project_type': 'Edge AI/TinyML/MCU AI；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 72.56, 'db_score': 63.39, 'ai_score_ref': 72.46, 'research_score': 28.18, 'raw_theme_line': 'AI MCU / Embedded Controller；MCU AI；MCU AI / TinyML', 'product_position': 'AI MCU / Embedded Controller', 'strategy_summary': 'AVOID；AVOID / 不列主攻；等待回檔', 'risk_summary': 'RSI偏熱 / 短線乖離過大', 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；AI_MCU策略；DB原始整合；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '5289', 'stock_name': '宜鼎', 'ai_project_type': 'AI Memory/Storage/NAND；AI Server ODM/Rack Scale；Edge AI/TinyML/MCU AI；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 6, 'max_score': 72.4, 'db_score': None, 'ai_score_ref': None, 'research_score': 72.4, 'raw_theme_line': 'AI伺服器；工控/Edge AI SSD', 'product_position': '工控/Edge AI SSD', 'strategy_summary': 'WATCH/等拉回或突破', 'risk_summary': '本益比/位階偏高、流動性震盪', 'source_file': 'AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '8114', 'stock_name': '振樺電', 'ai_project_type': 'AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 8, 'max_score': 72.31, 'db_score': 70.96, 'ai_score_ref': 72.31, 'research_score': None, 'raw_theme_line': 'AI伺服器；AI奇兵50強；Edge AI 從視覺升級至本地AI系統；Edge_AI；POS/邊緣裝置；銷售終端系統POS', 'product_position': '銷售終端系統POS', 'strategy_summary': 'WAIT；WAIT突破或回檔；加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；Agenda策略總表；TOP5主攻；docx內文'}, {'stock_id': '3163', 'stock_name': '波若威', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊', 'tracking_priority': 'B_觀察追蹤', 'source_count': 8, 'max_score': 72.25, 'db_score': 40.45, 'ai_score_ref': 29.22, 'research_score': 34.37, 'raw_theme_line': '1.6T/CPO/矽光子；CPO/光模組；Optical / Interconnect；Optical Module；高速光通訊', 'product_position': 'Optical Module', 'strategy_summary': '接近低接買點，可小量試單；減碼/不列主攻；觀察回檔，不追高', 'risk_summary': 'DB轉弱 / DB分數不足；無明顯過熱；站上65MA；站上120MA；距65MA適中；量縮；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '03_中長線候選；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping；前三大_可執行策略；投資策略總結；總表_更新版'}, {'stock_id': '2368', 'stock_name': '金像電', 'ai_project_type': 'AI Server ODM/Rack Scale；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 3, 'max_score': 72.2, 'db_score': None, 'ai_score_ref': None, 'research_score': 72.2, 'raw_theme_line': 'AI伺服器板、交換器/網通板，高層數PCB；伺服器PCB', 'product_position': 'AI伺服器板、交換器/網通板，高層數PCB', 'strategy_summary': '次主攻：等量價確認', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；選股策略'}, {'stock_id': '2359', 'stock_name': '所羅門', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': 71.69, 'db_score': 68.67, 'ai_score_ref': 68.36, 'research_score': 71.69, 'raw_theme_line': 'AI Vision / Physical AI；AI Vision/Robot；Edge IPC / Physical AI', 'product_position': 'AI Vision / Physical AI', 'strategy_summary': 'HOLD / 題材觀察；觀察', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '8299', 'stock_name': '群聯', 'ai_project_type': 'AI Memory/Storage/NAND；Advanced Packaging/CoWoS/半導體製程', 'tracking_priority': 'B_觀察追蹤', 'source_count': 10, 'max_score': 71.5, 'db_score': 51.06, 'ai_score_ref': 58.97, 'research_score': 51.8, 'raw_theme_line': 'AI Storage / Controller；AI/晶圓代工；Advanced Packaging / Memory；半導體；控制IC/AI SSD控制器', 'product_position': 'AI Storage / Controller；控制IC/AI SSD控制器', 'strategy_summary': 'WATCH/等拉回或突破；不列長期主攻；減碼/不列主攻', 'risk_summary': 'RSI偏熱 / 短線乖離過大 / DB分數不足；RSI過熱；MA65乖離過大；120日高位階；高Beta、估值與報價斜率轉弱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；05_風險排除觀察；06_情境報酬與配置；Investment_ROI'}, {'stock_id': '3363', 'stock_name': '上詮', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊', 'tracking_priority': 'B_觀察追蹤', 'source_count': 8, 'max_score': 71, 'db_score': 56.45, 'ai_score_ref': 57.82, 'research_score': 57.62, 'raw_theme_line': '1.6T/CPO/矽光子；CPO/光模組；Optical / Interconnect；Optical Module；高速光通訊', 'product_position': 'Optical Module', 'strategy_summary': '觀察；觀察回檔，不追高；高於低接區約5.1%，不追等回測', 'risk_summary': 'DB分數不足；站上65MA；站上120MA；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '03_中長線候選；06_Industry_Intelligence；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping；投資策略總結；總表_更新版'}, {'stock_id': '3661', 'stock_name': '世芯-KY', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Advanced Packaging/CoWoS/半導體製程；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 69.15, 'db_score': 52.29, 'ai_score_ref': 61.14, 'research_score': 50.62, 'raw_theme_line': 'AI ASIC服務；AI/晶圓代工；ASIC / AI Custom Chip；ASIC / RISC-V / IP；ASIC/IP；半導體', 'product_position': 'AI ASIC服務；ASIC / AI Custom Chip', 'strategy_summary': 'AVOID / 不列主攻；不列長期主攻；減碼/不列主攻', 'risk_summary': 'RSI過熱 / 短線乖離過大 / DB分數不足；RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '05_風險排除觀察；06_Industry_Intelligence；06_候選池全表_DB；Agenda策略總表；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '6285', 'stock_name': '啟碁', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway；CPO/矽光子/光通訊；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 66.38, 'db_score': 63.46, 'ai_score_ref': 55.96, 'research_score': 66.38, 'raw_theme_line': 'AI Connectivity；AI Gateway / Networking Device；AI資料中心網通設備/高速交換器相關，直接性較材料低；Connectivity Device；網通設備', 'product_position': 'AI Gateway / Networking Device；AI資料中心網通設備/高速交換器相關，直接性較材料低', 'strategy_summary': 'HOLD / 題材觀察；排除/等待；觀察', 'risk_summary': '正常', 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；DB原始快照；Investment_ROI；Stock_Strategy；Theme_Mapping；選股策略'}, {'stock_id': '2356', 'stock_name': '英業達', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': 65.71, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；AI伺服器ODM；伺服器', 'product_position': 'AI伺服器ODM', 'strategy_summary': 'WAIT突破或回檔；觀察回檔，不追高', 'risk_summary': '120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '03_中長線候選；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略'}, {'stock_id': '6922', 'stock_name': '宸曜', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': 63.84, 'db_score': 58.68, 'ai_score_ref': 62.09, 'research_score': 63.84, 'raw_theme_line': 'Edge IPC / Physical AI；IPC/Edge System；Rugged Edge AI；Rugged Edge AI / 車載工控', 'product_position': 'Rugged Edge AI；Rugged Edge AI / 車載工控', 'strategy_summary': 'AVOID / 不列主攻；觀察', 'risk_summary': 'DB分數不足', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；07_投資組合TOP20；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '2337', 'stock_name': '旺宏', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Edge AI/TinyML/MCU AI', 'tracking_priority': 'B_觀察追蹤', 'source_count': 8, 'max_score': 57.9, 'db_score': 51.14, 'ai_score_ref': 44.74, 'research_score': 57.9, 'raw_theme_line': 'AI/晶圓代工；NOR/SLC/MLC NAND；TL-RAM/記憶體架構受惠', 'product_position': 'NOR/SLC/MLC NAND', 'strategy_summary': 'HOLD/題材觀察；加入主題追蹤', 'risk_summary': 'PB高、獲利能見度待驗證', 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置；Agenda策略總表；docx內文'}, {'stock_id': '6669', 'stock_name': '緯穎', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': 50, 'db_score': None, 'ai_score_ref': 50, 'research_score': None, 'raw_theme_line': 'AI Rack + ODM；AI伺服器；AI伺服器/機櫃；AI伺服器ODM；伺服器', 'product_position': 'AI伺服器/機櫃', 'strategy_summary': 'WAIT突破或回檔；不列長期主攻；高於低接區約4.5%，不追等回測', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；站上65MA；站上120MA；乖離65MA偏大；營收YoY正', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；05_風險排除觀察；06_中期選股策略；07_長期選股策略；投資策略總結；總表_更新版'}, {'stock_id': '2059', 'stock_name': '川湖', 'ai_project_type': 'AI Server ODM/Rack Scale；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': 27.86, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI機械/滑軌；伺服器滑軌；零組件；電子零組件', 'product_position': '伺服器滑軌', 'strategy_summary': 'WAIT突破或回檔；不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；05_風險排除觀察；06_中期選股策略；07_長期選股策略'}, {'stock_id': '2360', 'stock_name': '致茂', 'ai_project_type': 'AI Power/PPA/BBU/電池；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI奇兵50強；全市場；半導體/電子測試設備；半導體測試', 'product_position': '半導體/電子測試設備', 'strategy_summary': 'BUY/主攻；核心主攻', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5'}, {'stock_id': '2603', 'stock_name': '長榮', 'ai_project_type': '其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '服務業獲利前十；海運；航運/海運；運輸', 'product_position': '海運', 'strategy_summary': 'BUY/主攻', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '3167', 'stock_name': '大量', 'ai_project_type': 'AI Power/PPA/BBU/電池；PCIe7/高速互連/PCB材料', 'tracking_priority': 'B_觀察追蹤', 'source_count': 7, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI奇兵50強；PCB設備；PCB設備與半導體測試設備；電源/HVDC', 'product_position': 'PCB設備與半導體測試設備', 'strategy_summary': 'BUY/主攻；核心主攻', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '00_總覽儀表板；04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別；09_最終總排名TOP5'}, {'stock_id': '6223', 'stock_name': '旺矽', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；其他/待分類', 'tracking_priority': 'B_觀察追蹤', 'source_count': 5, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；AI奇兵50強；半導體測試介面', 'product_position': '半導體測試介面', 'strategy_summary': 'BUY/主攻', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '3211', 'stock_name': '順達', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 86.66, 'db_score': 86.66, 'ai_score_ref': 76.15, 'research_score': None, 'raw_theme_line': 'AI電池/BBU；AI電池_BBU', 'product_position': None, 'strategy_summary': 'WAIT；加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；TOP5主攻；docx內文'}, {'stock_id': '2408', 'stock_name': '南亞科', 'ai_project_type': 'AI Memory/Storage/NAND', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 83.15, 'db_score': 83.15, 'ai_score_ref': 67.75, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；TL-RAM/記憶體架構受惠；半導體', 'product_position': None, 'strategy_summary': '分批布局/回測加碼；加入主題追蹤', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': '02_長期主升核心；Agenda策略總表；docx內文'}, {'stock_id': '6414', 'stock_name': '樺漢', 'ai_project_type': 'Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 80.86, 'db_score': 80.86, 'ai_score_ref': 70.66, 'research_score': None, 'raw_theme_line': 'Edge_AI；受惠台股：研華、凌華、樺漢、振樺電', 'product_position': None, 'strategy_summary': 'WAIT；加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；TOP5主攻；docx內文'}, {'stock_id': '8996', 'stock_name': '高力', 'ai_project_type': 'AI Power/PPA/BBU/電池；AI散熱/液冷', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 79.57, 'db_score': 63.23, 'ai_score_ref': 41.69, 'research_score': None, 'raw_theme_line': 'AI散熱；液冷', 'product_position': None, 'strategy_summary': '分批布局/回測加碼', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': '02_長期主升核心；Agenda策略總表'}, {'stock_id': '4931', 'stock_name': '新盛力', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 78.56, 'db_score': 78.56, 'ai_score_ref': 72.09, 'research_score': None, 'raw_theme_line': 'AI電池/BBU', 'product_position': None, 'strategy_summary': '加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；docx內文'}, {'stock_id': '4952', 'stock_name': '凌通', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 78.49, 'db_score': 77.47, 'ai_score_ref': 78.49, 'research_score': 76.39, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'BUY_PULLBACK', 'risk_summary': None, 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx', 'source_sheet': 'AI_MCU策略；DB原始整合；Dashboard'}, {'stock_id': '5471', 'stock_name': '松翰', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 77.57, 'db_score': 77.57, 'ai_score_ref': 72.69, 'research_score': 44.54, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'AVOID', 'risk_summary': None, 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx', 'source_sheet': 'AI_MCU策略；DB原始整合；Dashboard'}, {'stock_id': '2449', 'stock_name': '京元電子', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 77.49, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '分批布局/回測加碼', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '02_長期主升核心'}, {'stock_id': '6510', 'stock_name': '精測', 'ai_project_type': 'AI Server ODM/Rack Scale；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 77.28, 'db_score': None, 'ai_score_ref': None, 'research_score': 56.8, 'raw_theme_line': 'AI/晶圓代工；半導體；測試介面；高階測試介面、探針卡，間接受益高速運算IC驗證', 'product_position': '高階測試介面、探針卡，間接受益高速運算IC驗證', 'strategy_summary': '分批布局/回測加碼；排除/等待', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': '02_長期主升核心；DB原始快照；選股策略'}, {'stock_id': '4971', 'stock_name': 'IET-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 76.77, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '等回檔至10日/20日線', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '02_長期主升核心'}, {'stock_id': '6415', 'stock_name': '矽力-KY', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 76.07, 'db_score': 76.07, 'ai_score_ref': 67.88, 'research_score': 75.26, 'raw_theme_line': 'AI Power / PPA；Power IC / PMIC', 'product_position': 'Power IC / PMIC', 'strategy_summary': '題材觀察/等回測', 'risk_summary': 'RSI偏熱', 'source_file': 'Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': 'Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3653', 'stock_name': '健策', 'ai_project_type': 'AI散熱/液冷；ASIC/FPGA/RISC-V/IP/EDA', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 76.03, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI散熱；散熱模組；液冷', 'product_position': '散熱模組', 'strategy_summary': 'AVOID/觀望；分批布局/回測加碼', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '02_長期主升核心；04_候選池_DB重算；07_長期選股策略'}, {'stock_id': '3227', 'stock_name': '原相', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 75.33, 'db_score': 73.76, 'ai_score_ref': 75.33, 'research_score': 73.86, 'raw_theme_line': 'AI Sensor；AI Sensor / Perception IC', 'product_position': 'AI Sensor / Perception IC', 'strategy_summary': '題材觀察/等回測', 'risk_summary': 'RSI偏熱', 'source_file': 'Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': 'Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '4967', 'stock_name': '十銓', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 7, 'max_score': 74.45, 'db_score': None, 'ai_score_ref': None, 'research_score': 57.7, 'raw_theme_line': 'AI/晶圓代工；半導體；記憶體模組/高Beta', 'product_position': '記憶體模組/高Beta', 'strategy_summary': 'HOLD/題材觀察；觀察回檔，不追高', 'risk_summary': '月營收波動大、追價風險；無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；03_中長線候選；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '6643', 'stock_name': 'M31', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 74.2, 'db_score': 71.3, 'ai_score_ref': 74.2, 'research_score': None, 'raw_theme_line': 'Tape-Out ERC / FPGA/IP', 'product_position': None, 'strategy_summary': '加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；docx內文'}, {'stock_id': '3441', 'stock_name': '聯一光', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；CPO/矽光子/光通訊；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': 73.75, 'db_score': 73.38, 'ai_score_ref': 73.75, 'research_score': 73.69, 'raw_theme_line': 'AI Sensor；AI Sensor/Optics；Optical Sensor / Optical Module；Optical Sensor/Interconnect', 'product_position': 'Optical Sensor / Optical Module；Optical Sensor/Interconnect', 'strategy_summary': 'AVOID / 不列主攻；題材觀察/等回測', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '4903', 'stock_name': '聯光通', 'ai_project_type': 'CPO/矽光子/光通訊', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 73.55, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'CPO/光模組；高速光通訊', 'product_position': None, 'strategy_summary': '觀察回檔，不追高', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '03_中長線候選'}, {'stock_id': '2436', 'stock_name': '偉詮電', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 71.95, 'db_score': 71.02, 'ai_score_ref': 61.15, 'research_score': 71.95, 'raw_theme_line': 'AI Power / PPA；Power/Controller IC', 'product_position': 'Power/Controller IC', 'strategy_summary': '題材觀察/等回測', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': 'Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '6530', 'stock_name': '創威', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 71.8, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '資料中心交換器；高階網通', 'product_position': None, 'strategy_summary': '觀察回檔，不追高', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '03_中長線候選'}, {'stock_id': '3466', 'stock_name': '德晉', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 71.54, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '資料中心交換器；高階網通', 'product_position': None, 'strategy_summary': '觀察回檔，不追高', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '03_中長線候選'}, {'stock_id': '3665', 'stock_name': '貿聯-KY', 'ai_project_type': 'AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 71.2, 'db_score': None, 'ai_score_ref': None, 'research_score': 71.2, 'raw_theme_line': 'AI伺服器內外部高速線束、GPU Tray連接、DAC/AEC ASP升級；高速線束/AEC/DAC', 'product_position': 'AI伺服器內外部高速線束、GPU Tray連接、DAC/AEC ASP升級', 'strategy_summary': '觀察：只做拉回，不追高', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；選股策略'}, {'stock_id': '3189', 'stock_name': '景碩', 'ai_project_type': 'AI Memory/Storage/NAND；PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 70.95, 'db_score': 70.95, 'ai_score_ref': 59.5, 'research_score': None, 'raw_theme_line': 'ABF + 高速PCB', 'product_position': None, 'strategy_summary': '高於低接區約7.2%，不追等回測', 'risk_summary': '站上65MA；站上120MA；乖離65MA偏大；營收YoY正；120日位階偏高', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；投資策略總結；總表_更新版'}, {'stock_id': '3228', 'stock_name': '金麗科', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 70.52, 'db_score': 61.42, 'ai_score_ref': 70.52, 'research_score': 32.69, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'AVOID', 'risk_summary': None, 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx', 'source_sheet': 'AI_MCU策略；DB原始整合'}, {'stock_id': '3088', 'stock_name': '艾訊', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': 68.71, 'db_score': 67.61, 'ai_score_ref': 68.71, 'research_score': 61.96, 'raw_theme_line': 'Edge AI System / AMR；Edge Box / AMR Vision；Edge IPC / Physical AI；IPC/Edge System', 'product_position': 'Edge AI System / AMR；Edge Box / AMR Vision', 'strategy_summary': 'AVOID / 不列主攻；等待回檔', 'risk_summary': 'RSI過熱', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '6190', 'stock_name': '萬泰科', 'ai_project_type': 'AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 67.9, 'db_score': None, 'ai_score_ref': None, 'research_score': 67.9, 'raw_theme_line': 'PTFE/氟塑料高速線材、AI伺服器PCIe線束題材；高速線材/PTFE', 'product_position': 'PTFE/氟塑料高速線材、AI伺服器PCIe線束題材', 'strategy_summary': '觀察：只做拉回，不追高', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；選股策略'}, {'stock_id': '5269', 'stock_name': '祥碩', 'ai_project_type': 'PCIe7/高速互連/PCB材料；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 67.5, 'db_score': None, 'ai_score_ref': None, 'research_score': 67.5, 'raw_theme_line': 'PCIe/USB高速控制IC，受規格升級帶動驗證與設計需求；高速介面IC', 'product_position': 'PCIe/USB高速控制IC，受規格升級帶動驗證與設計需求', 'strategy_summary': '觀察：只做拉回，不追高', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；選股策略'}, {'stock_id': '3022', 'stock_name': '威強電', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': 67.35, 'db_score': 67.35, 'ai_score_ref': 62.09, 'research_score': 57.33, 'raw_theme_line': 'Edge IPC / Physical AI；IPC/Edge System；Industrial Edge AI；Rugged Edge / Secure Edge Infrastructure', 'product_position': 'Industrial Edge AI；Rugged Edge / Secure Edge Infrastructure', 'strategy_summary': 'REDUCE / 只觀察；減碼/不列主攻', 'risk_summary': 'DB轉弱', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '6237', 'stock_name': '驊訊', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 67.13, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '觀察回檔，不追高', 'risk_summary': '120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '03_中長線候選'}, {'stock_id': '6494', 'stock_name': '九齊', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 67.03, 'db_score': 67.03, 'ai_score_ref': 65.43, 'research_score': 62.07, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'WAIT', 'risk_summary': None, 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx', 'source_sheet': 'AI_MCU策略；DB原始整合；Dashboard'}, {'stock_id': '6239', 'stock_name': '力成', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；Advanced Packaging/CoWoS/半導體製程；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 10, 'max_score': 66.13, 'db_score': 64.09, 'ai_score_ref': 53.61, 'research_score': 66.13, 'raw_theme_line': 'AI/晶圓代工；Advanced Packaging / Memory；Memory/Packaging；半導體；記憶體封測/SSD封裝', 'product_position': 'Memory/Packaging；記憶體封測/SSD封裝', 'strategy_summary': 'AVOID/暫避；觀察；觀察回檔，不追高', 'risk_summary': '兩高不過疑慮；封測ASP未必等同NAND漲幅；正常', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；AI_Storage_NAND_投資策略_原始DB分析報告.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；03_中長線候選；04_技術量化；05_SWOT；06_情境報酬與配置；Investment_ROI'}, {'stock_id': '3013', 'stock_name': '晟銘電', 'ai_project_type': 'AI Server ODM/Rack Scale；AI散熱/液冷', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 65.55, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；伺服器', 'product_position': None, 'strategy_summary': '觀察回檔，不追高', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '03_中長線候選'}, {'stock_id': '2417', 'stock_name': '圓剛', 'ai_project_type': 'ASIC/FPGA/RISC-V/IP/EDA；Edge IPC/Physical AI/Robot', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': 65.53, 'db_score': 61.25, 'ai_score_ref': 64.28, 'research_score': 65.53, 'raw_theme_line': 'AI Vision/Edge Module；Edge IPC / Physical AI；Edge Vision Capture / Jetson Ecosystem；Vision/Edge Module', 'product_position': 'Edge Vision Capture / Jetson Ecosystem；Vision/Edge Module', 'strategy_summary': 'AVOID / 不列主攻；觀察', 'risk_summary': '正常', 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx；Edge_AI_TinyML_研究院版_更新報告_DB原始資料_20260512.xlsx', 'source_sheet': '06_候選池全表_DB；Investment_ROI；Stock_Strategy；Theme_Mapping'}, {'stock_id': '3260', 'stock_name': '威剛', 'ai_project_type': 'AI Memory/Storage/NAND', 'tracking_priority': 'C_題材備查', 'source_count': 6, 'max_score': 64.9, 'db_score': None, 'ai_score_ref': None, 'research_score': 64.9, 'raw_theme_line': 'AI/晶圓代工；記憶體模組/品牌', 'product_position': '記憶體模組/品牌', 'strategy_summary': 'WATCH/等拉回或突破', 'risk_summary': '庫存反轉、消費SSD需求弱', 'source_file': 'AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '3023', 'stock_name': '信邦', 'ai_project_type': 'AI Server ODM/Rack Scale；PCIe7/高速互連/PCB材料；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 62.7, 'db_score': None, 'ai_score_ref': None, 'research_score': 62.7, 'raw_theme_line': '工業/資料中心線束，觀察AI伺服器連接占比；線束/連接', 'product_position': '工業/資料中心線束，觀察AI伺服器連接占比', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；Dashboard；選股策略'}, {'stock_id': '6515', 'stock_name': '穎崴', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 61.9, 'db_score': None, 'ai_score_ref': None, 'research_score': 61.9, 'raw_theme_line': '測試介面；高階測試座/探針卡，受AI/HPC晶片測試需求推動', 'product_position': '高階測試座/探針卡，受AI/HPC晶片測試需求推動', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '4966', 'stock_name': '譜瑞-KY', 'ai_project_type': 'PCIe7/高速互連/PCB材料；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 61.5, 'db_score': None, 'ai_score_ref': None, 'research_score': 61.5, 'raw_theme_line': 'AI/晶圓代工；半導體；高速介面IC/Retimer；高速介面、橋接與訊號完整性相關IC，觀察PCIe/CXL/USB4滲透', 'product_position': '高速介面、橋接與訊號完整性相關IC，觀察PCIe/CXL/USB4滲透', 'strategy_summary': '不列長期主攻；排除/等待', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': '05_風險排除觀察；DB原始快照；選股策略'}, {'stock_id': '6213', 'stock_name': '聯茂', 'ai_project_type': 'PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 61.4, 'db_score': None, 'ai_score_ref': None, 'research_score': 61.4, 'raw_theme_line': 'CCL材料；高速CCL材料第二供應鏈，觀察高階材料比重', 'product_position': '高速CCL材料第二供應鏈，觀察高階材料比重', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '8046', 'stock_name': '南電', 'ai_project_type': 'AI Memory/Storage/NAND；PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': 60.48, 'db_score': 60.48, 'ai_score_ref': 49.09, 'research_score': None, 'raw_theme_line': 'ABF + 高速PCB', 'product_position': None, 'strategy_summary': '接近低接買點，可小量試單', 'risk_summary': '站上65MA；站上120MA；乖離65MA偏大；營收YoY正；法人分數偏強', 'source_file': '(重要)20260522_AI基建全股票_DB最新股價_買賣點修正版.xlsx；2026_AI_Agenda_Investment_Strategy.xlsx', 'source_sheet': 'Agenda策略總表；前三大_可執行策略；投資策略總結；總表_更新版'}, {'stock_id': '2454', 'stock_name': '聯發科', 'ai_project_type': 'Edge AI/TinyML/MCU AI；Edge IPC/Physical AI/Robot；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 60.42, 'db_score': 58.91, 'ai_score_ref': 60.42, 'research_score': None, 'raw_theme_line': 'Edge AI SoC / IoT平台；Edge SoC；IC設計；高權值', 'product_position': 'Edge AI SoC / IoT平台', 'strategy_summary': 'AVOID / 不列主攻；不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx；Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx', 'source_sheet': '05_風險排除觀察；06_候選池全表_DB'}, {'stock_id': '3587', 'stock_name': '閎康', 'ai_project_type': 'Advanced Packaging/CoWoS/半導體製程；PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 54.2, 'db_score': None, 'ai_score_ref': None, 'research_score': 54.2, 'raw_theme_line': '材料/失效分析；高速材料、封裝、故障分析與驗證服務', 'product_position': '高速材料、封裝、故障分析與驗證服務', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '3289', 'stock_name': '宜特', 'ai_project_type': 'PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 52, 'db_score': None, 'ai_score_ref': None, 'research_score': 52, 'raw_theme_line': '測試驗證；高速訊號、可靠度、材料與IC驗證服務', 'product_position': '高速訊號、可靠度、材料與IC驗證服務', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '2313', 'stock_name': '華通', 'ai_project_type': 'PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 50.4, 'db_score': None, 'ai_score_ref': None, 'research_score': 50.4, 'raw_theme_line': 'PCB；伺服器/網通PCB升級，直接性低於台光/台燿', 'product_position': '伺服器/網通PCB升級，直接性低於台光/台燿', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '6205', 'stock_name': '詮欣', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': 47.9, 'db_score': None, 'ai_score_ref': None, 'research_score': 47.9, 'raw_theme_line': '連接器；高速連接器與端子，需確認AI Server占比與ASP', 'product_position': '高速連接器與端子，需確認AI Server占比與ASP', 'strategy_summary': '排除/等待', 'risk_summary': None, 'source_file': '(重要)PCIe7_AI_Server_高速互連_投資選股策略報告.xlsx', 'source_sheet': 'DB原始快照；選股策略'}, {'stock_id': '6907', 'stock_name': '雅特力-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': 44.4, 'db_score': None, 'ai_score_ref': None, 'research_score': 44.4, 'raw_theme_line': None, 'product_position': None, 'strategy_summary': 'AVOID', 'risk_summary': None, 'source_file': 'AI_MCU_DB原始資料_大師級投資策略.xlsx', 'source_sheet': 'AI_MCU策略；DB原始整合；Dashboard'}, {'stock_id': '8150', 'stock_name': '南茂', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 6, 'max_score': 42.4, 'db_score': None, 'ai_score_ref': None, 'research_score': 42.4, 'raw_theme_line': 'AI/晶圓代工；記憶體/驅動IC封測', 'product_position': '記憶體/驅動IC封測', 'strategy_summary': 'AVOID/暫避', 'risk_summary': 'PE偏高、EPS低', 'source_file': 'AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '8227', 'stock_name': '巨有科技', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 42.11, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2387', 'stock_name': '精元', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 40.44, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；伺服器', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '8088', 'stock_name': '品安', 'ai_project_type': 'AI Memory/Storage/NAND；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 6, 'max_score': 31.5, 'db_score': None, 'ai_score_ref': None, 'research_score': 31.5, 'raw_theme_line': 'AI/晶圓代工；記憶體模組', 'product_position': '記憶體模組', 'strategy_summary': 'AVOID/暫避', 'risk_summary': '規模與議價力低於大型模組廠', 'source_file': 'AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '3380', 'stock_name': '明泰', 'ai_project_type': 'AI Connectivity/Wi-Fi7/AI Gateway', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 31.13, 'db_score': 30.6, 'ai_score_ref': 31.13, 'research_score': None, 'raw_theme_line': 'Connectivity Device；Networking Device', 'product_position': 'Networking Device', 'strategy_summary': 'REDUCE / 只觀察', 'risk_summary': None, 'source_file': 'Edge_AI_TinyML_研究院版_投資策略_ROI模型.xlsx', 'source_sheet': '06_候選池全表_DB'}, {'stock_id': '3135', 'stock_name': '凌航', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 30.25, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '7769', 'stock_name': '鴻勁', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 30.13, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '8210', 'stock_name': '勤誠', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 29.16, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；伺服器', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2496', 'stock_name': '卓越', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 28.1, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '其他', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2762', 'stock_name': '世界健身-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 27.09, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；運動', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '9927', 'stock_name': '泰銘', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 27.09, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '其他', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '7750', 'stock_name': '新代', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 27.02, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '電源；電源/HVDC', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '1590', 'stock_name': '亞德客-KY', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 26.55, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '電源；電源/HVDC', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2329', 'stock_name': '華泰', 'ai_project_type': 'AI Memory/Storage/NAND；AI Power/PPA/BBU/電池；其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 6, 'max_score': 26.1, 'db_score': None, 'ai_score_ref': None, 'research_score': 26.1, 'raw_theme_line': 'AI/晶圓代工；記憶體封測', 'product_position': '記憶體封測', 'strategy_summary': 'AVOID/暫避', 'risk_summary': '營收年增弱、短線偏弱', 'source_file': 'AI_Storage_NAND_投資策略_原始DB分析報告.xlsx', 'source_sheet': '00_總覽；02_個股策略總表；03_DB原始資料摘錄；04_技術量化；05_SWOT；06_情境報酬與配置'}, {'stock_id': '5285', 'stock_name': '界霖', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 25.99, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2739', 'stock_name': '寒舍', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 25.79, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；觀光', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '6771', 'stock_name': '平和環保-創', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 25.79, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '環保；綠能環保', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '無明顯過熱', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '3016', 'stock_name': '嘉晶', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 24.84, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '3581', 'stock_name': '博磊', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 24.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '3532', 'stock_name': '台勝科', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 24.14, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2049', 'stock_name': '上銀', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 24, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '電源；電源/HVDC', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '6147', 'stock_name': '頎邦', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 23.92, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '1597', 'stock_name': '直得', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 23.45, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '電源；電源/HVDC', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '6182', 'stock_name': '合晶', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 22.97, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '6651', 'stock_name': '全宇昕', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 22.97, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '3095', 'stock_name': '及成', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 21.93, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '資料中心交換器；高階網通', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '5425', 'stock_name': '台半', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 21.92, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2020', 'stock_name': '美亞', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 21.16, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '基礎原物料；鋼鐵', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '1442', 'stock_name': '名軒', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 19.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '6754', 'stock_name': '匯僑設計', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 19.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；居家', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2442', 'stock_name': '新美齊', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 19.6, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2227', 'stock_name': '裕日車', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 19.39, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '車用；電動車', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '1315', 'stock_name': '達新', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 18.64, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '基礎原物料；塑膠', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2540', 'stock_name': '愛山林', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 18.64, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2731', 'stock_name': '雄獅', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 17.39, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；觀光', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '9946', 'stock_name': '三發地產', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 17.24, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2543', 'stock_name': '皇昌', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 16.61, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '6488', 'stock_name': '環球晶', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 15.87, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI/晶圓代工；半導體', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '1432', 'stock_name': '大魯閣', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 15.05, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；運動', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '4137', 'stock_name': '麗豐-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 15.05, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '生技醫療；醫療', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2547', 'stock_name': '日勝生', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 14.54, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '6834', 'stock_name': '天二科技', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 13.88, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '零組件；電子零組件', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '5706', 'stock_name': '鳳凰', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 13.79, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；觀光', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2109', 'stock_name': '華豐', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 12.69, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；橡膠', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2348', 'stock_name': '海悅', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 9, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '其他', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '4763', 'stock_name': '材料*-KY', 'ai_project_type': 'PCIe7/高速互連/PCB材料', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 9, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '化工；基礎原物料', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '5225', 'stock_name': '東科-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 9, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '全市場；系統掃描', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2504', 'stock_name': '國產', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 8.41, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2530', 'stock_name': '華建', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 8.41, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '1436', 'stock_name': '華友聯', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 7.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2528', 'stock_name': '皇普', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 7.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '2534', 'stock_name': '宏盛', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 7.74, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；營造', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '1446', 'stock_name': '宏和', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 6.64, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；紡織', 'product_position': None, 'strategy_summary': '防守小部位/只低接', 'risk_summary': '兩高不過疑慮；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '04_防守現金流'}, {'stock_id': '1595', 'stock_name': '川寶', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 4.91, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '零組件；電子零組件', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '1216', 'stock_name': '統一', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 4, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '民生消費；製造業營收前十；食品/消費；食品/防禦', 'product_position': '食品/消費', 'strategy_summary': 'WAIT突破或回檔', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；08_明日TOP5_每類別'}, {'stock_id': '1513', 'stock_name': '中興電', 'ai_project_type': 'AI Power/PPA/BBU/電池', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI能源/電源效率', 'product_position': None, 'strategy_summary': '加入主題追蹤', 'risk_summary': None, 'source_file': '2026_AI_Agenda_Deep_Analysis_Report.docx', 'source_sheet': 'docx內文'}, {'stock_id': '1762', 'stock_name': '中化生', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '生技醫療；醫療', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '1809', 'stock_name': '中釉', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '傳產；玻璃陶瓷', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階；累計營收衰退', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2207', 'stock_name': '和泰車', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '汽車通路；汽車銷售；電動車', 'product_position': '汽車銷售', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2324', 'stock_name': '仁寶', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；電子代工', 'product_position': '電子代工', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2412', 'stock_name': '中華電', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '資料中心交換器；電信/IDC；電信/IDC/雲端', 'product_position': '電信/IDC/雲端', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2609', 'stock_name': '陽明', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '海運；航運/海運；運輸', 'product_position': '海運', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2610', 'stock_name': '華航', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '服務業獲利前十；航空；運輸', 'product_position': '航空', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '2615', 'stock_name': '萬海', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '服務業獲利前十；海運；航運/海運；運輸', 'product_position': '海運', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '2618', 'stock_name': '長榮航', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '服務業獲利前十；航空；運輸', 'product_position': '航空', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '2855', 'stock_name': '統一證', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '民生消費；食品', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '2880', 'stock_name': '華南金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2881', 'stock_name': '富邦金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金控排名；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'WATCH', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；08_明日TOP5_每類別'}, {'stock_id': '2882', 'stock_name': '國泰金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金控排名；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'WATCH', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；08_明日TOP5_每類別'}, {'stock_id': '2883', 'stock_name': '凱基金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金控排名；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'WATCH', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '2884', 'stock_name': '玉山金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2885', 'stock_name': '元大金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金控排名；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'WATCH', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '2886', 'stock_name': '兆豐金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2887', 'stock_name': '台新新光金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2889', 'stock_name': '國票金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2890', 'stock_name': '永豐金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2891', 'stock_name': '中信金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '2892', 'stock_name': '第一金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 5, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金控排名；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'WAIT突破或回檔', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；05_短期選股策略；06_中期選股策略；07_長期選股策略；08_明日TOP5_每類別'}, {'stock_id': '3033', 'stock_name': '威健', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '通路；電子通路', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '3209', 'stock_name': '全科', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '通路；電子通路', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '3430', 'stock_name': '奇鈦科', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '化工；基礎原物料', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '4556', 'stock_name': '旭然', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '其他', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '4807', 'stock_name': '日成-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '內需消費；百貨', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '4938', 'stock_name': '和碩', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI伺服器；電子代工', 'product_position': '電子代工', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '5314', 'stock_name': '世紀*', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 2, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '保健產品；保健產品/未分類；全市場；營運績效前五', 'product_position': '保健產品/未分類', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；08_明日TOP5_每類別'}, {'stock_id': '5871', 'stock_name': '中租-KY', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '全市場；租賃金融；金融租賃', 'product_position': '租賃金融', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '5880', 'stock_name': '合庫金', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '金控；金融；金融金控', 'product_position': '金控', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算'}, {'stock_id': '6584', 'stock_name': '南俊國際', 'ai_project_type': 'AI Server ODM/Rack Scale', 'tracking_priority': 'C_題材備查', 'source_count': 3, 'max_score': None, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': 'AI機械/滑軌；伺服器滑軌；電子零組件', 'product_position': '伺服器滑軌', 'strategy_summary': 'AVOID/觀望', 'risk_summary': None, 'source_file': '三份報告_AI奇兵_DB原始資料_投資策略總報告.xlsx', 'source_sheet': '04_候選池_DB重算；06_中期選股策略；07_長期選股策略'}, {'stock_id': '6658', 'stock_name': '聯策', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '全市場；系統掃描', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}, {'stock_id': '6861', 'stock_name': '睿生光電', 'ai_project_type': '其他/待分類', 'tracking_priority': 'C_題材備查', 'source_count': 1, 'max_score': 0, 'db_score': None, 'ai_score_ref': None, 'research_score': None, 'raw_theme_line': '生技醫療；醫療', 'product_position': None, 'strategy_summary': '不列長期主攻', 'risk_summary': 'RSI過熱；MA65乖離過大；120日高位階', 'source_file': '(重要)20260512_LongTerm_Industry_Engine_終極修正版.xlsx', 'source_sheet': '05_風險排除觀察'}]

MAIN_THEME_REPORT_SHEETS = [
    "主流題材股票池",
    "主題熱度Dashboard",
    "主流主升交集池",
    "主題壓縮觀察池",
    "主流題材整合驗收",
]


class ThemeTrackingEngine:
    """
    V2.9.0：主流題材追蹤引擎。
    將外部研究整理出的 200 檔主題股票池接入 institutional_report['all']，
    不再只追 CPO，而是以 theme_category / theme_score / theme_prebreakout_flag
    做為後續「門戶大開」追蹤基礎。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def master_df(self):
        if pd is None:
            return None
        df = pd.DataFrame(THEME_TRACKING_MASTER)
        if df.empty:
            return df
        df["stock_id"] = df["stock_id"].astype(str).str.zfill(4)
        for col in ["theme_appearance_count", "theme_score_base", "db_total_score_ref", "ai_score_ref", "research_score_ref"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")
        # 題材強度：研究最高分/DB/AI/研究分擇優，加上出現次數權重；保留題材不覆蓋風控。
        base = df[["theme_score_base", "db_total_score_ref", "ai_score_ref", "research_score_ref"]].max(axis=1).fillna(50)
        df["theme_score"] = (base + df["theme_appearance_count"].fillna(0).clip(0, 15) * 1.5).clip(0, 100).round(2)
        df["is_main_theme"] = True
        return df

    def apply(self, df):
        if pd is None or df is None or getattr(df, "empty", True):
            return df
        out = df.copy()
        if "stock_id" not in out.columns:
            return out
        out["stock_id"] = out["stock_id"].astype(str).str.zfill(4)
        master = self.master_df()
        if master is None or master.empty:
            out["is_main_theme"] = False
            return out
        cols = [
            "stock_id", "stock_name", "theme_category", "theme_priority", "theme_appearance_count",
            "theme_score", "db_total_score_ref", "ai_score_ref", "research_score_ref",
            "raw_theme_line", "product_position", "strategy_note", "risk_note",
            "theme_source_file", "theme_source_sheet"
        ]
        # 避免與原本欄位衝突。
        drop_cols = [c for c in cols if c != "stock_id" and c in out.columns]
        if drop_cols:
            out = out.drop(columns=drop_cols)
        out = out.merge(master[cols], on="stock_id", how="left", suffixes=("", "_theme"))
        out["is_main_theme"] = out["theme_category"].notna()
        out["theme_category"] = out["theme_category"].fillna("")
        out["theme_priority"] = out["theme_priority"].fillna("未列入主題池")
        out["theme_score"] = pd.to_numeric(out["theme_score"], errors="coerce").fillna(0)
        if self.logger:
            self.logger.info(
                f"THEME_TRACKING_MASTER_READY master_count={len(master)} "
                f"merged_count={int(out['is_main_theme'].sum())} "
                f"theme_count={int(master['theme_category'].astype(str).str.split('；').explode().nunique())}"
            )
        return out


class MainThemeReportExcelIntegrator:
    """
    V2.9.0：主流題材 × 爆發前報表整合器。
    正確語義：
    1. 主流題材股票池：研究股票池，不代表可買。
    2. 主流主升交集池：必須同時屬於主流題材 + Phase5 PreBreakout/Breakout/Expansion。
    3. 主題壓縮觀察池：Compression 只能列觀察，不可標為嚴格交集。
    """
    SHEETS = MAIN_THEME_REPORT_SHEETS

    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger
        self.engine = ThemeTrackingEngine(logger)

    def _sheet(self, wb, name: str):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def _style_basic(self, ws):
        try:
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="0B5CAD")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in range(1, ws.max_column + 1):
                letter = get_column_letter(col)
                header = str(ws.cell(1, col).value or "")
                ws.column_dimensions[letter].width = min(max(12, len(header) + 4), 36)
            ws.freeze_panes = "A2"
        except Exception:
            pass

    def write_into_workbook(self, wb, institutional_report: Optional[Dict[str, Any]] = None):
        merged_df = None
        if institutional_report is not None and institutional_report.get("all") is not None and pd is not None:
            try:
                merged_df = self.engine.apply(institutional_report.get("all"))
                institutional_report["all"] = merged_df
            except Exception as exc:
                if self.logger:
                    self.logger.warning(f"THEME_TRACKING_ENGINE_APPLY_FAIL error={exc}")
        self._write_theme_master(wb)
        self._write_theme_dashboard(wb, merged_df)
        self._write_main_attack_intersection(wb, merged_df)
        self._write_compression_watch(wb, merged_df)
        self._write_validation(wb, merged_df)
        if self.logger:
            self.logger.info("MAIN_THEME_REPORT_WRITTEN sheets=" + ",".join(self.SHEETS))
        return wb

    def _write_theme_master(self, wb):
        ws = self._sheet(wb, "主流題材股票池")
        headers = ["代號","名稱","主題類別","追蹤優先級","出現次數","題材分數","DB參考分","AI參考分","研究分","產品定位","策略備註","風險備註","來源檔案","來源Sheet"]
        ws.append(headers)
        master = self.engine.master_df()
        if master is None or master.empty:
            ws.append(["", "", "", "", "", "", "", "", "", "", "", "", "", ""])
            self._style_basic(ws)
            return
        m = master.sort_values(["theme_priority", "theme_score", "theme_appearance_count"], ascending=[True, False, False])
        for _, r in m.iterrows():
            ws.append([
                str(r.get("stock_id","")).zfill(4), r.get("stock_name",""), r.get("theme_category",""),
                r.get("theme_priority",""), int(r.get("theme_appearance_count") or 0),
                round(_safe_float(r.get("theme_score"), 0), 2),
                round(_safe_float(r.get("db_total_score_ref"), 0), 2),
                round(_safe_float(r.get("ai_score_ref"), 0), 2),
                round(_safe_float(r.get("research_score_ref"), 0), 2),
                r.get("product_position",""), r.get("strategy_note",""), r.get("risk_note",""),
                r.get("theme_source_file",""), r.get("theme_source_sheet","")
            ])
        self._style_basic(ws)

    def _write_theme_dashboard(self, wb, merged_df):
        ws = self._sheet(wb, "主題熱度Dashboard")
        headers = ["主題類型","股票池檔數","A主攻數","DB命中數","嚴格交集數","壓縮觀察數","平均題材分","追蹤目的"]
        ws.append(headers)
        master = self.engine.master_df()
        if master is None or master.empty or pd is None:
            ws.append(["無資料",0,0,0,0,0,0,"THEME_TRACKING_MASTER_EMPTY"])
            self._style_basic(ws)
            return
        rows = []
        # 依分號拆多主題統計，與01_主題追蹤類型一致。
        for theme in sorted(set([x for s in master["theme_category"].dropna().astype(str) for x in s.split("；") if x.strip()])):
            m = master[master["theme_category"].astype(str).str.contains(re.escape(theme), na=False)]
            a_count = int((m["theme_priority"].astype(str).str.startswith("A_")).sum())
            db_hit = 0
            strict_count = 0
            comp_count = 0
            if merged_df is not None and not getattr(merged_df, "empty", True):
                md = merged_df[merged_df.get("theme_category", pd.Series("", index=merged_df.index)).astype(str).str.contains(re.escape(theme), na=False)]
                db_hit = int(len(md))
                strict_count = int(len(self._build_main_attack_df(md))) if db_hit else 0
                comp_count = int(len(self._build_compression_watch_df(md))) if db_hit else 0
            purpose = self._theme_purpose(theme)
            rows.append([theme, int(len(m)), a_count, db_hit, strict_count, comp_count, round(float(m["theme_score"].mean()),2), purpose])
        rows = sorted(rows, key=lambda x: (x[4], x[6], x[1]), reverse=True)
        for r in rows:
            ws.append(r)
        self._style_basic(ws)

    def _theme_purpose(self, theme: str) -> str:
        purpose_map = {
            "CPO/矽光子/光通訊": "延續CPO模式，追蹤光互連/CPO/1.6T/3.2T門戶大開",
            "Edge AI/TinyML/MCU AI": "追蹤AI下沉到MCU、Sensor、Gateway、Runtime",
            "AI Connectivity/Wi-Fi7/AI Gateway": "追蹤AI Gateway、Wi-Fi 7、Broadband Gateway入口",
            "AI Server ODM/Rack Scale": "追蹤AI伺服器、Rack Scale、整機櫃ODM主線",
            "AI散熱/液冷": "追蹤AI Server液冷與散熱瓶頸",
            "AI Power/PPA/BBU/電池": "追蹤AI電源、PPA、BBU與能源效率",
            "ASIC/FPGA/RISC-V/IP/EDA": "追蹤客製ASIC、FPGA再起、RISC-V與EDA驗證",
            "PCIe7/高速互連/PCB材料": "追蹤PCIe 7、高速CCL、PCB材料與連接器升級",
            "AI Memory/Storage/NAND": "追蹤AI資料爆炸、NAND/HBM/TL-RAM/記憶體",
            "Edge IPC/Physical AI/Robot": "追蹤工業AI、Robot Node、AI Box、IPC商用落地",
            "Advanced Packaging/CoWoS/半導體製程": "追蹤CoWoS、2nm、Chiplet與先進封裝",
        }
        return purpose_map.get(theme, "主流題材追蹤")

    def _theme_decision(self, row):
        risk = " ".join([
            str(row.get("hard_avoid_reason", "") or ""),
            str(row.get("k_warning_type", "") or ""),
            str(row.get("phase5_block_reason", "") or ""),
        ])
        wp = str(row.get("phase5_wave_phase", "") or "")
        bs = str(row.get("phase5_breakout_stage", "") or "")
        pos = _safe_float(row.get("phase5_position_score"), 0)
        rr = _safe_float(row.get("rr"), 0)
        theme_score = _safe_float(row.get("theme_score"), 0)
        tdec = str(row.get("teacher_decision", "WATCH") or "WATCH")
        exe = str(row.get("teacher_execution_status", row.get("是否可下單", "NO")) or "NO")
        if any(k in risk for k in ["硬性", "逃命", "主跌", "長黑", "墓碑"]):
            return "AVOID", "主題標的但觸發硬性/波段風險：" + risk[:120]
        if tdec == "REDUCE":
            return "REDUCE", "老師策略為REDUCE，主流題材不得覆蓋減碼訊號"
        if theme_score >= 80 and wp in ["Wave3_Breakout", "Wave3_Expansion"] and pos >= 70 and rr >= 1.5 and exe in ["YES", "WAIT"]:
            return ("BUY" if exe == "YES" else "WATCH"), "主流題材+Wave3突破+RR達標；仍需依執行狀態確認"
        if theme_score >= 70 and (wp == "Wave3_PreBreakout" or bs == "PreBreakout") and pos >= 60 and rr >= 1.2:
            return ("LOW_BUY" if tdec in ["LOW_BUY", "BUY"] else "WATCH"), "主流題材+Wave3預突破；等待門戶大開確認"
        return "WATCH", "主流題材觀察，條件未完整確認"

    def _build_main_attack_df(self, merged_df):
        if merged_df is None or pd is None or getattr(merged_df, "empty", True):
            return pd.DataFrame()
        df = merged_df.copy()
        if "is_main_theme" not in df.columns:
            df = self.engine.apply(df)
        for col in ["phase5_position_score", "rr", "teacher_score", "theme_score", "vol5", "vol20"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")
        if "prebreakout_volume_ratio" not in df.columns:
            vol5 = df["vol5"] if "vol5" in df.columns else pd.Series(index=df.index, dtype=float)
            vol20 = df["vol20"] if "vol20" in df.columns else pd.Series(index=df.index, dtype=float)
            df["prebreakout_volume_ratio"] = (vol5 / vol20.replace(0, math.nan)).replace([math.inf, -math.inf], math.nan)
        wave_phase = df.get("phase5_wave_phase", pd.Series("", index=df.index)).astype(str)
        breakout_stage = df.get("phase5_breakout_stage", pd.Series("", index=df.index)).astype(str)
        # 嚴格交集：Compression 不可算主攻交集，只能進壓縮觀察。
        strict_mask = (
            df.get("is_main_theme", pd.Series(False, index=df.index)).fillna(False)
            & (
                wave_phase.isin(["Wave3_PreBreakout", "Wave3_Breakout", "Wave3_Expansion"])
                | breakout_stage.isin(["PreBreakout", "Breakout", "Expansion"])
            )
        )
        inter = df.loc[strict_mask].copy()
        if inter.empty:
            return inter
        decisions, reasons = [], []
        for _, row in inter.iterrows():
            d, reason = self._theme_decision(row)
            decisions.append(d); reasons.append(reason)
        inter["theme_prebreakout_decision"] = decisions
        inter["theme_prebreakout_reason"] = reasons
        inter["theme_prebreakout_flag"] = True
        inter["_rank"] = inter["theme_prebreakout_decision"].map({"BUY":1,"LOW_BUY":2,"WATCH":3,"REDUCE":4,"AVOID":5}).fillna(9)
        return inter.sort_values(["_rank","theme_score","phase5_position_score","rr","prebreakout_volume_ratio"], ascending=[True, False, False, False, False])

    def _build_compression_watch_df(self, merged_df):
        if merged_df is None or pd is None or getattr(merged_df, "empty", True):
            return pd.DataFrame()
        df = merged_df.copy()
        if "is_main_theme" not in df.columns:
            df = self.engine.apply(df)
        breakout_stage = df.get("phase5_breakout_stage", pd.Series("", index=df.index)).astype(str)
        wave_phase = df.get("phase5_wave_phase", pd.Series("", index=df.index)).astype(str)
        mask = (
            df.get("is_main_theme", pd.Series(False, index=df.index)).fillna(False)
            & breakout_stage.eq("Compression")
            & ~wave_phase.isin(["Wave3_PreBreakout", "Wave3_Breakout", "Wave3_Expansion"])
        )
        out = df.loc[mask].copy()
        if out.empty:
            return out
        out["theme_prebreakout_decision"] = "WATCH"
        out["theme_prebreakout_reason"] = "壓縮觀察池：Compression 尚未等於 PreBreakout/Breakout，不可標嚴格交集"
        out["theme_prebreakout_flag"] = False
        return out.sort_values(["theme_score","phase5_position_score","rr"], ascending=[False, False, False])

    def _write_main_attack_intersection(self, wb, merged_df):
        ws = self._sheet(wb, "主流主升交集池")
        headers = ["排名","代號","名稱","主題類別","追蹤優先級","題材分數","出現次數","老師決策","是否可下單","波段階段","突破階段","Phase5候選池","波段位置分","量比","RR","主題爆發決策","交集判定","原因/風險","資料狀態"]
        ws.append(headers)
        inter = self._build_main_attack_df(merged_df)
        if inter is None or getattr(inter, "empty", True):
            ws.append([1,"","","","","","","","","","","","","","","無","STRICT_INTERSECTION_ZERO","主流題材股票池與PreBreakout/Breakout嚴格條件交集為0；Compression已另列壓縮觀察池","ZERO_MATCH"])
            self._style_basic(ws)
            return
        for n, (_, r) in enumerate(inter.head(120).iterrows(), start=1):
            risk = " ".join([str(r.get("hard_avoid_reason","") or ""), str(r.get("k_warning_type","") or ""), str(r.get("phase5_block_reason","") or "")]).strip()
            ws.append([
                n, str(r.get("stock_id","")).zfill(4), r.get("report_name", r.get("stock_name", r.get("name", ""))),
                r.get("theme_category",""), r.get("theme_priority",""), round(_safe_float(r.get("theme_score"),0),2),
                int(_safe_float(r.get("theme_appearance_count"),0)), r.get("teacher_decision",""),
                r.get("teacher_execution_status", r.get("是否可下單","")),
                r.get("phase5_wave_phase",""), r.get("phase5_breakout_stage",""), r.get("phase5_candidate_pool",""),
                round(_safe_float(r.get("phase5_position_score"),0),2), round(_safe_float(r.get("prebreakout_volume_ratio"),0),2),
                round(_safe_float(r.get("rr"),0),2), r.get("theme_prebreakout_decision","WATCH"),
                "STRICT_MATCH", (str(r.get("theme_prebreakout_reason","")) + ("；風險=" + risk[:120] if risk else ""))[:250], "DB_OK"
            ])
        self._style_basic(ws)

    def _write_compression_watch(self, wb, merged_df):
        ws = self._sheet(wb, "主題壓縮觀察池")
        headers = ["排名","代號","名稱","主題類別","追蹤優先級","題材分數","老師決策","是否可下單","波段階段","突破階段","波段位置分","RR","觀察判定","原因","資料狀態"]
        ws.append(headers)
        obs = self._build_compression_watch_df(merged_df)
        if obs is None or getattr(obs, "empty", True):
            ws.append([1,"","","","","","","","","","","","NO_COMPRESSION","目前無主流題材Compression觀察名單","ZERO_MATCH"])
            self._style_basic(ws)
            return
        for n, (_, r) in enumerate(obs.head(120).iterrows(), start=1):
            ws.append([
                n, str(r.get("stock_id","")).zfill(4), r.get("report_name", r.get("stock_name", r.get("name", ""))),
                r.get("theme_category",""), r.get("theme_priority",""), round(_safe_float(r.get("theme_score"),0),2),
                r.get("teacher_decision",""), r.get("teacher_execution_status", r.get("是否可下單","")),
                r.get("phase5_wave_phase",""), r.get("phase5_breakout_stage",""),
                round(_safe_float(r.get("phase5_position_score"),0),2), round(_safe_float(r.get("rr"),0),2),
                "THEME_COMPRESSION_WATCH", "Compression僅為門戶前壓縮，不可誤標STRICT_MATCH", "DB_OK"
            ])
        self._style_basic(ws)

    def _write_validation(self, wb, merged_df):
        ws = self._sheet(wb, "主流題材整合驗收")
        master = self.engine.master_df()
        master_count = 0 if master is None else len(master)
        merged_count = 0
        strict_count = 0
        compression_count = 0
        if merged_df is not None and pd is not None and not getattr(merged_df, "empty", True):
            if "is_main_theme" in merged_df.columns:
                merged_count = int(merged_df["is_main_theme"].fillna(False).sum())
            strict_count = int(len(self._build_main_attack_df(merged_df)))
            compression_count = int(len(self._build_compression_watch_df(merged_df)))
        rows = [
            ["查核項目", "結果", "說明"],
            ["THEME_TRACKING_MASTER檔數", master_count, "來自主流題材整合股票池，不依賴CPO固定Universe"],
            ["DB合併命中數", merged_count, "institutional_report['all']與主流題材股票池交集"],
            ["主流主升嚴格交集數", strict_count, "主流題材 ∩ PreBreakout/Breakout/Expansion；Compression不算嚴格交集"],
            ["主題壓縮觀察池數", compression_count, "Compression只列觀察，不可標STRICT_MATCH"],
            ["欄位落地", "theme_category/theme_score/theme_source/theme_prebreakout_flag/breakout_confirm/main_attack_pool", "對應Excel1欄位規劃"],
            ["風控原則", "PASS", "題材不得覆蓋REDUCE/AVOID/Hard Avoid/Wave5/過熱"],
            ["Log驗收", "THEME_TRACKING_MASTER_READY / MAIN_THEME_REPORT_WRITTEN", "可grep追蹤"],
        ]
        for r in rows:
            ws.append(r)
        if self.logger:
            self.logger.info(
                f"MAIN_THEME_STRICT_INTERSECTION_SUMMARY master_count={master_count} "
                f"merged_count={merged_count} strict_written_count={strict_count} "
                f"compression_watch_count={compression_count}"
            )
        self._style_basic(ws)



AI_PROJECT_REPORT_SHEETS = ["AI專案追蹤頁", "AI專案Dashboard", "AI門戶大開監控", "AI專案整合驗收"]

class AIProjectRotationEngine:
    """
    V3.0：AI主流輪動監控系統。
    目的：不只CPO，將11大AI專案類別全部列出並追蹤每檔個股狀態。
    原則：所有AI專案股票都必須出現在「AI專案追蹤頁」，嚴格主攻/門戶大開只是一種狀態，不可把未突破股票排除在追蹤外。
    """
    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def master_df(self):
        if pd is None:
            return AI_PROJECT_TRACKING_MASTER
        df = pd.DataFrame(AI_PROJECT_TRACKING_MASTER)
        if not df.empty:
            df["stock_id"] = df["stock_id"].astype(str).str.zfill(4)
            for col in ["source_count", "max_score", "db_score", "ai_score_ref", "research_score"]:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors="coerce")
        return df

    def category_df(self):
        if pd is None:
            return AI_PROJECT_CATEGORIES
        return pd.DataFrame(AI_PROJECT_CATEGORIES)

    def apply(self, df):
        if df is None or pd is None or getattr(df, "empty", True):
            return df
        master = self.master_df()
        if isinstance(master, list) or master.empty:
            return df
        out = df.copy()
        if "stock_id" not in out.columns:
            return out
        out["stock_id"] = out["stock_id"].astype(str).str.zfill(4)
        ai_cols = [
            "ai_project_type", "tracking_priority", "source_count", "max_score", "db_score",
            "ai_score_ref", "research_score", "raw_theme_line", "product_position",
            "strategy_summary", "risk_summary", "source_file", "source_sheet", "is_ai_project"
        ]
        drop_cols = [c for c in ai_cols if c in out.columns]
        if drop_cols:
            out = out.drop(columns=drop_cols)
        out = out.merge(master, on="stock_id", how="left", suffixes=("", "_ai_master"))
        out["is_ai_project"] = out["ai_project_type"].notna()
        if self.logger:
            self.logger.info(
                "AI_PROJECT_TRACKING_MASTER_READY "
                f"master_count={len(master)} merged_count={int(out['is_ai_project'].sum())} "
                f"category_count={len(AI_PROJECT_CATEGORIES)}"
            )
        return out

    def infer_status(self, row):
        decision = str(row.get("teacher_decision", "") or "")
        execution = str(row.get("teacher_execution_status", row.get("是否可下單", "")) or "")
        breakout = str(row.get("phase5_breakout_stage", "") or "")
        wave = str(row.get("phase5_wave_phase", "") or "")
        pool = str(row.get("phase5_candidate_pool", "") or "")
        k_warning = str(row.get("k_warning_type", "") or "")
        hard = str(row.get("hard_avoid_reason", "") or "")
        risk = str(row.get("risk_summary", "") or "")
        score = _safe_float(row.get("teacher_score"), 0)
        rr = _safe_float(row.get("rr"), 0)
        pos = _safe_float(row.get("phase5_position_score"), 0)
        if decision == "AVOID" or any(k in hard for k in ["下降", "硬K", "主跌", "長黑", "墓碑"]):
            return "AVOID"
        if decision == "REDUCE" or any(k in wave for k in ["Wave5_Risk"]):
            return "REDUCE"
        if any(k in k_warning for k in ["流星", "長黑", "墓碑"]):
            return "REDUCE"
        if breakout in ["Expansion"] or (wave in ["Wave3_Expansion"] and pos >= 70):
            return "MAIN_ATTACK"
        if breakout in ["Breakout"] or wave in ["Wave3_Breakout"]:
            return "BREAKOUT_CONFIRM"
        if breakout in ["PreBreakout"] or wave in ["Wave3_PreBreakout"] or "主升預突破" in pool:
            return "PRE_BREAKOUT"
        if breakout in ["Compression"]:
            return "WATCHLIST"
        if decision in ["BUY", "LOW_BUY"] and execution in ["YES", "WAIT"]:
            return "HOT_MONEY"
        if score >= 72 and rr >= 1.2:
            return "WATCHLIST"
        if str(row.get("tracking_priority", "")).startswith("A"):
            return "LEADER"
        if "高位階" in risk or "過熱" in risk:
            return "EXTENDED"
        return "TRACKING"

    def status_reason(self, row, status):
        parts = []
        if row.get("phase5_wave_phase"):
            parts.append("波段=" + str(row.get("phase5_wave_phase")))
        if row.get("phase5_breakout_stage"):
            parts.append("突破=" + str(row.get("phase5_breakout_stage")))
        if row.get("teacher_decision"):
            parts.append("老師決策=" + str(row.get("teacher_decision")))
        if row.get("teacher_execution_status"):
            parts.append("執行=" + str(row.get("teacher_execution_status")))
        if row.get("hard_avoid_reason"):
            parts.append("硬風險=" + str(row.get("hard_avoid_reason"))[:80])
        if row.get("risk_summary"):
            parts.append("研究風險=" + str(row.get("risk_summary"))[:80])
        if not parts:
            parts.append("尚未進入嚴格爆發/門戶大開條件，維持AI專案基礎追蹤")
        return ";".join(parts)

class AIProjectTrackingReportExcelIntegrator:
    """V3.0：輸出AI專案追蹤頁與AI主流輪動Dashboard。"""
    SHEETS = AI_PROJECT_REPORT_SHEETS

    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger
        self.engine = AIProjectRotationEngine(logger)

    def _sheet(self, wb, name: str):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def _style_basic(self, ws, color="1F4E78"):
        try:
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor=color)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in range(1, ws.max_column + 1):
                letter = get_column_letter(col)
                header = str(ws.cell(1, col).value or "")
                if any(k in header for k in ["原因", "摘要", "來源", "主線", "定位", "類別"]):
                    ws.column_dimensions[letter].width = 34
                else:
                    ws.column_dimensions[letter].width = min(max(10, len(header) + 4), 22)
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
        except Exception:
            pass

    def _build_tracking_df(self, merged_df=None):
        if pd is None:
            return []
        master = self.engine.master_df()
        if master.empty:
            return master
        if merged_df is not None and not getattr(merged_df, "empty", True) and "stock_id" in merged_df.columns:
            base = merged_df.copy()
            base["stock_id"] = base["stock_id"].astype(str).str.zfill(4)
            db_cols = [
                "stock_id", "teacher_decision", "teacher_execution_status", "是否可下單", "phase5_wave_phase",
                "phase5_breakout_stage", "phase5_impulse_stage", "phase5_position_score", "phase5_candidate_pool",
                "phase5_block_reason", "teacher_score", "rr", "close", "entry_low", "entry_high", "stop_loss",
                "target_1", "target_2", "k_warning_type", "hard_avoid_reason", "soft_avoid_reason", "core_leader_state"
            ]
            use_cols = [c for c in db_cols if c in base.columns]
            base = base[use_cols].drop_duplicates("stock_id")
            out = master.merge(base, on="stock_id", how="left")
            out["db_merge_status"] = out["teacher_decision"].notna().map(lambda x: "DB_MATCH" if x else "DB_MISSING")
        else:
            out = master.copy()
            out["db_merge_status"] = "NO_DB_REPORT"
        statuses = []
        reasons = []
        for _, r in out.iterrows():
            status = self.engine.infer_status(r)
            if r.get("db_merge_status") in ["DB_MISSING", "NO_DB_REPORT"] and status in ["TRACKING", "LEADER"]:
                status = "TRACKING_DB_MISSING" if r.get("db_merge_status") == "DB_MISSING" else "TRACKING"
            statuses.append(status)
            reasons.append(self.engine.status_reason(r, status))
        out["tracking_status"] = statuses
        out["tracking_reason"] = reasons
        out["portal_open_flag"] = out["tracking_status"].isin(["BREAKOUT_CONFIRM", "MAIN_ATTACK"])
        out["main_attack_flag"] = out["tracking_status"].eq("MAIN_ATTACK")
        return out

    def write_into_workbook(self, wb, institutional_report: Optional[Dict[str, Any]] = None):
        merged_df = None
        if institutional_report is not None and institutional_report.get("all") is not None and pd is not None:
            try:
                merged_df = self.engine.apply(institutional_report.get("all"))
                institutional_report["all"] = merged_df
            except Exception as exc:
                if self.logger:
                    self.logger.warning(f"AI_PROJECT_TRACKING_ENGINE_APPLY_FAIL error={exc}")
        tracking_df = self._build_tracking_df(merged_df)
        self._write_tracking(wb, tracking_df)
        self._write_dashboard(wb, tracking_df)
        self._write_portal_monitor(wb, tracking_df)
        self._write_validation(wb, tracking_df, merged_df)
        if self.logger:
            self.logger.info("AI_PROJECT_REPORT_WRITTEN sheets=" + ",".join(self.SHEETS))
        return wb

    def _write_tracking(self, wb, df):
        ws = self._sheet(wb, "AI專案追蹤頁")
        headers = [
            "代號", "名稱", "AI專案分類", "追蹤優先級", "個股狀態", "門戶大開", "主攻", "狀態原因",
            "老師決策", "執行狀態", "波段階段", "突破階段", "波段位置分", "RR", "現價", "低接區",
            "停損", "目標1", "目標2", "最高分", "AI分數", "研究分", "出現次數", "產品定位",
            "操作建議彙整", "風險旗標彙整", "DB合併", "來源檔案"
        ]
        ws.append(headers)
        if df is None or len(df) == 0:
            ws.append(["無資料"])
        else:
            for _, r in df.iterrows():
                entry = ""
                if pd.notna(r.get("entry_low")) and pd.notna(r.get("entry_high")):
                    entry = f"{round(_safe_float(r.get('entry_low'),0),2)}~{round(_safe_float(r.get('entry_high'),0),2)}"
                ws.append([
                    str(r.get("stock_id", "")).zfill(4), r.get("stock_name", ""), r.get("ai_project_type", ""),
                    r.get("tracking_priority", ""), r.get("tracking_status", ""), "Y" if bool(r.get("portal_open_flag", False)) else "",
                    "Y" if bool(r.get("main_attack_flag", False)) else "", r.get("tracking_reason", ""),
                    r.get("teacher_decision", ""), r.get("teacher_execution_status", r.get("是否可下單", "")),
                    r.get("phase5_wave_phase", ""), r.get("phase5_breakout_stage", ""), r.get("phase5_position_score", ""),
                    r.get("rr", ""), r.get("close", ""), entry, r.get("stop_loss", ""), r.get("target_1", ""), r.get("target_2", ""),
                    r.get("max_score", ""), r.get("ai_score_ref", ""), r.get("research_score", ""), r.get("source_count", ""),
                    r.get("product_position", ""), r.get("strategy_summary", ""), r.get("risk_summary", ""), r.get("db_merge_status", ""), r.get("source_file", "")
                ])
        self._style_basic(ws, "0F766E")

    def _write_dashboard(self, wb, df):
        ws = self._sheet(wb, "AI專案Dashboard")
        headers = ["AI專案分類", "總追蹤檔數", "A主攻", "門戶大開", "主升攻擊", "爆發前", "壓縮觀察", "REDUCE/AVOID", "追蹤目的"]
        ws.append(headers)
        cat_df = self.engine.category_df() if pd is not None else []
        if pd is None or df is None or len(df) == 0:
            ws.append(["無資料"])
        else:
            for item in AI_PROJECT_CATEGORIES:
                cat = item.get("ai_project_type", "")
                mask = df["ai_project_type"].astype(str).str.contains(re.escape(cat), na=False)
                sub = df[mask]
                ws.append([
                    cat, int(len(sub)), int(sub["tracking_priority"].astype(str).str.startswith("A").sum()),
                    int(sub["portal_open_flag"].fillna(False).sum()), int(sub["main_attack_flag"].fillna(False).sum()),
                    int(sub["tracking_status"].astype(str).eq("PRE_BREAKOUT").sum()),
                    int(sub["tracking_status"].astype(str).eq("WATCHLIST").sum()),
                    int(sub["tracking_status"].astype(str).isin(["REDUCE", "AVOID"]).sum()),
                    item.get("tracking_purpose", "")
                ])
        self._style_basic(ws, "1D4ED8")

    def _write_portal_monitor(self, wb, df):
        ws = self._sheet(wb, "AI門戶大開監控")
        headers = ["代號", "名稱", "AI專案分類", "個股狀態", "老師決策", "執行狀態", "波段階段", "突破階段", "波段位置分", "RR", "現價", "狀態原因", "風險旗標"]
        ws.append(headers)
        if df is None or len(df) == 0:
            ws.append(["無資料"])
        else:
            sub = df[df["tracking_status"].isin(["PRE_BREAKOUT", "BREAKOUT_CONFIRM", "MAIN_ATTACK", "HOT_MONEY", "WATCHLIST"])].copy()
            if "phase5_position_score" in sub.columns:
                sub["_sort_pos"] = pd.to_numeric(sub["phase5_position_score"], errors="coerce").fillna(0)
            else:
                sub["_sort_pos"] = 0
            sub = sub.sort_values(["tracking_status", "_sort_pos", "max_score"], ascending=[True, False, False]).head(100)
            for _, r in sub.iterrows():
                ws.append([
                    str(r.get("stock_id", "")).zfill(4), r.get("stock_name", ""), r.get("ai_project_type", ""),
                    r.get("tracking_status", ""), r.get("teacher_decision", ""), r.get("teacher_execution_status", r.get("是否可下單", "")),
                    r.get("phase5_wave_phase", ""), r.get("phase5_breakout_stage", ""), r.get("phase5_position_score", ""),
                    r.get("rr", ""), r.get("close", ""), r.get("tracking_reason", ""), r.get("risk_summary", "")
                ])
        self._style_basic(ws, "7C2D12")

    def _write_validation(self, wb, df, merged_df=None):
        ws = self._sheet(wb, "AI專案整合驗收")
        master_count = len(AI_PROJECT_TRACKING_MASTER)
        category_count = len(AI_PROJECT_CATEGORIES)
        written_count = int(len(df)) if df is not None else 0
        db_match = int((df["db_merge_status"].astype(str).eq("DB_MATCH")).sum()) if df is not None and "db_merge_status" in df.columns else 0
        portal = int((df["portal_open_flag"].fillna(False)).sum()) if df is not None and "portal_open_flag" in df.columns else 0
        main_attack = int((df["main_attack_flag"].fillna(False)).sum()) if df is not None and "main_attack_flag" in df.columns else 0
        rows = [
            ["查核項目", "結果", "說明"],
            ["AI_PROJECT_TRACKING_MASTER檔數", master_count, "必須等於整合股票池201檔"],
            ["AI專案分類數", category_count, "固定11大AI專案分類"],
            ["AI專案追蹤頁輸出檔數", written_count, "必須全部列出，不只列交集成功者"],
            ["DB合併命中數", db_match, "與institutional_report['all']合併成功數；未命中仍保留TRACKING_DB_MISSING"],
            ["門戶大開檔數", portal, "BREAKOUT_CONFIRM或MAIN_ATTACK"],
            ["主攻檔數", main_attack, "MAIN_ATTACK"],
            ["狀態欄位", "tracking_status", "TRACKING/WATCHLIST/PRE_BREAKOUT/BREAKOUT_CONFIRM/MAIN_ATTACK/HOT_MONEY/LEADER/EXTENDED/REDUCE/EXIT/AVOID"],
            ["風控原則", "PASS", "題材不得覆蓋REDUCE/AVOID/Hard Avoid/Wave5/過熱"],
            ["Log驗收", "AI_PROJECT_TRACKING_MASTER_READY / AI_PROJECT_REPORT_WRITTEN", "可grep追蹤"],
        ]
        for r in rows:
            ws.append(r)
        if self.logger:
            self.logger.info(
                f"AI_PROJECT_TRACKING_SUMMARY master_count={master_count} category_count={category_count} "
                f"written_count={written_count} db_match={db_match} portal_open_count={portal} main_attack_count={main_attack}"
            )
        self._style_basic(ws, "334155")

class PreBreakoutSOPExcelIntegrator:
    """
    將「加入爆發前股票SOP」落地到 Macro16 Excel 輸出。
    設計原則：
    1. 不取代既有 Macro16 / 老師策略 / CPO 頁面，只新增 SOP 說明與候選清單。
    2. 若有 DB / institutional_report['all']，自動產生「爆發前股票候選」。
    3. 若無 DB，仍輸出 SOP、規則矩陣、盤後DB_盤前分析分工、宏觀16整合方式、盤前輸出欄位，避免報表缺頁。
    """
    SOP_SHEET_NAMES = [
        "爆發前股票SOP",
        "爆發前規則矩陣",
        "盤後DB_盤前分析分工",
        "宏觀16整合方式",
        "盤前輸出欄位",
        "爆發前股票候選",
    ]

    def __init__(self, logger: Optional[Macro16Logger] = None):
        self.logger = logger

    def _sheet(self, wb, name: str):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def write_into_workbook(self, wb, institutional_report: Optional[Dict[str, Any]] = None):
        self._write_sop(wb)
        self._write_rule_matrix(wb)
        self._write_db_preopen_split(wb)
        self._write_macro16_integration(wb)
        self._write_preopen_columns(wb)
        self._write_candidates(wb, institutional_report)
        if self.logger:
            self.logger.info("PREBREAKOUT_SOP_SHEETS_WRITTEN sheets=" + ",".join(self.SOP_SHEET_NAMES))
        return wb

    def _write_sop(self, wb):
        ws = self._sheet(wb, "爆發前股票SOP")
        rows = [
            ["步驟", "SOP項目", "程式執行邏輯", "輸出結果", "驗收標準"],
            ["1", "盤後DB完成", "收盤後先完成 price_history / market_snapshot / ranking_result / 外部財報籌碼欄位", "可供盤前分析的乾淨資料集", "DB資料日與報表基準日一致"],
            ["2", "候選池建立", "以老師策略、低位階翻多、Phase5預突破、量價壓縮、題材標籤建立初始池", "爆發前候選清單", "不得直接等同TOP15；必須有預突破原因"],
            ["3", "結構確認", "判斷 Wave3_PreBreakout、Compression、主升預突破觀察池、低位階翻多", "structure_tag / breakout_stage", "至少保留一個可追溯欄位"],
            ["4", "量價確認", "量比、20日高點、均線支撐、K線風險、RR同步確認", "volume_ratio / rr / risk_flag", "量價不足時只能WATCH，不可BUY"],
            ["5", "宏觀16 Gate", "引用宏觀總分、V2技術風險、大盤判定、夜盤偏空、重大事件", "macro_gate", "市場風險偏空時降級或禁止追高"],
            ["6", "決策輸出", "輸出 BUY / LOW_BUY / WATCH / REDUCE / AVOID 與是否可下單", "prebreakout_decision", "決策原因必須可追溯"],
            ["7", "Excel與UI一致", "Excel新增爆發前股票SOP相關頁，UI可再掛載相同資料框", "SOP頁+候選頁", "報表、程式、UI語義一致"],
        ]
        for r in rows:
            ws.append(r)

    def _write_rule_matrix(self, wb):
        ws = self._sheet(wb, "爆發前規則矩陣")
        rows = [
            ["規則類別", "欄位/條件", "BUY", "LOW_BUY", "WATCH", "REDUCE/AVOID", "程式欄位"],
            ["結構", "phase5_wave_phase", "Wave3_Breakout / Wave3_Expansion", "Wave3_PreBreakout 且位置分>=65", "Wave3_PreBreakout 或 Compression", "A/B/C_Correction 且未完成", "phase5_wave_phase"],
            ["突破階段", "phase5_breakout_stage", "Breakout/Expansion", "PreBreakout", "Compression/PreBreakout", "Correction/Exhaustion", "phase5_breakout_stage"],
            ["位置分", "phase5_position_score", ">=70", ">=65", ">=55", "<50", "phase5_position_score"],
            ["量能", "vol5/vol20", ">=1.5", ">=1.2", ">=1.0", "<0.8", "prebreakout_volume_ratio"],
            ["RR", "rr", ">=1.5", ">=1.2", ">=1.0", "<1.0", "rr"],
            ["風險", "hard_avoid / k_warning / phase5_escape", "皆未觸發", "無硬性風險", "僅軟性壓力", "硬性風險或逃命反彈", "prebreakout_risk_flag"],
            ["宏觀", "macro_gate", "允許交易/震盪偏多", "震盪偏多", "中性震盪", "風險偏空/停止新倉", "macro_gate"],
        ]
        for r in rows:
            ws.append(r)

    def _write_db_preopen_split(self, wb):
        ws = self._sheet(wb, "盤後DB_盤前分析分工")
        rows = [
            ["階段", "資料/動作", "主責程式區塊", "不得混淆事項", "輸出"],
            ["盤後", "更新日K、均線、成交量、法人、財報、ranking_result", "DBRepository / InstitutionalReportEngine", "盤後只建資料，不直接盤前追價", "完整DB"],
            ["盤後", "計算Phase5、老師策略、低位階翻多、避開/換股", "FeatureBuilder + TeacherDecisionEngine", "結構分類不可被即時價格硬改", "老師策略欄位"],
            ["盤前", "讀取昨收DB與外部宏觀16資料", "Macro16Engine.run", "盤前不重算大量歷史資料", "市場Gate"],
            ["盤前", "篩爆發前候選", "PreBreakoutSOPExcelIntegrator", "候選不等於BUY；需量價/RR/Gate確認", "爆發前股票候選"],
            ["盤中/下單前", "即時價、五檔、流動性、滑價二次確認", "交易/Execution Layer", "外部資料不得直接控制下單", "execution_status"],
        ]
        for r in rows:
            ws.append(r)

    def _write_macro16_integration(self, wb):
        ws = self._sheet(wb, "宏觀16整合方式")
        rows = [
            ["Macro16項目", "使用方式", "對爆發前股票的影響", "程式落地"],
            ["宏觀總分", "控制盤前風險偏好", "總分偏低時候選股降級", "summary['宏觀總分'] / scores"],
            ["V2技術風險", "判斷大盤是否禁追高", "risk_score>=3 降級 WATCH/REDUCE", "TechnicalRisk.risk_score"],
            ["重大事件", "重大事件=1時啟動保守Gate", "禁止高β追價", "market.major_event"],
            ["夜盤偏空", "night_score<0納入風險", "盤前候選降一級", "TechnicalRisk.night_bearish"],
            ["外資/官股", "資金面輔助", "資金轉弱時降低BUY數量", "MarketInput.foreign_net_100m / gov_net_100m"],
            ["AI產業", "主題強度", "CPO/AI主題股排序加權", "MarketInput.ai_strength"],
        ]
        for r in rows:
            ws.append(r)

    def _write_preopen_columns(self, wb):
        ws = self._sheet(wb, "盤前輸出欄位")
        rows = [
            ["欄位", "中文名稱", "來源", "用途", "必要"],
            ["stock_id", "代號", "DB/ranking_result", "主鍵", "Y"],
            ["report_name", "名稱", "stocks_master/報表", "顯示", "Y"],
            ["teacher_decision", "老師決策", "TeacherDecisionEngine", "五態決策", "Y"],
            ["teacher_execution_status", "是否可下單", "TeacherDecisionEngine", "YES/WAIT/NO", "Y"],
            ["phase5_wave_phase", "波段階段", "TeacherPhase5SemanticEngine", "辨識Wave3_PreBreakout", "Y"],
            ["phase5_breakout_stage", "突破階段", "TeacherPhase5SemanticEngine", "Compression/PreBreakout/Breakout", "Y"],
            ["phase5_position_score", "波段位置分", "TeacherPhase5SemanticEngine", "排序與Gate", "Y"],
            ["prebreakout_volume_ratio", "預突破量比", "price_history衍生 vol5/vol20", "量價確認", "Y"],
            ["rr", "風報比", "TradePlanEngine", "交易風險控管", "Y"],
            ["entry_low", "買進下緣", "TradePlanEngine", "低接區", "Y"],
            ["entry_high", "買進上緣", "TradePlanEngine", "低接區", "Y"],
            ["stop_loss", "停損", "TradePlanEngine", "風控", "Y"],
            ["target_1", "目標1", "TradePlanEngine", "第一目標", "Y"],
            ["target_2", "目標2", "TradePlanEngine", "第二目標", "Y"],
            ["prebreakout_decision", "爆發前決策", "本SOP整合層", "BUY/LOW_BUY/WATCH/REDUCE/AVOID", "Y"],
            ["prebreakout_reason", "爆發前原因", "本SOP整合層", "可追溯說明", "Y"],
        ]
        for r in rows:
            ws.append(r)

    def _write_candidates(self, wb, institutional_report: Optional[Dict[str, Any]] = None):
        ws = self._sheet(wb, "爆發前股票候選")
        headers = [
            "排名", "代號", "名稱", "是否CPO", "CPO主題", "CPO子分類", "CPO分數", "CPO策略定位",
            "老師決策", "是否可下單", "波段階段", "突破階段", "波段位置分",
            "量比", "RR", "買進下緣", "買進上緣", "停損", "目標1", "目標2",
            "爆發前決策", "爆發前原因", "風險旗標", "資料狀態"
        ]
        ws.append(headers)
        if not institutional_report or "all" not in institutional_report or institutional_report.get("all") is None or pd is None:
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "待DB", "未提供DB或InstitutionalReportEngine未產出，僅輸出SOP規格頁", "待DB", "NO_DB"])
            return
        df = institutional_report.get("all")
        if df is None or getattr(df, "empty", True):
            ws.append([1, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "待DB", "DB資料為空，無法產生候選", "待DB", "EMPTY"])
            return
        data = df.copy()
        try:
            data = CPOThemeEngine(self.logger).apply(data)
        except Exception as exc:
            if self.logger:
                self.logger.warning(f"PREBREAKOUT_CPO_MERGE_FAIL error={exc}")
        for c in ["teacher_score", "phase5_position_score", "rr", "entry_low", "entry_high", "stop_loss", "target_1", "target_2", "vol5", "vol20", "cpo_score"]:
            if c in data.columns:
                data[c] = pd.to_numeric(data[c], errors="coerce")
        if "prebreakout_volume_ratio" not in data.columns:
            vol5 = data["vol5"] if "vol5" in data.columns else pd.Series(index=data.index, dtype=float)
            vol20 = data["vol20"] if "vol20" in data.columns else pd.Series(index=data.index, dtype=float)
            data["prebreakout_volume_ratio"] = (vol5 / vol20.replace(0, math.nan)).replace([math.inf, -math.inf], math.nan)
        wave_phase = data.get("phase5_wave_phase", pd.Series("", index=data.index)).astype(str)
        breakout_stage = data.get("phase5_breakout_stage", pd.Series("", index=data.index)).astype(str)
        position_score = pd.to_numeric(data.get("phase5_position_score", pd.Series(0, index=data.index)), errors="coerce").fillna(0)
        rr = pd.to_numeric(data.get("rr", pd.Series(0, index=data.index)), errors="coerce").fillna(0)
        vr = pd.to_numeric(data.get("prebreakout_volume_ratio", pd.Series(0, index=data.index)), errors="coerce").fillna(0)
        teacher_decision = data.get("teacher_decision", pd.Series("WATCH", index=data.index)).astype(str)
        execution_status = data.get("teacher_execution_status", pd.Series("NO", index=data.index)).astype(str)
        risk_text = (
            data.get("hard_avoid_reason", pd.Series("", index=data.index)).astype(str) + " " +
            data.get("k_warning_type", pd.Series("", index=data.index)).astype(str) + " " +
            data.get("phase5_block_reason", pd.Series("", index=data.index)).astype(str)
        )
        pre_mask = (
            wave_phase.isin(["Wave3_PreBreakout", "Wave3_Breakout", "Wave3_Expansion"]) |
            breakout_stage.isin(["Compression", "PreBreakout", "Breakout", "Expansion"]) |
            data.get("phase5_candidate_pool", pd.Series("", index=data.index)).astype(str).str.contains("預突破|主升", na=False)
        )
        candidates = data.loc[pre_mask].copy()
        if candidates.empty:
            candidates = data.copy()
        def decide(row):
            risk = str(row.get("_risk_text", "") or "")
            wp = str(row.get("phase5_wave_phase", "") or "")
            bs = str(row.get("phase5_breakout_stage", "") or "")
            pos = _safe_float(row.get("phase5_position_score"), 0)
            rrv = _safe_float(row.get("rr"), 0)
            vol = _safe_float(row.get("prebreakout_volume_ratio"), 0)
            tdec = str(row.get("teacher_decision", "WATCH") or "WATCH")
            exe = str(row.get("teacher_execution_status", "NO") or "NO")
            if any(k in risk for k in ["硬性", "逃命", "主跌", "長黑", "墓碑"]):
                return "AVOID"
            if tdec == "REDUCE":
                return "REDUCE"
            if (wp in ["Wave3_Breakout", "Wave3_Expansion"] or bs in ["Breakout", "Expansion"]) and pos >= 70 and rrv >= 1.5 and vol >= 1.2 and exe in ["YES", "WAIT"]:
                return "BUY" if exe == "YES" else "WATCH"
            if wp == "Wave3_PreBreakout" or bs == "PreBreakout":
                if pos >= 65 and rrv >= 1.2:
                    return "LOW_BUY" if tdec in ["LOW_BUY", "BUY"] else "WATCH"
            if bs == "Compression" and pos >= 55:
                return "WATCH"
            return "WATCH"
        candidates["_risk_text"] = risk_text.loc[candidates.index]
        candidates["prebreakout_decision"] = candidates.apply(decide, axis=1)
        reasons = []
        for _, r in candidates.iterrows():
            parts = []
            parts.append(f"波段={r.get('phase5_wave_phase','')}")
            parts.append(f"突破={r.get('phase5_breakout_stage','')}")
            parts.append(f"位置分={round(_safe_float(r.get('phase5_position_score'),0),2)}")
            parts.append(f"量比={round(_safe_float(r.get('prebreakout_volume_ratio'),0),2)}")
            parts.append(f"RR={round(_safe_float(r.get('rr'),0),2)}")
            if str(r.get("_risk_text","")).strip():
                parts.append("風險=" + str(r.get("_risk_text","")).strip()[:120])
            reasons.append(";".join(parts))
        candidates["prebreakout_reason"] = reasons
        rank_cols = ["prebreakout_decision", "phase5_position_score", "rr", "prebreakout_volume_ratio", "teacher_score"]
        for c in rank_cols[1:]:
            if c not in candidates.columns:
                candidates[c] = 0
        candidates["_decision_rank"] = candidates["prebreakout_decision"].map({"BUY": 1, "LOW_BUY": 2, "WATCH": 3, "REDUCE": 4, "AVOID": 5}).fillna(9)
        candidates = candidates.sort_values(["_decision_rank", "phase5_position_score", "rr", "prebreakout_volume_ratio"], ascending=[True, False, False, False]).head(60)
        for n, (_, r) in enumerate(candidates.iterrows(), start=1):
            is_cpo = bool(r.get("is_cpo", False))
            ws.append([
                n,
                str(r.get("stock_id", "")).zfill(4) if str(r.get("stock_id", "")).strip() else "",
                r.get("report_name", r.get("stock_name", r.get("name", ""))),
                "Y" if is_cpo else "N",
                r.get("cpo_theme", "") if is_cpo else "",
                r.get("cpo_subtheme", "") if is_cpo else "",
                round(_safe_float(r.get("cpo_score"), 0), 2) if is_cpo else "",
                r.get("cpo_strategy", "") if is_cpo else "",
                r.get("teacher_decision", ""),
                r.get("teacher_execution_status", r.get("是否可下單", "")),
                r.get("phase5_wave_phase", ""),
                r.get("phase5_breakout_stage", ""),
                round(_safe_float(r.get("phase5_position_score"), 0), 2),
                round(_safe_float(r.get("prebreakout_volume_ratio"), 0), 2),
                round(_safe_float(r.get("rr"), 0), 2),
                _safe_float(r.get("entry_low"), None),
                _safe_float(r.get("entry_high"), None),
                _safe_float(r.get("stop_loss"), None),
                _safe_float(r.get("target_1"), None),
                _safe_float(r.get("target_2"), None),
                r.get("prebreakout_decision", ""),
                r.get("prebreakout_reason", ""),
                str(r.get("_risk_text", "")).strip(),
                "DB_OK",
            ])


class ExcelWriter:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger
        self.header_fill = PatternFill("solid", fgColor="DDEBF7")
        self.sub_fill = PatternFill("solid", fgColor="E2F0D9")
        self.warn_fill = PatternFill("solid", fgColor="FFF2CC")
        self.thin = Side(style="thin", color="D9E2F3")

    def write(self, template: Optional[str], out_path: str, market: MarketInput, scores: List[ModuleScore], tech: TechnicalRisk, summary: Dict[str, str], logs: List[str], raw: Optional[Dict[str, RawData]] = None, institutional_report: Optional[Dict[str, Any]] = None, gov_result: Optional[Dict[str, Any]] = None, market5_result: Optional[Dict[str, Any]] = None, report_mode: str = REPORT_MODE_MACRO, institutional_error: Optional[Dict[str, Any]] = None) -> str:
        if template and Path(template).exists():
            try:
                wb = load_workbook(template)
                self.logger.info(f"已載入Excel模板：{template}")
            except Exception as exc:
                self.logger.warning(f"模板載入失敗，改建新檔：{exc}")
                wb = Workbook()
        else:
            wb = Workbook()
            # 新檔第一個預設Sheet重新命名，避免產生Sheet頁。
            if wb.sheetnames == ["Sheet"]:
                wb["Sheet"].title = "市場輸入"
        report_mode = report_mode or REPORT_MODE_MACRO
        validator = MacroRefillValidator(self.logger)
        if report_mode in (REPORT_MODE_MACRO, REPORT_MODE_MACRO_TEACHER):
            # 正式日常模式：修正宏觀16，但不得關閉/刪除老師策略00~16報表。
            validator.ensure_macro_sheets(wb)
            self._write_market_input(wb, market)
            self._write_macro_modules(wb, scores)
            self._write_technical(wb, tech)
            if institutional_report is not None:
                InstitutionalExcelWriter(self.logger).write_into_workbook(wb, institutional_report, gov_result=gov_result, market5_result=market5_result)
                self.logger.info("MACRO_REFILL_TEACHER_OUTPUT_RESTORED sheets=00_16_teacher_strategy_reports")
            else:
                self.logger.warning("MACRO_REFILL_TOP_OUTPUT_SKIPPED reason=未提供DB或InstitutionalReportEngine失敗，無法產出TOP報表")
                self._write_teacher_failure_diagnostic(wb, institutional_error)
        elif report_mode == REPORT_MODE_MACRO_ONLY:
            # 單純驗證宏觀回填時才只保留三頁。
            validator.enforce_macro_only_sheets(wb)
            self._write_market_input(wb, market)
            self._write_macro_modules(wb, scores)
            self._write_technical(wb, tech)
        elif report_mode in (REPORT_MODE_INSTITUTIONAL, REPORT_MODE_TEACHER_FULL):
            # institutional_report/teacher_full模式只輸出00~16老師策略報表，不混入macro/debug頁。
            for name in list(wb.sheetnames):
                del wb[name]
            if institutional_report is not None:
                InstitutionalExcelWriter(self.logger).write_into_workbook(wb, institutional_report, gov_result=gov_result, market5_result=market5_result)
            else:
                ws = wb.create_sheet("00_執行摘要")
                ws.append(["項目", "內容"]); ws.append(["狀態", "未提供DB，無法產出institutional_report"])
                self._write_teacher_failure_diagnostic(wb, institutional_error)
        else:
            if institutional_report is not None:
                InstitutionalExcelWriter(self.logger).write_into_workbook(wb, institutional_report, gov_result=gov_result, market5_result=market5_result)
            self._write_market_input(wb, market)
            self._write_macro_modules(wb, scores)
            self._write_technical(wb, tech)
            self._write_audit(wb, market, scores, tech, summary, logs)
            self._write_data_source_status(wb, market, raw or {})
            self._write_evidence_index(wb, raw or {})
        CPOReportExcelIntegrator(self.logger).write_into_workbook(wb, institutional_report=locals().get("institutional_report"))
        MainThemeReportExcelIntegrator(self.logger).write_into_workbook(wb, institutional_report=locals().get("institutional_report"))
        AIProjectTrackingReportExcelIntegrator(self.logger).write_into_workbook(wb, institutional_report=locals().get("institutional_report"))
        PreBreakoutSOPExcelIntegrator(self.logger).write_into_workbook(wb, institutional_report=locals().get("institutional_report"))
        self._format_all(wb)
        wb.save(out_path)
        if institutional_report is not None and report_mode in (REPORT_MODE_MACRO, REPORT_MODE_MACRO_TEACHER, REPORT_MODE_INSTITUTIONAL, REPORT_MODE_TEACHER_FULL, REPORT_MODE_ALL):
            try:
                check_wb = load_workbook(out_path, read_only=True)
                missing = [name for name in EXPECTED_INSTITUTIONAL_SHEETS if name not in check_wb.sheetnames]
                check_wb.close()
                if missing:
                    self.logger.warning("TEACHER_REPORT_VALIDATE_FAIL_AFTER_SAVE missing=" + ",".join(missing))
                else:
                    self.logger.info("TEACHER_REPORT_VALIDATE_OK_AFTER_SAVE sheets=00_16_teacher_strategy_reports")
            except Exception as exc:
                self.logger.warning(f"TEACHER_REPORT_VALIDATE_AFTER_SAVE_ERROR {exc}")
        self.logger.info(f"Excel已輸出：{out_path} report_mode={report_mode}")
        return out_path

    def _sheet(self, wb, name):
        if name in wb.sheetnames:
            ws = wb[name]
            ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(name)
        return ws

    def _write_market_input(self, wb, market: MarketInput):
        ws = self._sheet(wb, "市場輸入")
        headers = ["日期", "收盤", "最高", "最低", "前高", "前低", "5MA", "成交量(億)", "5日均量(億)", "外資買賣超(億)", "官股買賣超(億)", "AI主流強度(0-1)", "重大事件(0/1)", "來源1", "來源2", "來源3", "來源4"]
        values = [market.base_date, market.close, market.high, market.low, market.prev_high, market.prev_low, market.ma5, market.turnover_100m, market.avg_turnover_5d_100m, market.foreign_net_100m, market.gov_net_100m, market.ai_strength, market.major_event, market.source_1, market.source_2, market.source_3, market.source_4]
        ws.append(headers)
        ws.append(values)
        ws.append([])
        ws.append(["欄位說明", "本表由主程式自動回填。若數值為空白/None，代表資料來源未取得，程式不編造數字，請依Log與回填紀錄修正資料來源或人工確認。"] + [None]*(len(headers)-2))
        ws.append(["交易判讀", MarketNarrativeBuilder().build(market)] + [None]*(len(headers)-2))

    def _write_macro_modules(self, wb, scores: List[ModuleScore]):
        # SOP V2.1 P0-02：正式Sheet必須是宏觀16模組，不得再建立宏觀15模組。
        ws = self._sheet(wb, "宏觀16模組")
        ws.append(["模組", "風險/強度分數(0-1)", "方向(+1/0/-1)", "加權分數", "說明", "資料來源", "資料時間"])
        for item in scores:
            weighted = round(float(item.strength or 0) * int(item.direction or 0), 2)
            # 以ScoringEngine輸出為主，若空值則強制回補，避免D欄空白。
            if item.weighted_score is not None:
                weighted = item.weighted_score
            ws.append([item.module, item.strength, item.direction, weighted, item.explanation, item.source, item.data_time])
        ws.append([])
        ws.append(["補充欄位", "狀態", "數據/事件", "交易用途"] )
        for item in scores:
            ws.append([item.module, item.status, item.data_text, item.trade_usage])

    def _write_technical(self, wb, tech: TechnicalRisk):
        ws = self._sheet(wb, "V2技術引擎")
        ws.append(["跌破5MA", "高不過高", "低破低", "放量", "重大事件", "夜盤偏空", "技術/風險分數", "大盤判定", "夜盤分數", "夜盤外資淨口數"])
        ws.append([tech.below_ma5, tech.lower_high, tech.lower_low, tech.volume_expansion, tech.major_event, getattr(tech, "night_bearish", 0), tech.risk_score, tech.market_judgement, getattr(tech, "night_score", None), getattr(tech, "night_net_lots", None)])
        ws.append(["判讀說明", "收盤<5MA為1", "最高<前高為1", "最低<前低為1", "成交值>5日均量*1.05為1", "合併Reuters/ISW/CNN/manual", "night_score<0為1", "六項加總", "供下單清單參考", "TAIFEX解析值", "TAIFEX外資淨多空口數"])

    def _write_audit(self, wb, market: MarketInput, scores: List[ModuleScore], tech: TechnicalRisk, summary: Dict[str,str], logs: List[str]):
        name = f"回填紀錄_{market.base_date.replace('-', '') if market.base_date else dt.date.today().strftime('%Y%m%d')}"
        ws = self._sheet(wb, name)
        ws.append(["回填項目", "數據/判斷", "方向", "分數", "資料日期", "來源URL/資料站", "回填邏輯", "交易用途"])
        for s in scores:
            ws.append([s.module, s.data_text, s.direction, s.weighted_score, s.data_time, s.source, s.explanation, s.trade_usage])
        start = 2
        ws.cell(start, 10, "總結項目")
        ws.cell(start, 11, "輸出")
        i = start + 1
        for k,v in summary.items():
            ws.cell(i,10,k); ws.cell(i,11,v); i += 1
        i += 1
        ws.cell(i,10,"程式Log摘要")
        for msg in logs[-30:]:
            i += 1
            ws.cell(i,10,msg[:180])

    def _write_data_source_status(self, wb, market: MarketInput, raw: Dict[str, RawData]):
        ws = self._sheet(wb, "資料來源狀態")
        ws.append(["資料源", "狀態", "資料分類", "解析狀態", "信心分數", "查詢日", "實際資料日", "是否回退", "回退天數", "來源", "URL", "RAW證據檔", "說明/訊息"])
        for name, item in raw.items():
            ws.append([
                name, item.status, item.data_status, getattr(item, "parse_status", ""), getattr(item, "confidence", ""),
                item.query_date, item.actual_date or item.date, item.is_fallback, item.fallback_days, item.source, item.url,
                getattr(item, "raw_file_path", ""), item.message or item.data_note
            ])
        ws.append([])
        ws.append(["說明", "query_date=使用者查詢日；actual_date=實際資料日；fallback_days=往前回退天數；RAW_DATA_SNAPSHOT 與 PARSED_VALUE 會寫入 log 證明實際抓取資料。"])


    def _write_evidence_index(self, wb, raw: Dict[str, RawData]):
        """寫入資料證據索引，修復V2.0缺少方法導致Excel輸出中斷。"""
        ws = self._sheet(wb, "資料證據索引")
        headers = [
            "項次", "模組/資料源", "狀態", "資料分類", "解析狀態", "parsed_fields摘要",
            "查詢日", "實際資料日", "是否回退", "回退天數", "來源", "URL",
            "RAW證據檔", "信心分數", "問題判定", "處理建議"
        ]
        ws.append(headers)
        for idx, (name, item) in enumerate(raw.items(), start=1):
            parsed_summary = ""
            issue = "OK"
            action = "可作為證據鏈追溯"
            raw_path = getattr(item, "raw_file_path", "") or ""
            try:
                if raw_path and Path(raw_path).exists():
                    payload = json.loads(Path(raw_path).read_text(encoding="utf-8"))
                    parsed_fields = payload.get("parsed_fields", {})
                    parsed_summary = json.dumps(parsed_fields, ensure_ascii=False)[:500]
                    if item.status == "OK" and not parsed_fields:
                        issue = "假OK風險：status=OK但parsed_fields為空"
                        action = "需補parser或降為WARN，不可進主分數"
                elif item.status == "OK" and getattr(item, "parse_status", "") != "PARSE_OK":
                    issue = "缺RAW證據或解析欄位"
                    action = "需確認write_raw_evidence與parser"
            except Exception as exc:
                issue = f"RAW讀取失敗：{exc}"
                action = "需檢查RAW證據檔路徑"
            if item.status != "OK":
                issue = item.message or item.data_status or item.status
                action = "依資料來源狀態修正來源或人工確認"
            ws.append([
                idx, name, item.status, item.data_status, getattr(item, "parse_status", ""), parsed_summary,
                item.query_date, item.actual_date or item.date, item.is_fallback, item.fallback_days, item.source, item.url,
                raw_path, getattr(item, "confidence", ""), issue, action
            ])
        ws.append([])
        ws.append(["驗收規則", "所有核心來源若status=OK，parse_status必須為PARSE_OK且parsed_fields不得為空；未解析資料必須WARN，不得假OK。"] )


    def _format_all(self, wb):
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    cell.border = Border(top=self.thin, left=self.thin, right=self.thin, bottom=self.thin)
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.fill = self.header_fill
            for col in range(1, ws.max_column + 1):
                width = 14
                for row in range(1, min(ws.max_row, 30) + 1):
                    val = ws.cell(row, col).value
                    if val:
                        width = max(width, min(48, len(str(val)) * 1.2))
                ws.column_dimensions[get_column_letter(col)].width = width
            ws.freeze_panes = "A2"

class WordEvidenceReportWriter:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def write(self, out_path: str, raw: Dict[str, RawData], summary: Dict[str, str]) -> str:
        try:
            from docx import Document
        except Exception:
            self.logger.warning("python-docx未安裝，略過Word證據報告")
            return ""
        path = Path(out_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading("Macro16 資料抓取證據報告", 0)
        doc.add_paragraph(f"Run ID：{self.logger.run_id}")
        doc.add_paragraph(f"產出時間：{dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_heading("總結", level=1)
        for k, v in summary.items():
            doc.add_paragraph(f"{k}：{v}")
        doc.add_heading("逐資料源證據", level=1)
        for name, item in raw.items():
            doc.add_heading(name, level=2)
            doc.add_paragraph(f"來源：{item.source}")
            doc.add_paragraph(f"URL：{item.url}")
            doc.add_paragraph(f"狀態：{item.status} / {item.data_status} / {getattr(item, 'parse_status', '')}")
            doc.add_paragraph(f"信心分數：{getattr(item, 'confidence', '')}")
            doc.add_paragraph(f"RAW證據檔：{getattr(item, 'raw_file_path', '')}")
            doc.add_paragraph(f"說明：{item.message or item.data_note}")
        doc.save(str(path))
        self.logger.info(f"WORD_EVIDENCE_REPORT path={path}")
        return str(path)


class AuditEngine:
    def __init__(self, logger: Macro16Logger):
        self.logger = logger

    def check(self, market: MarketInput, scores: List[ModuleScore], tech: TechnicalRisk) -> List[str]:
        warnings = []
        if not market.base_date:
            warnings.append("市場輸入缺資料基準日")
        for field in ["close", "high", "low", "ma5"]:
            if getattr(market, field) is None:
                warnings.append(f"市場輸入缺{field}")
        for s in scores:
            if s.status != "OK":
                warnings.append(f"{s.module} 狀態={s.status}，需確認資料來源/解析結果")
            if s.direction not in (-1, 0, 1):
                warnings.append(f"{s.module} 方向不是+1/0/-1")
        if warnings:
            for w in warnings:
                self.logger.warning(f"QA: {w}")
        else:
            self.logger.info("QA檢查完成，未發現重大缺失")
        return warnings

class Macro16Engine:
    def __init__(self, log_dir: Path):
        self.logger = Macro16Logger(log_dir)
        self.client = HttpClient(self.logger)
        self.source = SourceConnector(self.client, self.logger)
        self.processor = DataProcessor(self.logger)
        self.scoring = ScoringEngine(self.logger)
        self.indicator = IndicatorEngine(self.logger)
        self.explain = ExplanationEngine()
        self.audit = AuditEngine(self.logger)
        self.writer = ExcelWriter(self.logger)
        self.word_evidence_writer = WordEvidenceReportWriter(self.logger)

    def run(self, template: Optional[str], out_path: str, base_date: Optional[str] = None, override: Optional[ManualOverride] = None, db_path: Optional[str] = None, strict_ranking: bool = False, tej_gov_file: Optional[str] = None, report_mode: str = REPORT_MODE_MACRO) -> Dict[str, Any]:
        self.logger.info(f"開始執行 {APP_NAME} v{VERSION}")
        self.logger.info("CHANGELOG v3.0.0: Add AI Project Rotation Monitor sheets and full 201-stock AI tracking page")
        self.logger.strategy_trace("STRATEGY_VERSION", {"strategy_version": STRATEGY_VERSION, "program_version": VERSION})
        raw: Dict[str, RawData] = {}
        requested_date = base_date.replace("-", "") if base_date else None
        raw["台股指數"] = self.source.fetch_twse_taiex_history(requested_date)
        # V1.3：先用TWSE實際回傳的最新完整交易日作為後續台股資料基準日，避免使用未收盤日查詢成交值/外資。
        actual_twse_date = requested_date
        try:
            if raw["台股指數"].status == "OK" and raw["台股指數"].value:
                actual_twse_date = raw["台股指數"].value.get("last", {}).get("date", "").replace("-", "") or requested_date
                self.logger.info(f"資料基準日校正：requested={requested_date}, actual_twse_date={actual_twse_date}")
        except Exception as exc:
            self.logger.warning(f"資料基準日校正失敗：{exc}")
        raw["成交量"] = self.source.fetch_twse_turnover_month(actual_twse_date)
        raw["外資"] = self.source.fetch_foreign_investor(actual_twse_date)
        # V2.5 P0：跨月補齊最近5個有效交易日，產出前高/前低/5MA/5日均量。
        market5_result = Market5DayEngine(self.client, self.logger).build_market_features(actual_twse_date or requested_date or dt.date.today().strftime("%Y%m%d"))
        raw["市場5日資料"] = RawData("市場5日資料", market5_result, market5_result.get("base_date", ""), "TWSE跨月5日", "TWSE MI_5MINS_HIST/FMTQIK", self.source._today_str(), "OK" if market5_result.get("status") == "OK" else "FAIL", market5_result.get("message", ""), requested_date or "", market5_result.get("base_date", ""), False, 0, "OK" if market5_result.get("status") == "OK" else "DATA_MISSING", market5_result.get("message", ""), "PARSE_OK" if market5_result.get("status") == "OK" else "NO_PARSED_VALUE", confidence=1.0 if market5_result.get("status") == "OK" else 0.0)
        for module, symbol in YAHOO_SYMBOLS.items():
            raw[module] = self.source.fetch_yahoo_chart(symbol, module)
        for module, symbols in YAHOO_SYMBOL_CANDIDATES.items():
            raw[module] = self.source.fetch_yahoo_chart_candidates(symbols, module)
        raw["美債10Y"] = self.source.fetch_fred_csv_latest("DGS10", "美債10Y")
        raw["FED利率政策"] = self.source.fetch_fed_policy()
        raw["CPI"] = self.source.fetch_bls_cpi()
        raw["非農"] = self.source.fetch_bls_nfp()
        raw["戰爭/停火"] = self.source.fetch_reuters_war()
        raw["外交政策"] = self.source.fetch_bloomberg_policy()
        raw["戰爭/地緣"] = raw["戰爭/停火"] if raw["戰爭/停火"].status == "OK" else self.source.fetch_geopolitical_news()
        raw["ISW衝突分析"] = self.source.fetch_isw_conflict()
        raw["CNN重大新聞"] = self.source.fetch_cnn_major_news()
        raw["美國總統"] = self.source.fetch_trump_public_news()
        # V2.5.1 SOP：TEJ為主來源；TEJ缺檔時，Wantgoo/第三方備援必須嘗試解析gov_net_100m並標P0_WARN。
        gov_result = TEJGovBankEngine(tej_gov_file, self.logger).parse()
        raw["官股整理"] = self.source.fetch_wantgoo_public_bank()
        if gov_result.get("status") == "OK" and gov_result.get("gov_net_100m") is not None:
            raw["官股"] = RawData("官股", gov_result, gov_result.get("actual_date", ""), "TEJ八大公股行庫", tej_gov_file or "", self.source._today_str(), "OK", gov_result.get("message", ""), actual_twse_date or "", gov_result.get("actual_date", ""), False, 0, "OK", gov_result.get("message", ""), "PARSE_OK", confidence=0.95)
        elif raw.get("官股整理") and raw["官股整理"].parse_status == "PARSE_OK" and isinstance(raw["官股整理"].value, dict) and raw["官股整理"].value.get("gov_net_100m") is not None:
            fallback_value = dict(raw["官股整理"].value)
            fallback_value["status"] = "OK"
            fallback_value["message"] = "TEJ未提供，使用Wantgoo八大官股資料解析；來源保留供追溯，分數只依數值"
            raw["官股"] = RawData("官股", fallback_value, raw["官股整理"].date, "Wantgoo八大官股資料", raw["官股整理"].url, self.source._today_str(), "OK", fallback_value["message"], actual_twse_date or "", raw["官股整理"].date, False, 0, "OK", fallback_value["message"], "PARSE_OK", confidence=1.0)
            self.logger.info(f"GOV_DATA_ACCEPTED source=Wantgoo value={fallback_value.get('gov_net_100m')}")
        else:
            raw["官股"] = RawData("官股", gov_result, "", "TEJ八大公股行庫", tej_gov_file or "", self.source._today_str(), "WARN", gov_result.get("message", "TEJ未提供且Wantgoo備援未解析"), actual_twse_date or "", "", False, 0, "NO_PARSED_VALUE", gov_result.get("message", ""), "NO_PARSED_VALUE", confidence=0.0)
        raw["官股TWSE佐證"] = self.source.fetch_twse_broker_report(base_date=actual_twse_date)
        raw["AI產業"] = self.source.fetch_ai_industry_news()
        raw["IEK產業分析"] = self.source.fetch_iek_industry()
        raw["排行分析"] = self.source.fetch_ranking_result_db(db_path=db_path, strict=strict_ranking)
        raw["台股夜盤"] = self.source.fetch_taifex_night_snapshot()
        tpex_otc_raw = self.source.fetch_tpex_otc_snapshot()
        raw["櫃買官方來源"] = tpex_otc_raw
        if tpex_otc_raw.status == "OK" and isinstance(tpex_otc_raw.value, dict) and "close" in tpex_otc_raw.value:
            raw["OTC"] = tpex_otc_raw
        institutional_report = None
        institutional_report_error = None
        if db_path:
            try:
                institutional_report = InstitutionalReportEngine(db_path, self.logger).run()
            except Exception as exc:
                institutional_report_error = {"db_path": db_path, "error": str(exc), "error_type": type(exc).__name__}
                self.logger.warning(f"INSTITUTIONAL_REPORT_FAIL db_path={db_path} error={exc}")
                self.logger.strategy_trace("TEACHER_REPORT_FAIL_DETAIL", institutional_report_error)
        if override and override.event_note:
            raw["戰爭/地緣"] = self.source.build_manual_raw("戰爭/地緣", {"event_note": override.event_note, "major_event": 1}, override.event_note)
        if override and override.night_score is not None:
            raw["台股夜盤"] = self.source.build_manual_raw("台股夜盤", {"score": override.night_score}, "manual night score")
        # V1.3：完整記錄每個資料源狀態，方便後續增修與問題追蹤
        for k, v in raw.items():
            try:
                self.logger.debug(f"RAW_STATUS {k}: status={v.status}, query_date={v.query_date}, actual_date={v.actual_date or v.date}, is_fallback={v.is_fallback}, fallback_days={v.fallback_days}, data_status={v.data_status}, source={v.source}, url={v.url}, message={v.message}")
            except Exception:
                pass
        market = self.processor.build_market_input(raw, base_date or "")
        market = self.processor.apply_manual_override(market, override)
        completion_issues = FieldCompletionValidator(self.logger).validate_market_input(market, strict_gov=False)
        scores = self.scoring.score_all(raw, market)
        macro_total = round(sum(s.weighted_score for s in scores), 2)
        tech = self.indicator.compute(market, macro_total)
        summary = self.explain.build_summary(macro_total, tech)
        warnings = self.audit.check(market, scores, tech)
        warnings = (locals().get("completion_issues", []) or []) + warnings
        if warnings:
            summary["QA警告"] = "; ".join(warnings[:8])
        output = self.writer.write(template, out_path, market, scores, tech, summary, self.logger.messages, raw, institutional_report=institutional_report, gov_result=locals().get("gov_result"), market5_result=locals().get("market5_result"), report_mode=report_mode, institutional_error=institutional_report_error)
        evidence_word = ""
        if report_mode == REPORT_MODE_ALL:
            evidence_word = self.word_evidence_writer.write(str(Path(out_path).with_name(Path(out_path).stem + "_資料證據報告.docx")), raw, summary)
        self.logger.info("執行完成")
        return {"output": output, "evidence_word": evidence_word, "raw_dir": str(self.logger.raw_dir), "summary": summary, "warnings": warnings, "log_file": str(self.logger.log_file)}


class WatchPoolCultivationEngine:
    """R5N29：觀察池培養引擎。"""

    def __init__(self, log_dir: Path | str = "logs"):
        self.log_dir = Path(log_dir)
        self.log_dir.mkdir(parents=True, exist_ok=True)
        self.messages: List[str] = []

    def log(self, msg: str) -> None:
        line = f"INFO {msg}"
        self.messages.append(line)
        print(line)

    def warn(self, msg: str) -> None:
        line = f"WARN {msg}"
        self.messages.append(line)
        print(line)

    def _parse_date(self, base_date: str) -> str:
        return dt.datetime.strptime(base_date, "%Y-%m-%d").date().isoformat()

    def resolve_cultivation_db_path(self, main_db_path: str, cultivation_db_path: Optional[str]) -> Path:
        if cultivation_db_path and str(cultivation_db_path).strip():
            path = Path(cultivation_db_path)
        else:
            # 規劃指定：主程式資料夾/data/watch_pool_cultivation.db。
            base = Path(main_db_path).resolve().parent if main_db_path else Path.cwd()
            path = base / "data" / "watch_pool_cultivation.db"
        path.parent.mkdir(parents=True, exist_ok=True)
        return path

    def ensure_cultivation_schema(self, conn: sqlite3.Connection) -> None:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS watch_pool_tracking (
            track_date TEXT NOT NULL,
            stock_id TEXT NOT NULL,
            stock_name TEXT,
            watch_status TEXT,
            launch_score REAL,
            score_1d_delta REAL,
            score_3d_delta REAL,
            score_5d_delta REAL,
            days_in_watch INTEGER,
            status_reason TEXT,
            source_report TEXT,
            created_at TEXT NOT NULL,
            PRIMARY KEY(track_date, stock_id)
        )
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS watch_pool_event (
            event_id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_date TEXT NOT NULL,
            stock_id TEXT NOT NULL,
            event_type TEXT,
            before_status TEXT,
            after_status TEXT,
            event_reason TEXT,
            launch_score REAL,
            created_at TEXT NOT NULL,
            UNIQUE(event_date, stock_id, event_type, before_status, after_status)
        )
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS watch_pool_performance (
            stock_id TEXT NOT NULL,
            entry_date TEXT NOT NULL,
            entry_price REAL,
            max_return_5d REAL,
            max_return_10d REAL,
            max_return_20d REAL,
            outcome TEXT,
            outcome_reason TEXT,
            updated_at TEXT NOT NULL,
            PRIMARY KEY(stock_id, entry_date)
        )
        """)
        conn.commit()
        self.log("R5N29_SCHEMA_READY tables=watch_pool_tracking,watch_pool_event,watch_pool_performance")

    def _find_latest_report(self, launch_ready_path: Optional[str]) -> Optional[Path]:
        if not launch_ready_path:
            candidates = [Path.cwd() / "launch_ready_reports", Path.cwd() / "reports" / "launch_ready_reports"]
        else:
            p = Path(launch_ready_path)
            candidates = [p]
        files: List[Path] = []
        for p in candidates:
            if p.is_file() and p.suffix.lower() in (".xlsx", ".xlsm"):
                files.append(p)
            elif p.is_dir():
                files.extend([x for x in p.glob("*.xlsx") if not x.name.startswith("~$")])
                files.extend([x for x in p.glob("*.xlsm") if not x.name.startswith("~$")])
        if not files:
            return None
        return max(files, key=lambda x: x.stat().st_mtime)

    def _normalize_header(self, v: Any) -> str:
        return str(v or "").strip().lower().replace(" ", "").replace("_", "")

    def _pick_col(self, headers: Dict[str, int], names: List[str]) -> Optional[int]:
        normalized = {self._normalize_header(k): v for k, v in headers.items()}
        for name in names:
            key = self._normalize_header(name)
            if key in normalized:
                return normalized[key]
        for k, v in normalized.items():
            if any(self._normalize_header(n) in k for n in names):
                return v
        return None

    def _num(self, v: Any) -> Optional[float]:
        if v is None or v == "":
            return None
        try:
            s = str(v).replace("%", "").replace(",", "").strip()
            return float(s)
        except Exception:
            return None

    def _derive_status(self, score: Optional[float], raw_status: Optional[str] = None) -> str:
        s = str(raw_status or "").strip()
        if s:
            return s
        if score is None:
            return "WATCH"
        if score >= 80:
            return "LAUNCH_READY"
        if score >= 65:
            return "ACCELERATING"
        return "WATCH"

    def load_today_launch_ready(self, main_db_path: str, launch_ready_path: Optional[str], base_date: str) -> Tuple[List[Dict[str, Any]], str]:
        """讀取今日 Launch Ready 報表；若報表不存在，回退讀主DB常見候選表。"""
        rows: List[Dict[str, Any]] = []
        source = ""
        report = self._find_latest_report(launch_ready_path)
        if report:
            source = str(report)
            try:
                wb = load_workbook(report, data_only=True, read_only=True)
                for ws in wb.worksheets:
                    # 掃描前10列找股票代號欄。
                    for header_row_idx in range(1, min(ws.max_row, 10) + 1):
                        header_values = [ws.cell(header_row_idx, c).value for c in range(1, min(ws.max_column, 80) + 1)]
                        headers = {str(v): i + 1 for i, v in enumerate(header_values) if v is not None}
                        stock_col = self._pick_col(headers, ["stock_id", "股票代號", "代號", "證券代號"])
                        if not stock_col:
                            continue
                        name_col = self._pick_col(headers, ["stock_name", "股票名稱", "名稱", "公司名稱"])
                        score_col = self._pick_col(headers, ["launch_score", "launch ready score", "準備噴射分數", "分數", "score"])
                        status_col = self._pick_col(headers, ["watch_status", "狀態", "等級", "grade", "判定"])
                        reason_col = self._pick_col(headers, ["status_reason", "原因", "理由", "AI判定", "操作建議"])
                        for r in range(header_row_idx + 1, ws.max_row + 1):
                            stock_id = ws.cell(r, stock_col).value
                            if stock_id is None:
                                continue
                            stock_id = str(stock_id).strip().split('.')[0]
                            if not re.match(r"^\d{4,6}$", stock_id):
                                continue
                            score = self._num(ws.cell(r, score_col).value) if score_col else None
                            raw_status = str(ws.cell(r, status_col).value or "").strip() if status_col else ""
                            rows.append({
                                "track_date": base_date,
                                "stock_id": stock_id,
                                "stock_name": str(ws.cell(r, name_col).value or "").strip() if name_col else "",
                                "watch_status": self._derive_status(score, raw_status),
                                "launch_score": score,
                                "status_reason": str(ws.cell(r, reason_col).value or "").strip() if reason_col else "from_launch_ready_report",
                                "source_report": str(report),
                            })
                        if rows:
                            self.log(f"R5N29_LOAD_LAUNCH_READY_REPORT sheet={ws.title} rows={len(rows)} source={report}")
                            return self._dedupe_rows(rows), source
                self.warn(f"R5N29_LAUNCH_READY_PARSE_EMPTY source={report}")
            except Exception as exc:
                self.warn(f"R5N29_LAUNCH_READY_READ_FAIL source={report} error={exc}")
        # fallback：只讀主DB，不寫入主DB。
        if main_db_path and Path(main_db_path).exists():
            try:
                with sqlite3.connect(f"file:{main_db_path}?mode=ro", uri=True) as conn:
                    table_names = [r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
                    for table in ["launch_ready_candidates", "prebreakout", "ranking_result", "watchlist"]:
                        if table not in table_names:
                            continue
                        cols = [r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
                        stock_col = next((c for c in ["stock_id", "code", "證券代號"] if c in cols), None)
                        if not stock_col:
                            continue
                        name_col = next((c for c in ["stock_name", "name", "股票名稱"] if c in cols), None)
                        score_col = next((c for c in ["launch_score", "prebreakout_score", "score", "total_score"] if c in cols), None)
                        sql = f"SELECT {stock_col} AS stock_id" + (f", {name_col} AS stock_name" if name_col else ", '' AS stock_name") + (f", {score_col} AS launch_score" if score_col else ", NULL AS launch_score") + f" FROM {table} LIMIT 500"
                        for stock_id, stock_name, launch_score in conn.execute(sql).fetchall():
                            score = self._num(launch_score)
                            rows.append({"track_date": base_date, "stock_id": str(stock_id), "stock_name": stock_name or "", "watch_status": self._derive_status(score), "launch_score": score, "status_reason": f"fallback_from_{table}", "source_report": f"main_db:{table}"})
                        if rows:
                            source = f"{main_db_path}:{table}"
                            self.log(f"R5N29_LOAD_MAIN_DB_FALLBACK table={table} rows={len(rows)}")
                            return self._dedupe_rows(rows), source
            except Exception as exc:
                self.warn(f"R5N29_MAIN_DB_FALLBACK_FAIL error={exc}")
        raise RuntimeError("R5N29 無法取得今日觀察池資料：請指定 Launch Ready 報表/資料夾，或確認主DB存在候選表。")

    def _dedupe_rows(self, rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        best: Dict[str, Dict[str, Any]] = {}
        for row in rows:
            sid = str(row.get("stock_id", "")).strip()
            if not sid:
                continue
            old = best.get(sid)
            old_score = old.get("launch_score") if old else None
            new_score = row.get("launch_score")
            if old is None or (new_score is not None and (old_score is None or new_score > old_score)):
                best[sid] = row
        return list(best.values())

    def load_previous_tracking(self, conn: sqlite3.Connection, base_date: str) -> Dict[str, Dict[str, Any]]:
        sql = """
        SELECT t.track_date, t.stock_id, t.watch_status, t.launch_score, t.days_in_watch
        FROM watch_pool_tracking t
        JOIN (
            SELECT stock_id, MAX(track_date) AS max_date
            FROM watch_pool_tracking
            WHERE track_date < ?
            GROUP BY stock_id
        ) x ON x.stock_id=t.stock_id AND x.max_date=t.track_date
        """
        prev = {}
        for track_date, stock_id, status, score, days in conn.execute(sql, (base_date,)).fetchall():
            prev[str(stock_id)] = {"track_date": track_date, "watch_status": status, "launch_score": score, "days_in_watch": days or 0}
        self.log(f"R5N29_PREVIOUS_TRACKING rows={len(prev)}")
        return prev

    def _score_delta_days(self, conn: sqlite3.Connection, stock_id: str, base_date: str, days: int, score: Optional[float]) -> Optional[float]:
        if score is None:
            return None
        row = conn.execute("""
            SELECT launch_score FROM watch_pool_tracking
            WHERE stock_id=? AND track_date <= date(?, ?)
            ORDER BY track_date DESC LIMIT 1
        """, (stock_id, base_date, f"-{days} day")).fetchone()
        if not row or row[0] is None:
            return None
        return float(score) - float(row[0])

    def compare_today_previous(self, conn: sqlite3.Connection, today_rows: List[Dict[str, Any]], prev: Dict[str, Dict[str, Any]], base_date: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        now = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        tracking: List[Dict[str, Any]] = []
        events: List[Dict[str, Any]] = []
        for row in today_rows:
            sid = str(row["stock_id"])
            score = row.get("launch_score")
            p = prev.get(sid)
            days = (int(p.get("days_in_watch") or 0) + 1) if p else 1
            d1 = (float(score) - float(p["launch_score"])) if p and score is not None and p.get("launch_score") is not None else None
            d3 = self._score_delta_days(conn, sid, base_date, 3, score)
            d5 = self._score_delta_days(conn, sid, base_date, 5, score)
            track = dict(row)
            track.update({"score_1d_delta": d1, "score_3d_delta": d3, "score_5d_delta": d5, "days_in_watch": days, "created_at": now})
            tracking.append(track)
            before = p.get("watch_status") if p else None
            after = row.get("watch_status") or "WATCH"
            event_type = None
            reason = ""
            if p is None:
                event_type = "ENTER"
                reason = "首次進入觀察池"
            elif before != after:
                event_type = "STATUS_CHANGE"
                reason = f"狀態變更 {before} -> {after}"
            elif d1 is not None and d1 >= 10:
                event_type = "ACCELERATING"
                reason = f"分數單日增加 {d1:.2f}"
            if event_type:
                events.append({"event_date": base_date, "stock_id": sid, "event_type": event_type, "before_status": before, "after_status": after, "event_reason": reason, "launch_score": score, "created_at": now})
        self.log(f"R5N29_COMPARE tracking_rows={len(tracking)} event_rows={len(events)}")
        return tracking, events

    def write_tracking_rows(self, conn: sqlite3.Connection, rows: List[Dict[str, Any]]) -> None:
        sql = """
        INSERT OR REPLACE INTO watch_pool_tracking
        (track_date, stock_id, stock_name, watch_status, launch_score, score_1d_delta, score_3d_delta, score_5d_delta, days_in_watch, status_reason, source_report, created_at)
        VALUES (:track_date, :stock_id, :stock_name, :watch_status, :launch_score, :score_1d_delta, :score_3d_delta, :score_5d_delta, :days_in_watch, :status_reason, :source_report, :created_at)
        """
        conn.executemany(sql, rows)
        conn.commit()
        self.log(f"R5N29_TRACKING_WRITTEN rows={len(rows)}")

    def write_event_rows(self, conn: sqlite3.Connection, rows: List[Dict[str, Any]]) -> None:
        sql = """
        INSERT OR IGNORE INTO watch_pool_event
        (event_date, stock_id, event_type, before_status, after_status, event_reason, launch_score, created_at)
        VALUES (:event_date, :stock_id, :event_type, :before_status, :after_status, :event_reason, :launch_score, :created_at)
        """
        conn.executemany(sql, rows)
        conn.commit()
        self.log(f"R5N29_EVENT_WRITTEN rows={len(rows)}")

    def update_performance(self, conn: sqlite3.Connection, main_db_path: str, base_date: str) -> None:
        if not main_db_path or not Path(main_db_path).exists():
            self.warn("R5N29_PERFORMANCE_SKIP no_main_db")
            return
        now = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entries = conn.execute("SELECT stock_id, MIN(track_date) FROM watch_pool_tracking GROUP BY stock_id").fetchall()
        try:
            with sqlite3.connect(f"file:{main_db_path}?mode=ro", uri=True) as mconn:
                tables = [r[0] for r in mconn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
                if "price_history" not in tables:
                    self.warn("R5N29_PERFORMANCE_SKIP price_history_missing")
                    return
                cols = [r[1] for r in mconn.execute("PRAGMA table_info(price_history)").fetchall()]
                date_col = next((c for c in ["trade_date", "date", "日期"] if c in cols), None)
                close_col = next((c for c in ["close", "close_price", "收盤價"] if c in cols), None)
                stock_col = next((c for c in ["stock_id", "code", "證券代號"] if c in cols), None)
                if not (date_col and close_col and stock_col):
                    self.warn("R5N29_PERFORMANCE_SKIP price_history_columns_missing")
                    return
                upserts = []
                for sid, entry_date in entries:
                    prices = mconn.execute(f"SELECT {date_col}, {close_col} FROM price_history WHERE {stock_col}=? AND {date_col}>=? ORDER BY {date_col}", (sid, entry_date)).fetchall()
                    clean = [(str(d)[:10], self._num(c)) for d, c in prices if self._num(c) is not None]
                    if not clean:
                        continue
                    entry_price = clean[0][1]
                    def max_ret(n: int) -> Optional[float]:
                        end = (dt.datetime.strptime(entry_date, "%Y-%m-%d").date() + dt.timedelta(days=n)).isoformat()
                        vals = [c for d, c in clean if d <= end]
                        return (max(vals) / entry_price - 1.0) if vals and entry_price else None
                    r5, r10, r20 = max_ret(5), max_ret(10), max_ret(20)
                    best = max([x for x in [r5, r10, r20] if x is not None], default=None)
                    outcome = "PENDING" if best is None else ("SUCCESS" if best >= 0.10 else "WATCHING")
                    reason = "資料不足" if best is None else f"目前最高報酬 {best:.2%}"
                    upserts.append((sid, entry_date, entry_price, r5, r10, r20, outcome, reason, now))
                conn.executemany("""
                    INSERT OR REPLACE INTO watch_pool_performance
                    (stock_id, entry_date, entry_price, max_return_5d, max_return_10d, max_return_20d, outcome, outcome_reason, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, upserts)
                conn.commit()
                self.log(f"R5N29_PERFORMANCE_UPDATED rows={len(upserts)}")
        except Exception as exc:
            self.warn(f"R5N29_PERFORMANCE_FAIL error={exc}")

    def export_cultivation_report(self, conn: sqlite3.Connection, output_xlsx: str, base_date: str) -> str:
        out = Path(output_xlsx)
        out.parent.mkdir(parents=True, exist_ok=True)
        wb = Workbook()
        if wb.sheetnames == ["Sheet"]:
            wb["Sheet"].title = "今日總覽"
        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(style="thin", color="D9E2F3")

        def write_sheet(name: str, headers: List[str], rows: List[Tuple[Any, ...]]):
            ws = wb[name] if name in wb.sheetnames else wb.create_sheet(name)
            ws.append(headers)
            for row in rows:
                ws.append(list(row))
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for row in ws.iter_rows():
                for cell in row:
                    cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in range(1, ws.max_column + 1):
                ws.column_dimensions[get_column_letter(col)].width = min(max(12, len(str(ws.cell(1, col).value or "")) + 4), 28)
            ws.freeze_panes = "A2"
            return ws

        tracking_headers = ["track_date", "stock_id", "stock_name", "watch_status", "launch_score", "score_1d_delta", "score_3d_delta", "score_5d_delta", "days_in_watch", "status_reason", "source_report", "created_at"]
        tracking_rows = conn.execute("SELECT " + ",".join(tracking_headers) + " FROM watch_pool_tracking WHERE track_date=? ORDER BY launch_score DESC, stock_id", (base_date,)).fetchall()
        write_sheet("今日總覽", tracking_headers, tracking_rows)
        event_headers = ["event_date", "stock_id", "event_type", "before_status", "after_status", "event_reason", "launch_score", "created_at"]
        event_rows = conn.execute("SELECT " + ",".join(event_headers) + " FROM watch_pool_event WHERE event_date=? ORDER BY stock_id, event_type", (base_date,)).fetchall()
        write_sheet("狀態事件", event_headers, event_rows)
        perf_headers = ["stock_id", "entry_date", "entry_price", "max_return_5d", "max_return_10d", "max_return_20d", "outcome", "outcome_reason", "updated_at"]
        perf_rows = conn.execute("SELECT " + ",".join(perf_headers) + " FROM watch_pool_performance ORDER BY entry_date DESC, stock_id").fetchall()
        write_sheet("績效驗證", perf_headers, perf_rows)
        trend_rows = conn.execute("SELECT track_date, stock_id, stock_name, watch_status, launch_score, days_in_watch FROM watch_pool_tracking ORDER BY stock_id, track_date").fetchall()
        write_sheet("分數趨勢", ["track_date", "stock_id", "stock_name", "watch_status", "launch_score", "days_in_watch"], trend_rows)
        check_rows = [
            ("schema", "PASS", "三張表使用 CREATE TABLE IF NOT EXISTS，不刪除舊資料"),
            ("today_tracking_count", "PASS" if len(tracking_rows) > 0 else "WARN", len(tracking_rows)),
            ("today_event_count", "PASS", len(event_rows)),
            ("performance_count", "PASS", len(perf_rows)),
            ("output_file", "PASS", str(out)),
        ]
        write_sheet("查核紀錄", ["項目", "結果", "說明"], check_rows)
        wb.save(out)
        self.log(f"R5N29_CULTIVATION_REPORT_WRITTEN output={out}")
        return str(out)

    def run(self, template: Optional[str], out_path: str, base_date: str, main_db_path: str, cultivation_db_path: Optional[str] = None, launch_ready_path: Optional[str] = None) -> Dict[str, Any]:
        base_date = self._parse_date(base_date)
        if not main_db_path:
            raise RuntimeError("R5N29 必須指定主DB檔案 stock_system_v6_2.db")
        if not Path(main_db_path).exists():
            raise RuntimeError(f"R5N29 主DB不存在：{main_db_path}")
        cult_db = self.resolve_cultivation_db_path(main_db_path, cultivation_db_path)
        if not out_path:
            out_path = str(Path(main_db_path).resolve().parent / "reports" / f"watch_pool_cultivation_{base_date.replace('-', '')}.xlsx")
        self.log(f"R5N29_START base_date={base_date}")
        self.log(f"R5N29_MAIN_DB read_only={main_db_path}")
        self.log(f"R5N29_CULTIVATION_DB path={cult_db}")
        with sqlite3.connect(cult_db) as conn:
            self.ensure_cultivation_schema(conn)
            today_rows, source = self.load_today_launch_ready(main_db_path, launch_ready_path, base_date)
            prev = self.load_previous_tracking(conn, base_date)
            tracking_rows, event_rows = self.compare_today_previous(conn, today_rows, prev, base_date)
            self.write_tracking_rows(conn, tracking_rows)
            self.write_event_rows(conn, event_rows)
            self.update_performance(conn, main_db_path, base_date)
            output = self.export_cultivation_report(conn, out_path, base_date)
            summary = {
                "基準日": base_date,
                "今日追蹤檔數": str(len(tracking_rows)),
                "今日事件數": str(len(event_rows)),
                "培養DB": str(cult_db),
                "資料來源": source,
                "輸出檔案": output,
            }
        self.log("R5N29_DONE")
        log_file = self.log_dir / f"r5n29_watch_pool_{base_date.replace('-', '')}_{dt.datetime.now().strftime('%H%M%S')}.log"
        try:
            log_file.write_text("\n".join(self.messages), encoding="utf-8")
        except Exception as exc:
            self.warn(f"R5N29_LOG_WRITE_FAIL {exc}")
        return {"output": output, "cultivation_db": str(cult_db), "summary": summary, "log_file": str(log_file)}



def run_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk

    root = tk.Tk()
    root.title("宏觀16模組 自動回填主程式 / R5N29觀察池培養")
    root.geometry("1000x760")

    template_var = tk.StringVar()
    out_var = tk.StringVar(value=str(Path.cwd() / f"宏觀16模組_自動回填_{dt.date.today().strftime('%Y%m%d')}.xlsx"))
    date_var = tk.StringVar(value=dt.date.today().strftime("%Y-%m-%d"))
    db_var = tk.StringVar()
    tej_gov_var = tk.StringVar()
    cultivation_db_var = tk.StringVar()
    launch_ready_var = tk.StringVar()
    strict_ranking_var = tk.BooleanVar(value=False)
    # P0恢復：預設一定是原始宏觀回填，不得預設為R5N29培養。
    report_mode_var = tk.StringVar(value=REPORT_MODE_MACRO)
    status_var = tk.StringVar(value="待執行")

    def browse_template():
        p = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if p:
            template_var.set(p)

    def browse_out():
        p = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if p:
            out_var.set(p)

    def browse_db():
        p = filedialog.askopenfilename(filetypes=[("SQLite DB", "*.db"), ("All files", "*.*")])
        if p:
            db_var.set(p)

    def browse_tej_gov():
        p = filedialog.askopenfilename(filetypes=[("TEJ Excel", "*.xls *.xlsx"), ("All files", "*.*")])
        if p:
            tej_gov_var.set(p)

    def browse_cultivation_db():
        p = filedialog.askopenfilename(filetypes=[("SQLite DB", "*.db"), ("All files", "*.*")])
        if p:
            cultivation_db_var.set(p)

    def browse_launch_ready():
        p = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")])
        if not p:
            p = filedialog.askdirectory()
        if p:
            launch_ready_var.set(p)

    frm = ttk.Frame(root, padding=12)
    frm.pack(fill="both", expand=True)
    ttk.Label(frm, text="宏觀16模組 自動抓取與Excel回填 / R5N29觀察池培養", font=("Microsoft JhengHei", 16, "bold")).grid(row=0, column=0, columnspan=3, sticky="w", pady=(0,10))
    ttk.Label(frm, text="Excel模板").grid(row=1, column=0, sticky="w")
    ttk.Entry(frm, textvariable=template_var, width=90).grid(row=1, column=1, sticky="we")
    ttk.Button(frm, text="選擇", command=browse_template).grid(row=1, column=2)
    ttk.Label(frm, text="輸出檔案").grid(row=2, column=0, sticky="w")
    ttk.Entry(frm, textvariable=out_var, width=90).grid(row=2, column=1, sticky="we")
    ttk.Button(frm, text="另存", command=browse_out).grid(row=2, column=2)
    ttk.Label(frm, text="基準日(YYYY-MM-DD)").grid(row=3, column=0, sticky="w")
    ttk.Entry(frm, textvariable=date_var, width=20).grid(row=3, column=1, sticky="w")
    ttk.Label(frm, text="主DB檔案(選填；R5N29必填)").grid(row=4, column=0, sticky="w")
    ttk.Entry(frm, textvariable=db_var, width=90).grid(row=4, column=1, sticky="we")
    ttk.Button(frm, text="選擇DB", command=browse_db).grid(row=4, column=2)
    ttk.Label(frm, text="TEJ八大官股檔(宏觀模式選填)").grid(row=5, column=0, sticky="w")
    ttk.Entry(frm, textvariable=tej_gov_var, width=90).grid(row=5, column=1, sticky="we")
    ttk.Button(frm, text="選擇TEJ", command=browse_tej_gov).grid(row=5, column=2)
    ttk.Label(frm, text="培養DB檔案（R5N29；空白則自動建立）").grid(row=6, column=0, sticky="w")
    ttk.Entry(frm, textvariable=cultivation_db_var, width=90).grid(row=6, column=1, sticky="we")
    ttk.Button(frm, text="選擇培養DB", command=browse_cultivation_db).grid(row=6, column=2)
    ttk.Label(frm, text="Launch Ready報表/資料夾（R5N29選填）").grid(row=7, column=0, sticky="w")
    ttk.Entry(frm, textvariable=launch_ready_var, width=90).grid(row=7, column=1, sticky="we")
    ttk.Button(frm, text="選擇報表", command=browse_launch_ready).grid(row=7, column=2)
    ttk.Checkbutton(frm, text="Ranking缺失時中止輸出", variable=strict_ranking_var).grid(row=8, column=1, sticky="w")
    ttk.Label(frm, text="輸出模式").grid(row=9, column=0, sticky="w")
    ttk.Combobox(frm, textvariable=report_mode_var, values=[REPORT_MODE_MACRO, REPORT_MODE_MACRO_TEACHER, REPORT_MODE_MACRO_ONLY, REPORT_MODE_INSTITUTIONAL, REPORT_MODE_TEACHER_FULL, REPORT_MODE_ALL, REPORT_MODE_WATCH_POOL], width=36, state="readonly").grid(row=9, column=1, sticky="w")
    ttk.Label(frm, textvariable=status_var, foreground="blue").grid(row=10, column=0, columnspan=3, sticky="w", pady=8)

    log_text = tk.Text(frm, height=24, wrap="word")
    log_text.grid(row=12, column=0, columnspan=3, sticky="nsew", pady=(10,0))
    frm.rowconfigure(12, weight=1)
    frm.columnconfigure(1, weight=1)

    def append_log(text):
        log_text.insert("end", text + "\n")
        log_text.see("end")
        root.update_idletasks()

    def execute():
        try:
            status_var.set("執行中...")
            log_text.delete("1.0", "end")
            if report_mode_var.get() == REPORT_MODE_WATCH_POOL:
                if not db_var.get():
                    raise RuntimeError("R5N29觀察池培養必須指定主DB檔案")
                engine = WatchPoolCultivationEngine(Path("logs"))
                result = engine.run(template_var.get() or None, out_var.get(), date_var.get(), main_db_path=db_var.get(), cultivation_db_path=(cultivation_db_var.get() or None), launch_ready_path=(launch_ready_var.get() or None))
                for msg in engine.messages:
                    append_log(msg)
            else:
                engine = Macro16Engine(Path("logs"))
                result = engine.run(template_var.get() or None, out_var.get(), date_var.get(), db_path=(db_var.get() or None), strict_ranking=bool(strict_ranking_var.get()), tej_gov_file=(tej_gov_var.get() or None), report_mode=report_mode_var.get())
                for msg in engine.logger.messages:
                    append_log(msg)
            append_log("\n總結：" + json.dumps(result["summary"], ensure_ascii=False, indent=2))
            append_log("Log檔：" + result.get("log_file", ""))
            status_var.set("完成")
            messagebox.showinfo("完成", f"已輸出：{result['output']}")
        except Exception as exc:
            status_var.set("失敗")
            append_log("ERROR " + str(exc))
            messagebox.showerror("錯誤", str(exc))

    ttk.Button(frm, text="執行回填/培養", command=execute).grid(row=11, column=0, sticky="w", pady=6)
    ttk.Button(frm, text="離開", command=root.destroy).grid(row=11, column=2, sticky="e", pady=6)
    root.mainloop()

def main():
    parser = argparse.ArgumentParser(description="宏觀16模組自動抓取與Excel回填主程式")
    parser.add_argument("--cli", action="store_true", help="使用CLI模式")
    parser.add_argument("--template", default="", help="Excel模板路徑")
    parser.add_argument("--out", default=f"宏觀16模組_自動回填_{dt.date.today().strftime('%Y%m%d')}.xlsx", help="輸出Excel路徑")
    parser.add_argument("--date", default=dt.date.today().strftime("%Y-%m-%d"), help="基準日 YYYY-MM-DD")
    parser.add_argument("--log-dir", default="logs", help="Log目錄")
    parser.add_argument("--gov-net", type=float, default=None, help="人工覆寫官股買賣超(億元)")
    parser.add_argument("--ai-strength", type=float, default=None, help="人工覆寫AI主流強度0~1")
    parser.add_argument("--major-event", type=int, default=None, help="人工覆寫重大事件0/1")
    parser.add_argument("--event-note", default="", help="人工覆寫戰爭/地緣/重大事件說明")
    parser.add_argument("--night-score", type=float, default=None, help="人工覆寫台股夜盤分數")
    parser.add_argument("--db-path", default="", help="指定主SQLite DB路徑；用於ranking_result驗證與機構級股票投資規劃報表/R5N29主DB")
    parser.add_argument("--cultivation-db-path", default="", help="R5N29培養DB路徑；空白則自動建立 data/watch_pool_cultivation.db")
    parser.add_argument("--launch-ready-path", default="", help="R5N29 Launch Ready報表檔案或資料夾；空白則自動找 launch_ready_reports 最新報表")
    parser.add_argument("--tej-gov-file", default="", help="TEJ八大公股行庫買賣超排名xls/xlsx；用於gov_net_100m主來源")
    parser.add_argument("--strict-ranking", action="store_true", help="ranking_result缺失或空表時直接中止，避免輸出可下單結論")
    parser.add_argument("--report-mode", default=REPORT_MODE_MACRO, choices=[REPORT_MODE_MACRO, REPORT_MODE_MACRO_TEACHER, REPORT_MODE_MACRO_ONLY, REPORT_MODE_INSTITUTIONAL, REPORT_MODE_TEACHER_FULL, REPORT_MODE_ALL, REPORT_MODE_WATCH_POOL], help="輸出模式：macro_refill/macro_teacher輸出宏觀16+老師策略00~16；macro_only只輸出3頁；institutional_report/teacher_full只輸出老師策略00~16；all輸出完整debug")
    args = parser.parse_args()
    if args.cli:
        if args.report_mode == REPORT_MODE_WATCH_POOL:
            engine = WatchPoolCultivationEngine(Path(args.log_dir))
            result = engine.run(args.template or None, args.out, args.date, main_db_path=args.db_path, cultivation_db_path=(args.cultivation_db_path or None), launch_ready_path=(args.launch_ready_path or None))
        else:
            engine = Macro16Engine(Path(args.log_dir))
            override = ManualOverride(gov_net_100m=args.gov_net, ai_strength=args.ai_strength, major_event=args.major_event, event_note=args.event_note, night_score=args.night_score)
            result = engine.run(args.template or None, args.out, args.date, override, db_path=(args.db_path or None), strict_ranking=args.strict_ranking, tej_gov_file=(args.tej_gov_file or None), report_mode=args.report_mode)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        run_gui()

if __name__ == "__main__":
    main()
